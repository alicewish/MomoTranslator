import io
import os
import sys
from locale import getdefaultlocale
from pathlib import Path
import re
from shutil import copy2
from subprocess import Popen
import webbrowser

import cv2
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Inches
from itertools import chain
from loguru import logger
from natsort import natsorted
import numpy as np
from PIL import Image
from PyQt6.QtCore import QSettings, QSize, Qt, QTranslator, QCoreApplication
from PyQt6.QtGui import QAction, QActionGroup, QBrush, QDoubleValidator, QFont, QGuiApplication, QIcon, QImage, \
    QKeySequence, QPainter, QPixmap, QTransform
from PyQt6.QtWidgets import QAbstractItemView, QApplication, QButtonGroup, QComboBox, QDialog, QDockWidget, QFileDialog, \
    QGraphicsPixmapItem, QGraphicsScene, QGraphicsView, QHBoxLayout, QLabel, QLineEdit, QListView, QListWidget, \
    QListWidgetItem, QMainWindow, QMenu, QMessageBox, QProgressBar, QPushButton, QRadioButton, QStatusBar, QTabWidget, \
    QToolBar, QToolButton, QVBoxLayout, QWidget
from qtawesome import icon
import yaml

from storage import GlobalStorage
from config import GlobalConfig
from helper import timer_decorator

from core.color import Color, ColorGradient, ColorDouble, get_color_name
from core.bubble import get_colorful_bubbles, get_raw_bubbles, seg_bubbles
from core.frame import get_frame_grid_recursive, get_marked_frames, compute_frame_mask
from core.shape import Contr, PicGrid, Rect
from core.image import collect_valid_images_from, write_image, get_edge_pixels_rgba, find_dominant_color

from ioutils.f import make_dir

from OCR.Tesseract import Tesseract
from OCR.Vision import Vision

from core.torch import get_comictextdetector_mask

from matplotlib import colormaps
# 使用matplotlib的tab20颜色映射
colormap_tab20 = colormaps['tab20']

media_type_dict = {
    0: 'Comic',
    1: 'Manga',
}

media_language_dict = {
    0: 'English',
    1: 'Chinese Simplified',
    2: 'Chinese Traditional',
    3: 'Japanese',
    4: 'Korean',
}

color_patterns = [
        ['ffffff-15', '000000-60'],  # 白底黑字
    ]

image_format = 'jpeg'
max_chars = 5000

vertical = True

padding = 10

def capitalize_sentence(sentence: str) -> str:
    """
    首字母大写处理句子，同时考虑缩写。

    :param sentence: 需要处理的句子。
    :return: 首字母大写处理后的句子。
    """
    abbreviations = ['Dr.', 'Ms.', 'Mr.', 'Mrs.']
    for abbr in abbreviations:
        sentence = sentence.replace(abbr, abbr.lower())
    capitalized_sentence = sentence[0].upper() + sentence[1:].lower()
    for abbr in abbreviations:
        capitalized_sentence = capitalized_sentence.replace(abbr.lower(), abbr)
    return capitalized_sentence

def get_single_cnts(image_file, image_raw, all_masks):
    ih, iw = image_raw.shape[0:2]
    black_background = np.zeros((ih, iw), dtype=np.uint8)

    mask_pics = [x for x in all_masks if x.stem.startswith(image_file.stem)]
    single_cnts = []
    # ================针对每一张图================
    for m in range(len(mask_pics)):
        mask_pic = mask_pics[m]
        logger.warning(f'{mask_pic=}')

        # ================获取轮廓================
        transparent_image = cv2.imdecode(np.fromfile(mask_pic, dtype=np.uint8), -1)
        # 将半透明的像素变为透明
        transparent_image[transparent_image[..., 3] < 255] = 0
        # 获取所有完全不透明的像素的颜色
        non_transparent_pixels = transparent_image[transparent_image[..., 3] == 255, :3]
        # 获取所有不透明像素的唯一颜色
        unique_colors = np.unique(non_transparent_pixels, axis=0)
        # 对于每种颜色，提取对应的contour
        contour_list = []
        for color in unique_colors:
            mask = np.all(transparent_image[..., :3] == color, axis=-1)
            mask = (mask & (transparent_image[..., 3] == 255)).astype(np.uint8) * 255
            contours, _ = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            contour_list.extend(contours)

        # ================遍历每种气泡文字颜色组合================
        cp_str = mask_pic.stem.removeprefix(f'{image_file.stem}-Mask-')
        # ffffff-15-white~000000-60-black
        a_str, partition, b_str = cp_str.partition('~')
        a_parts = a_str.split('-')
        b_parts = b_str.split('-')
        color_pattern = [
            f'{a_parts[0]}-{a_parts[1]}',
            f'{b_parts[0]}-{b_parts[1]}',
        ]
        # print(f'{color_pattern=}')
        cp_bubble, cp_letter = color_pattern

        if isinstance(cp_bubble, list):
            # 气泡为渐变色
            color_bubble = ColorGradient(cp_bubble)
        else:
            # 气泡为单色
            color_bubble = Color(cp_bubble)

        if isinstance(cp_letter, list):
            # 文字为双色
            color_letter = ColorDouble(cp_letter)
            color_letter_big = color_letter
        else:
            # 文字为单
            color_letter = Color(cp_letter)
            color_letter_big = Color(f'{cp_letter[:6]}-{max(color_letter.padding + 120, 180)}')

        # ================获取对应的气泡文字蒙版================
        if color_bubble.type == 'Color':
            bubble_mask = color_bubble.get_img(image_raw)
        else:  # ColorGradient
            img_tuple = color_bubble.get_img(image_raw)
            bubble_mask, left_sample, right_sample = img_tuple
        letter_mask = color_letter_big.get_img(image_raw)

        for c in range(len(contour_list)):
            contour = contour_list[c]
            single_cnt = Contr(contour)

            # 在黑色背景上绘制单个轮廓的白色填充图像
            bit_white_bubble = cv2.drawContours(black_background.copy(), [single_cnt.contour], 0, 255, -1)

            # 获取原始图像在bit_white_bubble范围内的图像，其他部分为白色
            image_bubble_only = cv2.bitwise_and(image_raw, image_raw, mask=bit_white_bubble)
            image_bubble_only[bit_white_bubble == 0] = (255, 255, 255)

            # 通过将字符掩码与轮廓掩码相乘，得到只包含轮廓内部字符的图像
            letter_in_bubble = cv2.bitwise_and(letter_mask, letter_mask, mask=bit_white_bubble)
            # 获取所有非零像素的坐标
            px_pts = np.transpose(np.nonzero(letter_in_bubble))
            # 计算非零像素的数量
            px_area = px_pts.shape[0]
            # 获取所有非零像素的x和y坐标
            all_x = px_pts[:, 1]
            all_y = px_pts[:, 0]

            # 计算最小和最大的x、y坐标，并添加指定的padding
            min_x, max_x = np.min(all_x) - padding, np.max(all_x) + padding
            min_y, max_y = np.min(all_y) - padding, np.max(all_y) + padding

            # 限制padding后的坐标范围不超过原始图像的边界
            min_x, min_y = max(min_x, 0), max(min_y, 0)
            max_x, max_y = min(max_x, image_raw.shape[1]), min(max_y, image_raw.shape[0])

            # 从带有白色背景的图像中裁剪出带有padding的矩形区域
            if color_bubble.rgb_str == 'ffffff':
                cropped_image = image_bubble_only[min_y:max_y, min_x:max_x]
            else:
                cropped_image = letter_in_bubble[min_y:max_y, min_x:max_x]
            single_cnt.add_cropped_image(cropped_image)
            single_cnts.append(single_cnt)
    return single_cnts


class GUI(QMainWindow):
    def __init__(self):
        super().__init__()

        self._load_settings()

        # ui translator
        self._tr = QTranslator(self)
        QCoreApplication.instance().installTranslator(self._tr)
    
        self.a0_para()
        self.a1_initialize()
        self.a2_status_bar()
        self.a3_docks()
        self.a4_menubar()
        self.a5_toolbar()
        self.a9_setting()

        self.update_gui_language(self._settings.value('lang') or 'en_US')

    def _load_settings(self):
        self._settings = QSettings(GlobalStorage.GUI_SETTINGS.as_posix(), QSettings.IniFormat)
        # TODO load other settings

    def _save_settings(self):
        # TODO
        pass

    @staticmethod
    def collect_available_languages():
        available_languages = []

        file_names = os.listdir(GlobalStorage.GUI_TRANSLATION_FOLDER)
        for file_name in file_names:
            # TODO to fix hardcode
            result = re.match(r'MomoTranslator_(.*).qm$', file_name)
            if result:
                available_languages.append(result.group(1))

        return available_languages

    def update_gui_language(self, lang):
        self._settings.setValue('lang', lang)
        self.translate_widget_text()
        self.render_widget_text()
    
    def translate_widget_text(self):
        assert(self._settings)
        lang = self._settings.value('lang')
        qm = GlobalStorage.GUI_TRANSLATION_FOLDER / f'{GlobalConfig.APP_NAME}_{lang}.qm'
        if self._tr.load(qm.as_posix()):
            pass
        else:
            # TODO throw error
            pass

    def render_widget_text(self):
        self.file_menu.setTitle(self.tr('File'))
        self.open_folder_action.setText(self.tr('Open Folder'))
        self.save_image_action.setText(self.tr('Save Image'))
        self.recent_folders_menu.setTitle(self.tr('Recent Folders'))

        self.view_menu.setTitle(self.tr('View'))
        self.zoom_in_action.setText(self.tr('Zoom In'))
        self.zoom_out_action.setText(self.tr('Zoom Out'))
        self.fit_to_screen_action.setText(self.tr('Fit to Screen'))
        self.fit_to_width_action.setText(self.tr('Fit to Width'))
        self.fit_to_height_action.setText(self.tr('Fit to Height'))
        self.reset_zoom_action.setText(self.tr('Reset Zoom'))

        self.edit_menu.setTitle(self.tr('Edit'))

        self.nav_menu.setTitle(self.tr('Navigate'))
        self.prev_image_action.setText(self.tr('Previous Image'))
        self.next_image_action.setText(self.tr('Next Image'))
        self.first_image_action.setText(self.tr('First Image'))
        self.last_image_action.setText(self.tr('Last Image'))

        self.help_menu.setTitle(self.tr('Help'))
        self.about_action.setText(self.tr('About'))
        self.about_qt_action.setText(self.tr('About'))
        self.help_document_action.setText(f"{GlobalConfig.APP_NAME} {self.tr('Help')}")
        self.feedback_action.setText(self.tr('Bug Report'))
        self.update_action.setText(self.tr('Update Online'))

        while self.hb_step.count():
            child = self.hb_step.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
            else:
                del child

        self.pb_step = QPushButton(self.tr('Start'))
        self.pb_step.clicked.connect(self.start_task)
        self.pb_step.setText(self.tr('Start'))
        self.hb_step.addWidget(self.pb_step)

        self.task_names = [
            self.tr('Analyze Frames'),
            self.tr('Analyze Bubbles'),
            self.tr('OCR'),
            self.tr('Translate'),
            self.tr('Lettering'),
        ]

        for button in self.bg_step.buttons():
            self.bg_step.removeButton(button)
            button.deleteLater()
        
        for index, option in enumerate(self.task_names):
            self.rb_task = QRadioButton(option)
            self.hb_step.addWidget(self.rb_task)
            self.bg_step.addButton(self.rb_task, index)
            # 默认选中第一个 QRadioButton
            if index == 0:
                self.rb_task.setChecked(True)

        self.pb_task = QProgressBar()
        self.hb_step.addWidget(self.pb_task)

        self.pics_dock.setWindowTitle(self.tr('Image List'))
        self.setting_dock.setWindowTitle(self.tr('Setting'))

        while self.hb_media_type.count():
            child = self.hb_media_type.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
            else:
                del child
        
        self.lb_media_type = QLabel(self.tr('Layout'))
        self.lb_media_type.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.lb_media_type.setText(self.tr('Layout'))
        self.hb_media_type.addWidget(self.lb_media_type)

        self.media_type_names = [
            self.tr('Comic'),
            self.tr('Manga'),
        ]

        for button in self.bg_media_type.buttons():
            self.bg_media_type.removeButton(button)
            button.deleteLater()

        for index, option in enumerate(self.media_type_names):
            self.rb_media_type = QRadioButton(option)
            self.hb_media_type.addWidget(self.rb_media_type)
            self.bg_media_type.addButton(self.rb_media_type, index)
            if index == self.media_type_index:
                self.rb_media_type.setChecked(True)

        self.hb_media_type.addStretch(1)

        self.text_dock.setWindowTitle(self.tr('Text'))
        self.layer_dock.setWindowTitle(self.tr('Layer'))
        self.step_dock.setWindowTitle(self.tr('Step'))


        self.nav_tab.setTabText(0, self.tr('Thumbnails'))


    def a0_para(self):
        # ================初始化变量================
        self.screen_icon = icon('ei.screen')
        self.setWindowIcon(self.screen_icon)
        self.default_palette = QGuiApplication.palette()
        self.setAcceptDrops(True)

        self.setWindowTitle(f'{GlobalConfig.APP_NAME} {GlobalConfig.APP_VERSION}')
        self.resize(self._settings.value('size', QSize(1080, 768)))

        # 创建 QGraphicsScene 和 QGraphicsView 对象
        self.graphics_scene = QGraphicsScene(self)  # 图形场景对象
        self.graphics_view = QGraphicsView(self)  # 图形视图对象
        # 设置渲染、优化和视口更新模式
        # 设置渲染抗锯齿
        self.graphics_view.setRenderHint(QPainter.RenderHint.Antialiasing)
        # 设置优化标志以便不为抗锯齿调整
        self.graphics_view.setOptimizationFlag(QGraphicsView.OptimizationFlag.DontAdjustForAntialiasing, True)
        # 设置优化标志以便不保存画家状态
        self.graphics_view.setOptimizationFlag(QGraphicsView.OptimizationFlag.DontSavePainterState, True)
        # 设置视口更新模式为完整视口更新
        self.graphics_view.setViewportUpdateMode(QGraphicsView.ViewportUpdateMode.FullViewportUpdate)
        # 设置渲染提示为平滑图像变换，以提高图像的显示质量
        self.graphics_view.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)

    def a1_initialize(self):
        # ================图片列表================
        self.image_folder = None
        self.image_list = []
        self.image_file = None
        self.image_item = None
        self.image = QImage()
        self.pixmap_item = QGraphicsPixmapItem()
        self.image_index = -1
        self.filtered_image_index = -1
        self.view_mode = 0
        self.screen_scaling_factor = self.get_screen_scaling_factor()
        self.screen_scaling_factor_reciprocal = 1 / self.screen_scaling_factor
        # ================最近文件夹================
        self.recent_folders = []
        # ================设置================
        # 根据需要修改组织和应用名
        self.media_type_index = int(self._settings.value('media_type_index', 0))
        self.media_type = media_type_dict[self.media_type_index]
        self.media_language_index = int(self._settings.value('media_language_index', 0))
        self.media_language = media_language_dict[self.media_language_index]

    def a2_status_bar(self):
        # ================状态栏================
        self.status_bar = QStatusBar()
        # 设置状态栏，类似布局设置
        self.setStatusBar(self.status_bar)

    def a3_docks(self):
        # ================缩略图列表================
        self.thumbnails_widget = QListWidget(self)
        self.thumbnails_widget.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)  # 设置自定义右键菜单
        self.thumbnails_widget.customContextMenuRequested.connect(self.show_image_list_context_menu)  # 连接自定义右键菜单信号与函数
        self.thumbnails_widget.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)  # 设置单选模式
        self.thumbnails_widget.setViewMode(QListView.ViewMode.IconMode)  # 设置图标模式
        self.thumbnails_widget.setFlow(QListView.Flow.TopToBottom)  # 设置从上到下排列
        self.thumbnails_widget.setResizeMode(QListView.ResizeMode.Adjust)  # 设置自动调整大小
        self.thumbnails_widget.setWrapping(False)  # 关闭自动换行
        self.thumbnails_widget.setWordWrap(True)  # 开启单词换行
        self.thumbnails_widget.setIconSize(QSize(240, 240))  # 设置图标大小
        self.thumbnails_widget.setSpacing(5)  # 设置间距
        self.thumbnails_widget.itemSelectionChanged.connect(self.select_image_from_list)  # 连接图标选择信号与函数

        self.nav_tab = QTabWidget()  # 创建选项卡控件
        self.nav_tab.addTab(self.thumbnails_widget, self.tr('Thumbnails'))  # 添加缩略图列表到选项卡

        self.case_sensitive_button = QToolButton()  # 创建区分大小写按钮
        self.case_sensitive_button.setIcon(icon('msc.case-sensitive'))  # 设置图标
        self.case_sensitive_button.setCheckable(True)  # 设置可选中
        self.case_sensitive_button.setText(self.tr('Case Sensitive'))  # 设置文本
        self.case_sensitive_button.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonIconOnly)  # 设置仅显示图标

        self.whole_word_button = QToolButton()  # 创建全词匹配按钮
        self.whole_word_button.setIcon(icon('msc.whole-word'))  # 设置图标
        self.whole_word_button.setCheckable(True)  # 设置可选中
        self.whole_word_button.setText(self.tr('Whole Word'))  # 设置文本
        self.whole_word_button.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonIconOnly)  # 设置仅显示图标

        self.regex_button = QToolButton()  # 创建正则表达式按钮
        self.regex_button.setIcon(icon('mdi.regex'))  # 设置图标
        self.regex_button.setCheckable(True)  # 设置可选中
        self.regex_button.setText(self.tr('Use Regex'))  # 设置文本
        self.regex_button.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonIconOnly)  # 设置仅显示图标

        # 连接按钮的点击事件与搜索结果的刷新
        self.case_sensitive_button.clicked.connect(self.refresh_search_results)
        self.whole_word_button.clicked.connect(self.refresh_search_results)
        self.regex_button.clicked.connect(self.refresh_search_results)

        self.hb_search_bar = QHBoxLayout()  # 创建水平布局，用于将三个按钮放在一行
        self.hb_search_bar.setContentsMargins(0, 0, 0, 0)
        self.hb_search_bar.addStretch()
        self.hb_search_bar.addWidget(self.case_sensitive_button)
        self.hb_search_bar.addWidget(self.whole_word_button)
        self.hb_search_bar.addWidget(self.regex_button)

        self.search_bar = QLineEdit()  # 创建搜索栏
        self.search_bar.setPlaceholderText(self.tr('Search'))  # 设置占位符文本
        self.search_bar.textChanged.connect(self.filter_image_list)  # 连接文本改变信号与函数
        self.search_bar.setLayout(self.hb_search_bar)  # 将按钮添加到 QLineEdit 的右侧

        self.vb_search_nav = QVBoxLayout()  # 创建垂直布局，用于将搜索框和图片列表放在一列
        self.vb_search_nav.addWidget(self.search_bar)
        self.vb_search_nav.addWidget(self.nav_tab)

        self.pics_widget = QWidget()  # 创建 QWidget，用于容纳搜索框和图片列表
        self.pics_widget.setLayout(self.vb_search_nav)
        
        self.pics_dock = QDockWidget(self.tr('Image List'), self)  # 创建 QDockWidget，用于显示图片列表
        self.pics_dock.setObjectName('pics_dock')
        self.pics_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        self.pics_dock.setWidget(self.pics_widget)
        self.addDockWidget(Qt.DockWidgetArea.LeftDockWidgetArea, self.pics_dock)

        # ================创建设置工具================
        self.lb_setting = QLabel('翻译设置')
        self.lb_setting.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)

        self.hb_media_type = QHBoxLayout()
        self.lb_media_type = QLabel(self.tr('Layout'))
        self.lb_media_type.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.hb_media_type.addWidget(self.lb_media_type)

        # 创建选项组
        self.media_type_names = [
            self.tr('Comic'),
            self.tr('Manga'),
        ]
        self.bg_media_type = QButtonGroup()
        self.bg_media_type.buttonClicked.connect(self.update_media_type)
        for index, option in enumerate(self.media_type_names):
            self.rb_media_type = QRadioButton(option)
            self.hb_media_type.addWidget(self.rb_media_type)
            self.bg_media_type.addButton(self.rb_media_type, index)
            if index == self.media_type_index:
                self.rb_media_type.setChecked(True)

        self.hb_media_type.addStretch(1)

        # 创建语言选择下拉框
        self.cb_media_language = QComboBox()
        self.cb_media_language.addItems([
            self.tr("English"),
            self.tr("Chinese Simplified"),
            self.tr("Chinese Traditional"),
            self.tr("Japanese"),
            self.tr("Korean"),
        ])
        self.cb_media_language.setCurrentIndex(self.media_language_index)
        self.cb_media_language.currentIndexChanged.connect(self.update_media_language)

        # 在布局中添加下拉框
        self.hb_media_language = QHBoxLayout()
        self.lb_media_language = QLabel(self.tr('Language'))
        self.lb_media_language.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.hb_media_language.addWidget(self.lb_media_language)
        self.hb_media_language.addWidget(self.cb_media_language)
        self.hb_media_language.addStretch(1)

        self.vb_setting = QVBoxLayout()
        self.vb_setting.addWidget(self.lb_setting)
        self.vb_setting.addLayout(self.hb_media_type)
        self.vb_setting.addLayout(self.hb_media_language)
        self.vb_setting.addStretch(1)
        self.setting_tool = QWidget()
        self.setting_tool.setLayout(self.vb_setting)

        # ================创建文本工具================
        self.lb_text = QLabel('为LabelPlus提供兼容')
        self.lb_text.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.vb_text = QVBoxLayout()
        self.vb_text.addWidget(self.lb_text)
        self.vb_text.addStretch(1)
        self.text_tool = QWidget()
        self.text_tool.setLayout(self.vb_text)

        # ================创建图层工具================
        self.lb_layer = QLabel('为PS提供兼容')
        self.lb_layer.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.vb_layer = QVBoxLayout()
        self.vb_layer.addWidget(self.lb_layer)
        self.vb_layer.addStretch(1)
        self.layer_tool = QWidget()
        self.layer_tool.setLayout(self.vb_layer)

        # ================创建步骤工具================
        self.hb_step = QHBoxLayout()
        self.pb_step = QPushButton(self.tr('Start'))
        self.pb_step.clicked.connect(self.start_task)
        self.hb_step.addWidget(self.pb_step)

        # 创建选项组
        self.task_names = [
            self.tr('Analyze Frames'),
            self.tr('Analyze Bubbles'),
            self.tr('OCR'),
            self.tr('Translate'),
            self.tr('Lettering'),
        ]
        self.bg_step = QButtonGroup()
        for index, option in enumerate(self.task_names):
            self.rb_task = QRadioButton(option)
            self.hb_step.addWidget(self.rb_task)
            self.bg_step.addButton(self.rb_task, index)
            # 默认选中第一个 QRadioButton
            if index == 0:
                self.rb_task.setChecked(True)

        # 创建进度条
        self.pb_task = QProgressBar()
        self.hb_step.addWidget(self.pb_task)
        self.step_tool = QWidget()
        self.step_tool.setLayout(self.hb_step)

        self.setting_dock = QDockWidget(self.tr('Setting'), self)
        self.setting_dock.setObjectName("SettingDock")
        self.setting_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        self.setting_dock.setWidget(self.setting_tool)
        self.setting_dock.setMinimumWidth(200)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.setting_dock)

        self.text_dock = QDockWidget(self.tr('Text'), self)
        self.text_dock.setObjectName("TextDock")
        self.text_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        self.text_dock.setWidget(self.text_tool)
        self.text_dock.setMinimumWidth(200)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.text_dock)

        self.layer_dock = QDockWidget(self.tr('Layer'), self)
        self.layer_dock.setObjectName("LayerDock")
        self.layer_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        self.layer_dock.setWidget(self.layer_tool)
        self.layer_dock.setMinimumWidth(200)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.layer_dock)

        self.step_dock = QDockWidget(self.tr('Step'), self)
        self.step_dock.setObjectName("StepDock")
        self.step_dock.setAllowedAreas(Qt.DockWidgetArea.BottomDockWidgetArea | Qt.DockWidgetArea.TopDockWidgetArea)
        self.step_dock.setWidget(self.step_tool)
        self.step_dock.setMinimumWidth(200)
        self.addDockWidget(Qt.DockWidgetArea.BottomDockWidgetArea, self.step_dock)

        # 将 QGraphicsView 设置为中心窗口部件
        self.setCentralWidget(self.graphics_view)

        self.resizeDocks([self.pics_dock, self.setting_dock], [280, 250], Qt.Horizontal)


    def a4_menubar(self):
        # 文件菜单
        self.file_menu = self.menuBar().addMenu(self.tr('File'))

        # 添加打开文件夹操作
        self.open_folder_action = QAction(self.tr('Open Folder'), self)
        self.open_folder_action.setIcon(icon('ei.folder'))
        self.open_folder_action.setShortcut(QKeySequence("Ctrl+O"))
        self.open_folder_action.triggered.connect(self.open_folder_by_dialog)
        self.file_menu.addAction(self.open_folder_action)

        # 添加保存图片操作
        self.save_image_action = QAction(self.tr('Save Image'), self)
        self.save_image_action.setIcon(icon('ri.image-edit-fill'))
        self.save_image_action.setShortcut(QKeySequence.StandardKey.Save)
        self.save_image_action.triggered.connect(self.save_image)
        self.file_menu.addAction(self.save_image_action)

        self.file_menu.addSeparator()

        # 添加最近打开文件夹菜单项
        self.recent_folders_menu = self.file_menu.addMenu(self.tr('Recent Folders'))
        self.update_recent_folders_menu()

        # 显示菜单
        self.view_menu = self.menuBar().addMenu(self.tr('View'))

        # 添加视图菜单选项
        self.view_menu.addAction(self.text_dock.toggleViewAction())
        self.view_menu.addAction(self.layer_dock.toggleViewAction())
        self.view_menu.addAction(self.pics_dock.toggleViewAction())
        self.view_menu.addSeparator()

        # 添加缩放选项
        self.zoom_in_action = QAction(self.tr('Zoom In'), self)
        self.zoom_in_action.setIcon(icon('ei.zoom-in'))
        self.zoom_in_action.setShortcut(QKeySequence.StandardKey.ZoomIn)
        self.zoom_in_action.triggered.connect(self.zoom_in)
        self.view_menu.addAction(self.zoom_in_action)

        self.zoom_out_action = QAction(self.tr('Zoom Out'), self)
        self.zoom_out_action.setIcon(icon('ei.zoom-out'))
        self.zoom_out_action.setShortcut(QKeySequence.StandardKey.ZoomOut)
        self.zoom_out_action.triggered.connect(self.zoom_out)
        self.view_menu.addAction(self.zoom_out_action)

        self.view_menu.addSeparator()

        # 添加自适应屏幕选项
        self.fit_to_screen_action = QAction(self.tr('Fit to Screen'), self)
        self.fit_to_screen_action.setIcon(icon('mdi6.fit-to-screen-outline'))
        self.fit_to_screen_action.setShortcut(QKeySequence("Ctrl+F"))
        self.fit_to_screen_action.setCheckable(True)  # 将选项设置为可选
        self.fit_to_screen_action.toggled.connect(self.fit_to_view_toggled)
        self.view_menu.addAction(self.fit_to_screen_action)

        # 添加自适应宽度选项
        self.fit_to_width_action = QAction(self.tr('Fit to Width'), self)
        self.fit_to_width_action.setIcon(icon('ei.resize-horizontal'))
        self.fit_to_width_action.setShortcut(QKeySequence("Ctrl+W"))
        self.fit_to_width_action.setCheckable(True)  # 将选项设置为可选
        self.fit_to_width_action.toggled.connect(self.fit_to_view_toggled)
        self.view_menu.addAction(self.fit_to_width_action)

        # 添加自适应高度选项
        self.fit_to_height_action = QAction(self.tr('Fit to Height'), self)
        self.fit_to_height_action.setIcon(icon('ei.resize-vertical'))
        self.fit_to_height_action.setShortcut(QKeySequence("Ctrl+H"))
        self.fit_to_height_action.setCheckable(True)  # 设置选项为可选
        self.fit_to_height_action.toggled.connect(self.fit_to_view_toggled)
        self.view_menu.addAction(self.fit_to_height_action)

        # 添加重置缩放选项
        self.reset_zoom_action = QAction(self.tr('Reset Zoom'), self)
        self.reset_zoom_action.setIcon(icon('mdi6.backup-restore'))
        self.reset_zoom_action.setShortcut(QKeySequence("Ctrl+0"))
        self.reset_zoom_action.triggered.connect(self.reset_zoom)
        self.view_menu.addAction(self.reset_zoom_action)

        self.view_menu.addSeparator()

        # 添加显示选项
        self.display_modes = [(self.tr('Show Thumbnails'), 0),
                              (self.tr('Show Filenames'), 1),
                              (self.tr('Show Both'), 2)]
        self.display_mode_group = QActionGroup(self)
        for display_mode in self.display_modes:
            action = QAction(display_mode[0], self, checkable=True)
            action.triggered.connect(lambda _, mode=display_mode[1]: self.set_display_mode(mode))
            self.view_menu.addAction(action)
            self.display_mode_group.addAction(action)

        self.display_mode_group.actions()[2].setChecked(True)  # 默认选中 Show Both 选项

        # 编辑菜单
        self.edit_menu = self.menuBar().addMenu(self.tr('Edit'))

        # 导航菜单
        self.nav_menu = self.menuBar().addMenu(self.tr('Navigate'))

        # 添加上一张图片操作
        self.prev_image_action = QAction(self.tr('Previous Image'), self)
        self.prev_image_action.setIcon(icon('ei.arrow-left'))
        self.prev_image_action.setShortcut(QKeySequence("Ctrl+Left"))
        self.prev_image_action.triggered.connect(lambda: self.change_image(-1))
        self.nav_menu.addAction(self.prev_image_action)

        # 添加下一张图片操作
        self.next_image_action = QAction(self.tr('Next Image'), self)
        self.next_image_action.setIcon(icon('ei.arrow-right'))
        self.next_image_action.setShortcut(QKeySequence("Ctrl+Right"))
        self.next_image_action.triggered.connect(lambda: self.change_image(1))
        self.nav_menu.addAction(self.next_image_action)

        # 添加第一张图片操作
        self.first_image_action = QAction(self.tr('First Image'), self)
        self.first_image_action.setIcon(icon('ei.step-backward'))
        self.first_image_action.setShortcut(QKeySequence("Ctrl+Home"))
        self.first_image_action.triggered.connect(lambda: self.change_image("first"))
        self.nav_menu.addAction(self.first_image_action)

        # 添加最后一张图片操作
        self.last_image_action = QAction(self.tr('Last Image'), self)
        self.last_image_action.setIcon(icon('ei.step-forward'))
        self.last_image_action.setShortcut(QKeySequence("Ctrl+End"))
        self.last_image_action.triggered.connect(lambda: self.change_image("last"))
        self.nav_menu.addAction(self.last_image_action)

        # Configuration
        self.configuration_menu = self.menuBar().addMenu(self.tr('Configuration'))
        self.gui_languages_menu = self.configuration_menu.addMenu(self.tr('Languages'))
        available_languages = GUI.collect_available_languages()
        language_action_group = QActionGroup(self)
        for lang in available_languages:
            language_action = QAction(lang, self, checkable=True)
            if lang == self._settings.value('lang'):
                language_action.setChecked(True)
            language_action.triggered.connect(lambda _, l=lang: self.update_gui_language(l))
            language_action_group.addAction(language_action)
        self.gui_languages_menu.addActions(language_action_group.actions())


        # 帮助菜单
        self.help_menu = self.menuBar().addMenu(self.tr('Help'))

        self.about_action = QAction(f"{self.tr('About')} {GlobalConfig.APP_NAME}", self)
        self.about_action.triggered.connect(self.show_about_dialog)
        self.help_menu.addAction(self.about_action)

        self.about_qt_action = QAction(f"{self.tr('About')} Qt", self)
        self.about_qt_action.triggered.connect(QApplication.instance().aboutQt)
        self.help_menu.addAction(self.about_qt_action)

        self.help_document_action = QAction(f"{GlobalConfig.APP_NAME} {self.tr('Help')}", self)
        self.help_document_action.triggered.connect(self.show_help_document)
        self.help_menu.addAction(self.help_document_action)

        self.feedback_action = QAction(self.tr('Bug Report'), self)
        self.feedback_action.triggered.connect(self.show_feedback_dialog)
        self.help_menu.addAction(self.feedback_action)

        self.update_action = QAction(self.tr('Update Online'), self)
        self.update_action.triggered.connect(self.check_for_updates)
        self.help_menu.addAction(self.update_action)

    def a5_toolbar(self):
        # 添加顶部工具栏
        self.tool_bar = QToolBar(self.tr('Toolbar'), self)
        self.tool_bar.setObjectName("Toolbar")
        self.tool_bar.setIconSize(QSize(24, 24))
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea, self.tool_bar)
        self.tool_bar.setMovable(False)
        self.tool_bar.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextUnderIcon)
        self.tool_bar.addAction(self.open_folder_action)
        self.tool_bar.addAction(self.save_image_action)
        self.tool_bar.addSeparator()
        self.tool_bar.addAction(self.zoom_in_action)
        self.tool_bar.addAction(self.zoom_out_action)
        self.tool_bar.addAction(self.fit_to_screen_action)
        self.tool_bar.addAction(self.fit_to_width_action)
        self.tool_bar.addAction(self.fit_to_height_action)
        self.tool_bar.addAction(self.reset_zoom_action)
        self.tool_bar.addSeparator()
        self.tool_bar.addAction(self.first_image_action)
        self.tool_bar.addAction(self.prev_image_action)
        self.tool_bar.addAction(self.next_image_action)
        self.tool_bar.addAction(self.last_image_action)

        # 添加缩放百分比输入框
        self.scale_percentage_edit = QLineEdit(self)
        self.scale_percentage_edit.setFixedWidth(60)
        self.scale_percentage_edit.setValidator(QDoubleValidator(1, 1000, 2))
        self.scale_percentage_edit.setText('100.00')
        self.scale_percentage_edit.editingFinished.connect(self.scale_by_percentage)
        self.tool_bar.addWidget(self.scale_percentage_edit)
        self.tool_bar.addWidget(QLabel('%'))

    def a9_setting(self):
        # 读取上一次打开的文件夹、窗口位置和状态
        last_opened_folder = self._settings.value('last_opened_folder', '')
        geometry = self._settings.value('window_geometry')
        state = self._settings.value('window_state')
        # 如果上一次打开的文件夹存在，则打开它
        if last_opened_folder and os.path.exists(last_opened_folder) and os.path.isdir(last_opened_folder):
            self.open_folder_by_path(last_opened_folder)
        # 如果上一次有记录窗口位置，则恢复窗口位置
        if geometry is not None:
            self.restoreGeometry(geometry)
        # 如果上一次有记录窗口状态，则恢复窗口状态
        if state is not None:
            self.restoreState(state)

        # 显示窗口
        self.show()

    def _display_image(self, pixmap):
        self.image_size = pixmap.size()

        # ================清除之前的图像================
        if self.image_item:
            self.graphics_scene.removeItem(self.image_item)  # 移除之前的图片项

        # ================显示新图片================
        self.image_item = self.graphics_scene.addPixmap(pixmap)
        # 将视图大小设置为 pixmap 的大小，并将图像放入视图中
        self.graphics_view.setSceneRect(pixmap.rect().toRectF())
        self.graphics_view.setScene(self.graphics_scene)

        # 设置视图渲染选项
        self.graphics_view.setRenderHint(QPainter.RenderHint.Antialiasing)
        self.graphics_view.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)
        self.graphics_view.setOptimizationFlag(QGraphicsView.OptimizationFlag.DontAdjustForAntialiasing, True)
        self.graphics_view.setViewportUpdateMode(QGraphicsView.ViewportUpdateMode.FullViewportUpdate)
        self.graphics_view.setOptimizationFlag(QGraphicsView.OptimizationFlag.DontSavePainterState, True)
        self.graphics_view.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.graphics_view.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        # 设置视图转换和缩放选项
        self.graphics_view.setTransformationAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        self.graphics_view.setResizeAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        self.graphics_view.setBackgroundBrush(QBrush(Qt.GlobalColor.lightGray))
        self.graphics_view.setTransform(
            QTransform().scale(self.screen_scaling_factor_reciprocal, self.screen_scaling_factor_reciprocal))

        # 更新缩放比例
        self.update_scale_percentage()

        # 如果自适应屏幕选项已经选中，则按照自适应屏幕缩放图片
        if self.fit_to_screen_action.isChecked():
            self.fit_to_view("screen")
        # 如果自适应宽度选项已经选中，则按照自适应宽度缩放图片
        elif self.fit_to_width_action.isChecked():
            self.fit_to_view("width")
        # 如果自适应高度选项已经选中，则按照自适应宽度缩放图片
        elif self.fit_to_height_action.isChecked():
            self.fit_to_view("height")

    def open_image_from_data(self, image_data, bgr_order=False):
        # 如果输入是Pillow图像，将其转换为NumPy数组
        if isinstance(image_data, Image.Image):
            image_data = np.array(image_data)

        # 确保输入数据是NumPy数组
        if isinstance(image_data, np.ndarray):
            height, width, channel = image_data.shape
            bytes_per_line = channel * width

            # 如果输入图像使用BGR顺序，交换颜色通道以获得正确的RGB顺序
            if bgr_order:
                image_data = cv2.cvtColor(image_data, cv2.COLOR_BGR2RGB)

            # 如果输入图像有4个通道（带有Alpha通道），则使用QImage.Format_ARGB32
            if channel == 4:
                qimage_format = QImage.Format_ARGB32
            # 如果输入图像有3个通道，使用QImage.Format_RGB888
            elif channel == 3:
                qimage_format = QImage.Format_RGB888
            # 其他情况暂不支持
            else:
                raise ValueError("Unsupported number of channels in the input image.")

            # 将NumPy数组转换为QImage
            qimage = QImage(image_data.data, width, height, bytes_per_line, qimage_format)

            # 将QImage转换为QPixmap
            pixmap = QPixmap.fromImage(qimage)

            # 显示图像
            self._display_image(pixmap)
        else:
            raise ValueError("Input data must be a NumPy array or a Pillow Image.")
        QApplication.processEvents()

    def open_image_by_path(self, image_file):
        # 需要路径存在
        image_file = Path(image_file)
        if image_file.exists():
            self.image_file = image_file
            self.image_file_size = os.path.getsize(self.image_file)
            self.image_index = self.image_list.index(self.image_file)

            # ================显示新图片================
            pixmap = QPixmap(self.image_file.as_posix())
            self._display_image(pixmap)

            # 将当前图片项设为选中状态，更新状态栏信息
            self.thumbnails_widget.setCurrentRow(self.image_index)
            index_str = f'{self.image_index + 1}/{len(self.image_list)}'
            meta_str = f'{self.tr("Width")}: {self.image_size.width()} {self.tr("Height")}: {self.image_size.height()} | {self.tr("File Size")}: {self.image_file_size} bytes'
            status_text = f'{index_str} | {self.tr("Filnename")}: {self.image_file.name} | {meta_str}'
            self.status_bar.showMessage(status_text)

    def open_folder_by_path(self, folder_path):
        # 打开上次关闭程序时用到的文件夹
        # 打开最近的文件夹
        # 打开文件夹

        # 判断文件夹路径是否存在
        if os.path.exists(folder_path):
            # 获取所有图片文件的路径
            image_list = collect_valid_images_from(folder_path)
            if image_list:
                self.image_folder = folder_path
                self.image_list = image_list
                self.filtered_image_list = self.image_list
                self.image_index = 0
                self.filtered_image_index = 0
                self.image_file = self.image_list[self.image_index]

                # ================更新导航栏中的图片列表================
                self.thumbnails_widget.clear()
                for image_file in self.image_list:
                    # 将image的basename作为文本添加到thumbnail_item中
                    thumbnail_item = QListWidgetItem(image_file.name)
                    # 将image的图标设置为缩略图
                    thumbnail_item.setIcon(QIcon(image_file.as_posix()))
                    # 文本居中显示
                    thumbnail_item.setTextAlignment(Qt.AlignmentFlag.AlignHCenter)
                    thumbnail_item.setData(Qt.ItemDataRole.UserRole, image_file)
                    # 如果image文件存在，则将thumbnail_item添加到thumbnails_widget中
                    if image_file.exists():
                        self.thumbnails_widget.addItem(thumbnail_item)

                self.open_image_by_path(self.image_file)

                self.setWindowTitle(f'{GlobalConfig.APP_NAME} {GlobalConfig.APP_VERSION} - {self.image_folder}')
                # 将该文件夹添加到最近使用文件夹列表
                self.add_recent_folder(self.image_folder)

    def open_folder_by_dialog(self):
        # 如果self.image_folder已经设置，使用其上一级目录作为起始目录，否则使用当前目录
        self.image_folder = Path(self.image_folder) if self.image_folder else None
        start_directory = self.image_folder.parent.as_posix() if self.image_folder else "."
        image_folder = QFileDialog.getExistingDirectory(self, self.tr('Open Folder'), start_directory)
        if image_folder:
            self.image_folder = image_folder
            self.open_folder_by_path(self.image_folder)

    def update_recent_folders_menu(self):
        self.recent_folders_menu.clear()
        recent_folders = self._settings.value("recent_folders", [])
        for folder in recent_folders:
            action = QAction(str(folder), self)
            action.triggered.connect(lambda checked, p=folder: self.open_folder_by_path(p))
            self.recent_folders_menu.addAction(action)

    def add_recent_folder(self, folder_path):
        recent_folders = self._settings.value("recent_folders", [])
        if folder_path in recent_folders:
            recent_folders.remove(folder_path)
        recent_folders.insert(0, folder_path)
        recent_folders = recent_folders[:10]  # 保留最多10个最近文件夹

        self._settings.setValue("recent_folders", recent_folders)
        self.update_recent_folders_menu()

    def set_display_mode(self, mode):
        self.view_mode = mode
        self.thumbnails_widget.setUpdatesEnabled(False)
        for index in range(self.thumbnails_widget.count()):
            item = self.thumbnails_widget.item(index)
            data = item.data(Qt.ItemDataRole.UserRole)
            if self.view_mode == 0:  # Show thumbnails only
                item.setIcon(QIcon(data.as_posix()))
                item.setText('')
                self.thumbnails_widget.setWordWrap(False)
            elif self.view_mode == 1:  # Show filenames only
                item.setIcon(QIcon())  # Clear the icon
                item.setText(data.name)
                self.thumbnails_widget.setWordWrap(False)  # Make sure filenames don't wrap
            elif self.view_mode == 2:  # Show both thumbnails and filenames
                item.setIcon(QIcon(data.as_posix()))
                item.setText(data.name)
                self.thumbnails_widget.setWordWrap(True)

        thumbnail_size = QSize(240, 240)  # Set the thumbnail size to 240x240
        self.thumbnails_widget.setIconSize(thumbnail_size)  # Set the icon size

        if self.view_mode == 1:  # Show filenames only
            self.thumbnails_widget.setGridSize(QSize(-1, -1))  # Use default size for grid
        else:
            # Set grid size for thumbnails and both modes
            self.thumbnails_widget.setGridSize(
                QSize(thumbnail_size.width() + 30, -1))  # Add extra width for file names and spacing

        self.thumbnails_widget.setUpdatesEnabled(True)

    def select_image_from_list(self):
        # 通过键鼠点击获取的数据是真实序数image_index
        active_list_widget = self.nav_tab.currentWidget()
        selected_items = active_list_widget.selectedItems()
        if not selected_items:
            return
        current_item = selected_items[0]
        image_index = active_list_widget.row(current_item)
        if image_index != self.image_index:
            self.image_index = image_index
            image_file = current_item.data(Qt.ItemDataRole.UserRole)
            self.open_image_by_path(image_file)

    def open_image_in_viewer(self, file_path):
        if sys.platform == 'win32':
            os.startfile(os.path.normpath(file_path))
        elif sys.platform == 'darwin':
            Popen(['open', file_path])
        else:
            Popen(['xdg-open', file_path])

    def open_file_in_explorer(self, file_path):
        folder_path = os.path.dirname(file_path)
        if sys.platform == 'win32':
            Popen(f'explorer /select,"{os.path.normpath(file_path)}"')
        elif sys.platform == 'darwin':
            Popen(['open', '-R', file_path])
        else:
            Popen(['xdg-open', folder_path])

    def open_image_in_ps(self, file_path):
        if sys.platform == 'win32':
            photoshop_executable_path = "C:/Program Files/Adobe/Adobe Photoshop CC 2019/Photoshop.exe"  # 请根据您的Photoshop安装路径进行修改
            Popen([photoshop_executable_path, file_path])
        elif sys.platform == 'darwin':
            photoshop_executable_path = "/Applications/Adobe Photoshop 2021/Adobe Photoshop 2021.app"  # 修改此行
            Popen(['open', '-a', photoshop_executable_path, file_path])
        else:
            QMessageBox.warning(self, "Warning", "This feature is not supported on this platform.")

    def copy_to_clipboard(self, text):
        clipboard = QApplication.clipboard()
        clipboard.setText(text)

    def show_image_list_context_menu(self, point):
        item = self.thumbnails_widget.itemAt(point)
        if item:
            context_menu = QMenu(self)
            open_file_action = QAction(self.tr('Open in Explorer'), self)
            open_file_action.triggered.connect(
                lambda: self.open_file_in_explorer(item.data(Qt.ItemDataRole.UserRole)))
            open_image_action = QAction(self.tr('Open in Preview'), self)
            open_image_action.triggered.connect(
                lambda: self.open_image_in_viewer(item.data(Qt.ItemDataRole.UserRole)))
            context_menu.addAction(open_file_action)
            context_menu.addAction(open_image_action)

            # 添加一个打开方式子菜单
            open_with_menu = context_menu.addMenu(self.tr('Open with'))
            open_with_ps = QAction(self.tr('Photoshop'), self)
            open_with_menu.addAction(open_with_ps)

            # 添加拷贝图片路径、拷贝图片名的选项
            copy_image_path = QAction(self.tr('Copy Image Path'), self)
            copy_image_name = QAction(self.tr('Copy Image Name'), self)

            context_menu.addAction(copy_image_path)
            context_menu.addAction(copy_image_name)

            # 连接触发器
            open_with_ps.triggered.connect(lambda: self.open_image_in_ps(item.data(Qt.ItemDataRole.UserRole)))
            copy_image_path.triggered.connect(
                lambda: self.copy_to_clipboard(item.data(Qt.ItemDataRole.UserRole).as_posix()))
            copy_image_name.triggered.connect(lambda: self.copy_to_clipboard(item.text()))
            context_menu.exec(self.thumbnails_widget.mapToGlobal(point))

    def refresh_search_results(self):
        # 用于三个筛选按钮
        search_text = self.search_bar.text()
        self.filter_image_list(search_text)

    def filter_image_list(self, search_text):
        # 用于搜索框内容更新
        case_sensitive = self.case_sensitive_button.isChecked()
        whole_word = self.whole_word_button.isChecked()
        use_regex = self.regex_button.isChecked()

        flags = re.IGNORECASE if not case_sensitive else 0

        if use_regex:
            if whole_word:
                search_text = fr"\b{search_text}\b"
            try:
                regex = re.compile(search_text, flags)
            except re.error:
                return
        else:
            if whole_word:
                search_text = fr"\b{re.escape(search_text)}\b"
            else:
                search_text = re.escape(search_text)
            regex = re.compile(search_text, flags)

        # ================筛选列表更新================
        self.filtered_image_list = []
        for image_file in self.image_list:
            match = regex.search(image_file.name)
            if match:
                self.filtered_image_list.append(image_file)

        # ================缩略图列表更新================
        for index in range(self.thumbnails_widget.count()):
            item = self.thumbnails_widget.item(index)
            item_text = item.text()
            match = regex.search(item_text)
            if match:
                item.setHidden(False)
            else:
                item.setHidden(True)

    def save_image(self):
        file_name, _ = QFileDialog.getSaveFileName(self, self.tr('Save Image'), "", "Images (*.png *.xpm *.jpg)")

        if file_name:
            image = QImage(self.graphics_scene.sceneRect().size().toSize(), QImage.Format.Format_ARGB32)
            image.fill(Qt.GlobalColor.transparent)

            painter = QPainter(image)
            self.graphics_scene.render(painter)
            painter.end()

            image.save(file_name)

    def change_image(self, step):
        current_image_path = self.image_list[self.image_index]
        # 检查当前图片路径是否在过滤后的图片列表中
        if current_image_path not in self.filtered_image_list:
            return
        current_filtered_index = self.filtered_image_list.index(current_image_path)

        if step == "first":
            new_filtered_index = 0
        elif step == "last":
            new_filtered_index = len(self.filtered_image_list) - 1
        else:
            new_filtered_index = current_filtered_index + step

        if 0 <= new_filtered_index < len(self.filtered_image_list):
            new_image_path = self.filtered_image_list[new_filtered_index]
            self.open_image_by_path(new_image_path)

    def get_screen_scaling_factor(self):
        screen = QApplication.primaryScreen()
        if sys.platform == 'darwin':  # 如果是 MacOS 系统
            return screen.devicePixelRatio()
        else:
            return 1

    def zoom(self, scale_factor):
        self.graphics_view.setRenderHint(QPainter.RenderHint.Antialiasing)
        current_transform = self.graphics_view.transform()
        current_scale = current_transform.m11()
        new_scale = current_scale * scale_factor

        if 0.01 <= new_scale <= 10:
            current_transform.scale(scale_factor, scale_factor)
            self.graphics_view.setTransform(current_transform)
            self.update_scale_percentage()

    def zoom_in(self):
        self.zoom(1.2)

    def zoom_out(self):
        self.zoom(1 / 1.2)

    def fit_to_view(self, mode):
        if self.image_item:
            if mode == "screen":
                self.graphics_view.fitInView(self.image_item, Qt.AspectRatioMode.KeepAspectRatio)
            elif mode == "width":
                view_width = self.graphics_view.viewport().width()
                pixmap_width = self.image_item.pixmap().width()
                scale_factor = view_width / pixmap_width
                self.graphics_view.setTransform(QTransform().scale(scale_factor, scale_factor))
            elif mode == "height":
                view_height = self.graphics_view.viewport().height()
                pixmap_height = self.image_item.pixmap().height()
                scale_factor = view_height / pixmap_height
                self.graphics_view.setTransform(QTransform().scale(scale_factor, scale_factor))
            self.graphics_view.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)
            self.update_scale_percentage()

    def fit_to_view_toggled(self, checked):
        sender = self.sender()
        if sender == self.fit_to_screen_action:
            mode = 'screen'
            if checked:
                self.fit_to_width_action.setChecked(False)
                self.fit_to_height_action.setChecked(False)
        elif sender == self.fit_to_width_action:
            mode = 'width'
            if checked:
                self.fit_to_screen_action.setChecked(False)
                self.fit_to_height_action.setChecked(False)
        else:  # if sender == self.fit_to_height_action
            mode = 'height'
            if checked:
                self.fit_to_screen_action.setChecked(False)
                self.fit_to_width_action.setChecked(False)
        self.fit_to_view(mode)

    def reset_zoom(self):
        self.fit_to_screen_action.setChecked(False)
        self.fit_to_width_action.setChecked(False)
        self.fit_to_height_action.setChecked(False)
        # 使用 setTransform 代替 resetMatrix 并应用缩放因子
        self.graphics_view.setTransform(
            QTransform().scale(self.screen_scaling_factor_reciprocal, self.screen_scaling_factor_reciprocal))
        self.update_scale_percentage()

    def update_scale_percentage(self):
        current_scale = round(self.graphics_view.transform().m11() * 100, 2)
        self.scale_percentage_edit.setText(str(current_scale))

    def scale_by_percentage(self):
        scale_percentage = float(self.scale_percentage_edit.text())
        target_scale = scale_percentage / 100
        current_scale = self.graphics_view.transform().m11()

        scale_factor = target_scale / current_scale
        self.graphics_view.setRenderHint(QPainter.RenderHint.Antialiasing)
        self.graphics_view.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)
        self.graphics_view.scale(scale_factor, scale_factor)

    def show_about_dialog(self):
        about_dialog = QDialog(self)
        about_dialog.setWindowTitle(f"{self.tr('About')} {GlobalConfig.APP_NAME}")

        app_name_label = QLabel(GlobalConfig.APP_NAME)
        app_name_label.setFont(QFont("Arial", 20, QFont.Bold))

        version_label = QLabel(f"版本: {GlobalConfig.APP_VERSION}")
        libraries_label = QLabel("使用的库：PyQt6")
        license_label = QLabel("许可证: MIT License")

        main_layout = QVBoxLayout()
        main_layout.addWidget(app_name_label)
        main_layout.addWidget(version_label)
        main_layout.addWidget(libraries_label)
        main_layout.addWidget(license_label)
        about_dialog.setLayout(main_layout)
        about_dialog.exec_()

    def show_help_document(self):
        # 在此处显示帮助文档
        pass

    def show_feedback_dialog(self):
        recipient_email = "annebing@qq.com"  # 替换为您的电子邮件地址
        subject = "Bug Report - 漫画翻译工具"
        body = "请在此处详细描述您遇到的问题：\n\n\n\n软件版本：1.0.0"
        mailto_url = f"mailto:{recipient_email}?subject={subject}&body={body}"

        try:
            webbrowser.open(mailto_url)
        except webbrowser.Error as e:
            QMessageBox.warning(self, "错误", f"无法打开邮件客户端：{e}")

    def check_for_updates(self):
        # 在此处实现在线更新软件的功能
        pass

    @timer_decorator
    @logger.catch
    def step0_analyze_frames(self):
        if self.frame_yml_path.exists():
            with open(self.frame_yml_path, 'r') as yml_file:
                image_data = yaml.safe_load(yml_file)
        else:
            image_data = {}

        # ================分析画格================
        total_images = len(self.image_list)
        processed_images = 0

        for p, image_file in enumerate(self.image_list):
            logger.warning(f'{image_file=}')
            image_raw = cv2.imdecode(np.fromfile(image_file, dtype=np.uint8), -1)

            if image_file.name in image_data:
                frame_grid_strs = image_data[image_file.name]
            else:
                # 获取RGBA格式的边框像素
                edge_pixels_rgba = get_edge_pixels_rgba(image_raw)
                # 寻找主要颜色
                color_ratio, dominant_color = find_dominant_color(edge_pixels_rgba)
                # 获取颜色名称
                color_name = get_color_name(dominant_color)
                logger.info(
                    f"边框颜色中出现次数最多的颜色：{dominant_color}, 颜色名称：{color_name}, {color_ratio=}")
                
                # 计算主要颜色遮罩
                frame_mask = compute_frame_mask(image_raw, dominant_color)

                ih, iw = frame_mask.shape[0:2]
                rec0 = (0, 0, iw, ih)
                pg = PicGrid(frame_mask, self.media_type)
                grids = []
                if pg.basic:
                    rec0 = Rect(rec0, self.media_type)
                    rec0.get_devide(frame_mask)
                    grids.append(rec0)
                else:
                    rects = pg.rects_dic[pg.judgement]
                    for rect in rects:
                        get_frame_grid_recursive(rect, grids, frame_mask, self.media_type)
                frame_grid_strs = [rect.frame_grid_str for rect in grids]
                image_data[image_file.name] = frame_grid_strs

            processed_images += 1
            self.pb_task.setValue(int(processed_images / total_images * 100))
            QApplication.processEvents()

            if len(frame_grid_strs) >= 2:
                marked_frames = get_marked_frames(frame_grid_strs, image_raw)
                self.open_image_from_data(marked_frames, bgr_order=True)

        # Update progress bar to 100%
        self.pb_task.setValue(100)
        QApplication.processEvents()

        # Sort image_data dictionary by image file names
        image_data_sorted = {k: image_data[k] for k in natsorted(image_data)}

        # Save the image_data_sorted dictionary to the yml file
        with open(self.frame_yml_path, 'w') as yml_file:
            yaml.dump(image_data_sorted, yml_file)

    @timer_decorator
    @logger.catch
    def step1_analyze_bubbles(self):
        auto_subdir = GlobalStorage.Auto / self.image_folder.name
        make_dir(auto_subdir)

        if self.frame_yml_path.exists():
            with open(self.frame_yml_path, 'r') as yml_file:
                image_data = yaml.safe_load(yml_file)
        else:
            image_data = {}
        total_images = len(self.image_list)
        processed_images = 0
        all_masks_old = collect_valid_images_from(self.image_folder, mode='mask')

        for p, image_file in enumerate(self.image_list):
            logger.warning(f'{image_file=}')
            image_raw = cv2.imdecode(np.fromfile(image_file, dtype=np.uint8), -1)

            ih, iw = image_raw.shape[0:2]
            # ================矩形画格信息================
            if image_file.name in image_data:
                frame_grid_strs = image_data[image_file.name]
            else:
                frame_grid_strs = [f'0,0,{iw},{ih}~0,0,{iw},{ih}']
            # ================模型检测文字，文字显示为白色================
            if self._settings.value('torch') and GlobalStorage.DNNNet is not None:
                comictextdetector_mask = get_comictextdetector_mask(image_raw)
            else:
                comictextdetector_mask = None
            # ================针对每一张图================
            for c in range(len(color_patterns)):
                # ================遍历每种气泡文字颜色组合================
                color_pattern = color_patterns[c]
                cp_bubble, cp_letter = color_pattern

                if isinstance(cp_bubble, list):
                    # 气泡为渐变色
                    color_bubble = ColorGradient(cp_bubble)
                else:
                    # 气泡为单色
                    color_bubble = Color(cp_bubble)

                if isinstance(cp_letter, list):
                    # 文字为双色
                    color_letter = ColorDouble(cp_letter)
                    color_letter_big = color_letter
                else:
                    # 文字为单色
                    color_letter = Color(cp_letter)
                    color_letter_big = Color(f'{cp_letter[:6]}-{max(color_letter.padding + 120, 180)}')

                if color_bubble.type == 'Color':
                    bubble_mask = color_bubble.get_img(image_raw)
                else:  # ColorGradient
                    img_tuple = color_bubble.get_img(image_raw)
                    bubble_mask, left_sample, right_sample = img_tuple
                letter_mask = color_letter_big.get_img(image_raw)

                filtered_cnts = get_raw_bubbles(bubble_mask, letter_mask, comictextdetector_mask)
                colorful_raw_bubbles = get_colorful_bubbles(image_raw, filtered_cnts)

                # ================切割相连气泡================
                single_cnts = seg_bubbles(filtered_cnts, bubble_mask, letter_mask, self.media_type)
                # ================通过画格重新排序气泡框架================
                single_cnts_recs = []
                for f in range(len(frame_grid_strs)):
                    frame_grid_str = frame_grid_strs[f]
                    # 使用正则表达式提取字符串中的所有整数
                    int_values = list(map(int, re.findall(r'\d+', frame_grid_str)))
                    # 按顺序分配值
                    x, y, w, h, xx, yy, ww, hh = int_values
                    single_cnts_rec = []
                    for s in range(len(single_cnts)):
                        single_cnt = single_cnts[s]
                        if x <= single_cnt.cx <= x + w and y <= single_cnt.cy <= y + h:
                            single_cnts_rec.append(single_cnt)
                    if self.media_type == 'Manga':
                        ax = -1
                    else:
                        ax = 1
                    single_cnts_rec.sort(key=lambda x: ax * x.cx + x.cy)
                    single_cnts_recs.append(single_cnts_rec)
                single_cnts_ordered = list(chain(*single_cnts_recs))
                if len(single_cnts_ordered) >= 1:
                    colorful_single_bubbles = get_colorful_bubbles(image_raw, single_cnts_ordered)
                    self.open_image_from_data(colorful_single_bubbles, bgr_order=True)

                    alpha = 0.1
                    # 创建一个带有10%透明度原图背景的图像
                    transparent_image = np.zeros((image_raw.shape[0], image_raw.shape[1], 4), dtype=np.uint8)

                    # Bugfix: unlike jpeg format images, imdecode obtains rgba quads for png format
                    image = image_raw
                    if image_raw.shape[2] == 4:
                        image = np.delete(image_raw, 3, axis=2)

                    transparent_image[..., :3] = image
                    transparent_image[..., 3] = int(255 * alpha)
                    # 在透明图像上绘制contours，每个contour使用不同颜色

                    for s in range(len(single_cnts_ordered)):
                        bubble_cnt = single_cnts_ordered[s]
                        # 从 tab20 颜色映射中选择一个颜色
                        color = colormap_tab20(s % 20)[:3]
                        color_rgb = tuple(int(c * 255) for c in color)
                        color_bgra = color_rgb[::-1] + (255,)
                        cv2.drawContours(transparent_image, [bubble_cnt.contour], -1, color_bgra, -1)

                    self.open_image_from_data(transparent_image)
                    cp_preview_jpg = auto_subdir / f'{image_file.stem}-{color_bubble.rgb_str}-{color_bubble.color_name}~{color_letter.rgb_str}-{color_letter.color_name}.jpg'
                    cp_mask_cnt_pic = auto_subdir / f'{image_file.stem}-Mask-{color_bubble.rgb_str}-{color_bubble.color_name}~{color_letter.rgb_str}-{color_letter.color_name}.png'
                    write_image(cp_preview_jpg, colorful_single_bubbles)
                    write_image(cp_mask_cnt_pic, transparent_image)

            processed_images += 1
            self.pb_task.setValue(int(processed_images / total_images * 100))
            QApplication.processEvents()

        # ================搬运气泡蒙版================
        auto_all_masks = collect_valid_images_from(auto_subdir, mode='mask')
        # 如果步骤开始前没有气泡蒙版
        if not all_masks_old:
            for mask_src in auto_all_masks:
                mask_dst = self.image_folder / mask_src.name
                copy2(mask_src, mask_dst)

        # Update progress bar to 100%
        self.pb_task.setValue(100)
        QApplication.processEvents()

    def step2_OCR(self):
        if self.frame_yml_path.exists():
            with open(self.frame_yml_path, 'r') as yml_file:
                image_data = yaml.safe_load(yml_file)
        else:
            image_data = {}
        # ================分析画格================
        total_images = len(self.image_list)
        processed_images = 0
        # ================气泡蒙版================
        all_masks = collect_valid_images_from(self.image_folder, mode='mask')
        # ================DOCX文档================
        # 创建一个新的Document对象
        OCR_doc = Document()
        for i in range(len(self.image_list)):
            image_file = self.image_list[i]
            logger.warning(f'{image_file=}')

            OCR_doc.add_paragraph(image_file.stem)
            OCR_doc.add_paragraph('')

            image_raw = cv2.imdecode(np.fromfile(image_file, dtype=np.uint8), -1)
            ih, iw = image_raw.shape[0:2]

            # ================获取对应的文字图片================
            single_cnts = get_single_cnts(image_file, image_raw, all_masks)
            logger.debug(f'{len(single_cnts)=}')
            # ================矩形画格信息================
            if image_file.name in image_data:
                frame_grid_strs = image_data[image_file.name]
            else:
                frame_grid_strs = [f'0,0,{iw},{ih}~0,0,{iw},{ih}']
            # ================通过画格重新排序气泡框架================
            single_cnts_recs = []
            for f in range(len(frame_grid_strs)):
                frame_grid_str = frame_grid_strs[f]
                # 使用正则表达式提取字符串中的所有整数
                int_values = list(map(int, re.findall(r'\d+', frame_grid_str)))
                # 按顺序分配值
                x, y, w, h, xx, yy, ww, hh = int_values
                single_cnts_rec = []
                for s in range(len(single_cnts)):
                    single_cnt = single_cnts[s]
                    if x <= single_cnt.cx <= x + w and y <= single_cnt.cy <= y + h:
                        single_cnts_rec.append(single_cnt)
                if self.media_type == 'Manga':
                    ax = -1
                else:
                    ax = 1
                single_cnts_rec.sort(key=lambda x: ax * x.cx + x.cy)
                single_cnts_recs.append(single_cnts_rec)
            single_cnts_ordered = list(chain(*single_cnts_recs))

            if len(single_cnts_ordered) >= 1:
                colorful_single_bubbles = get_colorful_bubbles(image_raw, single_cnts_ordered)
                self.open_image_from_data(colorful_single_bubbles, bgr_order=True)

            for c in range(len(single_cnts_ordered)):
                single_cnt = single_cnts_ordered[c]
                # ================对裁剪后的图像进行文字识别================
                if GlobalConfig.SYSTEM in ['MAC', 'M1']:
                    # ================Vision================
                    v = Vision()
                    recognized_text = v.process(single_cnt.cropped_image, self.media_language)
                    lines = []
                    for text, (x, y, w, h), confidence in recognized_text:
                        # print(f"{text}[{confidence:.2f}] {x=:.2f}, {y=:.2f}, {w=:.2f}, {h=:.2f}")
                        lines.append(text)
                else:
                    # ================tesseract================
                    t = Tesseract()
                    recognized_text = t.process(single_cnt.cropped_image, self.media_language, vertical)
                    # 将recognized_text分成多行
                    lines = recognized_text.splitlines()
                # 去除空行
                non_empty_lines = [line for line in lines if line.strip()]
                # 将非空行合并为一个字符串
                cleaned_text = GlobalConfig.LINE_FEED.join(non_empty_lines)
                print(cleaned_text)
                print('-' * 24)

                # ================将图片添加到docx文件中================
                cropped_image_pil = Image.fromarray(single_cnt.cropped_image)
                image_dpi = cropped_image_pil.info.get('dpi', (96, 96))[0]  # 获取图片的dpi，如果没有dpi信息，则使用默认值96
                with io.BytesIO() as temp_buffer:
                    cropped_image_pil.save(temp_buffer, format=image_format.upper())
                    temp_buffer.seek(0)
                    pic_width_inches = cropped_image_pil.width / image_dpi
                    OCR_doc.add_picture(temp_buffer, width=Inches(pic_width_inches))
                # ================将识别出的文字添加到docx文件中================
                OCR_doc.add_paragraph(cleaned_text)
                # 在图片和文字之间添加一个空行
                OCR_doc.add_paragraph('')

            # Update progress bar
            self.pb_task.setValue(int(processed_images / total_images * 100))
            # Process pending events to keep the UI responsive
            QApplication.processEvents()

            # Increment processed images counter
            processed_images += 1

        # ================保存为DOCX================
        OCR_doc.save(self.OCR_docx_path)
        # Update progress bar to 100%
        self.pb_task.setValue(100)

    @logger.catch
    @timer_decorator
    def step3_translate(self):
        target_language = 'zh-CN'

        # 打开docx文件
        OCR_doc = Document(self.OCR_docx_path)

        # 读取并连接文档中的所有段落文本
        full_text = []
        for para in OCR_doc.paragraphs:
            full_text.append(para.text)
            # logger.debug(f'{para.text=}')

        index_dict = {}
        last_ind = 0
        inds = []
        for i in range(len(self.image_list)):
            image_file = self.image_list[i]
            logger.warning(f'{image_file=}')
            if image_file.stem in full_text[last_ind:]:
                ind = full_text[last_ind:].index(image_file.stem) + last_ind
                index_dict[image_file.stem] = ind
                inds.append(ind)
                last_ind = ind
                logger.debug(f'{ind=}')
        inds.append(len(full_text))
        # pprint(index_dict)

        pin = 0
        para_dict = {}
        for i in range(len(self.image_list)):
            image_file = self.image_list[i]
            if image_file.stem in index_dict:
                start_i = inds[pin] + 1
                end_i = inds[pin + 1]
                pin += 1
                para_dict[image_file.stem] = full_text[start_i:end_i]
        # pprint(para_dict)

        all_valid_paras = []
        for key in para_dict:
            paras = para_dict[key]
            valid_paras = [x for x in paras if x.strip() != '']
            if valid_paras:
                all_valid_paras.extend(valid_paras)

        simple_lines = []
        for a in range(len(all_valid_paras)):
            para = all_valid_paras[a]
            # 将多行文本替换为单行文本
            single_line_text = para.replace('\n', ' ')
            # 将文本分割成句子
            sentences = re.split(r' *[\.\?!][\'"\)\]]* *', single_line_text)
            # 对每个句子进行首字母大写处理，同时考虑缩写
            capitalized_sentences = [capitalize_sentence(sentence) for sentence in sentences if sentence]
            # 将处理后的句子连接成一个字符串
            processed_text = '. '.join(capitalized_sentences) + '.'
            simple_lines.append(processed_text)

        simple_text = GlobalConfig.LINE_FEED.join(simple_lines)
        print(simple_text)

        chunks = []
        current_chunk = ""
        for line in simple_lines:
            # 检查将当前行添加到当前块后的长度是否超过最大字符数
            if len(current_chunk) + len(line) + 1 > max_chars:  # 加1是为了考虑换行符
                chunks.append(current_chunk.strip())
                current_chunk = ""
            current_chunk += line + "\n"

        # 添加最后一个块（如果有内容的话）
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        for chunk in chunks:
            translated_text = GoogleTranslator(source='auto', target=target_language).translate(chunk)
            print(translated_text)

    def step4_lettering(self):
        pass

    def start_task(self, checked):
        # 当使用 QPushButton 的 clicked 信号连接到槽时，它会传递一个布尔值表示按钮是否被选中
        # 检查哪个 QRadioButton 被选中并执行相应的任务
        self.task_name = self.bg_step.checkedButton().text()
        self.task_ind = self.bg_step.checkedId()
        print(f'[{self.task_ind}]{self.task_name}')
        # 在这里执行您的任务逻辑
        if self.image_folder is not None:
            self.image_folder = Path(self.image_folder)
            if self.image_folder.exists():
                self.frame_yml_path = self.image_folder.parent / f'{self.image_folder.name}.yml'
                self.OCR_docx_path = self.image_folder.parent / f'{self.image_folder.name}-OCR.docx'
                if self.task_ind == 0:
                    self.step0_analyze_frames()
                elif self.task_ind == 1:
                    self.step1_analyze_bubbles()
                elif self.task_ind == 2:
                    self.step2_OCR()
                elif self.task_ind == 3:
                    self.step3_translate()
                elif self.task_ind == 4:
                    self.step4_lettering()

                # 自动选中下一个步骤，除非当前步骤是最后一个步骤
                if self.task_ind < self.bg_step.id(self.bg_step.buttons()[-1]):
                    next_task_ind = self.task_ind + 1
                    next_task_button = self.bg_step.button(next_task_ind)
                    if next_task_button:
                        next_task_button.setChecked(True)
                        self.task_ind = next_task_ind

        # Process pending events to keep the UI responsive
        QApplication.processEvents()

    def update_media_type(self, button):
        index = self.bg_media_type.id(button)
        self.media_type_index = index
        self.media_type = media_type_dict[index]

    def update_media_language(self, index):
        self.media_language_index = index
        self.media_language = media_language_dict[index]

    def closeEvent(self, event):
        # 如果有打开图片的文件夹，将其保存到程序设置中
        if self.image_folder:
            self._settings.setValue('last_opened_folder', self.image_folder)
        else:
            self._settings.setValue('last_opened_folder', '')
        # 保存窗口的几何和状态信息到程序设置中
        self._settings.setValue('window_geometry', self.saveGeometry())
        self._settings.setValue('window_state', self.saveState())
        self._settings.setValue('media_type_index', self.media_type_index)
        self._settings.setValue('media_language_index', self.media_language_index)
        self._settings.sync()
        event.accept()