from opencc import OpenCC

import codecs
import os
import os.path
import pickle
import re
import sys
from collections import OrderedDict
from csv import reader, writer, QUOTE_MINIMAL
from filecmp import cmp
from functools import partial, lru_cache, reduce, wraps
from getpass import getuser
from hashlib import md5
from io import StringIO
from locale import getdefaultlocale
from math import sqrt
from os.path import abspath, dirname, exists, expanduser, getsize, isfile, normpath
from pathlib import Path
from platform import machine, processor, python_version, system, uname
from re import I, IGNORECASE, escape, findall, match
from shutil import copy2
from subprocess import PIPE, Popen
from time import strftime, time
from traceback import print_exc
from unicodedata import normalize
from uuid import getnode

import yaml
from PIL import Image
from PyQt6.QtCore import QEventLoop, QTimer, QItemSelectionModel, QSize, Qt, pyqtSignal, QTranslator, QPoint
from PyQt6.QtGui import QAction, QActionGroup, QBrush, QDoubleValidator, QFont, QIcon, QImage, \
    QKeySequence, QPainter, QPixmap, QStandardItemModel, QColor, QPen, QStandardItem
from PyQt6.QtWidgets import QAbstractItemView, QApplication, QDockWidget, QFileDialog, \
    QGraphicsScene, QGraphicsView, QHBoxLayout, QLabel, QLineEdit, QListView, \
    QListWidget, QListWidgetItem, QMainWindow, QMenu, QStatusBar, QGraphicsSceneMouseEvent, \
    QTabWidget, QToolBar, QToolButton, QVBoxLayout, QWidget, QFrame, QHeaderView, QPlainTextEdit, QSizePolicy, \
    QTableView, QGraphicsTextItem, QGraphicsEllipseItem, QGraphicsItem, QGraphicsItemGroup, QStyledItemDelegate
from cv2 import COLOR_BGR2RGB, COLOR_BGRA2BGR, COLOR_GRAY2BGR, COLOR_RGB2BGR, cvtColor, imdecode, imencode
from loguru import logger
from matplotlib import colormaps
from natsort import natsorted
from numpy import array, clip, fromfile, ndarray, ones, sqrt, uint8
from psutil import virtual_memory
from qtawesome import icon as qicon
from ruamel.yaml import YAML
from webcolors import hex_to_rgb, name_to_rgb, rgb_to_name
from webcolors._definitions import _reversedict, _CSS3_NAMES_TO_HEX
from PyQt6.QtCore import QByteArray, QCommandLineOption, QCommandLineParser, QCoreApplication, QDate, \
    QEventLoop, QItemSelection, QItemSelectionModel, QModelIndex, QPoint, QRegularExpression, QSettings, QSize, \
    QThread, QTimer, QTranslator, QUrl, QVariant, Qt, pyqtSignal
from PyQt6.QtGui import QAction, QColor, QFont, QGuiApplication, QImage, QKeySequence, QPainter, QPalette, \
    QPixmap, QRegularExpressionValidator, QStandardItem, QStandardItemModel, QTextDocumentFragment, QWindow
from PyQt6.QtNetwork import QNetworkProxyFactory
from PyQt6.QtPrintSupport import QPrintDialog, QPrinter
from PyQt6.QtWidgets import QAbstractItemView, QApplication, QButtonGroup, QCalendarWidget, QCheckBox, \
    QComboBox, QDockWidget, QFontComboBox, QFontDialog, QFrame, QGridLayout, QGroupBox, QHBoxLayout, \
    QHeaderView, QLCDNumber, QLabel, QLineEdit, QListWidget, QMainWindow, QMenu, QMenuBar, QMessageBox, \
    QPlainTextEdit, QProgressBar, QPushButton, QRadioButton, QScrollArea, QSizePolicy, QSlider, QSpacerItem, \
    QSpinBox, QSplitter, QStatusBar, QStyleFactory, QTabBar, QTabWidget, QTableView, QTextEdit, QToolBar, \
    QToolButton, QTreeWidget, QVBoxLayout, QWidget, QStyledItemDelegate
from PyQt6.QtWebEngineWidgets import QWebEngineView
from manga_ocr import MangaOcr
from nltk.tokenize import sent_tokenize
from pandas import DataFrame
import asyncio
import builtins
import codecs
import json
import math
import os
import os.path
import pickle
import re
import string
import sys
import warnings
import webbrowser
from ast import AsyncFunctionDef, Attribute, ClassDef, FunctionDef, Import, ImportFrom, Name, NodeTransformer, \
    NodeVisitor, Pass, parse, unparse, walk
from collections import Counter, OrderedDict, defaultdict
from colorsys import hsv_to_rgb, rgb_to_hsv
from concurrent.futures import ThreadPoolExecutor, as_completed
from copy import deepcopy
from csv import reader, writer, QUOTE_MINIMAL
from datetime import datetime
from difflib import SequenceMatcher
from filecmp import cmp
from getpass import getuser
from hashlib import md5
from html import unescape
from io import BytesIO, StringIO
from itertools import chain, groupby, zip_longest
from locale import getdefaultlocale
from math import ceil, cos, floor, radians, sin, sqrt
from operator import mod
from os import getcwd, makedirs
from os.path import abspath, dirname, exists, expanduser, getmtime, getsize, isdir, isfile, normpath
from pathlib import Path
from platform import machine, platform, processor, python_version, release, system, uname, version
from pprint import pprint
from re import I, IGNORECASE, Pattern, escape, findall, finditer, match, search, sub, split
from re import compile as recompile
from shutil import copy2
from socket import gethostbyname, gethostname
from statistics import mean, median, multimode
from string import ascii_uppercase
from subprocess import PIPE, Popen, call
from textwrap import dedent, indent
from time import localtime, sleep, strftime, time
from tokenize import COMMENT, generate_tokens
from traceback import print_exc
from typing import List, Union
from unicodedata import normalize
from uuid import getnode, uuid4
from warnings import filterwarnings
import cv2
import numpy as np
import pkg_resources
import pyautogui
import pyperclip
import torch
import torch.nn as nn
import torch.nn.functional as F
import torchvision
import xmltodict
import yaml
from PIL import Image, ImageDraw, ImageFont
from PIL.Image import fromarray
from PyQt6.QtCore import QPointF, QRectF, QSettings, QSize, QTranslator, Qt, pyqtSignal
from PyQt6.QtGui import QAction, QActionGroup, QBrush, QColor, QDoubleValidator, QFont, QFontMetrics, QIcon, QImage, \
    QKeySequence, QPainter, QPainterPath, QPen, QPixmap, QPolygonF
from PyQt6.QtWidgets import QAbstractItemView, QApplication, QButtonGroup, QComboBox, QDialog, QDockWidget, QFileDialog, \
    QGraphicsDropShadowEffect, QGraphicsEllipseItem, QGraphicsItem, QGraphicsLineItem, QGraphicsPixmapItem, \
    QGraphicsPolygonItem, QGraphicsScene, QGraphicsTextItem, QGraphicsView, QHBoxLayout, QLabel, QLineEdit, QListView, \
    QListWidget, QListWidgetItem, QMainWindow, QMenu, QMessageBox, QProgressBar, QPushButton, QRadioButton, QStatusBar, \
    QTabWidget, QToolBar, QToolButton, QVBoxLayout, QWidget
from aip import AipOcr
from astor import to_source
from bs4 import BeautifulSoup, NavigableString, Tag
from cv2 import BORDER_CONSTANT, CHAIN_APPROX_SIMPLE, COLOR_BGR2BGRA, COLOR_BGR2GRAY, COLOR_BGR2RGB, COLOR_BGRA2BGR, \
    COLOR_BGRA2RGBA, COLOR_GRAY2BGR, COLOR_GRAY2BGRA, COLOR_RGB2BGR, CV_16U, FILLED, GaussianBlur, IMREAD_COLOR, \
    INPAINT_NS, INTER_LINEAR, MORPH_ELLIPSE, MORPH_RECT, RANSAC, RETR_EXTERNAL, RETR_LIST, RETR_TREE, \
    ROTATE_90_COUNTERCLOCKWISE, THRESH_BINARY, THRESH_OTSU, add, arcLength, bitwise_and, bitwise_not, bitwise_or, \
    bitwise_xor, boundingRect, boxPoints, circle, connectedComponentsWithStats, contourArea, copyMakeBorder, cvtColor, \
    dilate, dnn, drawContours, erode, fillPoly, findContours, findHomography, getStructuringElement, imdecode, imencode, \
    inRange, inpaint, line, minAreaRect, moments, pointPolygonTest, rectangle, resize, rotate, subtract, threshold, \
    warpPerspective
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Inches
from easyocr import Reader
from fontTools.ttLib import TTCollection, TTFont
from html2text import HTML2Text
from jieba import cut
from libcst import CSTTransformer, RemovalSentinel, parse_module
from loguru import logger
from mammoth import convert_to_html
from matplotlib import colormaps, font_manager
from natsort import natsorted
from nltk.corpus import names, words
from nltk.stem import WordNetLemmatizer
from numpy import arange, argmax, argmin, argsort, argwhere, array, asarray, ascontiguousarray, clip, diff, float16, \
    float32, float64, frombuffer, fromfile, greater, int16, int32, int64, linalg, maximum, minimum, mod, ndarray, \
    nonzero, ones, ones_like, percentile, sqrt, squeeze, std, subtract, transpose, uint8, unique, where, zeros, \
    zeros_like
from paddleocr import PaddleOCR
from pathvalidate import sanitize_filename
from prettytable import PrettyTable
from psd_tools import PSDImage
from psutil import virtual_memory
from pyautogui import locateOnScreen, locateAllOnScreen, center, click, press, keyDown, keyUp
from pyclipper import ET_CLOSEDPOLYGON, JT_ROUND, PyclipperOffset
from pypandoc import convert_text
from pytesseract import image_to_data, image_to_string
from pytz import UTC
from qtawesome import icon as qicon
from ruamel.yaml import YAML
from scipy import ndimage
from scipy.optimize import lsq_linear, nnls
from scipy.signal import argrelextrema
from shapely.geometry import LineString, MultiLineString, MultiPoint, Point, Polygon, box
from shapely.ops import nearest_points
from skimage.segmentation import watershed
from stdlib_list import stdlib_list
from unidecode import unidecode
from webcolors import hex_to_rgb, name_to_rgb, rgb_to_name
from webcolors._definitions import _reversedict, _CSS3_NAMES_TO_HEX

filterwarnings("ignore", category=DeprecationWarning)

use_torch = True
# use_torch = False
if use_torch:
    import torch

use_nlp = True
use_nlp = False
if use_nlp:
    import spacy

    nlp = spacy.load("en_core_web_sm")

# python3 -m spacy download en_core_web_sm #Mac
# python.exe -m spacy download en_core_web_sm #Win


# import nltk
# nltk.download('words')
# nltk.download('names')
# nltk.download('wordnet')
# nltk.download('omw-1.4')


# 合并常见单词和人名
good_words = set(words.words())
good_names = set(names.words())
good_words = good_words.union(good_names)

lemmatizer = WordNetLemmatizer()

python_ver = python_version()

# 创建转换实例，s2t表示简体到繁体，t2s表示繁体到简体
cc_s2t = OpenCC('s2t')  # 简体到繁体
cc_t2s = OpenCC('t2s')  # 繁体到简体


# ================================参数区================================
def a1_const():
    return


# Platforms
SYSTEM = ''
platform_system = system()
platform_uname = uname()
os_kernal = platform_uname.machine
if os_kernal in ['x86_64', 'AMD64']:
    if platform_system == 'Windows':
        SYSTEM = 'WINDOWS'
    elif platform_system == 'Linux':
        SYSTEM = 'LINUX'
    else:  # 'Darwin'
        SYSTEM = 'MAC'
else:  # os_kernal = 'arm64'
    if platform_system == 'Windows':
        SYSTEM = 'WINDOWS'
    elif platform_system == 'Darwin':
        SYSTEM = 'M1'
    else:
        SYSTEM = 'PI'

locale_tup = getdefaultlocale()
lang_code = locale_tup[0]

username = getuser()
homedir = expanduser("~")
homedir = Path(homedir)
DOWNLOADS = homedir / 'Downloads'
DOCUMENTS = homedir / 'Documents'

mac_address = ':'.join(findall('..', '%012x' % getnode()))
node_name = platform_uname.node

current_dir = dirname(abspath(__file__))
current_dir = Path(current_dir)

dirpath = os.getcwd()
ProgramFolder = Path(dirpath)
UserDataFolder = ProgramFolder / 'MomoHanhuaUserData'

python_vs = f"{sys.version_info.major}.{sys.version_info.minor}"

APP_NAME = 'MomoTranslator'
MAJOR_VERSION = 2
MINOR_VERSION = 0
PATCH_VERSION = 0
APP_VERSION = f'v{MAJOR_VERSION}.{MINOR_VERSION}.{PATCH_VERSION}'

APP_AUTHOR = '墨问非名'

if SYSTEM == 'WINDOWS':
    encoding = 'gbk'
    line_feed = '\n'
    cmct = 'ctrl'
else:
    encoding = 'utf-8'
    line_feed = '\n'
    cmct = 'command'

if SYSTEM in ['MAC', 'M1']:
    import applescript
    from Vision import VNImageRequestHandler, VNRecognizeTextRequest
    from Quartz import CGEventCreateMouseEvent, CGEventPost, kCGEventLeftMouseDown, kCGEventLeftMouseUp, \
        kCGMouseButtonLeft, kCGHIDEventTap, CGEventSetType, kCGRenderingIntentDefault
    from Quartz.CoreGraphics import CGDataProviderCreateWithData, CGColorSpaceCreateDeviceRGB, CGImageCreate, CGPoint

    processor_name = processor()
else:
    processor_name = machine()

if SYSTEM == 'WINDOWS':
    import pytesseract

    # 如果PATH中没有tesseract可执行文件，请指定tesseract路径
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

    from winocr import recognize_pil

line_feeds = line_feed * 2

lf = line_feed
lfs = line_feeds

ignores = ('~$', '._')

type_dic = {
    'xlsx': '.xlsx',
    'csv': '.csv',
    'pr': '.prproj',
    'psd': '.psd',
    'tor': '.torrent',
    'xml': '.xml',
    'audio': ('.aif', '.mp3', '.wav', '.flac', '.m4a', '.ogg'),
    'video': ('.mp4', '.mkv', '.avi', '.flv', '.mov', '.wmv'),
    'compressed': ('.zip', '.rar', '.7z', '.tar', '.gz', '.bz2'),
    'font': ('.ttc', '.ttf', '.otf'),
    'comic': ('.cbr', '.cbz', '.rar', '.zip', '.pdf', '.txt'),
    'pic': ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'),
    'log': '.log',
    'json': '.json',
    'pickle': '.pkl',
    'python': '.py',
    'txt': '.txt',
    'doc': ('.doc', '.docx'),
    'ppt': ('.ppt', '.pptx'),
    'pdf': '.pdf',
    'html': ('.html', '.htm'),
    'css': '.css',
    'js': '.js',
    'markdown': ('.md', '.markdown'),
    'yml': ('.yml', '.yaml'),
}

ram = str(round(virtual_memory().total / (1024.0 ** 3)))

video_width = 1920
video_height = 1080
video_size = (video_width, video_height)

pylupdate = 'pylupdate6'
lrelease = 'lrelease'

window_title_prefix = f'{APP_NAME} {APP_VERSION}'

py_path = Path(__file__).resolve()
py_dev_path = py_path.parent / f'{py_path.stem}_dev.py'
py_i18n_path = py_path.parent / f'{py_path.stem}_i18n.py'
py_frag_path = py_path.parent / f'{py_path.stem}_frag.py'
py_lite_path = py_path.parent / f'{py_path.stem}_lite.py'
py_struct_path = py_path.parent / f'{py_path.stem}_struct.py'

google_max_chars = 5000

pictures_exclude = '加框,分框,框,涂白,填字,修图,-,copy,副本,拷贝,顺序,打码,测试,标注,边缘,标志,伪造'
pic_tuple = tuple(pictures_exclude.split(','))

pre_tuple = (
    'zzz,'
)

scan_tuple = (
    'zSoU-Nerd',
    'zWater',
    'ZZZZZ',
    'zzzDQzzz',
    'zzz DQ zzz',
    'zzz LDK6 zzz',
    'zzz-mephisto',
    'zzz MollK6 zzz',
    'z',
    'zzz empire',
    'zzdelirium_dargh',
    'zzTLK',
    'zzz6 (Darkness-Empire)',
    'zfire',
)

spacing_ratio = 0.2

pos_0 = (0, 0)
pos_c = ('center', 'center')

color_blue = (255, 0, 0)
color_green = (0, 255, 0)
color_red = (0, 0, 255)
color_white = (255, 255, 255)
color_black = (0, 0, 0)

color_yellow = (0, 255, 255)  # 黄色
color_cyan = (255, 255, 0)  # 青色
color_magenta = (255, 0, 255)  # 洋红色
color_silver = (192, 192, 192)  # 银色
color_gray = (128, 128, 128)  # 灰色
color_maroon = (0, 0, 128)  # 褐红色
color_olive = (0, 128, 128)  # 橄榄色
color_purple = (128, 0, 128)  # 紫色
color_teal = (128, 128, 0)  # 蓝绿色
color_navy = (128, 0, 0)  # 海军蓝色
color_orange = (0, 165, 255)  # 橙色
color_pink = (203, 192, 255)  # 粉色
color_brown = (42, 42, 165)  # 棕色
color_gold = (0, 215, 255)  # 金色
color_lavender = (250, 230, 230)  # 薰衣草色
color_beige = (220, 245, 245)  # 米色
color_mint_green = (189, 255, 172)  # 薄荷绿
color_turquoise = (208, 224, 64)  # 绿松石色
color_indigo = (130, 0, 75)  # 靛蓝色
color_coral = (80, 127, 255)  # 珊瑚色
color_salmon = (114, 128, 250)  # 鲑鱼色
color_chocolate = (30, 105, 210)  # 巧克力色
color_tomato = (71, 99, 255)  # 番茄色
color_violet = (226, 43, 138)  # 紫罗兰色
color_goldenrod = (32, 165, 218)  # 金菊色
color_fuchsia = (255, 0, 255)  # 紫红色
color_crimson = (60, 20, 220)  # 深红色
color_dark_orchid = (204, 50, 153)  # 暗兰花色
color_slate_blue = (205, 90, 106)  # 石板蓝色
color_medium_sea_green = (113, 179, 60)  # 中等海洋绿色

rgba_white = (255, 255, 255, 255)
trans_white = (255, 255, 255, 127)
q_trans_white = QColor(255, 255, 255, 127)
rgba_zero = (0, 0, 0, 0)
rgba_black = (0, 0, 0, 255)

index_color = (0, 0, 255, 255)

trans_red = (255, 0, 0, 128)  # 半透明红色
trans_green = (0, 255, 0, 128)  # 半透明绿色
trans_purple = (128, 0, 128, 168)  # 半透明紫色
trans_yellow = (255, 255, 0, 128)  # 半透明黄色
trans_blue = (0, 255, 255, 128)  # 半透明蓝色
trans_olive = (0, 128, 128, 128)  # 半透明橄榄色

mark_px = [0, 0, 0, 1]

pad = 5
lower_bound_white = array([255 - pad, 255 - pad, 255 - pad])
upper_bound_white = array([255, 255, 255])
lower_bound_black = array([0, 0, 0])
upper_bound_black = array([pad, pad, pad])

# 定义一个字符串，包含常见的结束标点符号，不能出现在断句最前
not_start_punct = ',，.。;；:：?？!！”’·-》>:】【]、)）…'
# 定义一个字符串，包含常见的起始标点符号，不能出现在断句最后
not_end_punct = '“‘《<【[(（'
# 定义一个字符串，包含常见的结束字，不能出现在断句最前
not_start_char = '上中下内出完的地得了么呢吗嘛呗吧着个就前世里图们来'
# 定义一个字符串，包含常见的起始字，不能出现在断句最后
not_end_char = '太每帮跟另向'

valid_chars = set('.,;!?-—`~!@#$%^&*()_+={}[]|\\:;"\'<>,.?/')

punc_table_full = str.maketrans(r'：；，。！？“”‘’（）', r""":;,.!?""''()""")
punc_table_simple = str.maketrans(r'：；，。！？（）', r""":;,.!?()""")

proper_nouns = {
    'Insomnia',
}

sep_word = "supercalifragilisticexpialidocious"

fine_words = [
    'jai',
    'jin',
    'jor',
    'fer',
    'binnu',
]

replacements = {
    '‘': "'",
    '’': "'",
    '“': '"',
    '”': '"',
}

corrections = [
    ('LI', 'U'),
    ('X', 'Y'),
    # ('P', 'D'),
    # ('R', 'D'),
    ('ther', 'them'),
]

punct_rep_rules = {
    '.': ['…', '… ', '--'],
    '. ': ['…', '… ', '--'],
    ',': ['…', '.'],
    "'": ['“', '”', '"'],
    'I': [','],
    'l': ['j'],
    '‡': ['I'],
}

# 定义替换规则表
replace_rules = {
    '…': '...',
    'GET! T': 'GET IT',
    'WH!': 'NH!',
    'al/': 'all',
    '/T': 'IT',
    # '|': 'I',
    # '/': '!',
}

better_abbrs = {
    "YOL'RE": "YOU'RE",
    "YOL'LL": "YOU'LL",
    "YOL'VE": "YOU'VE",
    "WE'YE": "WE'VE",
    "WEIVE": "WE'VE",
    "WEVE": "WE'VE",
    "WHOVE": "WHO'VE",
    "L'VE": "I'VE",
    "L've": "I've",
    "Ijust": "I just",
    "IT'SA": "IT'S A",
    "IT'5": "IT'S",
    "IM": "I'M",
    "T'm": "I'M",
    "they'l": "they'll",
    "THEYIRE": "THEY'RE",
    "DONIT": "DON'T",
    "BOUT": "ABOUT",
    "ABOLIT": "ABOUT",
    "JLIST": "JUST",
    "SONLIVA": "SONUVA",
    "WORR'ED": "WORRIED",
    "SISTERIS": "SISTER'S",
    "LUCIFERIS": "LUCIFER'S",
    "CALLERIS": "CALLER'S",
    "IE": "IF",
    "I6": "IS",
    "SPELIEL": "SPECIAL",
    "LACIES": "LADIES",
    "LBT": "LBJ",
    'ALOCA': 'A LOCA',
    # 'WH': 'NH',
    'YOW': 'YOU',
    'OKAX': 'OKAY',
    'TO-': 'TO--',
    'Daramit': 'Dammit',
}

# ================语气词================
interjections = ('huh', 'hn')

src_head_2c = ['Source', '目标语言']

media_type_dict = {
    0: 'Comic',
    1: 'Manga',
}

media_lang_dict = {
    0: 'English',
    1: 'Chinese Simplified',
    2: 'Chinese Traditional',
    3: 'Japanese',
    4: 'Korean',
}

easyocr_language_map = {
    'English': 'en',
    'Chinese Simplified': 'ch_sim',
    'Japanese': 'ja',
}

# 南非语（af），阿塞拜疆语（az），波斯尼亚语（bs），简体中文（ch_sim），繁体中文（ch_tra），捷克语（cs），威尔士语（cy），丹麦语（da），德语（de），英语（en） ），西班牙语（es），爱沙尼亚语（et），法语（fr），爱尔兰语（ga），克罗地亚语（hr），匈牙利语（hu），印度尼西亚语（id），冰岛语（is），意大利语（it），日语（ja） ），韩文（ko），库尔德（ku），拉丁语（la），立陶宛语（lt），拉脱维亚语（lv），毛利人（mi），马来语（ms），马耳他语（mt），荷兰语（nl），挪威语（否），奥克西唐（oc），波兰语（pl），葡萄牙语（pt），罗马尼亚语（ro），塞尔维亚语（拉丁语）（rs_latin），斯洛伐克语（sk）（需要重新访问），斯洛文尼亚语（sl），阿尔巴尼亚语（sq），瑞典语（sv），斯瓦希里语（sw），泰语（th），他加禄语（tl），土耳其语（tr），乌兹别克（uz），越南语（vi）

paddleocr_language_map = {
    'English': 'en',
    'Chinese Simplified': 'ch',
    'Chinese Traditional': 'chinese_cht',
    'Japanese': 'japan',
}
paddle_langs = [
    'ch', 'en', 'korean', 'japan', 'chinese_cht', 'ta', 'te', 'ka', 'latin', 'arabic', 'cyrillic', 'devanagari',
]

similar_chars_map = {
    'а': 'a',
    'А': 'A',
    'Ä': 'A',
    'д': 'A',
    'в': 'B',
    'В': 'B',
    'ь': 'b',
    'Ь': 'B',
    'е': 'e',
    'Е': 'E',
    'н': 'H',
    'Н': 'H',
    'Ы': 'H',
    'І': 'I',
    'к': 'k',
    'К': 'K',
    'М': 'M',
    'П': 'N',
    'о': 'o',
    'О': 'O',
    'р': 'p',
    'Р': 'P',
    'с': 'c',
    'С': 'C',
    'т': 'T',
    'Т': 'T',
    'и': 'u',
    'И': 'U',
    'х': 'x',
    'Х': 'X',
    'у': 'y',
    'У': 'Y',
    'ч': '4',
    'Ч': '4',
}

special_keywords = [
    'setIcon', 'icon',
    'setShortcut', 'QKeySequence',
    'QSettings', 'value', 'setValue',
    'triggered',
    'setWindowTitle', 'windowTitle',
]

language_tuples = [
    # ================支持的语言================
    ('zh_CN', 'Simplified Chinese', '简体中文', '简体中文'),
    ('zh_TW', 'Traditional Chinese', '繁体中文', '繁體中文'),
    ('en_US', 'English', '英语', 'English'),
    ('ja_JP', 'Japanese', '日语', '日本語'),
    ('ko_KR', 'Korean', '韩语', '한국어'),
    # ================未来支持的语言================
    # ('es_ES', 'Spanish', '西班牙语', 'Español'),
    # ('fr_FR', 'French', '法语', 'Français'),
    # ('de_DE', 'German', '德语', 'Deutsch'),
    # ('it_IT', 'Italian', '意大利语', 'Italiano'),
    # ('pt_PT', 'Portuguese', '葡萄牙语', 'Português'),
    # ('ru_RU', 'Russian', '俄语', 'Русский'),
    # ('ar_AR', 'Arabic', '阿拉伯语', 'العربية'),
    # ('nl_NL', 'Dutch', '荷兰语', 'Nederlands'),
    # ('sv_SE', 'Swedish', '瑞典语', 'Svenska'),
    # ('tr_TR', 'Turkish', '土耳其语', 'Türkçe'),
    # ('pl_PL', 'Polish', '波兰语', 'Polski'),
    # ('he_IL', 'Hebrew', '希伯来语', 'עברית'),
    # ('da_DK', 'Danish', '丹麦语', 'Dansk'),
    # ('fi_FI', 'Finnish', '芬兰语', 'Suomi'),
    # ('no_NO', 'Norwegian', '挪威语', 'Norsk'),
    # ('hu_HU', 'Hungarian', '匈牙利语', 'Magyar'),
    # ('cs_CZ', 'Czech', '捷克语', 'Čeština'),
    # ('ro_RO', 'Romanian', '罗马尼亚语', 'Română'),
    # ('el_GR', 'Greek', '希腊语', 'Ελληνικά'),
    # ('id_ID', 'Indonesian', '印度尼西亚语', 'Bahasa Indonesia'),
    # ('th_TH', 'Thai', '泰语', 'ภาษาไทย'),
]

input_size = 1024
input_h = 1024
input_w = 1024
device = 'cpu'
scaleup = True
stride = 64

p_zh_char = re.compile(r'[^\u4e00-\u9fffA-Za-z，。、,\. ]')
p_zh = re.compile(r'[\u4e00-\u9fff]')
p_en = re.compile(r'\b[A-Za-z]+\b')
p_color = re.compile(r'([a-fA-F0-9]{6})-?(\d{0,3})', I)
p_issue_w_dot = re.compile(r'(.+?)(?!\d) (\d{2,5})', I)
p_num_chara = re.compile(r'(\d+)(\D+)')
p_comment = re.compile(r'^(\*|[①-⑨])')
p_lp_coor = re.compile(r'----------------\[(\d+)\]----------------\[(\d+\.\d+),(\d+\.\d+),(\d+)\]', I)

# 使用matplotlib的tab20颜色映射
colormap_tab20 = colormaps['tab20']

tesseract_language_options = {
    'Chinese Simplified': 'chi_sim',
    'Chinese Traditional': 'chi_tra',
    'English': 'eng',
    'Japanese': 'jpn',
    'Korean': 'kor',
}

vision_language_options = {
    'Chinese Simplified': 'zh-Hans',
    'Chinese Traditional': 'zh-Hant',
    'English': 'en',
    'Japanese': 'ja',
    'Korean': 'ko',
}

winocr_language_options = {
    'Chinese Simplified': 'zh-CN',
    'Chinese Traditional': 'zh-TW',
    'English': 'en',
    'Japanese': 'ja',
    'Korean': 'ko',
}

baidu_language_options = {
    'Chinese Simplified': 'CHN_ENG',
    'Chinese Traditional': 'CHN_ENG',
    'English': 'ENG',
    'Japanese': 'JPN',
    'Korean': 'KOR',
}

font2ps_dic = {
    '时尚中黑': 'TRENDS',
    '苹方': 'PingFangSC-Regular',
    '手札': 'HannotateSC-W5',
    '微软雅黑': 'MicrosoftYaHei',
}

# 创建一个新的 KeyboardEvent，模拟按下回车键。
# 'keydown' 是事件类型，表示一个按键被按下。
# 'key' 和 'code' 属性分别表示按下的按键和按键代码。
# 'which' 属性表示按键的字符编码。
# 'bubbles' 和 'cancelable' 属性表示事件是否可以冒泡和被取消。
press_enter_js = '''
var event = new KeyboardEvent('keydown', {
    'key': 'Enter',
    'code': 'Enter',
    'which': 13,
    'bubbles': true,
    'cancelable': true
});
Array.from(document.querySelectorAll('textarea'))[0].dispatchEvent(event);
'''

as_paste = f'''
delay 0.5
tell application "System Events"
    key code 9 using command down
end tell
'''

as_paste_n_enter = f'''
delay 0.5
tell application "System Events"
    key code 9 using command down
    delay 0.5
    key code 36
end tell
'''

as_enter = f'''
delay 0.5
tell application "System Events"
    key code 36
end tell
'''

as_funk = """
set soundName to "Funk"
do shell script "afplay /System/Library/Sounds/" & soundName & ".aiff"
"""

as_submarine = """
set soundName to "Submarine"
do shell script "afplay /System/Library/Sounds/" & soundName & ".aiff"
"""

as_Tingting_uploaded = f"""
say "全部上传完毕" speaking rate 180
"""

button_js_code = """
var buttons = Array.from(document.querySelectorAll('button[aria-label=\'附加文件\']'));
if (buttons.length > 0) {
    buttons[0].click();
}
console.log('找到按钮数量：', buttons.length);
"""

chatGPT4o_prompt = """请分画格详细描述以下内容：
1. 描述每个画格的场景和剧情。
2. 描述每个画格中的角色的外貌、服装、情绪，和角色之间的互动和动作。
3. 描述漫画的视觉效果，例如颜色、光线、阴影等。
4. 提供每个画格中出现的文本，标注说话人，并将文本翻译成中文。"""

chatGPT4o_prompt = """请分画格详细描述以下内容：
1. 描述每个画格的场景和剧情。
2. 提供每个画格中出现的文本，标注说话人，并将文本翻译成中文。"""

chatGPT4o_prompt = """请分画格详细描述每个画格的场景和剧情，提供每个画格中出现的文本，标注说话人，并将文本翻译成中文。"""

prompt_1st_line = chatGPT4o_prompt.splitlines()[0]

chatgpt_prefix = 'https://chatgpt.com/'

gpt4o_spec_str = '这张图片是来自漫画'


# ================================基础函数区================================
def a2_base():
    return


def kernel(size):
    return ones((size, size), uint8)


def kernel_hw(h, w):
    return ones((h, w), uint8)


kernel1 = kernel(1)
kernel2 = kernel(2)
kernel3 = kernel(3)
kernel4 = kernel(4)
kernel5 = kernel(5)
kernel6 = kernel(6)
kernel7 = kernel(7)
kernel8 = kernel(8)
kernel9 = kernel(9)
kernel10 = kernel(10)
kernel11 = kernel(11)
kernel12 = kernel(12)
kernel13 = kernel(13)
kernel14 = kernel(14)
kernel15 = kernel(15)
kernel20 = kernel(20)
kernel25 = kernel(25)
kernel30 = kernel(30)
kernel35 = kernel(35)
kernel40 = kernel(40)
kernel50 = kernel(50)
kernel60 = kernel(60)
kernalh1w5 = kernel_hw(1, 5)
kernalh1w6 = kernel_hw(1, 6)
kernalh1w7 = kernel_hw(1, 7)
kernalh1w8 = kernel_hw(1, 8)
kernalh1w9 = kernel_hw(1, 9)
kernalh1w10 = kernel_hw(1, 10)
kernalh1w11 = kernel_hw(1, 11)
kernalh1w12 = kernel_hw(1, 12)
kernalh1w13 = kernel_hw(1, 13)
kernalh1w14 = kernel_hw(1, 14)
kernalh1w15 = kernel_hw(1, 15)
kernalh1w16 = kernel_hw(1, 16)
kernalh1w17 = kernel_hw(1, 17)
kernalh1w18 = kernel_hw(1, 18)
kernalh1w19 = kernel_hw(1, 19)
kernalh1w20 = kernel_hw(1, 20)
kernalh5w1 = kernel_hw(5, 1)
kernalh8w1 = kernel_hw(8, 1)
kernalh10w1 = kernel_hw(10, 1)

CSS3_HEX_TO_NAMES = _reversedict(_CSS3_NAMES_TO_HEX)


def timer_decorator(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time()
        result = func(*args, **kwargs)
        elapsed_time = time() - start_time

        hours, rem = divmod(elapsed_time, 3600)
        minutes, seconds = divmod(rem, 60)

        if hours > 0:
            show_run_time = f"{int(hours)}时{int(minutes)}分{seconds:.2f}秒"
        elif minutes > 0:
            show_run_time = f"{int(minutes)}分{seconds:.2f}秒"
        else:
            show_run_time = f"{seconds:.2f}秒"

        logger.debug(f"{func.__name__} took: {show_run_time}")
        return result

    return wrapper


def is_decimal_or_comma(s):
    pattern = r'^\d*\.?\d*$|^\d*[,]?\d*$'
    return bool(match(pattern, s))


def is_valid_file(file_path, suffixes):
    if not file_path.is_file():
        return False
    if not file_path.stem.startswith(ignores):
        if suffixes:
            return file_path.suffix.lower() in suffixes
        else:
            return True
    return False


def printe(e):
    print(e)
    logger.error(e)
    print_exc()


def clamp(value, min_value, max_value):
    return max(min_value, min(value, max_value))


def reduce_list(input_list):
    try:
        # 尝试使用dict.fromkeys方法
        output_list = list(OrderedDict.fromkeys(input_list))
    except TypeError:
        # 如果发生TypeError（可能是因为列表包含不可哈希的对象），
        # 则改用更慢的方法
        output_list = []
        for input in input_list:
            if input not in output_list:
                output_list.append(input)
    return output_list


# ================创建目录================
def make_dir(file_path):
    if not exists(file_path):
        try:
            os.makedirs(file_path)
        except BaseException as e:
            print(e)


# ================获取文件夹列表================
def get_dirs(rootdir):
    dirs_list = []
    if rootdir and rootdir.exists():
        # 列出目录下的所有文件和目录
        lines = os.listdir(rootdir)
        for line in lines:
            filepath = Path(rootdir) / line
            if filepath.is_dir():
                dirs_list.append(filepath)
    dirs_list.sort()
    return dirs_list


def get_files(rootdir, file_type=None, direct=False):
    rootdir = Path(rootdir)
    file_paths = []

    # 获取文件类型的后缀
    # 默认为所有文件
    suffixes = type_dic.get(file_type, file_type)
    if isinstance(suffixes, str):
        suffixes = (suffixes,)

    # 如果根目录存在
    if rootdir and rootdir.exists():
        # 只读取当前文件夹下的文件
        if direct:
            files = os.listdir(rootdir)
            for file in files:
                file_path = Path(rootdir) / file
                if is_valid_file(file_path, suffixes):
                    file_paths.append(file_path)
        # 读取所有文件
        else:
            for root, dirs, files in os.walk(rootdir):
                for file in files:
                    file_path = Path(root) / file
                    if is_valid_file(file_path, suffixes):
                        file_paths.append(file_path)

    # 使用natsorted()进行自然排序，
    # 使列表中的字符串按照数字顺序进行排序
    file_paths = natsorted(file_paths)
    return file_paths


def filter_items(old_list, prefix=pre_tuple, infix=scan_tuple, suffix=pic_tuple, item_attr='stem'):
    """
    这个函数用于过滤一个列表，根据指定的前缀、中缀和后缀来排除不需要的元素。
    可以根据文件的全名或者文件名（不包括扩展名）来进行过滤。

    :param old_list: 原始列表。
    :param prefix: 要排除的前缀元组。
    :param infix: 要排除的中间文本元组。
    :param suffix: 要排除的后缀元组。
    :param item_attr: 'name' 或 'stem'，基于文件全名或仅基于文件主名进行过滤。
    :return: 过滤后的新列表，不包含任何匹配前缀、中缀或后缀的元素。
    """

    # 定义一个内部函数来判断一个元素是否应该被排除
    def is_excluded(item):
        # 检查元素是否以任何给定的前缀开始
        for p in prefix:
            if item.startswith(p):
                return True
        # 检查元素的名字是否包含任何给定的中缀
        for i in infix:
            if i == item:
                return True
        # 检查元素是否以任何给定的后缀结束
        for s in suffix:
            if item.endswith(s):
                return True
        # 如果元素不匹配任何排除规则，则不应该排除
        return False

    # 使用列表推导式来过滤原始列表
    # 对于列表中的每一个元素，我们先获取其指定的属性（'name'或'stem'），然后检查是否应该排除
    filtered_list = [item for item in old_list if not is_excluded(getattr(item, item_attr))]
    return filtered_list


@logger.catch
def get_valid_imgs(rootdir, vmode='raw'):
    all_pics = get_files(rootdir, 'pic', True)
    jpgs = [x for x in all_pics if x.suffix.lower() in ('.jpg', '.jpeg')]
    pngs = [x for x in all_pics if x.suffix.lower() == '.png']

    all_masks = [x for x in pngs if '-Mask-' in x.stem]
    all_Whitens = [x for x in pngs if x.stem.endswith('-Whiten')]
    exclude_pngs = all_masks + all_Whitens
    no_masks = [x for x in pngs if '-Mask-' not in x.stem]

    valid_jpgs = filter_items(jpgs)
    valid_pngs = filter_items(no_masks)

    valid_img_list = [x for x in all_pics if x not in exclude_pngs]
    valid_img_list = filter_items(valid_img_list)

    if vmode == 'raw':
        valid_imgs = valid_img_list
    else:
        valid_imgs = all_masks
    return valid_imgs


# @lru_cache
def iload_data(file_path):
    data_dic = {}
    if file_path.exists():
        if file_path.suffix == '.yml':
            with open(file_path, 'r', encoding='utf-8') as file:
                data_dic = yaml.safe_load(file)
        elif file_path.suffix == '.pkl':
            with open(file_path, 'rb') as file:
                data_dic = pickle.load(file)
    return data_dic


# ================读取文本================
def read_txt(file_path, encoding='utf-8'):
    """
    读取指定文件路径的文本内容。

    :param file_path: 文件路径
    :param encoding: 文件编码，默认为'utf-8'
    :return: 返回读取到的文件内容，如果文件不存在则返回None
    """
    file_content = None
    if file_path.exists():
        with open(file_path, mode='r', encoding=encoding) as file_object:
            file_content = file_object.read()
    return file_content


# ================写入文件================
def write_txt(file_path, text_input, encoding='utf-8', ignore_empty=True):
    """
    将文本内容写入指定的文件路径。

    :param file_path: 文件路径
    :param text_input: 要写入的文本内容，可以是字符串或字符串列表
    :param encoding: 文件编码，默认为'utf-8'
    :param ignore_empty: 是否忽略空内容，默认为True
    """
    if text_input:
        save_text = True
        if isinstance(text_input, list):
            otext = lf.join(text_input)
        else:
            otext = text_input
        file_content = read_txt(file_path, encoding)
        if file_content == otext or (ignore_empty and otext == ''):
            save_text = False
        if save_text:
            with open(file_path, mode='w', encoding=encoding, errors='ignore') as f:
                f.write(otext)


def generate_md5(img_array):
    img_data = imencode('.png', img_array)[1].tostring()
    file_hash = md5()
    file_hash.update(img_data)
    return file_hash.hexdigest()


# ================对文件算MD5================
def md5_w_size(path, blksize=2 ** 20):
    if isfile(path) and exists(path):  # 判断目标是否文件,及是否存在
        file_size = getsize(path)
        if file_size <= 256 * 1024 * 1024:  # 512MB
            with open(path, 'rb') as f:
                cont = f.read()
            hash_object = md5(cont)
            t_md5 = hash_object.hexdigest()
            return t_md5, file_size
        else:
            m = md5()
            with open(path, 'rb') as f:
                while True:
                    buf = f.read(blksize)
                    if not buf:
                        break
                    m.update(buf)
            t_md5 = m.hexdigest()
            return t_md5, file_size
    else:
        return None


def write_csv(csv_path, data_input, headers=None):
    temp_csv = csv_path.parent / 'temp.csv'

    try:
        if isinstance(data_input, list):
            if len(data_input) >= 1:
                if csv_path.exists():
                    with codecs.open(temp_csv, 'w', 'utf_8_sig') as f:
                        f_csv = writer(f, delimiter=',', quotechar='"', quoting=QUOTE_MINIMAL, escapechar='\\')
                        if headers:
                            f_csv.writerow(headers)
                        f_csv.writerows(data_input)
                    if md5_w_size(temp_csv) != md5_w_size(csv_path):
                        copy2(temp_csv, csv_path)
                    if temp_csv.exists():
                        os.remove(temp_csv)
                else:
                    with codecs.open(csv_path, 'w', 'utf_8_sig') as f:
                        f_csv = writer(f, delimiter=',', quotechar='"', quoting=QUOTE_MINIMAL, escapechar='\\')
                        if headers:
                            f_csv.writerow(headers)
                        f_csv.writerows(data_input)
        else:  # DataFrame
            if csv_path.exists():
                data_input.to_csv(temp_csv, encoding='utf-8', index=False)
                if md5_w_size(temp_csv) != md5_w_size(csv_path):
                    copy2(temp_csv, csv_path)
                if temp_csv.exists():
                    os.remove(temp_csv)
            else:
                data_input.to_csv(csv_path, encoding='utf-8', index=False)
    except BaseException as e:
        printe(e)


def conv_img(img, target_format='PIL'):
    """
    将图像转换为指定的格式。

    :param img: 输入图像，可以是 NumPy 数组或 PIL 图像。
    :param target_format: 目标格式，可以是 'PIL' 或 'CV'。
    :return: 转换后的图像。
    """
    if target_format == 'PIL':
        if isinstance(img, ndarray):
            # 转换 NumPy 数组为 PIL 图像
            if len(img.shape) == 2:  # 灰度或黑白图像
                cimg = Image.fromarray(img, 'L')
            else:  # if len(img.shape) == 3:  # 彩色图像
                cimg = Image.fromarray(img, 'RGB')
        else:  # isinstance(img, Image.Image)
            cimg = img
    else:
        # 如果是PIL图像，转换为NumPy数组
        if isinstance(img, Image.Image):
            cimg = array(img)
            # 如果图像有三个维度，并且颜色为三通道，则进行颜色空间的转换
            if cimg.ndim == 3 and cimg.shape[2] == 3:
                cimg = cvtColor(cimg, COLOR_RGB2BGR)
        else:  # isinstance(img, ndarray)
            cimg = img
    return cimg


@logger.catch
def write_pic(pic_path, picimg):
    pic_path = Path(pic_path)
    ext = pic_path.suffix
    temp_pic = pic_path.parent / f'{pic_path.stem}-temp{ext}'

    # 检查输入图像的类型
    if isinstance(picimg, bytes):
        # 如果是字节对象，直接写入文件
        with open(temp_pic, 'wb') as f:
            f.write(picimg)
    else:
        # 如果是PIL图像，转换为NumPy数组
        if isinstance(picimg, Image.Image):
            picimg = array(picimg)

            # 如果图像有三个维度，并且颜色为三通道，则进行颜色空间的转换
            if picimg.ndim == 3 and picimg.shape[2] == 3:
                picimg = cvtColor(picimg, COLOR_RGB2BGR)

        # 检查图像是否为空
        if picimg is None or picimg.size == 0:
            logger.error(f'{pic_path=}')
            # raise ValueError("The input image is empty.")
            return pic_path

        # 保存临时图像
        imencode(ext, picimg)[1].tofile(temp_pic)
    # 检查临时图像和目标图像的md5哈希和大小是否相同
    if not pic_path.exists() or md5_w_size(temp_pic) != md5_w_size(pic_path):
        copy2(temp_pic, pic_path)
    # 删除临时图像
    if temp_pic.exists():
        os.remove(temp_pic)
    return pic_path


# #@logger.catch
def write_docx(docx_path, docu):
    temp_docx = docx_path.parent / 'temp.docx'
    if docx_path.exists():
        docu.save(temp_docx)
        if md5_w_size(temp_docx) != md5_w_size(docx_path):
            copy2(temp_docx, docx_path)
        if temp_docx.exists():
            os.remove(temp_docx)
    else:
        docu.save(docx_path)


def write_yml(yml_path, data):
    temp_yml = yml_path.parent / 'temp.yml'
    if yml_path.exists():
        with open(temp_yml, 'w', encoding='utf-8') as temp_file:
            yaml.dump(data, temp_file, default_flow_style=False, allow_unicode=True)
        if not cmp(temp_yml, yml_path):
            copy2(temp_yml, yml_path)
        if temp_yml.exists():
            os.remove(temp_yml)
    else:
        with open(yml_path, 'w', encoding='utf-8') as yaml_file:
            yaml.dump(data, yaml_file, default_flow_style=False, allow_unicode=True)


def common_prefix(strings):
    """
    返回字符串列表中的共同前缀。

    :param strings: 字符串列表
    :return: 共同前缀
    """
    # pprint(strings)
    if not strings:
        return ""
    common_prefix = strings[0]
    for s in strings[1:]:
        while not s.startswith(common_prefix):
            common_prefix = common_prefix[:-1]
            if not common_prefix:
                return ""
    return common_prefix


def common_suffix(strings):
    """
    返回字符串列表中的共同后缀。

    :param strings: 字符串列表
    :return: 共同后缀
    """
    if not strings:
        return ""
    common_suffix = strings[0]
    for s in strings[1:]:
        while not s.endswith(common_suffix):
            common_suffix = common_suffix[1:]
            if not common_suffix:
                return ""
    return common_suffix


# @logger.catch
def parse_range(range_str):
    # 使用"~"分割字符串，并转化为浮点数列表
    # logger.warning(f'{range_str=}')
    range_strs = range_str.split('~')
    ranges = [float(x) for x in range_strs]
    return ranges


@logger.catch
def find_nth_largest(nums, n):
    if len(nums) < n:
        return None, None
    # 使用enumerate获取元素及其索引，并按值排序
    sorted_nums = sorted(enumerate(nums), key=lambda x: x[1], reverse=True)
    # 获取第N大的元素（注意列表索引从0开始，所以要用n-1）
    nth_largest = sorted_nums[n - 1]
    # nth_largest是一个元组，其中第一个元素是原始索引，第二个元素是值
    original_index, value = nth_largest
    return value, original_index


def lcs(X, Y):
    """
    计算两个字符串X和Y的最长公共子序列（Longest Common Subsequence, LCS）。

    :param X: 第一个需要比较的字符串。
    :param Y: 第二个需要比较的字符串。
    :return: 返回最长公共子序列。
    """
    m = len(X)  # X的长度
    n = len(Y)  # Y的长度

    # 初始化一个二维列表L，用于存储子问题的解
    L = [[0] * (n + 1) for i in range(m + 1)]

    # 动态规划填表
    for i in range(m + 1):
        for j in range(n + 1):
            if i == 0 or j == 0:
                L[i][j] = 0
            elif X[i - 1] == Y[j - 1]:
                L[i][j] = L[i - 1][j - 1] + 1
            else:
                L[i][j] = max(L[i - 1][j], L[i][j - 1])

    # 从填好的表中构造LCS
    index = L[m][n]
    lcs = [''] * (index + 1)
    lcs[index] = ''
    i = m
    j = n
    while i > 0 and j > 0:
        if X[i - 1] == Y[j - 1]:
            lcs[index - 1] = X[i - 1]
            i -= 1
            j -= 1
            index -= 1
        elif L[i - 1][j] > L[i][j - 1]:
            i -= 1
        else:
            j -= 1

    return ''.join(lcs)


def color_diff(str1, str2):
    """
    对比两个字符串，并用颜色标出第一个字符串中与第二个字符串不同的字符。

    :param str1: 第一个需要对比的字符串。
    :param str2: 第二个需要对比的字符串。
    :return: 返回第一个字符串，其中与第二个字符串不同的字符将被标红。
    """
    common = lcs(str1, str2)  # 计算最长公共子序列
    i = 0
    colored_str = ""
    for c in str1:
        if i < len(common) and c == common[i]:
            colored_str += c  # 相同字符不变
            i += 1
        else:
            colored_str += f"\033[91m{c}\033[0m"  # 不同字符标红

    return colored_str


def convert_str2num(s):
    if '.' in s:
        return float(s)
    else:
        return int(s)


# @logger.catch
def form2data(ocr_zdata_form):
    ocr_zdata = []
    for t in range(len(ocr_zdata_form)):
        row = ocr_zdata_form[t]
        if row and '|' in row:
            row_nums_str, par, row_text = row.partition('|')
            row_nums = row_nums_str.split(',')
            # 使用列表推导式将字符串转换为相应的整数或小数
            row_nums = [convert_str2num(x) for x in row_nums]
            tess_result = row_nums + [row_text]
            ocr_zdata.append(tess_result)
        else:
            logger.error(f'[{t}]{row=}')
    return ocr_zdata


def iread_csv(csv_file, pop_head=True, get_head=False):
    # 使用'rb'模式读取文件内容，然后解码为字符串，同时移除NUL字符
    with open(csv_file, 'rb') as file:
        file_content = file.read().decode('utf-8').replace('\x00', '')
    # 使用字符串IO模拟文件对象
    f = StringIO(file_content)
    f_csv = reader(f)
    if pop_head:
        # 获取首行并在需要时将其从数据中删除
        head = next(f_csv, [])
    else:
        head = []
    # 使用列表推导式简化数据读取
    idata = [tuple(row) for row in f_csv]
    if get_head:
        return idata, head
    else:
        return idata


# ================================基础图像函数区================================
def a3_pic():
    return


def rect2poly(x, y, w, h):
    # 四个顶点为：左上，左下，右下，右上
    points = [
        (x, y),  # 左上
        (x, y + h),  # 左下
        (x + w, y + h),  # 右下
        (x + w, y),  # 右上
    ]
    return points


def hex2int(hex_num):
    hex_num = f'0x{hex_num}'
    int_num = int(hex_num, 16)
    return int_num


def rgb2str(rgb_tuple):
    r, g, b = rgb_tuple
    color_str = f'{r:02x}{g:02x}{b:02x}'
    return color_str


def toBGR(img_raw):
    # 检查图像的维度（颜色通道数）
    if len(img_raw.shape) == 2:
        # 图像是灰度图（只有一个颜色通道），将其转换为BGR
        img_raw = cvtColor(img_raw, COLOR_GRAY2BGR)
    elif img_raw.shape[2] == 3:
        # 图像已经是BGR格式（有三个颜色通道），不需要转换
        pass
    elif img_raw.shape[2] == 4:
        # 图像是BGRA格式（有四个颜色通道），将其转换为BGR，移除Alpha通道
        img_raw = cvtColor(img_raw, COLOR_BGRA2BGR)
    return img_raw


def color2rgb(color):
    if isinstance(color, str):
        if color.startswith("#"):
            # 十六进制颜色
            return hex_to_rgb(color)
        else:
            # 颜色名称
            return name_to_rgb(color)
    elif isinstance(color, (tuple, list, ndarray)):
        color = tuple(color)  # 如果输入是NumPy数组，将其转换为元组
        if len(color) == 3:
            # RGB颜色
            return color
        elif len(color) == 4:
            # RGBA颜色，忽略透明度
            return color[:3]
    else:
        raise ValueError("不支持的颜色格式")


def closest_color(requested_color):
    min_colors = {}
    for key, name in CSS3_HEX_TO_NAMES.items():
        r_c, g_c, b_c = hex_to_rgb(key)
        rd = (r_c - requested_color[0]) ** 2
        gd = (g_c - requested_color[1]) ** 2
        bd = (b_c - requested_color[2]) ** 2
        min_colors[(rd + gd + bd)] = name
    return min_colors[min(min_colors.keys())]


# @logger.catch
def get_color_name(color):
    rgb_color = color2rgb(color)
    try:
        color_name = rgb_to_name(rgb_color)
    except ValueError:
        color_name = closest_color(rgb_color)
    return color_name


def print_colored_text(text, color_str):
    # 将16进制颜色转换为RGB
    rgb_color = color2rgb(color_str)
    r, g, b = rgb_color

    # 使用ANSI转义序列设置文本颜色
    ansi_color = f"\033[38;2;{r};{g};{b}m"
    reset_color = "\033[0m"

    # 打印彩色文本
    print(f"{ansi_color}{text}{reset_color}")


def highlight_diffs(content0, content):
    table = PrettyTable()
    table.field_names = ["Index", "Original Data", "Modified Data"]

    # 遍历content0和content的元素
    for idx, (c0, c) in enumerate(zip(content0, content), 0):
        # 如果元素不同，则添加高亮
        if c0 != c:
            table.add_row([idx, f"\033[1;31m{c0}\033[0m", f"\033[1;31m{c}\033[0m"])
        else:
            table.add_row([idx, c0, c])


def idx2label(idx):
    if 0 <= idx < 26:
        return string.ascii_uppercase[idx]
    elif 26 <= idx < 52:
        return string.ascii_lowercase[idx - 26]
    else:
        return str(idx)


def pt2tup(point):
    return (int(point.x), int(point.y))


def get_dist2rect(point, rect):
    x, y, w, h = rect
    dx = max(x - point.x, 0, point.x - (x + w))
    dy = max(y - point.y, 0, point.y - (y + h))
    return sqrt(dx ** 2 + dy ** 2)


def get_poly_by_cnt(contour_polys, cnt):
    # 这里假设每个 cnt 在 contour_polys 中只对应一个 poly
    for poly in contour_polys:
        if poly.cnt == cnt:
            return poly
    return None


def crop_img(src_img, br, pad=0):
    # 输入参数:
    # src_img: 原始图像(numpy array)
    # br: 裁剪矩形(x, y, w, h)，分别代表左上角坐标(x, y)以及宽度和高度
    # pad: 额外填充，默认值为0

    x, y, w, h = br
    ih, iw = src_img.shape[0:2]

    # 计算裁剪区域的边界坐标，并确保它们不超过图像范围
    y_min = clip(y - pad, 0, ih - 1)
    y_max = clip(y + h + pad, 0, ih - 1)
    x_min = clip(x - pad, 0, iw - 1)
    x_max = clip(x + w + pad, 0, iw - 1)

    # 使用numpy的切片功能对图像进行裁剪
    cropped = src_img[y_min:y_max, x_min:x_max]
    return cropped


def get_clipboard_data():
    """
    从剪贴板获取数据
    :return 如果没有错误，返回剪贴板数据；否则返回 None
    """
    process = Popen(['pbpaste'], stdout=PIPE, stderr=PIPE)
    output, error = process.communicate()

    if not error:
        return output.decode('utf-8')
    else:
        print("Error:", error)
        return None


def get_ext(zipped):
    lower = [min(x) for x in zipped]
    upper = [max(x) for x in zipped]
    ext_lower = [round(max(x - ext_padding, 0)) for x in lower]
    ext_upper = [round(min(x + ext_padding, 255)) for x in upper]
    ext_lower = tuple(reversed(ext_lower))
    ext_upper = tuple(reversed(ext_upper))
    ext = (ext_lower, ext_upper)
    return ext


def get_sample(video_img, left_rgb, right_rgb):
    zipped = list(zip(left_rgb, right_rgb))
    ext_lower, ext_upper = get_ext(zipped)
    logger.debug(f'{ext_lower=}, {ext_upper=}')
    sample_mask = inRange(video_img, array(ext_lower), array(ext_upper))
    return sample_mask


# ================================图像函数区================================
def a4_apple_script():
    return


def remove_common_indent(script):
    """
    删除脚本中每行开头的相同长度的多余空格
    :param script: 要处理的脚本
    :return: 删除多余空格后的脚本
    """
    lines = script.split('\n')
    min_indent = min(len(line) - len(line.lstrip()) for line in lines if line.strip())
    return '\n'.join(line[min_indent:] for line in lines)


@logger.catch
@timer_decorator
def run_apple_script(script):
    """
    执行 AppleScript 脚本
    :param script：要执行的 AppleScript 脚本
    :return 如果执行成功，返回执行结果；否则返回 None
    """
    if script:
        script = remove_common_indent(script)
        logger.debug(f'{script}')
        result = applescript.run(script)

        if result.code == 0:
            return result.out
        else:
            print(f'{result.err=}')


def get_browser_current_tab_url(browser):
    """
    获取浏览器当前标签页的 URL
    :param browser：浏览器名称，可以是 'Safari' 或 'Google Chrome'
    :return 当前标签页的 URL
    """
    if browser == 'Safari':
        apple_script = f'''
        tell application "{browser}"
            set current_url to URL of front document
            return current_url
        end tell
        '''
    elif browser.startswith('Google Chrome'):
        apple_script = f'''
        tell application "{browser}"
            set current_url to URL of active tab of front window
            return current_url
        end tell
        '''
    else:
        print(f"Error: Unsupported browser {browser}.")
        return None
    return run_apple_script(apple_script)


def get_browser_current_tab_title(browser):
    """
    获取浏览器当前标签页的标题
    :param browser：浏览器名称，可以是 'Safari' 或 'Chrome'
    :return 当前标签页的标题
    """
    if browser == 'Safari':
        apple_script = f'''
        tell application "{browser}"
            set current_title to name of front document
            return current_title
        end tell
        '''
    elif browser.startswith('Google Chrome'):
        apple_script = f'''
        tell application "{browser}"
            set current_title to title of active tab of front window
            return current_title
        end tell
        '''
    else:
        print(f"Error: Unsupported browser {browser}.")
        return None

    return run_apple_script(apple_script)


def get_browser_current_tab_html(browser, activate_browser=False):
    """
    获取浏览器当前标签页的 HTML 内容
    :param browser：浏览器名称，可以是 'Safari' 或 'Chrome'
    :return 当前标签页的 HTML 内容
    """
    js_code = "document.documentElement.outerHTML;"

    # 如果需要激活浏览器，则设置 activate_command 为 "activate"，否则为空字符串。
    activate_command = "activate" if activate_browser else ""

    if browser == 'Safari':
        apple_script = f'''
        tell application "{browser}"
            {activate_command}
            set curr_tab to current tab of front window
            do JavaScript "{js_code}" in curr_tab
            set the clipboard to result
            return (the clipboard as text)
        end tell
        '''
    elif browser.startswith('Google Chrome'):
        apple_script = f'''
        tell application "{browser}"
            {activate_command}
            set curr_tab to active tab of front window
            execute curr_tab javascript "{js_code}"
            set the clipboard to result
            return (the clipboard as text)
        end tell
        '''
    else:
        print(f"Error: Unsupported browser {browser}.")
        return None

    return run_apple_script(apple_script)


def open_html_file_in_browser(file_path, browser):
    """
    在浏览器中打开 HTML 文件
    :param file_path：HTML 文件的路径
    :param browser：浏览器名称，可以是 'Safari' 或 'Chrome'
    """
    apple_script = f'''
    tell application "{browser}"
        activate
        open POSIX file "{file_path}"
    end tell
    '''
    return run_apple_script(apple_script)


@timer_decorator
@logger.catch
def save_from_browser(browser: str):
    """
    主函数
    :param browser：浏览器名称，可以是 'Safari' 或 'Chrome'
    """
    current_url = get_browser_current_tab_url(browser)
    title = get_browser_current_tab_title(browser)
    logger.debug(f'{current_url=}')
    logger.warning(f'{title=}')
    if current_url and title:
        if current_url.startswith(chatgpt_prefix):
            safe_title = title.replace('/', '／').replace('\\', '＼')
            chatgpt_html = ChatGPT / f'{sanitize_filename(safe_title)}-{Path(current_url).stem}.html'
            content = get_browser_current_tab_html(browser)
            # logger.info(f'{content}')
            soup = BeautifulSoup(content, 'html.parser')
            pretty_html = soup.prettify()
            write_txt(chatgpt_html, pretty_html)
    return chatgpt_html


# ================================图像函数区================================
def a5_frame():
    return


def get_edge_pxs(img_raw):
    h, w = img_raw.shape[:2]
    # 转换为RGBA格式
    img_rgba = cvtColor(img_raw, COLOR_BGR2BGRA)
    # 将非边框像素的alpha设置为0
    mask = ones((h, w), dtype=bool)
    mask[edge_size:-edge_size, edge_size:-edge_size] = False
    img_rgba[~mask] = [0, 0, 0, 0]
    return img_rgba


def find_dominant_color(edge_pixels_rgba):
    # 获取边框像素
    border_pixels = edge_pixels_rgba[where(edge_pixels_rgba[..., 3] != 0)]
    # 计算每种颜色的出现次数
    colors, counts = unique(border_pixels, axis=0, return_counts=True)
    # 获取出现次数最多的颜色的索引
    dominant_color_index = argmax(counts)
    # 计算占据边框面积的比例
    h, w = edge_pixels_rgba.shape[:2]
    frame_pxs = 2 * (h + w) * edge_size - 4 * edge_size ** 2
    color_ratio = counts[dominant_color_index] / frame_pxs
    dominant_color = colors[dominant_color_index]
    return color_ratio, dominant_color


def find_dominant_colors(edge_pixels_rgba, tolerance=10, top_colors=10):
    # 获取边框像素
    border_pixels = edge_pixels_rgba[where(edge_pixels_rgba[..., 3] != 0)]
    # 计算每种颜色的出现次数
    colors, counts = np.unique(border_pixels, axis=0, return_counts=True)

    h, w = edge_pixels_rgba.shape[:2]
    frame_pxs = 2 * (h + w) * edge_size - 4 * edge_size ** 2

    # 计算 color_ratios
    color_ratios = counts / frame_pxs

    # 创建包含颜色和 color_ratio 的元组列表
    color_and_ratios = sorted([(color, ratio) for color, ratio in zip(colors, color_ratios)], key=lambda x: -x[1])

    # 只保留前 10 种出现次数最多的颜色
    color_and_ratios = color_and_ratios[:top_colors]
    color_and_ratios_raw = deepcopy(color_and_ratios)

    combined_color_and_ratios = []
    while len(color_and_ratios) > 0:
        current_color, current_ratio = color_and_ratios.pop(0)
        combined_color_and_ratios.append((current_color, current_ratio))

        similar_color_indices = []
        for i, (color, ratio) in enumerate(color_and_ratios):
            if np.all(np.abs(current_color[:3] - color[:3]) <= tolerance):
                current_ratio += ratio
                similar_color_indices.append(i)

        # 移除已归类的相似颜色
        for index in sorted(similar_color_indices, reverse=True):
            color_and_ratios.pop(index)

    return color_and_ratios_raw, combined_color_and_ratios


def compute_frame_mask_single(img_raw, dominant_color, tolerance):
    dominant_color_int32 = array(dominant_color[:3]).astype(int32)
    lower_bound = maximum(0, dominant_color_int32 - tolerance)
    upper_bound = minimum(255, dominant_color_int32 + tolerance)
    mask = inRange(img_raw, lower_bound, upper_bound)
    return mask


def examine_white_line(row, min_dist, top_n, min_length):
    white_segments = np.split(row, where(diff(row) != 1)[0] + 1)
    sorted_segments = sorted(white_segments, key=len, reverse=True)

    top_segments = sorted_segments[:top_n]
    top_lengths = [len(segment) for segment in top_segments]

    valid_segments = []
    for idx, length in enumerate(top_lengths):
        if length >= min_length:
            valid_segments.append(top_segments[idx])

    if len(valid_segments) >= 2:
        distance = valid_segments[0][-1] - valid_segments[1][0]
        if distance >= min_dist:
            return sum(top_lengths[:2])
        else:
            return top_lengths[0]
    elif len(valid_segments) == 1:
        return top_lengths[0]
    else:
        return 0


@logger.catch
def get_white_lines(binary_mask, method, media_type):
    """
    从给定的二值化图像中，提取白线。

    :param binary_mask: 二值化图像
    :param method: str，切分方法，可以是 'horizontal' 或 'vertical'
    :param media_type: str，媒体类型，例如 'Manga'

    :return: 经过校正的白线坐标列表。
    """
    if media_type in grid_ratio_dic:
        grid_ratio = grid_ratio_dic.get(media_type)
    else:
        grid_ratio = grid_ratio_dic.get('Comic')

    ih, iw = binary_mask.shape[0:2]

    if method == 'horizontal':  # 水平线
        dep = ih
        leng = iw
        picture = binary_mask
    else:  # method=='vertical' # 竖直线
        dep = iw
        leng = ih
        picture = binary_mask.T

    # ================粗筛================
    # 框架白线最短长度
    min_leng = grid_ratio * leng

    # 投影
    columns = (picture != 0).sum(0)
    rows = (picture != 0).sum(1)

    # 找出白线
    ori_gray = where(rows >= min_leng)[0]

    # ================细筛================
    # 对相邻白线分组
    cons = np.split(ori_gray, where(diff(ori_gray) != 1)[0] + 1)

    # 合并相邻的分组，如果它们之间的距离在 1～3 的范围内
    merged_cons = []
    cur_group = cons[0]

    for next_group in cons[1:]:
        if next_group[0] - cur_group[-1] <= 3:
            cur_group = np.concatenate((cur_group, next_group))
        else:
            merged_cons.append(cur_group)
            cur_group = next_group

    # 添加最后一个处理过的分组
    merged_cons.append(cur_group)

    # 转化为投影线段
    white_lines = [(x[0], x[-1] - x[0]) for x in merged_cons if len(x) >= 1]

    if merge_nearby:
        merged_lines = []
        if white_lines:
            cur_start, cur_len = white_lines[0]
            for next_start, next_len in white_lines[1:]:
                if next_start - (cur_start + cur_len) <= merge_thres:
                    cur_len = next_start + next_len - cur_start
                else:
                    merged_line = (cur_start, cur_len)
                    merged_lines.append(merged_line)
                    cur_start, cur_len = next_start, next_len
            merged_line = (cur_start, cur_len)
            merged_lines.append(merged_line)
    else:
        merged_lines = white_lines

    # ================最小厚度检验================
    white_lines_normal = [x for x in merged_lines if x[1] > min_frame_thickness]

    logger.debug(f'{white_lines_normal=}')
    return white_lines_normal


@logger.catch
def get_recs(method, white_lines, iw, ih, offset_x, offset_y, media_type):
    """
    根据给定的方法和白色分割线，计算矩形区域列表。

    :param method: str，切分方法，可以是 'horizontal' 或 'vertical'
    :param white_lines: list of tuples，表示白色分割线的起点和长度，例如 [(start_pt1, length1), (start_pt2, length2), ...]
    :param iw: int，图像的宽度
    :param ih: int，图像的高度
    :param offset_x: int，横向偏移量
    :param offset_y: int，纵向偏移量
    :param media_type: str，媒体类型，例如 'Manga'

    :return: list of tuples，返回计算得到的矩形区域列表，每个矩形表示为 (x, y, w, h)
    """
    y_min, y_max, x_min, x_max = 0, 0, 0, 0
    main_rec = (0, 0, iw, ih)

    if method == 'horizontal':
        x_min = 0
        x_max = iw
        dep = ih
    else:  # method=='vertical'
        y_min = 0
        y_max = ih
        dep = iw

    true_dividers = [0]

    for pind in range(len(white_lines)):
        tu = white_lines[pind]
        start_pt, leng = tu
        if edge_ratio * dep <= start_pt < start_pt + leng <= (1 - edge_ratio) * dep:
            true_dividers.append(start_pt)

    true_dividers.append(dep)
    rectangles = []
    valid = True

    small_box_num = 0
    for pind in range(len(true_dividers) - 1):
        upper = true_dividers[pind]
        lower = true_dividers[pind + 1]

        if method == 'horizontal':
            y_min = upper
            y_max = lower
        else:  # method=='vertical'
            x_min = upper
            x_max = lower

        w = x_max - x_min
        h = y_max - y_min

        x = x_min + offset_x
        y = y_min + offset_y

        rectangle = (x, y, w, h)

        # 如果切分时有矩形太小则不放入划分列表
        if w <= grid_width_min or h <= grid_height_min:
            print(f'{rectangle=}')
            small_box_num += 1
        else:
            rectangles.append(rectangle)

    # 如果切分时有多个矩形太小则不合理
    if small_box_num >= 2:
        valid = False

    if not valid:
        rectangles = [main_rec]

    # 日漫从右到左，其他从左到右
    if media_type == 'Manga' and method == 'vertical':
        rectangles.reverse()

    return rectangles


# ================================qt函数区================================
def a6_pyqt():
    return


def copy2clipboard(text):
    clipboard = QApplication.clipboard()
    clipboard.setText(text)


def iact(text, icon=None, shortcut=None, checkable=False, toggled_func=None, trig=None):
    """创建并返回一个QAction对象"""
    action = QAction(text)
    if icon:  # 检查icon_name是否不为None
        action.setIcon(qicon(icon))
    if shortcut:
        action.setShortcut(QKeySequence(shortcut))
    if checkable:
        action.setCheckable(True)
    if toggled_func:
        action.toggled.connect(toggled_func)
    if trig:
        action.triggered.connect(trig)
    return action


def ibut(text, icon):
    button = QToolButton()
    button.setIcon(qicon(icon))
    button.setCheckable(True)
    button.setText(text)
    button.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonIconOnly)
    return button


def get_search_regex(search_text, case_sensitive, whole_word, use_regex):
    """根据搜索条件，返回对应的正则表达式模式对象。"""

    # 如果不区分大小写，设置正则标志为忽略大小写
    flags = IGNORECASE if not case_sensitive else 0

    # 如果不使用正则表达式，对搜索文本进行转义，避免特殊字符影响匹配
    if not use_regex:
        search_text = escape(search_text)

    # 如果选择全词匹配，为搜索文本添加边界匹配符
    if whole_word:
        search_text = fr"\b{search_text}\b"

    # 尝试编译正则表达式，如果失败返回None
    try:
        return re.compile(search_text, flags)
    except re.error:
        return None


# ================================图像函数区================================
def a9_dev():
    return


def open_in_viewer(file_path):
    if sys.platform == 'win32':
        os.startfile(normpath(file_path))
    elif sys.platform == 'darwin':
        Popen(['open', file_path])
    else:
        Popen(['xdg-open', file_path])


def open_in_explorer(file_path):
    folder_path = dirname(file_path)
    if sys.platform == 'win32':
        Popen(f'explorer /select,"{normpath(file_path)}"')
    elif sys.platform == 'darwin':
        Popen(['open', '-R', file_path])
    else:
        Popen(['xdg-open', folder_path])


def open_in_ps(file_path):
    if sys.platform == 'win32':
        photoshop_executable_path = "C:/Program Files/Adobe/Adobe Photoshop CC 2019/Photoshop.exe"  # 请根据您的Photoshop安装路径进行修改
        Popen([photoshop_executable_path, file_path])
    elif sys.platform == 'darwin':
        photoshop_executable_path = "/Applications/Adobe Photoshop 2021/Adobe Photoshop 2021.app"  # 修改此行
        Popen(['open', '-a', photoshop_executable_path, file_path])
    else:
        logger.warning("This feature is not supported on this platform.")


# 处理单个字体
# @logger.catch
def process_font(font):
    name_table = font['name']
    # 中文名的langID
    zh_lang_id = 2052
    en_lang_id = 1033
    # nameID 4 对应的是完整的字体名称
    name_id = 4
    # 首先查找是否有中文名称
    font_name = name_table.getName(name_id, 3, 1, zh_lang_id)
    if font_name is None:
        # 如果没有中文名称，查找英文名称
        font_name = name_table.getName(name_id, 3, 1, en_lang_id)
    if font_name is None:
        return None
    return font_name.toUnicode()


@logger.catch
def update_font_metadata(font_meta_csv, font_head):
    font_meta_list = []
    if font_meta_csv.exists():
        font_meta_list = iread_csv(font_meta_csv, True, False)
    else:
        for i in range(len(system_fonts)):
            font = system_fonts[i]
            font_path = Path(font)
            if font_path.exists():
                file_size = os.path.getsize(font) / 1024  # size in KB
                imes = f'Font {i + 1}: {font}, Size: {file_size:.2f} KB'
                # 读取字体文件
                meta_list = []
                if font_path.suffix.lower() == '.ttc':
                    font_collection = TTCollection(font_path)
                    # 对于TTC文件，处理每个字体
                    fonts = font_collection.fonts
                    for f in range(len(fonts)):
                        font = fonts[f]
                        font_name = process_font(font)
                        postscript_name = font['name'].getName(6, 3, 1, 0x409)
                        if postscript_name:
                            postscript_name = postscript_name.toStr()
                        meta = (font_name, postscript_name)
                        meta_list.append(meta)
                else:
                    font = TTFont(font_path)
                    # 对于TTF文件，处理单个字体
                    font_name = process_font(font)
                    postscript_name = font['name'].getName(6, 3, 1, 0x409)
                    if postscript_name:
                        postscript_name = postscript_name.toStr()
                    meta = (font_name, postscript_name)
                    meta_list.append(meta)
                meta_list = [x for x in meta_list if x[0] and x[1]]
                if meta_list:
                    imes += f', Meta: {meta_list=}'
                    for meta_ent in meta_list:
                        font_meta = (font_path.name,) + meta_ent
                        font_meta_list.append(font_meta)
                logger.debug(imes)
        font_meta_list = reduce_list(font_meta_list)
        font_meta_list.sort()
        write_csv(font_meta_csv, font_meta_list, font_head)
    return font_meta_list


def wrap_strings_with_tr(py_path):
    """
    该函数用于在 Python 文件中自动将从 QWidget 或 QMainWindow 继承的类中的字符串包装在 self.tr() 中。
    这有助于支持字符串的多语言翻译。

    :param py_path: 要处理的 Python 文件的路径。
    """

    # 从文件中读取内容
    content = read_txt(py_path)

    # 匹配从 QWidget 或 QMainWindow 继承的类
    class_pattern = r'class\s+\w+\s*\((?:\s*\w+\s*,\s*)*(QWidget|QMainWindow)(?:\s*,\s*\w+)*\s*\)\s*:'
    class_matches = finditer(class_pattern, content)

    # 在这些类中查找没有包裹在 self.tr() 中的字符串
    string_pattern = r'(?<!self\.tr\()\s*(".*?"|\'.*?\')(?!\s*\.(?:format|replace)|\s*\%|settings\.value)(?!\s*(?:setIcon|icon|setShortcut|QKeySequence|QSettings)\s*\()'

    # 特殊关键字，我们不希望替换包含这些关键字的字符串
    special_keywords = [...]  # 请在这里定义特殊关键字列表

    for class_match in class_matches:
        class_start = class_match.start()
        next_match = search(r'class\s+\w+\s*\(', content[class_start + 1:])
        if next_match:
            class_end = next_match.start()
        else:
            class_end = len(content)

        class_content = content[class_start:class_end]

        matches = finditer(string_pattern, class_content)
        new_class_cont = class_content
        for match in reversed(list(matches)):
            # 检查是否需要替换文本
            str_value = match.group(1)[1:-1]

            # 获取字符串所在的整行
            line_start = class_content.rfind('\n', 0, match.start()) + 1
            line_end = class_content.find('\n', match.end())
            line = class_content[line_start:line_end]

            # 检查字符串是否为十进制数或包含逗号，若是则跳过
            if is_decimal_or_comma(str_value) or str_value.endswith('%'):
                continue

            # 检查行中是否包含我们不需要替换的关键字
            if any(keyword in line for keyword in special_keywords):
                continue

            start = match.start(1)
            end = match.end(1)
            logger.info(f'正在修改: {match.group(1)}')
            # 使用单引号包裹字符串
            new_class_cont = new_class_cont[:start] + f'self.tr(\'{match.group(1)[1:-1]}\')' + new_class_cont[end:]
        content = content[:class_start] + new_class_cont + content[class_start + class_end:]
    updated_content = content

    print(updated_content)
    write_txt(py_i18n_path, updated_content)


@logger.catch
def self_translate_qt(py_path, language_tuples):
    """
    本函数用于提取可翻译字符串，生成 .ts 文件，更新翻译文件并将 .ts 文件转换为 .qm 文件。

    :param py_path: 一个包含 Python 文件的路径对象，用于从中提取可翻译字符串。
    :param language_tuples: 一个包含语言信息的元组列表，每个元组包括语言代码、英文名称、中文名称和自称名称。
    """
    for i in range(len(language_tuples)):
        language_tuple = language_tuples[i]
        lang_code, en_name, cn_name, self_name = language_tuple

        ts_file = UserDataFolder / f'{APP_NAME}_{lang_code}.ts'
        csv_file = UserDataFolder / f'{APP_NAME}_{lang_code}.csv'

        en2dst, dst2en = {}, {}
        existing_sources = set()
        src = []
        if csv_file.exists():
            src = iread_csv(csv_file, True, False)
            for pind in range(len(src)):
                entry = src[pind]
                en_str = entry[0]
                dst_str = entry[-1]
                existing_sources.add(en_str)
                if en_str != '' and dst_str != '':
                    en2dst[en_str] = dst_str
                    dst2en[dst_str] = en_str

        # 提取可翻译的字符串生成ts文件
        cmd = f'{pylupdate} {py_path.as_posix()} -ts {ts_file.as_posix()}'
        logger.debug(f'{cmd}')
        call(cmd, shell=True)
        # 解析 ts 文件
        xml_text = read_txt(ts_file)
        doc_parse = xmltodict.parse(xml_text)
        TS = doc_parse['TS']

        contexts = TS['context']
        updated_contexts = deepcopy(contexts)
        missing_sources = []

        for c in range(len(contexts)):
            context = contexts[c]
            updated_context = deepcopy(context)
            messages = context['message']
            updated_messages = deepcopy(messages)
            for m in range(len(updated_messages)):
                message = updated_messages[m]
                source = message['source']
                if source not in existing_sources:
                    logger.debug(f'{source=}')
                    missing_sources.append(source)
                if source in en2dst:
                    message['translation'] = en2dst[source]
            updated_context['message'] = updated_messages
            updated_contexts[c] = updated_context

        doc_parse['TS']['context'] = updated_contexts
        if missing_sources:
            new_src = deepcopy(src)
            for missing_source in missing_sources:
                new_src.append([missing_source, ''])
            write_csv(csv_file, new_src, src_head_2c)

        # 保存更新后的 ts 文件
        with ts_file.open('w', encoding='utf-8') as f:
            f.write(xmltodict.unparse(doc_parse))

        # 生成 qm
        cmd = f'{lrelease} {ts_file.as_posix()}'
        call(cmd, shell=True)


def get_mini_boxes(contour):
    # 计算轮廓的最小面积矩形
    bounding_box = minAreaRect(contour)
    # 获取最小面积矩形的四个顶点，并按x坐标排序
    points = sorted(list(boxPoints(bounding_box)), key=lambda x: x[0])
    # 根据y坐标确定顶点的排序
    if points[1][1] > points[0][1]:
        index_1, index_4 = 0, 1
    else:
        index_1, index_4 = 1, 0
    if points[3][1] > points[2][1]:
        index_2, index_3 = 2, 3
    else:
        index_2, index_3 = 3, 2
    # 组成最小面积矩形的顶点列表
    box = [points[index_1], points[index_2], points[index_3], points[index_4]]
    # 返回顶点列表和最小边的长度
    return box, min(bounding_box[1])


def box_score_fast(bitmap, _box):
    # 获取图像的高度和宽度
    h, w = bitmap.shape[:2]
    box = _box.copy()
    # 计算边界框在图像中的位置
    xmin = clip(np.floor(box[:, 0].min()).astype(int), 0, w - 1)
    xmax = clip(np.ceil(box[:, 0].max()).astype(int), 0, w - 1)
    ymin = clip(np.floor(box[:, 1].min()).astype(int), 0, h - 1)
    ymax = clip(np.ceil(box[:, 1].max()).astype(int), 0, h - 1)
    # 在边界框内创建掩码
    mask = zeros((ymax - ymin + 1, xmax - xmin + 1), dtype=uint8)
    # 调整边界框坐标以适应掩码
    box[:, 0] -= xmin
    box[:, 1] -= ymin
    # 填充掩码以表示边界框区域
    fillPoly(mask, box.reshape(1, -1, 2).astype(int32), 1)
    # 如果bitmap的类型是float16，转换为float32以避免后续操作中的问题
    if bitmap.dtype == float16:
        bitmap = bitmap.astype(float32)
    # 返回边界框内部区域的平均得分
    box_score = cv2.mean(bitmap[ymin:ymax + 1, xmin:xmax + 1], mask)[0]
    return box_score


def unclip(box, unclip_ratio=1.5):
    # 将边界框转换为多边形
    poly = Polygon(box)
    # 计算需要扩大的距离
    distance = poly.area * unclip_ratio / poly.length
    # 使用pyclipper进行多边形的扩大操作
    offset = PyclipperOffset()
    offset.AddPath(box, JT_ROUND, ET_CLOSEDPOLYGON)
    # 执行扩大操作并返回结果
    expanded = array(offset.Execute(distance))
    return expanded


def box_area(box):
    """
    计算边界框的面积。

    :param box: 边界框的坐标，格式为4xn。
    :return: 边界框的面积。
    """
    return (box[2] - box[0]) * (box[3] - box[1])


def xywh2xyxy(x):
    """
    将边界框从 [x, y, w, h] 格式转换为 [x1, y1, x2, y2] 格式。
    其中x1,y1是边界框的左上角坐标，x2,y2是右下角坐标。

    :param x: 原始边界框坐标，格式为nx4。
    :return: 转换后的边界框坐标。
    """
    # Convert nx4 boxes from [x, y, w, h] to [x1, y1, x2, y2] where xy1=top-left, xy2=bottom-right
    y = x.clone() if isinstance(x, torch.Tensor) else np.copy(x)
    y[:, 0] = x[:, 0] - x[:, 2] / 2  # top left x
    y[:, 1] = x[:, 1] - x[:, 3] / 2  # top left y
    y[:, 2] = x[:, 0] + x[:, 2] / 2  # bottom right x
    y[:, 3] = x[:, 1] + x[:, 3] / 2  # bottom right y
    return y


def box_iou(box1, box2):
    # https://github.com/pytorch/vision/blob/master/torchvision/ops/boxes.py
    """
    Return intersection-over-union (Jaccard index) of boxes.
    Both sets of boxes are expected to be in (x1, y1, x2, y2) format.
    Arguments:
        box1 (Tensor[N, 4])
        box2 (Tensor[M, 4])
    Returns:
        iou (Tensor[N, M]): the NxM matrix containing the pairwise
            IoU values for every element in boxes1 and boxes2

    计算两组边界框之间的交并比（IoU）。

    :param box1: 第一组边界框，格式为Tensor[N, 4]。
    :param box2: 第二组边界框，格式为Tensor[M, 4]。
    :return: 交并比（IoU）矩阵，大小为[N, M]。
    """

    area1 = box_area(box1.T)
    area2 = box_area(box2.T)

    # inter(N,M) = (rb(N,M,2) - lt(N,M,2)).clamp(0).prod(2)
    # 计算交集
    inter = (torch.min(box1[:, None, 2:], box2[:, 2:]) - torch.max(box1[:, None, :2], box2[:, :2])).clamp(0).prod(2)
    # 计算IoU
    return inter / (area1[:, None] + area2 - inter)  # iou = inter / (area1 + area2 - inter)


def non_max_suppression(prediction, conf_thres=0.25, iou_thres=0.45, classes=None, agnostic=False, multi_label=False,
                        labels=(), max_det=300):
    """
    Runs Non-Maximum Suppression (NMS) on inference results

    Returns:
         list of detections, on (n,6) tensor per image [xyxy, conf, cls]

    对推理结果运行非最大抑制（NMS）。

    :param prediction: 推理结果，每张图像一个(n,6)张量 [xyxy, conf, cls]。
    :param conf_thres: 置信度阈值，用于过滤。
    :param iou_thres: IoU阈值，用于非最大抑制。
    :param classes: 仅保留特定类别的检测结果。
    :param agnostic: 是否忽略类别在非最大抑制中。
    :param multi_label: 是否为每个边界框考虑多个标签。
    :param labels: 额外的标签数据。
    :param max_det: 每张图片最大检测数量。
    :return: 每张图片的检测结果列表。
    """

    # 如果prediction是numpy数组，则转换为torch.Tensor
    if isinstance(prediction, ndarray):
        prediction = torch.from_numpy(prediction)

    # 计算类别数
    # number of classes
    nc = prediction.shape[2] - 5  # 类别数 = 预测张量的第三维大小 - 5
    # 筛选出满足置信度阈值的候选检测
    xc = prediction[..., 4] > conf_thres  # candidates

    print(torch.max(prediction[..., 4]))
    # 基本检查
    # Checks
    assert 0 <= conf_thres <= 1, f'Invalid Confidence threshold {conf_thres}, valid values are between 0.0 and 1.0'
    assert 0 <= iou_thres <= 1, f'Invalid IoU {iou_thres}, valid values are between 0.0 and 1.0'

    # 设置
    # Settings
    # 边界框的最小和最大宽高
    min_wh, max_wh = 2, 4096  # (pixels) minimum and maximum box width and height
    # torchvision.ops.nms()中的最大边界框数
    max_nms = 30000  # maximum number of boxes into torchvision.ops.nms()
    # 超时时间限制
    time_limit = 10.0  # seconds to quit after
    # 是否要求冗余检测
    redundant = True  # require redundant detections
    # 如果类别数大于1，则启用多标签
    multi_label &= nc > 1  # multiple labels per box (adds 0.5ms/img)
    # 是否使用合并NMS
    merge = False  # use merge-NMS

    t = time()
    output = [torch.zeros((0, 6), device=prediction.device)] * prediction.shape[0]
    for xi, x in enumerate(prediction):  # image index, image inference
        # Apply constraints
        # x[((x[..., 2:4] < min_wh) | (x[..., 2:4] > max_wh)).any(1), 4] = 0  # width-height
        x = x[xc[xi]]  # confidence

        # Cat apriori labels if autolabelling
        if labels and len(labels[xi]):
            l = labels[xi]
            v = torch.zeros((len(l), nc + 5), device=x.device)
            v[:, :4] = l[:, 1:5]  # box
            v[:, 4] = 1.0  # conf
            v[range(len(l)), l[:, 0].long() + 5] = 1.0  # cls
            x = torch.cat((x, v), 0)

        # If none remain process next image
        if not x.shape[0]:
            continue

        # Compute conf
        x[:, 5:] *= x[:, 4:5]  # conf = obj_conf * cls_conf

        # Box (center x, center y, width, height) to (x1, y1, x2, y2)
        box = xywh2xyxy(x[:, :4])

        # Detections matrix nx6 (xyxy, conf, cls)
        if multi_label:
            i, j = (x[:, 5:] > conf_thres).nonzero(as_tuple=False).T
            x = torch.cat((box[i], x[i, j + 5, None], j[:, None].float()), 1)
        else:  # best class only
            conf, j = x[:, 5:].max(1, keepdim=True)
            x = torch.cat((box, conf, j.float()), 1)[conf.view(-1) > conf_thres]

        # Filter by class
        if classes is not None:
            x = x[(x[:, 5:6] == torch.tensor(classes, device=x.device)).any(1)]

        # Apply finite constraint
        # if not torch.isfinite(x).all():
        #     x = x[torch.isfinite(x).all(1)]

        # Check shape
        n = x.shape[0]  # number of boxes
        if not n:  # no boxes
            continue
        elif n > max_nms:  # excess boxes
            x = x[x[:, 4].argsort(descending=True)[:max_nms]]  # sort by confidence

        # Batched NMS
        c = x[:, 5:6] * (0 if agnostic else max_wh)  # classes
        boxes, scores = x[:, :4] + c, x[:, 4]  # boxes (offset by class), scores
        i = torchvision.ops.nms(boxes, scores, iou_thres)  # NMS
        if i.shape[0] > max_det:  # limit detections
            i = i[:max_det]
        if merge and (1 < n < 3E3):  # Merge NMS (boxes merged using weighted mean)
            # update boxes as boxes(i,4) = weights(i,n) * boxes(n,4)
            iou = box_iou(boxes[i], boxes) > iou_thres  # iou matrix
            weights = iou * scores[None]  # box weights
            x[i, :4] = torch.mm(weights, x[:, :4]).float() / weights.sum(1, keepdim=True)  # merged boxes
            if redundant:
                i = i[iou.sum(1) > 1]  # require redundancy

        output[xi] = x[i]
        if (time() - t) > time_limit:
            print(f'WARNING: NMS time limit {time_limit}s exceeded')
            break  # time limit exceeded

    return output


@timer_decorator
@logger.catch
def CTD2data(blks, lines_map, resize_ratio):
    # 使用非极大抑制处理检测结果，过滤掉重叠度高于nms_thresh的边界框
    blks = non_max_suppression(blks, conf_thresh, nms_thresh)[0]
    # 如果检测结果不在CPU上，将其移动到CPU并转换为NumPy数组
    if blks.device != 'cpu':
        blks = blks.detach_().cpu().numpy()
    # 根据提供的缩放比例调整边界框的宽度和高度
    # 调整宽度坐标
    blks[..., [0, 2]] = blks[..., [0, 2]] * resize_ratio[0]
    # 调整高度坐标
    blks[..., [1, 3]] = blks[..., [1, 3]] * resize_ratio[1]
    # 提取边界框坐标并转换为整数类型
    blines = blks[..., 0:4].astype(int32)
    # 提取置信度并四舍五入到小数点后三位
    confs = np.round(blks[..., 4], 3)
    # 提取对应的类别
    cls = blks[..., 5].astype(int32)
    bccs = (blines, cls, confs)
    # 处理lines_map，只取第一个通道的内容
    lines_map = lines_map[:, 0, :, :]
    # 通过给定的阈值进行二值化处理，以区分前景和背景
    segmentation = lines_map > seg_thresh
    # 初始化存储边界框和得分的列表
    lines = []
    scores_batch = []
    # 根据lines_map的数据类型获取批次大小
    if isinstance(lines_map, torch.Tensor):
        batch_size = lines_map.size(0)
    else:
        batch_size = lines_map.shape[0]
    # 获取当前样本的高度和宽度
    dest_height = lines_map.shape[1]
    dest_width = lines_map.shape[2]
    if not isinstance(dest_width, int):
        dest_width = dest_width.item()
        dest_height = dest_height.item()
    # 遍历每个样本
    for batch_i in range(batch_size):
        pred = lines_map[batch_i]
        _bitmap = segmentation[batch_i]

        # 使用最小边界框从位图中提取边界框和得分
        # 确保_bitmap是二维的
        assert len(_bitmap.shape) == 2
        # 如果pred是torch.Tensor，就将其转换为numpy数组，便于处理
        if isinstance(pred, torch.Tensor):
            bitmap = _bitmap.cpu().numpy()  # 第一个通道
            pred = pred.cpu().detach().numpy()
        else:
            bitmap = _bitmap
        height, width = bitmap.shape
        bmask = (bitmap * 255).astype(uint8)
        contours, _ = findContours(bmask, RETR_LIST, CHAIN_APPROX_SIMPLE)
        num_contours = min(len(contours), max_candidates)
        # 初始化存储边界框的数组
        boxes = zeros((num_contours, 4, 2), dtype=int16)
        # 初始化存储边界框得分的数组
        scores = zeros((num_contours,), dtype=float32)

        for index in range(num_contours):
            # 获取单个轮廓
            contour = contours[index].squeeze(1)
            # 获取轮廓的最小边界框及其最短边长
            points, sside = get_mini_boxes(contour)
            if sside < 2:  # 如果最短边长小于2，忽略该轮廓
                continue
            points = array(points)
            # 计算边界框的得分
            score = box_score_fast(pred, contour)
            # 使用unclip方法扩大边界框
            box = unclip(points, unclip_ratio=unclip_ratio).reshape(-1, 1, 2)
            # 再次获取扩大后的最小边界框及其最短边长
            box, sside = get_mini_boxes(box)
            box = array(box)
            # 将边界框的坐标调整到目标图像尺寸
            box[:, 0] = clip(np.round(box[:, 0] / width * dest_width), 0, dest_width)
            box[:, 1] = clip(np.round(box[:, 1] / height * dest_height), 0, dest_height)
            boxes[index, :, :] = box.astype(int16)  # 将调整后的边界框存入数组
            # 存入边界框得分
            scores[index] = score
        # 将提取的边界框和得分添加到批次列表中
        lines.append(boxes)
        scores_batch.append(scores)
    # 将提取的边界框和得分列表分别赋值给lines和scores
    # 根据得分过滤
    idx = where(scores_batch[0] > box_thresh)
    lines, scores_batch = lines[0][idx], scores_batch[0][idx]

    # 如果检测到的线条为空，则将lines数组置为空列表
    if lines.size == 0:
        lines = []
    else:
        # 否则，将lines数组中的数据类型转换为float64
        lines = lines.astype(float64)
        # 调整线条坐标，使其匹配到原图的尺寸
        lines[..., 0] *= resize_ratio[0]  # 调整X坐标
        lines[..., 1] *= resize_ratio[1]  # 调整Y坐标
        # 将调整后的线条坐标转换回int32类型
        lines = lines.astype(int32)
    return bccs, lines


def get_CTD_data(image):
    # ================缩放并填充图像，同时满足步长倍数约束================
    ih, iw = image.shape[:2]
    # 缩放比例 (新 / 旧)
    r = min(input_h / ih, input_w / iw)

    # 计算缩放后的图像尺寸
    new_w = int(round(iw * r))
    new_h = int(round(ih * r))
    # 计算需要填充的宽度和高度，以满足模型输入尺寸
    dw = int(input_w - new_w)
    dh = int(input_h - new_h)

    # 计算图像缩放比例，用于将检测结果调整回原图尺寸
    resize_ratio = (iw / (input_w - dw), ih / (input_h - dh))

    if (iw, ih) != (new_w, new_h):
        # 需调整大小
        image = resize(image, (new_w, new_h), interpolation=INTER_LINEAR)
        # logger.warning(f'{iw}*{ih}->{new_w}*{new_h}')

    # 对图像进行边界填充，以达到模型输入尺寸
    image = copyMakeBorder(image, 0, dh, 0, dw, BORDER_CONSTANT, value=color_black)

    # 将处理后的图像转换为模型输入所需的blob格式
    blob = dnn.blobFromImage(image, scalefactor=1 / 255.0, size=(input_size, input_size))
    # 设置模型输入并进行前向传播，得到掩膜和线图
    CTD_model.setInput(blob)
    blks, mask, lines_map = CTD_model.forward(uoln)
    if mask.shape[1] == 2:  # 一些OpenCV版本的输出结果是颠倒的
        tmp = mask
        mask = lines_map
        lines_map = tmp

    # 去除填充部分
    mask = mask.squeeze()
    mask = mask[..., :mask.shape[0] - dh, :mask.shape[1] - dw]
    lines_map = lines_map[..., :lines_map.shape[2] - dh, :lines_map.shape[3] - dw]

    # 如果掩膜是torch张量，将其转换为numpy数组，并调整尺寸回原图尺寸
    if isinstance(mask, torch.Tensor):
        mask = mask.squeeze_()
        if mask.device != 'cpu':
            mask = mask.detach().cpu()
        mask = mask.numpy()
    else:
        mask = mask.squeeze()

    # 将掩膜值从概率转换为0-255的像素值，并转换为uint8类型
    mask = mask * 255
    mask = mask.astype(uint8)
    # 将掩膜调整回原始图像尺寸
    mask = resize(mask, (iw, ih), interpolation=INTER_LINEAR)
    return mask, blks, lines_map, resize_ratio


@timer_decorator
@logger.catch
def get_CTD_mask(img_file):
    """
    根据指定的切分模式处理图像，生成CTD掩码图像。

    :param img_file: 输入图像的路径，Path对象。
    """
    # 构建CTD掩码图像的保存路径
    ctd_png = auto_subdir / f'{img_file.stem}-CTD.png'
    # 如果CTD掩码图像已存在且不需要更新，则直接从文件加载
    if ctd_png.exists() and not renew_CTD:
        CTD_mask = imdecode(fromfile(ctd_png, dtype=uint8), -1)
        # 转换为灰度图
        if len(CTD_mask.shape) == 3:
            CTD_mask = cvtColor(CTD_mask, COLOR_BGR2GRAY)
    else:
        image = imdecode(fromfile(img_file, dtype=uint8), -1)
        image = toBGR(image)
        height, width = image.shape[:2]
        black_bg = zeros((height, width), dtype=uint8)
        if CTD_cut == 'whole':
            # 'whole'模式：直接处理整个图像
            CTD_mask = get_CTD_data(image)[0]
        else:
            # 计算切分尺寸
            if CTD_cut == 'square':
                # 'square'模式：取宽高中的最小值
                sq_size = min(width, height)
            else:  # CTD_cut == 'tile'
                # 'tile'模式：固定尺寸1024x1024
                sq_size = 1024
            overlap = int(overlap_ratio * sq_size)

            # 计算切分次数
            cols = max(1, ceil((width - overlap) / (sq_size - overlap)))
            rows = max(1, ceil((height - overlap) / (sq_size - overlap)))

            logger.warning(f'{cols=}, {rows=}')

            masks = []
            poses = []

            # 切分图像
            for row in range(rows):
                for col in range(cols):
                    left = max(0, min(width - sq_size, (sq_size - overlap) * col))
                    upper = max(0, min(height - sq_size, (sq_size - overlap) * row))
                    right = left + sq_size
                    lower = upper + sq_size

                    # 切分图像
                    cropped_image = image[upper:lower, left:right]
                    mask = get_CTD_data(cropped_image)[0]
                    masks.append(mask)
                    poses.append((upper, lower, left, right))

            CTD_mask = black_bg.copy()
            for m in range(len(masks)):
                mask = masks[m]
                pos = poses[m]
                upper, lower, left, right = pos
                # 叠加图像，保留最亮的像素
                CTD_mask[upper:lower, left:right] = maximum(CTD_mask[upper:lower, left:right], mask)

        write_pic(ctd_png, CTD_mask)
    return CTD_mask


def get_filter_inds(good_inds, hierarchy):
    # 对有包含关系的轮廓，排除父轮廓
    filter_inds = []
    for i in range(len(good_inds)):
        contour_index = good_inds[i]
        parent_index = hierarchy[0][contour_index][3]
        is_parent_contour = False

        # 检查是否有任何 good_index 轮廓将此轮廓作为其父轮廓
        for other_index in good_inds:
            hier = hierarchy[0][other_index]
            nextt, prev, first_child, parent = hier
            current_parent_index = parent

            # 检查所有上层父轮廓
            while current_parent_index != -1:
                if current_parent_index == contour_index:
                    is_parent_contour = True
                    break
                current_parent_index = hierarchy[0][current_parent_index][3]

            if is_parent_contour:
                break

        # 如果此轮廓不是任何 good_index 轮廓的父轮廓，将其添加到 filter_inds
        if not is_parent_contour:
            filter_inds.append(contour_index)
    return filter_inds


# @logger.catch
def get_colorful_bubbles(img_raw, bubble_cnts):
    img_pil = fromarray(cvtColor(img_raw, COLOR_BGR2RGB))
    overlay = Image.new('RGBA', img_pil.size, rgba_zero)
    draw = ImageDraw.Draw(overlay)

    for f in range(len(bubble_cnts)):
        bubble_cnt = bubble_cnts[f]
        contour_points = [tuple(point[0]) for point in bubble_cnt.contour]
        # 如果坐标点数量少于2，那么跳过这个轮廓
        if len(contour_points) < 2:
            continue
        # 从 tab20 颜色映射中选择一个颜色
        color = colormap_tab20(f % 20)[:3]
        color_rgb = tuple(int(c * 255) for c in color)
        color_rgba = color_rgb + (int(255 * bubble_alpha),)
        draw.polygon(contour_points, fill=color_rgba)

    for f in range(len(bubble_cnts)):
        bubble_cnt = bubble_cnts[f]
        # 在轮廓中心位置添加序数
        text = str(f + 1)
        text_bbox = draw.textbbox(bubble_cnt.cxy, text, font=msyh_font100, anchor="mm")
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]
        xy_pos = (bubble_cnt.cx - text_width // 2, bubble_cnt.cy - text_height // 2)
        draw.text(xy_pos, text, font=msyh_font100, fill=trans_purple)

    # 将带有半透明颜色轮廓的透明图像与原图混合
    img_pil.paste(overlay, mask=overlay)
    # 将 PIL 图像转换回 OpenCV 图像
    blended_img = cvtColor(array(img_pil), COLOR_RGB2BGR)
    return blended_img


# @logger.catch
def get_textblock_bubbles(img_raw, all_textblocks):
    ih, iw = img_raw.shape[0:2]
    black_bg = zeros((ih, iw), dtype=uint8)

    # 将原始图像转换为PIL图像
    img_pil = fromarray(cvtColor(img_raw, COLOR_BGR2RGB))
    overlay = Image.new('RGBA', img_pil.size, rgba_zero)
    draw = ImageDraw.Draw(overlay)

    all_core_brp_coords = []
    for b in range(len(all_textblocks)):
        textblock = all_textblocks[b]
        textlines = textblock.textlines
        textblock_colorf = colormap_tab20(b % 20)[:3]
        textblock_rgb = tuple(int(c * 255) for c in textblock_colorf)
        textblock_rgba = textblock_rgb + (int(255 * textblock_alpha),)
        # 将RGB颜色转换为HSV
        r, g, b = textblock_rgb
        h, s, v = rgb_to_hsv(r / 255.0, g / 255.0, b / 255.0)
        contour_points = [tuple(point[0]) for point in textblock.block_contour]
        if len(textlines) > 0:
            mask = drawContours(black_bg.copy(), [array(contour_points)], -1, 255, FILLED)
            custom_kernel = kernel5
            # custom_kernel = kernel1
            mask_dc = dilate(mask, custom_kernel, iterations=1)
            contours, _ = findContours(mask_dc, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
            contour = contours[0]
            points = [tuple(point[0]) for point in contour]
            draw.polygon(points, fill=textblock_rgba)

            for l in range(len(textlines)):
                textline = textlines[l]
                textwords = textline.textwords
                # 对颜色进行微调以生成新的颜色
                textline_color = [(c + (l % 5) * 0.03) % 1.0 for c in textblock_colorf]
                # 将颜色值限定在0-255的范围内
                textline_rgb = tuple(min(max(int(c * 255), 0), 255) for c in textline_color)
                textline_rgba = textline_rgb + (int(255 * textline_alpha),)
                draw.rectangle(textline.br_uv, fill=textline_rgba)

                for w in range(len(textwords)):
                    textword = textwords[w]
                    # 改变色相值以得到不同的颜色
                    h = (h + (w % 5) * 0.2) % 1.0
                    textword_color = hsv_to_rgb(h, s, v)
                    textword_rgb = tuple(min(max(int(c * 255), 0), 255) for c in textword_color)
                    textword_rgba = textword_rgb + (int(255 * textword_alpha),)
                    draw.rectangle(textword.br_uv, fill=textword_rgba)

            contour_pairs = zip(contour_points, contour_points[1:] + [contour_points[0]])
            for start_pt, end_point in contour_pairs:
                draw.line([start_pt, end_point], fill=textblock_rgb, width=1)
            if len(textlines) == 1:
                core_brp_coords = list(textblock.core_brp.exterior.coords)
                all_core_brp_coords.append(core_brp_coords)
    for a in range(len(all_core_brp_coords)):
        core_brp_coords = all_core_brp_coords[a]
        draw.polygon(core_brp_coords, outline=color_navy)

    # 将带有半透明颜色轮廓的透明图像与原图混合
    img_pil.paste(overlay, mask=overlay)
    # 将 PIL 图像转换回 OpenCV 图像
    blended_img = cvtColor(array(img_pil), COLOR_RGB2BGR)
    return blended_img


# @logger.catch
def get_raw_bubbles(bubble_mask, letter_mask, left_sample, right_sample, CTD_mask, bboxes=None):
    ih, iw = bubble_mask.shape[0:2]
    black_bg = zeros((ih, iw), dtype=uint8)
    all_contours, hierarchy = findContours(bubble_mask, RETR_TREE, CHAIN_APPROX_SIMPLE)

    all_cnts = []
    for a in range(len(all_contours)):
        contour = all_contours[a]
        cnt = Contr(contour)
        all_cnts.append(cnt)

    good_inds = []
    for a in range(len(all_cnts)):
        cnt = all_cnts[a]
        br_w_ratio = cnt.br_w / iw
        br_h_ratio = cnt.br_h / ih
        br_ratio = cnt.area / (cnt.br_w * cnt.br_h)
        portion_ratio = cnt.area / (iw * ih)
        # ================对轮廓数值进行初筛================
        condition_a1 = area_min <= cnt.area <= area_max
        condition_a2 = perimeter_min <= cnt.perimeter <= perimeter_max
        if cnt.area >= 500:
            # 狭长气泡
            condition_a3 = thickness_min - 2 <= cnt.thickness <= thickness_max
        else:
            condition_a3 = thickness_min <= cnt.thickness <= thickness_max
        condition_a4 = br_w_min <= cnt.br_w <= br_w_max
        condition_a5 = br_h_min <= cnt.br_h <= br_h_max
        condition_a6 = br_wh_min <= min(cnt.br_w, cnt.br_h) <= max(cnt.br_w, cnt.br_h) <= br_wh_max
        condition_a7 = br_wnh_min <= cnt.br_w + cnt.br_h <= br_wnh_max
        condition_a8 = br_w_ratio_min <= br_w_ratio <= br_w_ratio_max
        condition_a9 = br_h_ratio_min <= br_h_ratio <= br_h_ratio_max
        condition_a10 = br_ratio_min <= br_ratio <= br_ratio_max
        condition_a11 = br_w_percent_min * iw <= cnt.br_x <= cnt.br_u <= br_w_percent_max * iw
        condition_a12 = br_h_percent_min * ih <= cnt.br_y <= cnt.br_v <= br_h_percent_max * ih
        condition_a13 = portion_ratio_min <= portion_ratio <= portion_ratio_max
        condition_a14 = area_perimeter_ratio_min <= cnt.area_perimeter_ratio <= area_perimeter_ratio_max

        condition_as = [
            condition_a1,
            condition_a2,
            condition_a3,
            condition_a4,
            condition_a5,
            condition_a6,
            condition_a7,
            # condition_a8,
            # condition_a9,
            # condition_a10,
            # condition_a11,
            # condition_a12,
            # condition_a13,
            # condition_a14,
        ]

        if all(condition_as):
            # ================结合文字图像进一步筛选================
            # 使用白色填充轮廓区域
            filled_contour = drawContours(black_bg.copy(), [cnt.contour], 0, 255, -1)

            # 使用位运算将 mask 和 filled_contour 相交，得到轮廓内的白色像素
            bubble_in_contour = bitwise_and(bubble_mask, filled_contour)
            letter_in_contour = bitwise_and(letter_mask, filled_contour)
            bubble_px = np.sum(bubble_in_contour == 255)
            letter_px = np.sum(letter_in_contour == 255)

            # ================检测文字轮廓================
            letter_contours, letter_hier = findContours(letter_in_contour, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
            letter_cnts = []
            for q in range(len(letter_contours)):
                letter_contour = letter_contours[q]
                letter_cnt = Contr(letter_contour)
                letter_cnts.append(letter_cnt)
            dot_cnts = [x for x in letter_cnts if x.area <= 5]

            # ================对轮廓数值进行初筛================
            bubble_px_ratio = bubble_px / cnt.area  # 气泡像素占比
            letter_px_ratio = letter_px / cnt.area  # 文字像素占比
            BnL_px_ratio = bubble_px_ratio + letter_px_ratio

            # ================计算气泡缩小一圈（5像素）内文字像素数量================
            # 步骤1：使用erode函数缩小filled_contour，得到一个缩小的轮廓
            contour_e5 = erode(filled_contour, kernel5, iterations=1)
            # 步骤2：从原始filled_contour中减去contour_e5，得到轮廓的边缘
            contour_edges = subtract(filled_contour, contour_e5)
            # 步骤3：在letter_mask上统计边缘上的白色像素数量
            edge_pixels = where(contour_edges == 255)
            letter_inner_px_cnt = 0
            for y, x in zip(*edge_pixels):
                if letter_mask[y, x] == 255:
                    letter_inner_px_cnt += 1

            # ================计算气泡扩大一圈（5像素）内文字像素数量================
            # 步骤1：使用dilate函数扩大filled_contour，得到一个扩大的轮廓
            contour_d5 = dilate(filled_contour, kernel5, iterations=1)
            # 步骤2：从扩大的contour_d5中减去原始filled_contour，得到轮廓的外缘
            contour_outer_edges = subtract(contour_d5, filled_contour)
            # 步骤3：在letter_mask上统计外缘上的白色像素数量
            outer_edge_pixels = where(contour_outer_edges == 255)
            letter_outer_px_cnt = 0
            for y, x in zip(*outer_edge_pixels):
                if letter_mask[y, x] == 255:
                    letter_outer_px_cnt += 1
            border_thickness = letter_outer_px_cnt / cnt.perimeter
            # logger.debug(f'{border_thickness=:.4f}')

            # ================检测文字轮廓================
            letter_contours, letter_hier = findContours(letter_in_contour, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
            letter_cnts = []
            for q in range(len(letter_contours)):
                letter_contour = letter_contours[q]
                letter_cnt = Contr(letter_contour)
                letter_cnts.append(letter_cnt)
            max_out_letter_cnts = [x for x in letter_cnts if x.br_h > max_font_size or x.br_w > 2 * max_font_size]
            # logger.warning(f'{max_out_letter_cnts=}')

            # ================对轮廓数值进行细筛================
            CTD_mask_px_ratio = 0.1
            if CTD_mask is not None:
                mask_in_contour = bitwise_and(CTD_mask, filled_contour)
                CTD_mask_px = np.sum(mask_in_contour >= CTD_thres)
                # CTD文字像素占比
                CTD_mask_px_ratio = CTD_mask_px / cnt.area
                if CTD_mask_px_ratio >= 0.01 and better_info:
                    logger.debug(f'{CTD_mask_px_ratio=}')

            condition_b12 = True
            if left_sample is not None:
                # 使用filled_contour和left_sample的位运算，计算交集
                left_sample_in_contour = bitwise_and(left_sample, filled_contour)
                # 计算交集内的白色像素数量
                left_sample_px = np.sum(left_sample_in_contour == 255)
                # 计算filled_contour的面积
                filled_contour_area = np.sum(filled_contour == 255)
                # 检查left_sample在filled_contour内部的面积是否不小于filled_contour面积的30%
                condition_b12 = left_sample_px >= 0.1 * filled_contour_area
                if do_dev_pic:
                    left_sample_png = debug_dir / f'{media_type}_left_sample.png'
                    left_sample_in_contour_png = debug_dir / f'{media_type}_left_sample_in_contour.png'
                    write_pic(left_sample_png, left_sample)
                    if condition_b12:
                        write_pic(left_sample_in_contour_png, left_sample_in_contour)

            if bboxes:
                condition_b0 = False
                if cnt.polygon:
                    # 确保轮廓足以构成多边形
                    for bbox in bboxes:
                        bbox_polygon = box(*bbox)
                        # 计算两个多边形的交集
                        intersection = cnt.polygon.intersection(bbox_polygon)
                        # 返回重叠区域的面积
                        overlap_area = intersection.area
                        min_area = min(100, int(0.5 * bbox_polygon.area))
                        if overlap_area >= min_area:
                            condition_b0 = True
                            # 如果找到至少一个满足条件的bbox，就不再继续检查这个cnt
                            break
            else:
                condition_b0 = True

            condition_b1 = bubble_px_ratio_min <= bubble_px_ratio <= bubble_px_ratio_max
            condition_b2 = letter_px_ratio_min <= letter_px_ratio <= letter_px_ratio_max
            condition_b3 = bubble_px_min <= bubble_px <= bubble_px_max
            condition_b4 = letter_px_min <= letter_px <= letter_px_max
            condition_b5 = BnL_px_ratio_min <= BnL_px_ratio <= BnL_px_ratio_max
            condition_b6 = edge_px_count_min <= letter_inner_px_cnt <= edge_px_count_max
            condition_b7 = border_thickness_min <= border_thickness <= border_thickness_max
            condition_b8 = (len(max_out_letter_cnts) == 0)
            condition_b9 = CTD_mask_px_ratio_min <= CTD_mask_px_ratio <= CTD_mask_px_ratio_max
            condition_b10 = len(dot_cnts) <= 50
            condition_b11 = letter_cnts_min <= len(letter_cnts) <= letter_cnts_max
            condition_bs = [
                condition_b0,
                condition_b1,
                condition_b2,
                condition_b3,
                condition_b4,
                condition_b5,
                condition_b6,
                condition_b7,
                # condition_b8,
                condition_b9,
                # condition_b10,
                # condition_b11,
                # condition_b12,
            ]

            if all(condition_bs):
                good_inds.append(a)
                # logger.warning(f'{a=}')
                logger.debug(f'{bubble_px_ratio=}')
                logger.debug(f'{letter_px_ratio=}')
                logger.warning(f'{border_thickness=}')

    filter_inds = get_filter_inds(good_inds, hierarchy)

    filter_cnts = []
    for g in range(len(filter_inds)):
        filter_index = filter_inds[g]
        filter_cnt = all_cnts[filter_index]
        filter_cnts.append(filter_cnt)
        # print(f'[{g + 1}]{filter_cnt.br=},{filter_cnt.area=:.0f},{filter_cnt.perimeter=:.2f}')
        print(f'[{g + 1}]{filter_cnt=}')

    return filter_cnts


@logger.catch
def order2yaml(order_yml, ordered_cnts, custom_cnts, img_file):
    order_data = iload_data(order_yml)
    content0 = [cnt.cxy_str for cnt in ordered_cnts]
    content = [cnt.cxy_str for cnt in custom_cnts]
    highlight_diffs(content0, content)
    if content != content0:
        order_data[img_file.name] = content
        bubble_data_sorted = {k: order_data[k] for k in natsorted(order_data)}
        write_yml(order_yml, bubble_data_sorted)
        logger.debug(f"已保存到{order_yml}")


def get_paint_colors(group):
    # 根据group选择颜色
    if group == 1:
        dot_color = QColor('red')
        text_color = QColor('purple')
    elif group == 2:
        dot_color = QColor('blue')
        text_color = QColor('navy')
    else:
        dot_color = QColor('green')
        text_color = QColor('maroon')
    return dot_color, text_color


@logger.catch
def get_formatted_stem(file_stem, format='txt'):
    if format == 'doc':
        formatted_stem = file_stem
    elif format == 'html':
        formatted_stem = f'<p>{file_stem}</p>'
    else:  # format == 'txt':
        formatted_stem = f'>>>>>>>>[{file_stem.name}]<<<<<<<<'
    formatted_stem = normalize('NFC', formatted_stem)
    # logger.debug(f'{formatted_stem=}')
    return formatted_stem


@logger.catch
def find_bubbles(text_input):
    bubbles = []
    if isinstance(text_input, str):
        textlines = text_input.splitlines()
    else:
        textlines = text_input
    coor_inds = []
    for t in range(len(textlines)):
        textline = textlines[t]
        m_lp_coor = p_lp_coor.match(textline)
        if m_lp_coor:
            coor_inds.append(t)
    coor_inds.append(len(textlines) - 1)

    for c in range(len(coor_inds) - 1):
        coor_ind = coor_inds[c]
        next_coor_ind = coor_inds[c + 1]
        textline = textlines[coor_ind]
        m_lp_coor = p_lp_coor.match(textline)
        if m_lp_coor:
            bubble_id = int(m_lp_coor.group(1))
            coor_x = float(m_lp_coor.group(2))
            coor_y = float(m_lp_coor.group(3))
            group = int(m_lp_coor.group(4))
            content_lines = textlines[coor_ind + 1:next_coor_ind - 1]
            content = lf.join(content_lines)
            bubble = {
                'id': bubble_id,
                'coor_x': coor_x,
                'coor_y': coor_y,
                'group': group,
                'content': content
            }
            bubbles.append(bubble)
    return bubbles


@logger.catch
def create_index_dict(file_stems, full_paragraphs, format='doc'):
    """根据图片文件名（不包括扩展名）和段落列表创建索引字典和索引列表"""
    full_paragraphs_normalized = [normalize('NFC', paragraph) for paragraph in full_paragraphs]
    index_dict = {}
    last_ind = 0
    indexes = []
    for i, file_stem in enumerate(file_stems):
        formatted_stem = get_formatted_stem(file_stem, format)
        if formatted_stem in full_paragraphs[last_ind:]:
            ind = full_paragraphs[last_ind:].index(formatted_stem) + last_ind
            index_dict[formatted_stem] = ind
            indexes.append(ind)
            last_ind = ind
    indexes.append(len(full_paragraphs))
    return index_dict, indexes


@timer_decorator
def create_para_dic(file_stems, index_dict, indexes, lines, format='doc'):
    """根据图片文件名、索引字典、索引列表和行列表创建段落字典"""
    pin = 0
    para_dic = {}
    start_inds = []
    # 初始化'head'部分
    if indexes:
        start_ind = 0
        end_ind = indexes[0]
        if end_ind > 0:
            stem_text_list = lines[start_ind:end_ind]
            para_dic['head'] = stem_text_list
    for file_stem in file_stems:
        formatted_stem = get_formatted_stem(file_stem, format)
        if formatted_stem in index_dict:
            start_ind = indexes[pin] + 1
            start_inds.append(start_ind)
            end_ind = indexes[pin + 1]
            pin += 1
            stem_text_list = lines[start_ind:end_ind]
            if format == 'txt':
                rlp_pic_bubbles = find_bubbles(stem_text_list)
                para_dic[formatted_stem] = rlp_pic_bubbles
            else:
                para_dic[formatted_stem] = stem_text_list
    return para_dic


lp_head = """
1,0
-
框内
框外
-
Default Comment
You can edit me
"""


@logger.catch
def read_rlp(rlp_txt, img_list):
    if rlp_txt.exists():
        rlp_text = read_txt(rlp_txt, encoding='utf-8-sig')
        rlp_lines = rlp_text.splitlines()
        rlp_index_dict, rlp_inds = create_index_dict(img_list, rlp_lines, 'txt')
        rlp_para_dic = create_para_dic(img_list, rlp_index_dict, rlp_inds, rlp_lines, 'txt')
    else:
        rlp_para_dic = {}
        head = lp_head.strip().splitlines()
        head += ['', '']
        rlp_para_dic['head'] = head
        for i in range(len(img_list)):
            img_file = img_list[i]
            formatted_stem = get_formatted_stem(img_file)
            rlp_para_dic[formatted_stem] = []
    return rlp_para_dic


@logger.catch
def save_rlp(rlp_txt, rlp_para_dic, img_list):
    rlp_lines = []
    rlp_lines += rlp_para_dic['head']
    for i in range(len(img_list)):
        img_file = img_list[i]
        formatted_stem = get_formatted_stem(img_file)
        rlp_lines.append(formatted_stem)
        rlp_pic_bubbles = rlp_para_dic.get(formatted_stem, [])
        for r in range(len(rlp_pic_bubbles)):
            # ================针对每一个气泡================
            rlp_pic_bubble = rlp_pic_bubbles[r]
            id = rlp_pic_bubble['id']
            coor_x = rlp_pic_bubble['coor_x']
            coor_y = rlp_pic_bubble['coor_y']
            group = rlp_pic_bubble['group']
            content = rlp_pic_bubble['content']
            meta_line = f'----------------[{id}]----------------[{coor_x:.3f},{coor_y:.3f},{group}]'
            rlp_lines.append(meta_line)
            rlp_lines.extend(content.splitlines())
            rlp_lines.append('')
        rlp_lines.append('')
    rlp_text = lf.join(rlp_lines)
    write_txt(rlp_txt, rlp_text)


class AppConfig:
    # 设置默认配置文件路径和用户配置文件路径
    default_config_yml = UserDataFolder / f'{APP_NAME}_config.yml'
    if python_ver in ['3.11.8']:
        user_config_yml = UserDataFolder / f'{APP_NAME}_{processor_name}_{ram}GB_anaconda_config.yml'
    else:
        user_config_yml = UserDataFolder / f'{APP_NAME}_{processor_name}_{ram}GB_config.yml'
    master_config_yml = ProgramFolder / f'{APP_NAME}_master_config.yml'
    logger.debug(f'当前用户配置文件名为：{user_config_yml.name}')

    def __init__(self, config_file, config_data):
        # 初始化时设置配置文件路径，并加载配置数据
        self.config_file = Path(config_file)
        self.yaml = YAML()
        self.yaml.indent(mapping=2, sequence=4, offset=2)
        self.config_data = config_data

    @classmethod
    def load(cls):
        # 加载配置文件。如果用户配置文件不存在，则复制默认配置文件
        config_file = cls.user_config_yml
        if not config_file.exists():
            copy2(cls.default_config_yml, config_file)

        # 使用 ruamel.yaml 来读取配置文件，因为它可以保留文件中的注释和顺序
        yaml = YAML()
        yaml.indent(mapping=2, sequence=4, offset=2)
        with open(cls.user_config_yml, mode='r', encoding='utf-8') as yf:
            config_data = yaml.load(yf)

        if cls.master_config_yml.exists():
            with open(cls.master_config_yml, mode='r', encoding='utf-8') as yf:
                master_cfg = yaml.load(yf)
            config_data.update(master_cfg)
        return cls(config_file, config_data)

    def get(self, key, default_value=None):
        # 获取指定配置项的值，如果没有找到，则返回默认值
        return self.config_data.get(key, default_value)

    def set(self, key, value):
        # 设置指定配置项的值，并保存到配置文件中
        self.config_data[key] = value
        self.save()

    def save(self):
        # 保存配置数据到配置文件
        try:
            with self.config_file.open('w') as f:
                self.yaml.dump(self.config_data, f)
        except IOError as e:
            logger.error(f'保存配置文件时出错：{e}')
        else:
            logger.info(f'成功保存配置文件到 {self.config_file}')


# 普通颜色与范围
class Color:

    def __init__(self, color_str):
        self.type = self.__class__.__name__
        self.color_str = color_str
        self.m_color = p_color.match(color_str)
        self.rgb_str = self.m_color.group(1)
        self.padding = self.m_color.group(2)

        self.white = (self.rgb_str == 'ffffff')
        self.black = (self.rgb_str == '000000')

        self.rr = self.color_str[:2]
        self.gg = self.color_str[2:4]
        self.bb = self.color_str[4:6]

        self.r = hex2int(self.rr)
        self.g = hex2int(self.gg)
        self.b = hex2int(self.bb)
        self.a = 255
        self.rgb = (self.r, self.g, self.b)
        self.bgr = (self.b, self.g, self.r)
        self.rgba = (self.r, self.g, self.b, self.a)
        self.bgra = (self.b, self.g, self.r, self.a)

        # 计算亮度
        self.luminance = 0.299 * self.r + 0.587 * self.g + 0.114 * self.b

        if self.padding == '':
            if self.black:
                self.padding = black_padding
            elif self.white:
                self.padding = white_padding
            elif self.luminance > 128:  # 浅色
                self.padding = light_padding
            elif self.luminance < 64:  # 深色
                self.padding = dark_padding
            else:
                self.padding = normal_padding
        else:
            self.padding = int(self.padding)

        self.ext_lower = [int(max(x - self.padding, 0)) for x in self.bgr]
        # self.ext_lower = [int(max(x - 10, 0)) for x in self.bgr]
        self.ext_upper = [int(min(x + self.padding, 255)) for x in self.bgr]

        self.color_name = get_color_name(self.rgb)

        if self.padding == 15:
            self.descripts = [self.rgb_str, self.color_name]
        else:
            self.descripts = [self.rgb_str, f'{self.padding}', self.color_name]
        self.descript = '-'.join(self.descripts)

    def put_padding(self, padding):
        self.padding = padding

    def get_range_img(self, task_img):
        task_img = toBGR(task_img)
        frame_mask = inRange(task_img, array(self.ext_lower), array(self.ext_upper))
        return frame_mask

    def __str__(self):
        return f"{self.type}('{self.rgb_str}', '{self.color_name}', {self.padding})"

    def __repr__(self):
        return f'{self}'


# 渐变颜色与范围
class ColorGradient:

    def __init__(self, color_str_list):
        # 获取类的名字
        self.type = self.__class__.__name__
        # 获取输入的颜色列表
        self.color_str_list = color_str_list
        # 将输入的颜色列表连接成字符串
        self.color_str = '-'.join(color_str_list)

        # 获取输入的颜色列表中的两种颜色
        self.left_color_str, self.right_color_str = color_str_list
        # 将输入的颜色转换为Color对象
        self.left_color = Color(self.left_color_str)
        self.right_color = Color(self.right_color_str)

        # 获取Color对象的rgb字符串表示
        self.rgbs = [self.left_color.rgb_str, self.right_color.rgb_str]
        # 将Color对象的rgb字符串表示连接成一个字符串
        self.rgb_str = '-'.join(self.rgbs)

        # 获取Color对象的颜色名字
        self.color_name = f'{self.left_color.color_name}-{self.right_color.color_name}'

        # 获取Color对象的描述
        self.descripts = [self.left_color.descript, self.right_color.descript]
        # 将Color对象的描述连接成一个字符串
        self.descript = '-'.join(self.descripts)

        # 获取Color对象的padding属性的最大值
        self.padding = int(max(self.left_color.padding, self.right_color.padding))

        # 将两种颜色的rgb值压缩成一个列表
        self.zipped = list(zip(self.left_color.rgb, self.right_color.rgb))
        logger.warning(f'{self.zipped=}')
        # 计算两种颜色的rgb值的平均值
        self.middle = tuple([int(mean(x)) for x in self.zipped])

        # 将中间颜色的rgb值转换为字符串
        self.middle_str = rgb2str(self.middle)
        # 判断中间颜色是否为白色
        self.white = (self.middle_str == 'ffffff')
        # 判断中间颜色是否为黑色
        self.black = (self.middle_str == '000000')

        # 获取中间颜色的rgb值
        self.r, self.g, self.b = self.middle
        # 设置颜色的透明度为255
        self.a = 255
        # 获取颜色的rgba值
        self.rgb = (self.r, self.g, self.b)
        self.bgr = (self.b, self.g, self.r)
        self.rgba = (self.r, self.g, self.b, self.a)
        self.bgra = (self.b, self.g, self.r, self.a)

        # 获取颜色的扩展下限和上限
        self.ext_lower, self.ext_upper = get_ext(self.zipped)
        # 记录颜色的扩展下限和上限
        logger.debug(f'{self.ext_lower=}, {self.ext_upper=}')

    def put_padding(self, padding):
        self.padding = padding

    def get_range_img(self, task_img):
        # 获取在颜色扩展范围内的图像
        frame_mask = inRange(task_img, array(self.ext_lower), array(self.ext_upper))
        # 获取在左边颜色和中间颜色之间的图像
        left_sample = get_sample(task_img, self.left_color.rgb, self.middle)
        # 获取在右边颜色和中间颜色之间的图像
        right_sample = get_sample(task_img, self.right_color.rgb, self.middle)
        # 将以上三种图像打包成元组
        img_tuple = (frame_mask, left_sample, right_sample)
        # 返回包含三种图像的元组
        return img_tuple

    def __str__(self):
        return f"{self.type}('{self.rgb_str}', '{self.color_name}', {self.padding})"

    def __repr__(self):
        return f'{self}'


# 两种颜色与范围
class ColorDouble:
    def __init__(self, color_str_list):
        self.type = self.__class__.__name__
        self.color_str_list = color_str_list
        self.rgb_str = '-'.join(color_str_list)
        self.color_str = '-'.join(color_str_list)

        self.left_color_str, self.right_color_str = color_str_list
        self.left_color = Color(self.left_color_str)
        self.right_color = Color(self.right_color_str)

        self.rgbs = [self.left_color.rgb_str, self.right_color.rgb_str]
        self.rgb_str = '-'.join(self.rgbs)

        self.color_name = f'{self.left_color.color_name}-{self.right_color.color_name}'

        self.descripts = [self.left_color.descript, self.right_color.descript]
        self.descript = '-'.join(self.descripts)

        self.padding = int(max(self.left_color.padding, self.right_color.padding))

    def put_padding(self, padding):
        self.padding = padding

    def get_range_img(self, video_img):
        left_color_mask = self.left_color.get_range_img(video_img)
        right_color_mask = self.right_color.get_range_img(video_img)
        frame_mask = add(left_color_mask, right_color_mask)
        return frame_mask

    def __str__(self):
        return f"{self.type}('{self.rgb_str}', '{self.color_name}', {self.padding})"

    def __repr__(self):
        return f'{self}'


class Contr:

    def __init__(self, contour):
        self.type = self.__class__.__name__
        self.contour = contour
        # 如果轮廓点集的长度大于等于3，可以构成多边形
        if len(contour) >= 3:
            self.polygon = Polygon(np.vstack((contour[:, 0, :], contour[0, 0, :])))
            if not self.polygon.is_valid:
                self.polygon = self.polygon.buffer(0)
        else:
            # 如果点集不足以构成多边形，则设置为None
            self.polygon = None

        # 计算轮廓的面积和周长
        self.area = contourArea(self.contour)
        self.perimeter = arcLength(self.contour, True)
        # 根据周长和面积计算厚度和面积周长比
        if self.perimeter != 0:
            self.thickness = self.area / self.perimeter
            self.area_perimeter_ratio = self.area / self.perimeter / self.perimeter
        else:
            self.thickness = 0
            self.area_perimeter_ratio = 0

        # 轮廓的外接矩形
        self.br = boundingRect(self.contour)
        self.br_x, self.br_y, self.br_w, self.br_h = self.br
        self.br_u = self.br_x + self.br_w
        self.br_v = self.br_y + self.br_h
        self.br_m = int(self.br_x + 0.5 * self.br_w)
        self.br_n = int(self.br_y + 0.5 * self.br_h)
        self.br_area = self.br_w * self.br_h
        self.br_rt = self.area / self.br_area
        # 将外接矩形转换为多边形点集
        self.br_pts = rect2poly(*self.br)
        self.brp = Polygon(self.br_pts)
        self.br_xy = (self.br_x, self.br_y)
        self.br_uv = (self.br_u, self.br_v)
        # 轮廓的平均宽度和高度
        self.avg_w = self.area / self.br_h
        self.avg_h = self.area / self.br_w
        # 计算轮廓的质心坐标
        self.M = moments(self.contour)
        if self.M['m00'] != 0:
            self.cx = int(self.M['m10'] / self.M['m00'])
            self.cy = int(self.M['m01'] / self.M['m00'])
        else:
            self.cx, self.cy = 0, 0
        self.cxy = (self.cx, self.cy)
        self.cyx = (self.cy, self.cx)
        self.cxy_pt = Point(self.cxy)
        self.cxy_str = f'{self.cx},{self.cy}'
        self.cyx_str = f'{self.cy},{self.cx}'
        self.cxv = (self.cx, self.br_v - 3)
        self.order_pt = self.cxy
        # self.order_pt = self.cxv

        self.strokes = []

    def add_stroke(self, stroke):
        self.strokes.append(stroke)

    def add_line_polygon(self, nw, nh):
        if text_direction == 'Horizontal':
            self.line_polygon = Polygon([
                (0, self.br_y),
                (nw, self.br_y),
                (nw, self.br_v),
                (0, self.br_v)
            ])
            self.line_br = [0, self.br_y, nw, self.br_h]
            self.line_br_xy = (0, self.br_y)
            self.line_br_uv = (nw, self.br_v)
        else:
            self.line_polygon = Polygon([
                (self.br_x, 0),
                (self.br_u, 0),
                (self.br_u, nh),
                (self.br_x, nh)
            ])
            self.line_br = [self.br_x, 0, self.br_w, nh]
            self.line_br_xy = (self.br_x, 0)
            self.line_br_uv = (self.br_u, nh)

    def add_cropped_img(self, cropped_img, letter_coors, color_pattern):
        self.cropped_img = cropped_img
        self.letter_coors = letter_coors
        self.center_pt = letter_coors[-1]
        self.color_pattern = color_pattern

    def get_core_br(self, black_bg):
        # ================获取不受气泡尖影响的外接矩形================
        binary = drawContours(black_bg.copy(), [self.contour], -1, 255, FILLED)
        if self.br_rt >= 0.7:
            # ================本来就接近圆形================
            self.core_br = self.br
            self.core_contour0 = self.contour
        else:
            if self.area >= 4000:
                kernel = kernel40
            elif self.area >= 3000:
                kernel = kernel30
            elif self.area >= 2000:
                kernel = kernel20
            else:
                kernel = kernel10
            erosion = erode(binary, kernel, iterations=1)
            dilation = dilate(erosion, kernel, iterations=1)
            self.core_bulk = bitwise_and(dilation, binary)
            self.core_contours, _ = findContours(self.core_bulk, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
            if self.core_contours:
                # 根据新的轮廓找到外接矩形的左上角点
                self.core_contour0 = self.core_contours[0]
                self.core_br = boundingRect(self.core_contour0)
            else:
                if do_dev_pic:
                    core_preview_png = debug_dir / f'{media_type}_core_preview.png'
                    write_pic(core_preview_png, binary)
                self.core_br = self.br
                self.core_contour0 = self.contour
        self.core_br_x, self.core_br_y, self.core_br_w, self.core_br_h = self.core_br

    def is_inside(self, another_cnt):
        """检查当前轮廓是否在另一个轮廓内部"""
        # 使用轮廓的中心点来检查
        result = pointPolygonTest(another_cnt.contour, self.cxy, False)
        return result > 0  # 如果返回的距离为正，当前轮廓在另一个轮廓内部

    def __str__(self):
        return f"{self.type}([A]{self.area:.1f}, [P]{self.perimeter:.2f})[BR]{self.br}"

    def __repr__(self):
        return f'{self}'


class TextWord:
    def __init__(self, letter_cnt, media_type):
        self.type = self.__class__.__name__
        self.letter_cnts = []
        self.media_type = media_type
        self.word_ext_px = word_ext_px
        if isinstance(letter_cnt, list):
            self.letter_cnt = letter_cnt[0]
            self.letter_cnts = letter_cnt
        else:
            self.letter_cnt = letter_cnt
            self.letter_cnts.append(self.letter_cnt)
        self.letter_count = len(self.letter_cnts)
        self.calc_brs()

    def calc_brs(self):
        self.br_x = min(letter.br_x for letter in self.letter_cnts)
        self.br_u = max(letter.br_x + letter.br_w for letter in self.letter_cnts)
        self.br_y = min(letter.br_y for letter in self.letter_cnts)
        self.br_v = max(letter.br_y + letter.br_h for letter in self.letter_cnts)
        self.br_w = self.br_u - self.br_x
        self.br_h = self.br_v - self.br_y
        self.br_m = int(self.br_x + 0.5 * self.br_w)
        self.br_n = int(self.br_y + 0.5 * self.br_h)
        self.br_center = (self.br_m, self.br_n)
        self.br_center_pt = Point(self.br_m, self.br_n)
        self.br_area = self.br_w * self.br_h
        self.br = (self.br_x, self.br_y, self.br_w, self.br_h)
        self.br_uv = (self.br_x, self.br_y, self.br_u, self.br_v)
        self.left_top = (self.br_x, self.br_y)
        self.left_bottom = (self.br_x, self.br_v)
        self.right_top = (self.br_u, self.br_y)
        self.right_bottom = (self.br_u, self.br_v)
        self.br_pts = rect2poly(*self.br)
        self.brp = Polygon(self.br_pts)
        if text_direction == 'Horizontal':
            # ================往右端略微延长================
            self.ext_br = (self.br_x, self.br_y, self.br_w + self.word_ext_px, self.br_h)
        else:  # media_type in ['Manga']
            # ================往下端略微延长================
            self.ext_br = (self.br_x, self.br_y, self.br_w, self.br_h + self.word_ext_px)
        self.ext_br_pts = rect2poly(*self.ext_br)
        self.ext_brp = Polygon(self.ext_br_pts)

    def add_letter_cnt(self, letter_cnt):
        self.letter_cnts.append(letter_cnt)
        if text_direction == 'Horizontal':
            self.letter_cnts.sort(key=lambda x: x.cx + x.cy)
        else:  # media_type in ['Manga']
            self.letter_cnts.sort(key=lambda x: -x.cx + x.cy)
        self.letter_count = len(self.letter_cnts)
        self.calc_brs()

    def add_letter_mask(self, letter_mask):
        self.letter_mask = letter_mask
        self.ih, self.iw = self.letter_mask.shape[0:2]
        self.black_bg = zeros((self.ih, self.iw), dtype=uint8)
        self.all_contours = [x.contour for x in self.letter_cnts]
        self.all_letters = drawContours(self.black_bg.copy(), self.all_contours, -1, 255, FILLED)
        self.letter_in_word = bitwise_and(self.letter_mask, self.all_letters)
        # 获取所有非零像素的坐标
        self.px_pts = transpose(nonzero(self.letter_in_word))
        # 计算非零像素的数量
        self.px_area = self.px_pts.shape[0]
        self.letter_area = self.px_area / len(self.letter_cnts)

    def __str__(self):
        return f"{self.type}[{len(self.letter_cnts)}]br{self.br}center{self.br_center}"

    def __repr__(self):
        return f'{self}'


class TextLine:

    def __init__(self, textword, media_type):
        self.type = self.__class__.__name__
        self.textwords = []
        self.textword = textword
        self.media_type = media_type
        self.line_ext_px = line_ext_px
        self.textwords.append(self.textword)
        self.letter_count = sum([word.letter_count for word in self.textwords])
        self.word_count = len(self.textwords)
        self.calc_brs()

    def calc_brs(self):
        self.br_x = min(word.br_x for word in self.textwords)
        self.br_u = max(word.br_x + word.br_w for word in self.textwords)
        self.br_y = min(word.br_y for word in self.textwords)
        self.br_v = max(word.br_y + word.br_h for word in self.textwords)
        self.br_w = self.br_u - self.br_x
        self.br_h = self.br_v - self.br_y
        # ================处理超大字号================
        if adapt_big_letter and self.br_h >= 2 * line_ext_px:
            self.line_ext_px = max(self.line_ext_px, int(0.8 * self.br_h))
            logger.warning(f'{self.br_h=}, {self.line_ext_px=}')
        self.br_m = int(self.br_x + 0.5 * self.br_w)
        self.br_n = int(self.br_y + 0.5 * self.br_h)
        self.br_center = (self.br_m, self.br_n)
        self.br_center_pt = Point(self.br_m, self.br_n)
        self.br_area = self.br_w * self.br_h
        self.br = (self.br_x, self.br_y, self.br_w, self.br_h)
        self.br_uv = (self.br_x, self.br_y, self.br_u, self.br_v)
        self.left_top = (self.br_x, self.br_y)
        self.left_bottom = (self.br_x, self.br_v)
        self.right_top = (self.br_u, self.br_y)
        self.right_bottom = (self.br_u, self.br_v)
        self.br_pts = rect2poly(*self.br)
        self.brp = Polygon(self.br_pts)
        if text_direction == 'Horizontal':
            # ================往右端略微延长================
            self.ext_br = (self.br_x, self.br_y, self.br_w + self.line_ext_px, self.br_h)
        else:  # media_type in ['Manga']
            # ================往下端略微延长================
            self.ext_br = (self.br_x, self.br_y, self.br_w, self.br_h + self.line_ext_px)
        self.ext_br_pts = rect2poly(*self.ext_br)
        self.ext_brp = Polygon(self.ext_br_pts)

    def add_textword(self, textword):
        self.textwords.append(textword)
        if text_direction == 'Horizontal':
            self.textwords.sort(key=lambda x: x.br_x + x.br_y)
        else:  # media_type in ['Manga']
            self.textwords.sort(key=lambda x: -x.br_x + x.br_y)
        self.textwords.sort(key=lambda x: x.br_x + x.br_y)
        self.letter_count += textword.letter_count
        self.word_count = len(self.textwords)
        self.calc_brs()

    def __str__(self):
        return f"{self.type}[{len(self.textwords)}]br{self.br}center{self.br_center}"

    def __repr__(self):
        return f'{self}'


class iTextBlock:

    def __init__(self, textline, media_type):
        self.type = self.__class__.__name__
        self.textlines = []
        self.textline = textline
        self.media_type = media_type
        self.block_ext_px = block_ext_px
        self.textlines.append(self.textline)
        self.letter_count = sum([line.letter_count for line in self.textlines])
        self.word_count = sum([line.word_count for line in self.textlines])
        self.line_count = len(self.textlines)
        self.calc_brs()

    def calc_brs(self):
        self.br_x = min(textline.br_x for textline in self.textlines)
        self.br_u = max(textline.br_u for textline in self.textlines)
        self.br_y = min(textline.br_y for textline in self.textlines)
        self.br_v = max(textline.br_v for textline in self.textlines)
        self.br_w = self.br_u - self.br_x
        self.br_h = self.br_v - self.br_y
        self.br_m = int(self.br_x + 0.5 * self.br_w)
        self.br_n = int(self.br_y + 0.5 * self.br_h)
        self.br_center = (self.br_m, self.br_n)
        self.br_center_pt = Point(self.br_m, self.br_n)
        self.br_area = self.br_w * self.br_h
        self.br = (self.br_x, self.br_y, self.br_w, self.br_h)
        self.br_uv = (self.br_x, self.br_y, self.br_u, self.br_v)
        self.br_pts = rect2poly(*self.br)
        self.brp = Polygon(self.br_pts)

        # ================柱钩================
        if text_direction == 'Horizontal':
            self.core_br = (
                self.br_m - block_ratio * self.block_ext_px,
                self.br_y - self.block_ext_px,
                2 * block_ratio * self.block_ext_px,
                self.block_ext_px,
            )
        else:  # media_type in ['Manga']
            self.core_br = (
                self.br_x - self.block_ext_px,
                self.br_n - block_ratio * self.block_ext_px,
                self.block_ext_px,
                2 * block_ratio * self.block_ext_px,
            )

        self.core_br_pts = rect2poly(*self.core_br)
        self.core_brp = Polygon(self.core_br_pts)

        # ================包络多边形================
        # 根据文本块的所有文本行的外接矩形生成最小包络多边形
        # 先上后下，先左后右的顺序收集所有的外接矩形点
        all_rect_pts = []
        if text_direction == 'Horizontal':
            for textline in self.textlines:
                all_rect_pts.append((textline.left_top))  # 左上
                all_rect_pts.append((textline.left_bottom))  # 左下
            for textline in reversed(self.textlines):
                all_rect_pts.append((textline.right_bottom))  # 右下
                all_rect_pts.append((textline.right_top))  # 右上
        else:
            for textline in self.textlines:
                all_rect_pts.append((textline.right_top))  # 右上
                all_rect_pts.append((textline.left_top))  # 左上
            for textline in reversed(self.textlines):
                all_rect_pts.append((textline.left_bottom))  # 左下
                all_rect_pts.append((textline.right_bottom))  # 右下
        # 利用所有的外接矩形点构建包络多边形
        self.block_poly = Polygon(all_rect_pts)
        # 直接将这些点转化为OpenCV可以使用的轮廓格式，即(N,1,2)的NumPy数组
        self.block_contour = array(all_rect_pts).reshape((-1, 1, 2)).astype(int32)
        self.block_cnt = Contr(self.block_contour)
        x, y, w, h = boundingRect(self.block_contour)
        # 扩展矩形的四个边界
        x -= rec_pad_w
        y -= rec_pad_h
        w += 2 * rec_pad_w
        h += 2 * rec_pad_h
        # 从扩展后的矩形获取新的轮廓
        self.expanded_contour = array([
            [x, y],
            [x + w, y],
            [x + w, y + h],
            [x, y + h]
        ]).reshape((-1, 1, 2)).astype(int32)
        if use_rec:
            self.expanded_cnt = Contr(self.expanded_contour)
        else:
            self.expanded_cnt = self.block_cnt

    def add_textline(self, textline):
        self.textlines.append(textline)
        if text_direction == 'Horizontal':
            self.textlines.sort(key=lambda x: lower_ratio * x.br_x + x.br_y)
        else:  # media_type in ['Manga']
            self.textlines.sort(key=lambda x: -x.br_x + lower_ratio * x.br_y)
        # ================处理超大字号================
        if textline.line_ext_px >= 50:
            self.block_ext_px = max(self.block_ext_px, textline.line_ext_px)
        self.letter_count += textline.letter_count
        self.word_count += textline.word_count
        self.line_count = len(self.textlines)
        self.calc_brs()

    def __str__(self):
        return f"{self.type}[{len(self.textlines)}]br{self.br}center{self.br_center}"

    def __repr__(self):
        return f'{self}'


class Rect:

    def __init__(self, rect_tup, media_type):
        self.type = self.__class__.__name__  # 获取类名作为类型
        self.rect_tup = rect_tup  # 矩形四元组（x, y, 宽, 高）
        self.media_type = media_type  # 媒体类型

        self.x, self.y, self.w, self.h = self.rect_tup  # 解构四元组到x, y, 宽, 高
        self.xx, self.yy, self.ww, self.hh = self.x, self.y, self.w, self.h
        self.rect_inner_tuple = (self.xx, self.yy, self.ww, self.hh)  # 内部矩形四元组

        self.x_max = self.x + self.w  # 矩形最大x值
        self.y_max = self.y + self.h  # 矩形最大y值

        self.area = self.w * self.h  # 矩形面积
        self.basic = False  # 是否为不可分割的基本方块，初始化为False

    def put_frame_mask(self, frame_mask_group):
        # 根据矩形位置和大小从原图中切割子图
        self.sub_frame_mask_group = []
        for frame_mask in frame_mask_group:
            self.sub_frame_mask = frame_mask[self.y:self.y_max, self.x:self.x_max]
            self.sub_frame_mask_group.append(self.sub_frame_mask)
        self.pg = PicGrid(
            self.sub_frame_mask_group,  # 子图
            self.media_type,  # 媒体类型
            self.x,  # 矩形x值
            self.y,  # 矩形y值
        )
        self.get_rect_inner()

    def get_rect_inner(self):
        # 在四条边上添加用户指定宽度的白色边框
        border_width = 3
        clean_kernel = kernel15
        # clean_kernel = kernel5
        for sub_frame_mask in self.sub_frame_mask_group:
            self.sub_frame_mask = sub_frame_mask
            # 复制子图像
            self.cleaned_mask = self.sub_frame_mask.copy()

            if optimize_inner:
                # 定义一个字典，用于映射方向和对应的数组切片
                direction_slices = {
                    'top': np.s_[:border_width, :],
                    'bottom': np.s_[-border_width:, :],
                    'left': np.s_[:, :border_width],
                    'right': np.s_[:, -border_width:]
                }

                # 遍历四个方向，检查黑色像素数量并在满足条件时添加边框
                for direction, slices in direction_slices.items():
                    # 计算指定方向的黑色像素数量
                    black_pxs = np.sum(self.cleaned_mask[slices] == 0)
                    # 如果黑色像素数量小于等于阈值，则在该方向添加边框
                    if black_pxs <= black_pxs_thres:
                        self.cleaned_mask[slices] = 255

                # 使用形态学操作移除噪声
                # self.cleaned_mask = morphologyEx(self.cleaned_mask, MORPH_OPEN, kernel15)
                # 进行膨胀操作
                self.cleaned_mask = dilate(self.cleaned_mask, clean_kernel, iterations=1)
                # 进行腐蚀操作
                self.cleaned_mask = erode(self.cleaned_mask, clean_kernel, iterations=1)

            # 找到子图中所有黑色像素的坐标
            black_pxs = where(self.cleaned_mask == 0)
            black_y, black_x = black_pxs

            if not black_x.size or not black_y.size:
                # 如果没有黑色像素，那么内部矩形就是整个矩形
                pass
            else:
                # 计算所有黑色像素的外接矩形
                min_x, max_x = np.min(black_x), np.max(black_x)
                min_y, max_y = np.min(black_y), np.max(black_y)

                # 更新内部矩形的坐标和大小
                self.xx = self.x + min_x
                self.yy = self.y + min_y
                self.ww = max_x - min_x + 1
                self.hh = max_y - min_y + 1

            self.rect_inner_tuple = (self.xx, self.yy, self.ww, self.hh)  # 内部矩形四元组
            self.frame_grid = [self.rect_tup, self.rect_inner_tuple]  # 矩形网格，包括原矩形和内部矩形
            self.frame_grid_str = '~'.join([','.join([f'{y}' for y in x]) for x in self.frame_grid])  # 矩形网格字符串表示
            if self.rect_inner_tuple != self.rect_tup:
                break

    def __str__(self):
        return f"{self.type}({self.x}, {self.y}, {self.w}, {self.h})"

    def __repr__(self):
        return f'{self}'


class PicGrid:

    def __init__(self, frame_mask_group, media_type, offset_x=0, offset_y=0):
        self.type = self.__class__.__name__  # 类型名称
        self.media_type = media_type  # 媒体类型
        self.offset_x = offset_x  # 水平偏移量
        self.offset_y = offset_y  # 垂直偏移量
        self.frame_mask_group = frame_mask_group
        self.frame_mask_ind = 0
        self.find_recs(self.frame_mask_ind)

    def find_recs(self, frame_mask_ind):
        self.frame_mask = self.frame_mask_group[frame_mask_ind]  # 二值化的掩码图
        self.basic = False  # 是否是基本单元格
        self.judgement = None  # 判断结果

        # if do_dev_pic:
        #     frame_mask_grid_png = debug_dir / f'{media_type}_frame_mask_grid.png'
        #     write_pic(frame_mask_grid_png, self.frame_mask)

        # 图片长宽
        self.ih, self.iw = self.frame_mask.shape[0:2]  # 图片高度和宽度
        self.px_pts = transpose(nonzero(self.frame_mask))  # 非零像素点的坐标
        self.px_area = self.px_pts.shape[0]  # 非零像素点的数量

        # 计算白色像素的比例
        total_pixels = self.ih * self.iw  # 图片的总像素数量
        white_ratio = self.px_area / total_pixels  # 白色像素的比例

        # 水平和垂直白线
        self.white_lines_horizontal = get_white_lines(self.frame_mask, 'horizontal', media_type)  # 水平白线
        self.white_lines_vertical = get_white_lines(self.frame_mask, 'vertical', media_type)  # 垂直白线
        self.white_lines = {
            'horizontal': self.white_lines_horizontal,  # 水平白线
            'vertical': self.white_lines_vertical  # 垂直白线
        }

        # 水平和垂直矩形
        self.rects_horizontal = get_recs(
            'horizontal',
            self.white_lines_horizontal,
            self.iw,
            self.ih,
            self.offset_x,
            self.offset_y,
            self.media_type,
        )  # 水平矩形
        self.rects_vertical = get_recs(
            'vertical',
            self.white_lines_vertical,
            self.iw,
            self.ih,
            self.offset_x,
            self.offset_y,
            self.media_type,
        )  # 垂直矩形
        self.rects_dic = {
            'horizontal': self.rects_horizontal,  # 水平矩形
            'vertical': self.rects_vertical  # 垂直矩形
        }

        # 判断矩形排列方向
        self.rects = None
        if len(self.rects_horizontal) == len(self.rects_vertical) == 1 or white_ratio >= max_white_ratio:
            self.basic = True  # 如果只有一个矩形，设为基本单元格
        elif len(self.rects_horizontal) >= 2:
            # 如果水平矩形数量大于等于2，判断结果为'horizontal'
            self.judgement = 'horizontal'
        else:
            # 否则判断结果为'vertical'
            self.judgement = 'vertical'
        if self.judgement:
            self.rects = self.rects_dic[self.judgement]  # 根据判断结果，选择对应的矩形

        # 检查self.basic和self.frame_mask_ind
        if self.basic and self.frame_mask_ind < len(self.frame_mask_group) - 1:
            self.frame_mask_ind += 1  # 增加self.frame_mask_ind
            self.find_recs(self.frame_mask_ind)  # 重新运行find_recs方法

    def __str__(self):
        return f"{self.type}({self.iw}, {self.ih}, {self.px_area}, '{self.judgement}')"

    def __repr__(self):
        return f'{self}'


def get_frame_grid_recursive(rect, grids, frame_mask_group, media_type):
    rec = Rect(rect, media_type)
    rec.put_frame_mask(frame_mask_group)  # 对当前Rect进行分割

    if rec.pg.basic:  # 如果当前Rect对象是基本单元格
        grids.append(rec)  # 将其添加到格子列表中
    else:  # 如果当前Rect对象不是基本单元格
        rects = rec.pg.rects  # 获取当前Rect对象的子矩形列表
        for r in rects:  # 对子矩形列表中的每个矩形
            get_frame_grid_recursive(r, grids, frame_mask_group, media_type)  # 递归进行处理


# @logger.catch
def get_grids(frame_mask_group, media_type):
    frame_mask = frame_mask_group[0]
    ih, iw = frame_mask.shape[0:2]  # 获取图像的高度和宽度
    rec0 = (0, 0, iw, ih)  # 定义初始矩形的坐标和大小
    pg = PicGrid(frame_mask_group, media_type)
    grids = []  # 初始化格子列表
    if pg.basic:  # 如果PicGrid对象是基本单元格
        # ================如果主图不可分割================
        rec0 = Rect(rec0, media_type)
        rec0.put_frame_mask(frame_mask_group)  # 对Rect对象进行分割
        grids.append(rec0)  # 将Rect对象添加到格子列表中
    else:
        # ================如果主图可以分割================
        rects = pg.rects_dic[pg.judgement]  # 获取PicGrid对象的子矩形列表
        for rect in rects:  # 对子矩形列表中的每个矩形
            get_frame_grid_recursive(rect, grids, frame_mask_group, media_type)  # 递归进行处理
    return grids  # 返回格子列表


class Line2Prev(QGraphicsLineItem):
    def __init__(self, parent: QGraphicsItem = None):
        super().__init__(parent)
        self.type = self.__class__.__name__
        self.setPen(QPen(QColor(128, 0, 128, 128), 5))

    def update_line(self, end_poly):
        self.setLine(self.parentItem().cnt.cx, self.parentItem().cnt.cy, end_poly.cnt.cx, end_poly.cnt.cy)

    def __str__(self):
        return f"{self.type}"

    def __repr__(self):
        return f'{self}'


class LabelOrd(QGraphicsItem):
    def __init__(self, text, parent=None):
        super().__init__(parent)
        self.type = self.__class__.__name__
        self.text = text
        self.font = QFont('Arial', 50, QFont.Weight.Normal)
        self.textColor = QColor('#696969')
        self.outlineColor = QColor(255, 255, 255)

    def boundingRect(self):
        fm = QFontMetrics(self.font)
        return QRectF(fm.boundingRect(self.text))

    def paint(self, painter, option, widget):
        path = QPainterPath()
        path.addText(0, 0, self.font, self.text)
        painter.setBrush(self.textColor)
        painter.setPen(QPen(self.outlineColor, 2, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin))
        painter.drawPath(path)


class ContourOrd(QGraphicsTextItem):
    def __init__(self, cnt, idx, order_window, parent=None, *args, **kwargs):
        super().__init__(parent, *args, **kwargs)
        self.type = self.__class__.__name__
        self.cnt = cnt
        self.poly_idx = idx
        self.parent_window = order_window
        self.swap_color = QColor(*[int(255 * c) for c in colormap_tab20(idx % 20)[:3]])
        self.swap_color.setAlpha(128)
        self.darker_color = self.swap_color.darker(300)
        self.darker_color.setAlpha(200)
        self.setPlainText(str(idx))
        self.setDefaultTextColor(self.darker_color)
        # 设置字体
        self.setFont(QFont('Arial', 100, QFont.Weight.Normal))
        # 设置位置
        text_rect = self.boundingRect()
        self.setPos(cnt.cx - text_rect.width() / 2, cnt.cy - text_rect.height() / 2)
        self.setGraphicsEffect(
            QGraphicsDropShadowEffect(blurRadius=5, xOffset=0, yOffset=0, color=q_trans_white))

    def update_ord(self):
        idx = self.parentItem().mode_indices[self.parent_window.order_mode]
        if idx is not None:
            self.setPlainText(str(idx))
        else:
            self.setPlainText('')


class ContourPoly(QGraphicsPolygonItem):

    def __init__(self, cnt, idx, order_window, *args, **kwargs):
        super().__init__(QPolygonF([QPointF(pt[0], pt[1]) for pt in cnt.contour.squeeze()]), *args, **kwargs)
        self.type = self.__class__.__name__
        self.cnt = cnt
        self.poly_idx = idx  # 不变
        self.idx_label = idx2label(self.poly_idx)  # 不变
        self.parent_window = order_window
        self.mode_indices = {'swap': idx, 'manual': None}  # 可变
        self.swap_indice = self.mode_indices['swap']
        self.manual_indice = self.mode_indices['manual']
        self.sel_states = {'swap': False, 'manual': False}
        self.swap_color = QColor(*[int(255 * c) for c in colormap_tab20(idx % 20)[:3]])
        self.swap_color.setAlpha(128)
        self.setAcceptHoverEvents(True)
        self.setFlag(QGraphicsItem.GraphicsItemFlag.ItemIsSelectable)
        self.setPen(QPen(self.swap_color, 5))
        self.setBrush(QBrush(self.swap_color, Qt.BrushStyle.SolidPattern))
        self.centroid_dot = QGraphicsEllipseItem(self.cnt.cx - r_dot, self.cnt.cy - r_dot, 2 * r_dot, 2 * r_dot,
                                                 parent=self)
        self.centroid_dot.setBrush(QBrush(QColor(0, 0, 255, 128)))
        self.line2prev = Line2Prev(parent=self)
        self.label_ord = LabelOrd(self.idx_label, parent=self)
        self.contour_ord = ContourOrd(cnt, idx, order_window, parent=self)

    def setup_default_ord(self):
        bounding_rect = self.boundingRect()
        scene_rect = self.scene().sceneRect()
        ord_width = self.label_ord.boundingRect().width()
        ord_height = self.label_ord.boundingRect().height()
        ay = bounding_rect.top() + ord_height
        if bounding_rect.left() - ord_width >= scene_rect.left():
            ax = bounding_rect.left() - ord_width
        else:
            ax = bounding_rect.right()
        self.label_ord.setPos(ax, ay)

    def update_poly(self):
        # ================序数================
        self.mode_indices['swap'] = self.parent_window.custom_cnts.index(self.cnt)
        if self.cnt in self.parent_window.manual_cnts:
            self.mode_indices['manual'] = self.parent_window.manual_cnts.index(self.cnt)
        else:
            self.mode_indices['manual'] = None
        self.swap_indice = self.mode_indices['swap']
        self.manual_indice = self.mode_indices['manual']
        self.contour_ord.update_ord()
        # ================选中与否================
        if self.sel_states[self.parent_window.order_mode]:
            self.setPen(QPen(Qt.red, 3))
        else:
            self.setPen(QPen(self.swap_color, 5))
        # ================连线================
        if self.parent_window.order_mode == 'swap':
            if self.swap_indice == 0:
                self.line2prev.update_line(self)
            else:
                prev_polygon = next((x for x in self.parent_window.contour_polys
                                     if x.swap_indice == self.swap_indice - 1), None)
                if prev_polygon:
                    self.line2prev.update_line(prev_polygon)
        elif self.parent_window.order_mode == 'manual':
            self.line2prev.update_line(self)

    def toggle_selected(self):
        current_mode = self.parent_window.order_mode
        self.sel_states[current_mode] = not self.sel_states[current_mode]
        self.update_poly()

    def mouseDoubleClickEvent(self, event):
        if self.parent_window.order_mode == 'manual' and self.cnt in self.parent_window.manual_cnts:
            self.parent_window.manual_cnts.remove(self.cnt)
            self.parent_window.update_contourpolys()
        else:
            super().mouseDoubleClickEvent(event)

    def mousePressEvent(self, event):
        if self.parent_window.order_mode == 'swap':
            imes = f"点击了轮廓{self.poly_idx}{self.idx_label}"
            logger.debug(imes)
            self.parent_window.status_bar.showMessage(imes)
            self.toggle_selected()
            self.sel_polygons = [x for x in self.parent_window.contour_polys if x.sel_states['swap']]
            self.sel_poly_idxs = [x.poly_idx for x in self.sel_polygons]
            # logger.debug(f"所有选中轮廓: {self.sel_poly_idxs}")
            if self.sel_states['swap'] and len(self.sel_poly_idxs) == 2:
                logger.debug(f"尝试交换{self.sel_poly_idxs[0]}和{self.sel_poly_idxs[1]}")
                self.parent_window.swap_contours(self.sel_poly_idxs[0], self.sel_poly_idxs[1])
                for x in self.sel_polygons:
                    x.toggle_selected()
        elif self.parent_window.order_mode == 'manual':
            if self.cnt not in self.parent_window.manual_cnts:
                self.parent_window.manual_cnts.append(self.cnt)
                self.parent_window.update_contourpolys()


class SearchLine(QLineEdit):
    def __init__(self, parent=None):
        super(SearchLine, self).__init__(parent)

        self.type = self.__class__.__name__
        self.parent_window = parent
        # 区分大小写按钮
        self.case_sensitive_button = ibut(self.tr('Case Sensitive'), 'msc.case-sensitive')
        # 全词匹配按钮
        self.whole_word_button = ibut(self.tr('Whole Word'), 'msc.whole-word')
        # 正则表达式按钮
        self.regex_button = ibut(self.tr('Use Regex'), 'mdi.regex')

        # 创建水平布局，用于将三个按钮放在一行
        self.hb_search_bar = QHBoxLayout()
        self.hb_search_bar.setContentsMargins(0, 0, 0, 0)
        self.hb_search_bar.setSpacing(2)
        self.hb_search_bar.addStretch()
        self.hb_search_bar.addWidget(self.case_sensitive_button)
        self.hb_search_bar.addWidget(self.whole_word_button)
        self.hb_search_bar.addWidget(self.regex_button)

        self.case_sensitive_button.clicked.connect(lambda: self.parent_window.filter_imgs(self.text()))
        self.whole_word_button.clicked.connect(lambda: self.parent_window.filter_imgs(self.text()))
        self.regex_button.clicked.connect(lambda: self.parent_window.filter_imgs(self.text()))
        self.textChanged.connect(self.parent_window.filter_imgs)

        # 设置占位符文本
        self.setPlaceholderText(self.tr('Search'))
        # 将按钮添加到 QLineEdit 的右侧
        self.setLayout(self.hb_search_bar)


class CustImageList(QListWidget):
    def __init__(self, parent=None):
        super().__init__(parent)

        self.type = self.__class__.__name__
        self.parent_window = parent
        self.display_mode = 0
        self.font = QFont()
        self.font.setWordSpacing(0)
        # 设置图标模式
        self.setViewMode(QListView.IconMode)
        self.setIconSize(QSize(thumb_size, thumb_size))  # 设置图标大小
        self.setResizeMode(QListView.ResizeMode.Adjust)  # 设置自动调整大小
        self.setSpacing(5)  # 设置间距
        self.setWordWrap(True)  # 开启单词换行
        self.setWrapping(False)  # 关闭自动换行
        self.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)  # 设置单选模式
        self.setFlow(QListView.Flow.TopToBottom)  # 设置从上到下排列
        self.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)  # 设置自定义右键菜单
        self.customContextMenuRequested.connect(self.show_context_menu)
        self.itemSelectionChanged.connect(self.on_img_selected)
        self.load_img_list()

    def load_img_list(self):
        # 加载图片列表
        self.clear()
        for i in range(len(self.parent_window.img_list)):
            img = self.parent_window.img_list[i]
            pixmap = QPixmap(str(img)).scaled(thumb_size, thumb_size, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            item = QListWidgetItem(QIcon(pixmap), img.name)
            item.setFont(self.font)
            item.setData(Qt.ItemDataRole.UserRole, img)
            item.setTextAlignment(Qt.AlignBottom | Qt.AlignHCenter)
            self.addItem(item)

    def on_img_selected(self):
        # 通过键鼠点击获取的数据是真实序数img_ind
        selected_items = self.selectedItems()
        if selected_items:
            current_item = selected_items[0]
            img_ind = self.row(current_item)
            if img_ind != self.parent_window.img_ind:
                img_file = current_item.data(Qt.ItemDataRole.UserRole)
                self.parent_window.open_img_by_path(img_file)

    def set_display_mode(self, display_mode):
        self.display_mode = display_mode
        self.setUpdatesEnabled(False)
        for index in range(self.count()):
            item = self.item(index)
            data = item.data(Qt.ItemDataRole.UserRole)
            if self.display_mode == 0:  # 仅显示缩略图
                item.setIcon(QIcon(data.as_posix()))
                item.setText('')
                self.setWordWrap(False)
            elif self.display_mode == 1:  # 仅显示文件名
                item.setIcon(QIcon())  # 清除图标
                item.setText(data.name)
                self.setWordWrap(False)  # 确保文件名不换行
            elif self.display_mode == 2:  # 同时显示缩略图和文件名
                item.setIcon(QIcon(data.as_posix()))
                item.setText(data.name)
                self.setWordWrap(True)

        self.setIconSize(QSize(thumb_size, thumb_size))  # 设置图标大小
        if self.display_mode == 1:  # 仅显示文件名
            self.setGridSize(QSize(-1, -1))  # 使用默认大小的网格
        else:
            # 为缩略图和两种模式设置网格大小
            # 为文件名和间距增加额外的宽度
            self.setGridSize(QSize(thumb_size + 30, -1))

        self.setUpdatesEnabled(True)

    def show_context_menu(self, point):
        item = self.itemAt(point)
        if item:
            context_menu = QMenu(self)
            # 创建并添加菜单项
            open_file_action = QAction(self.tr('Open in Explorer'), self)
            open_img_action = QAction(self.tr('Open in Preview'), self)
            open_with_ps = QAction(self.tr('Photoshop'), self)
            # 添加拷贝图片路径、拷贝图片名的选项
            copy_img_path = QAction(self.tr('Copy Image Path'), self)
            copy_img_name = QAction(self.tr('Copy Image Name'), self)

            context_menu.addAction(open_file_action)
            context_menu.addAction(open_img_action)
            # 添加一个打开方式子菜单
            open_with_menu = context_menu.addMenu(self.tr('Open with'))
            open_with_menu.addAction(open_with_ps)
            context_menu.addAction(copy_img_path)
            context_menu.addAction(copy_img_name)

            open_file_action.triggered.connect(lambda: open_in_explorer(item.data(Qt.ItemDataRole.UserRole)))
            open_img_action.triggered.connect(lambda: open_in_viewer(item.data(Qt.ItemDataRole.UserRole)))
            open_with_ps.triggered.connect(lambda: open_in_ps(item.data(Qt.ItemDataRole.UserRole)))
            copy_img_path.triggered.connect(lambda: copy2clipboard(item.data(Qt.ItemDataRole.UserRole).as_posix()))
            copy_img_name.triggered.connect(lambda: copy2clipboard(item.text()))

            context_menu.exec(self.mapToGlobal(point))


class CustGraphicsScene(QGraphicsScene):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.type = self.__class__.__name__
        self.parent_window = parent
        self.img_file = None

    def load_qimg(self, img_data, img_file=None):
        self.img_file = img_file
        # 如果输入是Pillow图像，将其转换为NumPy数组
        if isinstance(img_data, Image.Image):
            img_data = array(img_data)

        # 确保输入数据是NumPy数组
        if isinstance(img_data, ndarray):
            # 检查图像是否是灰度图
            if len(img_data.shape) == 2:  # 灰度图，只有高度和宽度
                height, width = img_data.shape
                bytes_per_line = width  # 灰度图像的每行字节数
                qimg_format = QImage.Format_Grayscale8  # 灰度图像格式
                # 将NumPy数组转换为QImage
                qimage = QImage(img_data.data, width, height, bytes_per_line, qimg_format)
            else:  # 彩色图像
                height, width, channel = img_data.shape
                bytes_per_line = channel * width
                if channel == 4:
                    # 如果输入图像有4个通道（带有Alpha通道）
                    qimg_format = QImage.Format_ARGB32
                elif channel == 3:
                    # 如果输入图像有3个通道
                    # 如果输入图像使用BGR顺序，交换颜色通道以获得正确的RGB顺序
                    img_data = cvtColor(img_data, COLOR_BGR2RGB)
                    qimg_format = QImage.Format_RGB888
                # 将NumPy数组转换为QImage
                qimage = QImage(img_data.data, width, height, bytes_per_line, qimg_format)
            # 将QImage转换为QPixmap
            pixmap = QPixmap.fromImage(qimage)
            # ================清除之前的图像================
            self.clear()
            # ================显示新图片================
            self.addPixmap(pixmap)
            # 将视图大小设置为 pixmap 的大小，并将图像放入视图中
            self.setSceneRect(pixmap.rect().toRectF())

    def mousePressEvent(self, event: QGraphicsSceneMouseEvent):
        if event.button() == Qt.MouseButton.RightButton and hasattr(self.parent_window, 'on_create'):
            x = int(event.scenePos().x())
            y = int(event.scenePos().y())
            self.parent_window.on_create(x, y)
        super().mousePressEvent(event)


class CustGraphicsView(QGraphicsView):
    zoomChanged = pyqtSignal(float)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.type = self.__class__.__name__
        self.parent_window = parent
        # 默认缩放级别，表示原始大小
        self.zoom_level = scaling_factor_reci
        self.start_pt = None
        self.selected_contours = []
        self.setViewportUpdateMode(QGraphicsView.ViewportUpdateMode.FullViewportUpdate)
        self.setTransformationAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        self.setResizeAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setInteractive(True)
        self.setMouseTracking(True)
        # 设置渲染、优化和视口更新模式
        # 设置渲染抗锯齿
        self.setRenderHint(QPainter.RenderHint.Antialiasing)
        # 设置优化标志以便不为抗锯齿调整
        self.setOptimizationFlag(QGraphicsView.OptimizationFlag.DontAdjustForAntialiasing, True)
        # 设置优化标志以便不保存画家状态
        self.setOptimizationFlag(QGraphicsView.OptimizationFlag.DontSavePainterState, True)
        # 设置视口更新模式为完整视口更新
        self.setViewportUpdateMode(QGraphicsView.ViewportUpdateMode.FullViewportUpdate)
        # 设置渲染提示为平滑图像变换，以提高图像的显示质量
        self.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)
        self.setBackgroundBrush(QBrush(Qt.GlobalColor.lightGray))

    def cust_zoom(self, factor):
        self.scale(factor, factor)
        self.zoom_level *= factor
        # 在主窗口中更新缩放输入
        self.parent_window.update_zoom_label()
        self.zoomChanged.emit(self.zoom_level)

    def cust_zoom_in(self):
        self.cust_zoom(1.25)

    def cust_zoom_out(self):
        self.cust_zoom(0.8)

    def fit2view(self, fmode):
        rect = self.scene().itemsBoundingRect()
        view_rect = self.viewport().rect()
        scale_factor_x = view_rect.width() / rect.width()
        scale_factor_y = view_rect.height() / rect.height()
        if fmode == 'width':
            scale_factor = scale_factor_x
        elif fmode == 'height':
            scale_factor = scale_factor_y
        elif fmode == 'screen':
            scale_factor = min(scale_factor_x, scale_factor_y)
        elif fmode == 'original':
            scale_factor = scaling_factor_reci
        self.zoom_level = scale_factor
        self.resetTransform()
        self.scale(scale_factor, scale_factor)
        self.zoomChanged.emit(self.zoom_level)

    def mousePressEvent(self, event):
        self.start_pt = event.pos()
        scene_pt = self.mapToScene(self.start_pt)
        if self.scene().sceneRect().contains(scene_pt) and hasattr(self.parent_window,
                                                                   'order_mode') and self.parent_window.order_mode == 'swap':
            if event.modifiers() == Qt.ShiftModifier:
                self.setDragMode(QGraphicsView.RubberBandDrag)
            else:
                self.setDragMode(QGraphicsView.ScrollHandDrag)
        super().mousePressEvent(event)

    def mouseReleaseEvent(self, event):
        if hasattr(self.parent_window,
                   'order_mode') and self.parent_window.order_mode == 'swap' and self.dragMode() == QGraphicsView.RubberBandDrag:
            selected_items = self.scene().selectedItems()
            for item in selected_items:
                if isinstance(item, ContourPoly):
                    self.selected_contours.append(item)
                    item.toggle_selected()
        self.setDragMode(QGraphicsView.ScrollHandDrag)
        super().mouseReleaseEvent(event)


class OrderWindow(QMainWindow):

    def __init__(self):
        super().__init__()
        self.type = self.__class__.__name__
        self.a0_para()
        self.a1_initialize()
        self.a2_status_bar()
        self.a3_docks()
        self.a4_actions()
        self.a5_menubar()
        self.a6_toolbar()
        self.a9_setting()

    def b1_window(self):
        return

    def a0_para(self):
        # ================初始化变量================
        self.screen_icon = qicon('ei.screen')
        self.setWindowIcon(self.screen_icon)
        self.cgs = CustGraphicsScene(self)
        self.cgv = CustGraphicsView(self)
        self.cgv.setScene(self.cgs)
        self.cgv.zoomChanged.connect(self.update_zoom_label)
        self.resize(window_w, window_h)

    def a1_initialize(self):
        # ================图片列表================
        self.img_folder = img_folder
        self.auto_subdir = Auto / self.img_folder.name
        make_dir(self.auto_subdir)
        self.frame_yml = self.img_folder.parent / f'{self.img_folder.name}.yml'
        self.order_yml = self.img_folder.parent / f'{self.img_folder.name}-气泡排序.yml'
        self.cnts_dic_pkl = self.img_folder.parent / f'{self.img_folder.name}.pkl'
        self.img_list = get_valid_imgs(self.img_folder)
        self.all_masks = get_valid_imgs(self.img_folder, vmode='mask')
        self.frame_data = iload_data(self.frame_yml)
        self.order_data = iload_data(self.order_yml)
        self.cnts_dic = iload_data(self.cnts_dic_pkl)
        self.filter_img_list = self.img_list
        self.img_ind = clamp(img_ind, 0, len(self.img_list) - 1)
        self.img_file = self.img_list[self.img_ind]
        self.contour_polys = []
        self.setWindowTitle(self.img_file.name)

    def a2_status_bar(self):
        # ================状态栏================
        self.status_bar = QStatusBar()
        # 设置状态栏，类似布局设置
        self.setStatusBar(self.status_bar)

    def a3_docks(self):
        self.nav_tab = QTabWidget(self)
        self.cil = CustImageList(self)
        self.nav_tab.addTab(self.cil, self.tr('Thumbnails'))

        self.search_line = SearchLine(self)
        self.vb_search_nav = QVBoxLayout(self)
        self.vb_search_nav.addWidget(self.search_line)
        self.vb_search_nav.addWidget(self.nav_tab)
        self.pics_widget = QWidget()
        self.pics_widget.setLayout(self.vb_search_nav)

        self.pics_dock = QDockWidget(self.tr('Image List'), self)
        self.pics_dock.setObjectName("PicsDock")
        self.pics_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        self.pics_dock.setWidget(self.pics_widget)
        self.addDockWidget(Qt.DockWidgetArea.LeftDockWidgetArea, self.pics_dock)

    def a4_actions(self):
        self.le_scale_percent = QLineEdit(self)
        self.le_scale_percent.setFixedWidth(50)
        self.le_scale_percent.setValidator(QDoubleValidator(1, 1000, 2))

        self.open_folder_action = iact(self.tr('Open Folder'), 'ei.folder', QKeySequence.StandardKey.Open,
                                       trig=self.open_folder_by_dialog)
        self.zoom_in_action = iact(self.tr('Zoom In'), 'ei.zoom-in', QKeySequence.StandardKey.ZoomIn,
                                   trig=self.cgv.cust_zoom_in)
        self.zoom_out_action = iact(self.tr('Zoom Out'), 'ei.zoom-out', QKeySequence.StandardKey.ZoomOut,
                                    trig=self.cgv.cust_zoom_out)
        self.fit2screen_action = iact(self.tr('Fit to Screen'), 'mdi6.fit-to-screen-outline', "Alt+F",
                                      trig=lambda: self.cgv.fit2view("screen"))
        self.fit2width_action = iact(self.tr('Fit to Width'), 'ei.resize-horizontal', "Alt+W",
                                     trig=lambda: self.cgv.fit2view("width"))
        self.fit2height_action = iact(self.tr('Fit to Height'), 'ei.resize-vertical', "Alt+H",
                                      trig=lambda: self.cgv.fit2view("height"))
        self.reset_zoom_action = iact(self.tr('Reset Zoom'), 'mdi6.backup-restore', "Ctrl+0",
                                      trig=lambda: self.cgv.fit2view("original"))
        self.prev_img_action = iact(self.tr('Previous Image'), 'ei.arrow-left', "Ctrl+Left",
                                    trig=lambda: self.nav_img(-1))
        self.next_img_action = iact(self.tr('Next Image'), 'ei.arrow-right', "Ctrl+Right",
                                    trig=lambda: self.nav_img(1))
        self.first_img_action = iact(self.tr('First Image'), 'ei.step-backward', "Ctrl+Home",
                                     trig=lambda: self.nav_img("first"))
        self.last_img_action = iact(self.tr('Last Image'), 'ei.step-forward', "Ctrl+End",
                                    trig=lambda: self.nav_img("last"))
        self.swap_order_action = iact(self.tr('Swap Ordering'), 'mdi.swap-horizontal', checkable=True,
                                      trig=lambda: self.set_order_mode('swap'))
        self.manual_order_action = iact(self.tr('Manual Ordering'), 'fa.hand-grab-o', checkable=True,
                                        trig=lambda: self.set_order_mode('manual'))
        self.undo_action = iact(self.tr('Undo'), 'fa5s.undo', QKeySequence.StandardKey.Undo, trig=self.undo)
        self.redo_action = iact(self.tr('Redo'), 'fa5s.redo', QKeySequence.StandardKey.Redo, trig=self.redo)
        self.save_action = iact(self.tr('Save'), 'msc.save', QKeySequence.StandardKey.Save, trig=self.save2yaml)

        self.ag_sorting = QActionGroup(self)
        self.ag_sorting.addAction(self.swap_order_action)
        self.ag_sorting.addAction(self.manual_order_action)
        self.swap_order_action.setChecked(True)  # 默认选中
        self.undo_action.setEnabled(False)
        self.redo_action.setEnabled(False)
        self.le_scale_percent.editingFinished.connect(self.scale_by_percent)
        self.update_zoom_label()

    def a5_menubar(self):
        # 文件菜单
        self.file_menu = self.menuBar().addMenu(self.tr('File'))
        self.file_menu.addAction(self.open_folder_action)
        self.file_menu.addAction(self.save_action)

        # 显示菜单
        self.view_menu = self.menuBar().addMenu(self.tr('View'))
        # 视图菜单选项
        self.view_menu.addAction(self.pics_dock.toggleViewAction())
        self.view_menu.addSeparator()
        # 缩放选项
        self.view_menu.addAction(self.zoom_in_action)
        self.view_menu.addAction(self.zoom_out_action)
        self.view_menu.addAction(self.fit2screen_action)
        self.view_menu.addAction(self.fit2width_action)
        self.view_menu.addAction(self.fit2height_action)
        self.view_menu.addAction(self.reset_zoom_action)
        self.view_menu.addSeparator()

        # 显示选项
        self.display_modes = [(self.tr('Show Thumbnails'), 0),
                              (self.tr('Show Filenames'), 1),
                              (self.tr('Show Both'), 2)]
        self.display_mode_group = QActionGroup(self)
        for display_mode in self.display_modes:
            action = QAction(display_mode[0], self, checkable=True)
            action.triggered.connect(lambda _, dmode=display_mode[1]: self.cil.set_display_mode(dmode))
            self.view_menu.addAction(action)
            self.display_mode_group.addAction(action)

        # 默认选中 Show Both 选项
        self.display_mode_group.actions()[2].setChecked(True)

        # 编辑菜单
        self.edit_menu = self.menuBar().addMenu(self.tr('Edit'))
        self.edit_menu.addAction(self.swap_order_action)
        self.edit_menu.addAction(self.manual_order_action)
        self.edit_menu.addAction(self.undo_action)
        self.edit_menu.addAction(self.redo_action)

        # 导航菜单
        self.nav_menu = self.menuBar().addMenu(self.tr('Navigate'))
        self.nav_menu.addAction(self.prev_img_action)
        self.nav_menu.addAction(self.next_img_action)
        self.nav_menu.addAction(self.first_img_action)
        self.nav_menu.addAction(self.last_img_action)

    def a6_toolbar(self):
        self.tool_bar = QToolBar(self)
        self.tool_bar.setObjectName("Toolbar")
        self.tool_bar.setIconSize(QSize(24, 24))
        self.tool_bar.setMovable(False)
        self.tool_bar.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextUnderIcon)
        self.addToolBar(self.tool_bar)
        self.tool_bar.addAction(self.open_folder_action)
        self.tool_bar.addSeparator()
        self.tool_bar.addAction(self.zoom_in_action)
        self.tool_bar.addAction(self.zoom_out_action)
        self.tool_bar.addAction(self.fit2screen_action)
        self.tool_bar.addAction(self.fit2width_action)
        self.tool_bar.addAction(self.fit2height_action)
        self.tool_bar.addAction(self.reset_zoom_action)
        self.tool_bar.addSeparator()
        self.tool_bar.addAction(self.first_img_action)
        self.tool_bar.addAction(self.prev_img_action)
        self.tool_bar.addAction(self.next_img_action)
        self.tool_bar.addAction(self.last_img_action)
        self.tool_bar.addSeparator()
        self.tool_bar.addAction(self.swap_order_action)
        self.tool_bar.addAction(self.manual_order_action)
        self.tool_bar.addAction(self.undo_action)
        self.tool_bar.addAction(self.redo_action)
        self.tool_bar.addAction(self.save_action)
        self.tool_bar.addWidget(self.le_scale_percent)
        self.tool_bar.addWidget(QLabel('%'))

    def a9_setting(self):
        self.open_img_by_path(self.img_file)
        self.setCentralWidget(self.cgv)
        self.show()

    def open_img_by_path(self, img_file):
        img_file = Path(img_file)
        if img_file.exists() and img_file != self.cgs.img_file:
            self.img_file = img_file
            self.img_file_size = getsize(self.img_file)
            self.img_ind = self.img_list.index(self.img_file)
            self.setWindowTitle(self.img_file.name)
            # ================显示图片================
            self.img_raw = imdecode(fromfile(self.img_file, dtype=uint8), -1)
            self.ih, self.iw = self.img_raw.shape[0:2]
            self.cgs.load_qimg(self.img_raw, self.img_file)
            self.scale_by_percent()
            self.update_zoom_label()
            # ================将当前图片项设为选中状态================
            self.cil.blockSignals(True)  # 阻止信号
            self.cil.setCurrentRow(self.img_ind)
            self.cil.blockSignals(False)  # 恢复信号
            # ================更新状态栏信息================
            index_str = f'{self.img_ind + 1}/{len(self.img_list)}'
            meta_str = f'{self.tr("Width")}: {self.iw} {self.tr("Height")}: {self.ih} | {self.tr("File Size")}: {self.img_file_size} bytes'
            status_text = f'{index_str} | {self.tr("Filnename")}: {self.img_file.name} | {meta_str}'
            self.status_bar.showMessage(status_text)
            QApplication.processEvents()
            self.open_contour_polys()
        QApplication.processEvents()

    def open_contour_polys(self):
        self.contour_polys.clear()
        self.manual_cnts = []
        self.undo_stack = []
        self.redo_stack = []
        self.order_mode = 'swap'
        self.swap_order_action.setChecked(True)
        if self.img_file in self.cnts_dic:
            self.ordered_cnts = self.cnts_dic[self.img_file]
        else:
            img_file, self.ordered_cnts = order1pic(self.img_folder, self.img_ind, media_type)
        self.custom_cnts = deepcopy(self.ordered_cnts)
        if self.custom_cnts:
            for idx, cnt in enumerate(self.custom_cnts):
                self.contour_poly = ContourPoly(cnt, idx, self)
                # 加入到场景中
                self.cgs.addItem(self.contour_poly)
                self.contour_poly.setup_default_ord()
                self.contour_polys.append(self.contour_poly)
                # 更新显示效果
                self.contour_poly.update_poly()

    def open_folder_by_path(self, folder_path):
        # 判断文件夹路径是否存在
        folder_path = Path(folder_path)
        if folder_path.exists():
            # 获取所有图片文件的路径
            img_list = get_valid_imgs(folder_path)
            if img_list and folder_path != self.img_folder:
                self.img_folder = folder_path
                self.auto_subdir = Auto / self.img_folder.name
                make_dir(self.auto_subdir)
                self.frame_yml = self.img_folder.parent / f'{self.img_folder.name}.yml'
                self.order_yml = self.img_folder.parent / f'{self.img_folder.name}-气泡排序.yml'
                self.cnts_dic_pkl = self.img_folder.parent / f'{self.img_folder.name}.pkl'
                self.img_list = img_list
                self.all_masks = get_valid_imgs(self.img_folder, vmode='mask')
                self.frame_data = iload_data(self.frame_yml)
                self.order_data = iload_data(self.order_yml)
                self.cnts_dic = iload_data(self.cnts_dic_pkl)
                self.filter_img_list = self.img_list
                self.img_ind = 0
                self.img_file = self.img_list[self.img_ind]
                # ================更新导航栏中的图片列表================
                self.cil.load_img_list()
                self.open_img_by_path(self.img_file)

    def open_folder_by_dialog(self):
        # 如果self.img_folder已经设置，使用其上一级目录作为起始目录，否则使用当前目录
        self.img_folder = Path(self.img_folder) if self.img_folder else None
        start_directory = self.img_folder.parent.as_posix() if self.img_folder else "."
        img_folder = QFileDialog.getExistingDirectory(self, self.tr('Open Folder'), start_directory)
        if img_folder:
            self.open_folder_by_path(img_folder)

    def nav_img(self, nav_step):
        cur_img_path = self.img_list[self.img_ind]
        # 检查当前图片路径是否在过滤后的图片列表中
        if cur_img_path not in self.filter_img_list:
            return
        cur_filter_ind = self.filter_img_list.index(cur_img_path)
        if nav_step == "first":
            new_filter_ind = 0
        elif nav_step == "last":
            new_filter_ind = len(self.filter_img_list) - 1
        else:
            new_filter_ind = cur_filter_ind + nav_step
        if 0 <= new_filter_ind < len(self.filter_img_list):
            new_img_file = self.filter_img_list[new_filter_ind]
            self.open_img_by_path(new_img_file)

    def filter_imgs(self, search_text: str):
        # 获取搜索框的三个条件：是否区分大小写、是否全词匹配、是否使用正则表达式
        case_sensitive = self.search_line.case_sensitive_button.isChecked()
        whole_word = self.search_line.whole_word_button.isChecked()
        use_regex = self.search_line.regex_button.isChecked()

        # 获取对应的正则表达式模式对象
        regex = get_search_regex(search_text, case_sensitive, whole_word, use_regex)
        if not regex:
            return

        # 根据正则表达式筛选图片列表
        self.filter_img_list = [img_file for img_file in self.img_list if regex.search(img_file.name)]

        # 更新缩略图列表：如果图片名匹配正则表达式，显示该项，否则隐藏
        for index in range(self.cil.count()):
            item = self.cil.item(index)
            item_text = item.text()
            item.setHidden(not bool(regex.search(item_text)))

    def update_contourpolys(self):
        for c in range(len(self.contour_polys)):
            self.contour_poly = self.contour_polys[c]
            self.contour_poly.update_poly()

    def update_zoom_label(self):
        self.le_scale_percent.setText(f'{self.cgv.zoom_level * 100:.2f}')

    def scale_by_percent(self):
        target_scale = float(self.le_scale_percent.text()) / 100
        current_scale = self.cgv.transform().m11()
        scale_factor = target_scale / current_scale
        self.cgv.scale(scale_factor, scale_factor)

    def set_order_mode(self, order_mode):
        self.order_mode = order_mode
        self.update_contourpolys()

    def swap_contours(self, poly_idx1, poly_idx2):
        poly1 = [x for x in self.contour_polys if x.poly_idx == poly_idx1][0]
        poly2 = [x for x in self.contour_polys if x.poly_idx == poly_idx2][0]

        # 保存当前的轮廓顺序，以便于后续的撤销操作
        self.undo_stack.append(self.custom_cnts.copy())
        self.redo_stack.clear()
        # 交换
        idx1 = self.custom_cnts.index(poly1.cnt)
        idx2 = self.custom_cnts.index(poly2.cnt)
        self.custom_cnts[idx1], self.custom_cnts[idx2] = self.custom_cnts[idx2], self.custom_cnts[idx1]

        # 在状态栏显示交换信息
        imes = f'交换了轮廓{poly1.idx_label}和轮廓{poly2.idx_label}'
        # logger.debug(imes)
        self.status_bar.showMessage(imes)
        # 更新撤销和重做按钮的状态
        self.undo_action.setEnabled(True)
        self.redo_action.setEnabled(False)
        self.update_contourpolys()

    def undo(self):
        if not self.undo_stack:
            return

        self.redo_stack.append(self.custom_cnts.copy())
        self.custom_cnts = self.undo_stack.pop()
        self.update_contourpolys()

        self.undo_action.setEnabled(len(self.undo_stack) > 0)
        self.redo_action.setEnabled(True)

    def redo(self):
        if not self.redo_stack:
            return

        self.undo_stack.append(self.custom_cnts.copy())
        self.custom_cnts = self.redo_stack.pop()
        self.update_contourpolys()

        self.redo_action.setEnabled(len(self.redo_stack) > 0)
        self.undo_action.setEnabled(True)

    def keyPressEvent(self, event):
        if self.order_mode == 'swap':
            if event.key() == Qt.Key_Escape:
                self.sel_polygons = [x for x in self.contour_polys if x.sel_states['swap']]
                for s in range(len(self.sel_polygons)):
                    self.contour_poly = self.sel_polygons[s]
                    self.contour_poly.toggle_selected()
                # 更新轮廓的高亮显示
                self.update_contourpolys()
                self.status_bar.showMessage('已取消所有选中的轮廓')
            else:
                self.sel_polygons = [x for x in self.contour_polys if x.sel_states['swap']]
                if self.sel_polygons:
                    direction = None
                    if event.key() == Qt.Key_Minus:
                        direction = -1
                    elif event.key() == Qt.Key_Equal:
                        direction = 1
                    if direction:
                        self.sel_polygons.sort(key=lambda x: x.swap_indice, reverse=(direction == 1))
                        boundary_check = (0 if direction == -1 else len(self.custom_cnts) - 1)
                        if self.sel_polygons[0].swap_indice != boundary_check:
                            for poly in self.sel_polygons:
                                cnt2 = self.custom_cnts[poly.swap_indice + direction]
                                poly2 = get_poly_by_cnt(self.contour_polys, cnt2)
                                self.swap_contours(poly.poly_idx, poly2.poly_idx)
                            self.update_contourpolys()
            self.sel_polygons = [x for x in self.contour_polys if x.sel_states['swap']]

    def save2yaml(self):
        if self.order_mode == 'swap':
            order2yaml(self.order_yml, self.ordered_cnts, self.custom_cnts, self.img_file)
        elif self.order_mode == 'manual':
            logger.debug(f'{self.manual_cnts=}')
            if len(self.manual_cnts) == len(self.custom_cnts):
                order2yaml(self.order_yml, self.ordered_cnts, self.manual_cnts, self.img_file)
            else:
                self.status_bar.showMessage("在手动模式下，必须排序所有轮廓才能保存")


class ColorTextItemDelegate(QStyledItemDelegate):
    def __init__(self, parent=None):
        super().__init__(parent)

    def paint(self, painter, option, index):
        # 设置画笔颜色基于单元格内容
        text = index.model().data(index, Qt.ItemDataRole.DisplayRole)
        if text == "G1框内":
            painter.setPen(QPen(QColor("red")))
        elif text == "G2框外":
            painter.setPen(QPen(QColor("blue")))
        else:
            painter.setPen(QPen(QColor("black")))  # 默认颜色

        # 绘制文本
        painter.drawText(option.rect, Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter, text)


class CustTableView(QTableView):
    def __init__(self, parent=None):
        super(CustTableView, self).__init__(parent)
        self.parent = parent

        self.setItemDelegateForColumn(2, ColorTextItemDelegate(self))  # 假设你想在第三列应用颜色
        # 设置表格为不可编辑状态
        self.setEditTriggers(QTableView.EditTrigger.NoEditTriggers)

        # 设置选择模式和选择行为
        self.setSelectionMode(QTableView.SelectionMode.SingleSelection)
        self.setSelectionBehavior(QTableView.SelectionBehavior.SelectRows)

        # 表格宽度的自适应调整
        self.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        # 水平方向标签拓展剩下的窗口部分，填满表格
        self.horizontalHeader().setStretchLastSection(True)

        # 自定义右键菜单
        self.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.customContextMenuRequested.connect(self.onContextMenu)

        # self.setDragDropMode(QAbstractItemView.InternalMove)
        self.setEditTriggers(QAbstractItemView.NoEditTriggers)  # 设置表格不可更改

        # 水平方向，表格大小拓展到适当的尺寸
        self.setSelectionMode(QAbstractItemView.SingleSelection)  # 单选
        self.setSelectionBehavior(QAbstractItemView.SelectRows)  # 单元格选择
        self.setFrameShape(QFrame.NoFrame)  # 设置无表格的外框

        self.setStyleSheet("""
        QTableView {
            selection-background-color: lightblue;
            selection-color: black;
        }
        """)

    def onContextMenu(self, point: QPoint):
        contextMenu = QMenu(self)
        contextMenu.addAction(self.parent.group1Action)
        contextMenu.addAction(self.parent.group2Action)
        contextMenu.addAction(self.parent.deleteAction)
        contextMenu.exec(self.mapToGlobal(point))


class DraggableGroup(QGraphicsItemGroup):
    def __init__(self, x, y, id, group, parent=None):
        super().__init__()
        # 创建椭圆
        self.id = int(id)
        self.group = int(group)
        diameter = 20
        dot_color, text_color = get_paint_colors(group)
        self.ellipse = QGraphicsEllipseItem(x - diameter / 2, y - diameter / 2, diameter, diameter, self)
        self.ellipse.setBrush(QBrush(dot_color))
        self.ellipse.setPen(QPen(Qt.NoPen))

        # 创建文本
        self.text_item = QGraphicsTextItem(str(self.id), self)
        self.text_item.setDefaultTextColor(text_color)
        self.text_item.setFont(QFont('Arial', 100, QFont.Weight.Normal))
        self.text_item.setPos(x, y)  # 相对于组的位置

        # 设置组的拖动属性
        self.setFlags(QGraphicsItem.ItemIsMovable | QGraphicsItem.ItemSendsGeometryChanges)
        # 连接父窗口
        self.parent = parent

    def change_id(self, id):
        self.id = id
        self.text_item.setPlainText(str(self.id))

    def change_group(self, group):
        self.group = group
        dot_color, text_color = get_paint_colors(group)
        self.ellipse.setBrush(QBrush(dot_color))
        self.text_item.setDefaultTextColor(text_color)

    def itemChange(self, change, value):
        if change == QGraphicsItem.ItemPositionHasChanged:
            # 获取椭圆在场景中的新位置
            new_scene_pos = self.ellipse.mapToScene(self.ellipse.rect().center())
            new_x = int(new_scene_pos.x())
            new_y = int(new_scene_pos.y())
            # 更新表格坐标
            # logger.debug(f'{new_x},{new_y}')
            self.parent.update_coordinates(self, new_x, new_y)
        return super().itemChange(change, value)


class CustTableModel(QStandardItemModel):
    def __init__(self, rows, columns, parent=None):
        super().__init__(rows, columns, parent)

    def data(self, index, role=Qt.ItemDataRole.DisplayRole):
        if role == Qt.ItemDataRole.DisplayRole:
            text = super().data(index, role)
            if text is not None:
                return text.replace('\n', ' ')  # 将换行替换为空格用于显示
        return super().data(index, role)

    def setData(self, index, value, role=Qt.ItemDataRole.EditRole):
        if role == Qt.ItemDataRole.EditRole:
            # 存储原始多行文本
            return super().setData(index, value, role)
        return False

    def addData(self, row, column, value):
        item = QStandardItem(value)
        self.setItem(row, column, item)


class LabelPlusWindow(QMainWindow):

    def __init__(self):
        super().__init__()
        self.type = self.__class__.__name__
        self.a0_para()
        self.a1_initialize()
        self.a2_status_bar()
        self.a3_docks()
        self.a4_actions()
        self.a5_menubar()
        self.a6_toolbar()
        self.a9_setting()

    def b1_window(self):
        return

    def a0_para(self):
        # ================初始化变量================
        self.screen_icon = qicon('ei.screen')
        self.setWindowIcon(self.screen_icon)
        self.cgs = CustGraphicsScene(self)
        self.cgv = CustGraphicsView(self)
        self.cgv.setScene(self.cgs)
        self.cgv.zoomChanged.connect(self.update_zoom_label)
        self.resize(window_w, window_h)

    def a1_initialize(self):
        # ================图片列表================
        self.img_folder = img_folder
        self.auto_subdir = Auto / self.img_folder.name
        make_dir(self.auto_subdir)
        self.img_list = get_valid_imgs(self.img_folder)
        self.rlp_txt = self.img_folder.parent / f'{self.img_folder.name}翻译_0.txt'
        self.rlp_para_dic = read_rlp(self.rlp_txt, self.img_list)
        self.filter_img_list = self.img_list
        self.img_ind = clamp(img_ind, 0, len(self.img_list) - 1)
        self.bubble_ellipses = []
        if self.img_ind < len(self.img_list):
            self.img_file = self.img_list[self.img_ind]
            self.setWindowTitle(self.img_file.name)
        else:
            self.img_file = None

    def a2_status_bar(self):
        # ================状态栏================
        self.status_bar = QStatusBar()
        # 设置状态栏，类似布局设置
        self.setStatusBar(self.status_bar)

    def a3_docks(self):
        self.nav_tab = QTabWidget(self)
        self.cil = CustImageList(self)
        self.nav_tab.addTab(self.cil, self.tr('Thumbnails'))

        self.search_line = SearchLine(self)

        # 构建Model/View
        table_header = [self.tr('ID'), self.tr('Content'), self.tr('Group'), self.tr('Coordinates'), ]
        self.plus_tv_im = CustTableModel(10, 4, self)  # 数据模型,10行4列
        self.plus_tv_im.setHorizontalHeaderLabels(table_header)
        self.plus_tv_sm = QItemSelectionModel(self.plus_tv_im)  # Item选择模型
        self.plus_tv_sm.selectionChanged.connect(self.update_plus_pte)

        # 设置表格属性
        self.plus_tv = CustTableView(self)
        self.plus_tv.setModel(self.plus_tv_im)  # 设置数据模型
        self.plus_tv.setSelectionModel(self.plus_tv_sm)  # 设置选择模型
        if hide_extra:
            self.plus_tv.setColumnHidden(0, True)  # ID
            self.plus_tv.setColumnHidden(2, True)  # 组别
            self.plus_tv.setColumnHidden(3, True)  # 坐标

        self.plus_pte = QPlainTextEdit()
        self.plus_pte.setLineWrapMode(QPlainTextEdit.WidgetWidth)
        self.plus_pte.setObjectName("plainTextEdit")
        self.plus_pte.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        self.plus_pte.textChanged.connect(self.on_text_changed)

        self.pics_widget = QWidget()
        self.vb_search_nav = QVBoxLayout(self)
        self.vb_search_nav.addWidget(self.search_line)
        self.vb_search_nav.addWidget(self.nav_tab)
        self.pics_widget.setLayout(self.vb_search_nav)

        self.lp_widget = QWidget()
        self.lp_vb = QVBoxLayout(self)
        self.lp_vb.addWidget(self.plus_tv)
        self.lp_vb.addWidget(self.plus_pte)
        self.lp_widget.setLayout(self.lp_vb)

        self.pics_dock = QDockWidget(self.tr('Image List'), self)
        self.pics_dock.setObjectName("PicsDock")
        self.pics_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        self.pics_dock.setWidget(self.pics_widget)
        self.addDockWidget(Qt.DockWidgetArea.LeftDockWidgetArea, self.pics_dock)

        self.lp_dock = QDockWidget(self.tr('LabelPlus'), self)
        self.lp_dock.setObjectName("LabelplusDock")
        self.lp_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        self.lp_dock.setWidget(self.lp_widget)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.lp_dock)

    def a4_actions(self):
        self.le_scale_percent = QLineEdit(self)
        self.le_scale_percent.setFixedWidth(50)
        self.le_scale_percent.setValidator(QDoubleValidator(1, 1000, 2))

        self.open_folder_action = iact(self.tr('Open Folder'), 'ei.folder', QKeySequence.StandardKey.Open,
                                       trig=self.open_folder_by_dialog)
        self.zoom_in_action = iact(self.tr('Zoom In'), 'ei.zoom-in', QKeySequence.StandardKey.ZoomIn,
                                   trig=self.cgv.cust_zoom_in)
        self.zoom_out_action = iact(self.tr('Zoom Out'), 'ei.zoom-out', QKeySequence.StandardKey.ZoomOut,
                                    trig=self.cgv.cust_zoom_out)
        self.fit2screen_action = iact(self.tr('Fit to Screen'), 'mdi6.fit-to-screen-outline', "Alt+F",
                                      trig=lambda: self.cgv.fit2view("screen"))
        self.fit2width_action = iact(self.tr('Fit to Width'), 'ei.resize-horizontal', "Alt+W",
                                     trig=lambda: self.cgv.fit2view("width"))
        self.fit2height_action = iact(self.tr('Fit to Height'), 'ei.resize-vertical', "Alt+H",
                                      trig=lambda: self.cgv.fit2view("height"))
        self.reset_zoom_action = iact(self.tr('Reset Zoom'), 'mdi6.backup-restore', "Ctrl+0",
                                      trig=lambda: self.cgv.fit2view("original"))
        self.prev_img_action = iact(self.tr('Previous Image'), 'ei.arrow-left', "Ctrl+Left",
                                    trig=lambda: self.nav_img(-1))
        self.next_img_action = iact(self.tr('Next Image'), 'ei.arrow-right', "Ctrl+Right",
                                    trig=lambda: self.nav_img(1))
        self.first_img_action = iact(self.tr('First Image'), 'ei.step-backward', "Ctrl+Home",
                                     trig=lambda: self.nav_img("first"))
        self.last_img_action = iact(self.tr('Last Image'), 'ei.step-forward', "Ctrl+End",
                                    trig=lambda: self.nav_img("last"))

        self.deleteAction = iact(self.tr('Delete'), 'ri.delete-bin-7-line', QKeySequence.StandardKey.Delete,
                                 trig=self.on_delete)
        self.group1Action = iact(self.tr('Inside'), 'ph.number-circle-one-fill', "Ctrl+1",
                                 trig=partial(self.on_group, 1))
        self.group2Action = iact(self.tr('Outside'), 'ph.number-circle-two-fill', "Ctrl+2",
                                 trig=partial(self.on_group, 2))
        self.up_action = iact(self.tr('Up'), 'ei.chevron-up', "Ctrl+Up", trig=partial(self.on_move, -1))
        self.down_action = iact(self.tr('Down'), 'ei.chevron-down', "Ctrl+Down", trig=partial(self.on_move, 1))
        self.top_action = iact(self.tr('Top'), 'mdi.arrow-collapse-up', "Alt+Up", trig=partial(self.on_move, '1'))
        self.bottom_action = iact(self.tr('Bottom'), 'mdi.arrow-collapse-down', "Alt+Down",
                                  trig=partial(self.on_move, '-1'))
        self.undo_action = iact(self.tr('Undo'), 'fa5s.undo', QKeySequence.StandardKey.Undo, trig=self.undo)
        self.redo_action = iact(self.tr('Redo'), 'fa5s.redo', QKeySequence.StandardKey.Redo, trig=self.redo)
        self.save_action = iact(self.tr('Save'), 'msc.save', QKeySequence.StandardKey.Save, trig=self.save2lp)

        self.undo_action.setEnabled(False)
        self.redo_action.setEnabled(False)

        self.le_scale_percent.editingFinished.connect(self.scale_by_percent)
        self.update_zoom_label()

    def a5_menubar(self):
        # 文件菜单
        self.file_menu = self.menuBar().addMenu(self.tr('File'))
        self.file_menu.addAction(self.open_folder_action)

        # 显示菜单
        self.view_menu = self.menuBar().addMenu(self.tr('View'))
        # 视图菜单选项
        self.view_menu.addAction(self.pics_dock.toggleViewAction())
        self.view_menu.addSeparator()
        # 缩放选项
        self.view_menu.addAction(self.zoom_in_action)
        self.view_menu.addAction(self.zoom_out_action)
        self.view_menu.addAction(self.fit2screen_action)
        self.view_menu.addAction(self.fit2width_action)
        self.view_menu.addAction(self.fit2height_action)
        self.view_menu.addAction(self.reset_zoom_action)
        self.view_menu.addSeparator()

        # 显示选项
        self.display_modes = [(self.tr('Show Thumbnails'), 0),
                              (self.tr('Show Filenames'), 1),
                              (self.tr('Show Both'), 2)]
        self.display_mode_group = QActionGroup(self)
        for display_mode in self.display_modes:
            action = QAction(display_mode[0], self, checkable=True)
            action.triggered.connect(lambda _, dmode=display_mode[1]: self.cil.set_display_mode(dmode))
            self.view_menu.addAction(action)
            self.display_mode_group.addAction(action)

        # 默认选中 Show Both 选项
        self.display_mode_group.actions()[2].setChecked(True)

        # 导航菜单
        self.nav_menu = self.menuBar().addMenu(self.tr('Navigate'))
        self.nav_menu.addAction(self.prev_img_action)
        self.nav_menu.addAction(self.next_img_action)
        self.nav_menu.addAction(self.first_img_action)
        self.nav_menu.addAction(self.last_img_action)

        # 编辑菜单
        self.edit_menu = self.menuBar().addMenu(self.tr('Edit'))
        self.edit_menu.addAction(self.deleteAction)
        self.edit_menu.addAction(self.group1Action)
        self.edit_menu.addAction(self.group2Action)
        self.edit_menu.addAction(self.up_action)
        self.edit_menu.addAction(self.down_action)
        self.edit_menu.addAction(self.top_action)
        self.edit_menu.addAction(self.bottom_action)
        self.edit_menu.addAction(self.undo_action)
        self.edit_menu.addAction(self.redo_action)

    def a6_toolbar(self):
        self.tool_bar = QToolBar(self)
        self.tool_bar.setObjectName("Toolbar")
        self.tool_bar.setIconSize(QSize(24, 24))
        self.tool_bar.setMovable(False)
        self.tool_bar.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextUnderIcon)
        self.addToolBar(self.tool_bar)
        self.tool_bar.addAction(self.open_folder_action)
        self.tool_bar.addSeparator()
        self.tool_bar.addAction(self.zoom_in_action)
        self.tool_bar.addAction(self.zoom_out_action)
        self.tool_bar.addAction(self.fit2screen_action)
        self.tool_bar.addAction(self.fit2width_action)
        self.tool_bar.addAction(self.fit2height_action)
        self.tool_bar.addAction(self.reset_zoom_action)
        self.tool_bar.addSeparator()
        self.tool_bar.addAction(self.first_img_action)
        self.tool_bar.addAction(self.prev_img_action)
        self.tool_bar.addAction(self.next_img_action)
        self.tool_bar.addAction(self.last_img_action)
        self.tool_bar.addSeparator()
        self.tool_bar.addAction(self.deleteAction)
        self.tool_bar.addAction(self.group1Action)
        self.tool_bar.addAction(self.group2Action)
        self.tool_bar.addAction(self.up_action)
        self.tool_bar.addAction(self.down_action)
        self.tool_bar.addAction(self.top_action)
        self.tool_bar.addAction(self.bottom_action)
        self.tool_bar.addAction(self.undo_action)
        self.tool_bar.addAction(self.redo_action)
        self.tool_bar.addAction(self.save_action)
        self.tool_bar.addSeparator()
        self.tool_bar.addWidget(self.le_scale_percent)
        self.tool_bar.addWidget(QLabel('%'))

    def a9_setting(self):
        self.open_img_by_path(self.img_file)
        self.setCentralWidget(self.cgv)
        self.show()

    def loop_time(self, leng):
        loop = QEventLoop()
        QTimer.singleShot(leng, loop.quit)
        loop.exec_()

    def open_img_by_path(self, img_file):
        if img_file is not None:
            img_file = Path(img_file)
            if img_file.exists() and img_file != self.cgs.img_file:
                self.img_file = img_file
                self.formatted_stem = get_formatted_stem(self.img_file)
                self.img_file_size = getsize(self.img_file)
                self.img_ind = self.img_list.index(self.img_file)
                self.setWindowTitle(self.img_file.name)
                # ================显示图片================
                self.img_raw = imdecode(fromfile(self.img_file, dtype=uint8), -1)
                self.ih, self.iw = self.img_raw.shape[0:2]
                self.cgs.load_qimg(self.img_raw, self.img_file)
                self.scale_by_percent()
                self.update_zoom_label()
                # ================将当前图片项设为选中状态================
                self.cil.blockSignals(True)  # 阻止信号
                self.cil.setCurrentRow(self.img_ind)
                self.cil.blockSignals(False)  # 恢复信号
                # ================更新状态栏信息================
                index_str = f'{self.img_ind + 1}/{len(self.img_list)}'
                meta_str = f'{self.tr("Width")}: {self.iw} {self.tr("Height")}: {self.ih} | {self.tr("File Size")}: {self.img_file_size} bytes'
                status_text = f'{index_str} | {self.tr("Filnename")}: {self.img_file.name} | {meta_str}'
                self.status_bar.showMessage(status_text)
                QApplication.processEvents()
                self.open_lp_bubbles()
        QApplication.processEvents()

    def open_lp_bubbles(self):
        self.undo_stack = []
        self.redo_stack = []
        self.bubble_ellipses.clear()
        rlp_pic_bubbles = self.rlp_para_dic.get(self.formatted_stem, [])
        # ================设置表格行数================
        self.plus_tv_im.setRowCount(len(rlp_pic_bubbles))
        for r in range(len(rlp_pic_bubbles)):
            # ================针对每一个气泡================
            rlp_pic_bubble = rlp_pic_bubbles[r]
            id = rlp_pic_bubble['id']
            coor_x = rlp_pic_bubble['coor_x']
            coor_y = rlp_pic_bubble['coor_y']
            group = rlp_pic_bubble['group']
            content = rlp_pic_bubble['content']
            x = int(self.iw * coor_x)
            y = int(self.ih * coor_y)
            pt_xy = (x, y)
            group_str = ''
            if group == 1:
                group_str = 'G1框内'
            elif group == 2:
                group_str = 'G2框外'

            # 创建一个可拖动的组
            bubble_ellipse = DraggableGroup(x, y, id, group, self)
            self.bubble_ellipses.append(bubble_ellipse)

            item0 = QStandardItem(str(id))
            item1 = QStandardItem(content)
            item2 = QStandardItem(group_str)
            item3 = QStandardItem(f'{x},{y}')
            self.plus_tv_im.setItem(r, 0, item0)
            self.plus_tv_im.setItem(r, 1, item1)
            self.plus_tv_im.setItem(r, 2, item2)
            self.plus_tv_im.setItem(r, 3, item3)

        for idx, bubble_ellipse in enumerate(self.bubble_ellipses):
            self.cgs.addItem(bubble_ellipse)

        # 检查是否有行已经被选中
        if not self.plus_tv.selectionModel().hasSelection() and len(rlp_pic_bubbles) > 0:
            # 选中第一行
            self.plus_tv.selectRow(0)

        self.update_plus_pte()

    @logger.catch
    def update_coordinates(self, item, x, y):
        # 更新表格中的坐标信息
        row_id = item.id - 1
        item3 = QStandardItem(f'{x},{y}')
        self.plus_tv_im.setItem(row_id, 3, item3)

    def update_plus_pte(self):
        current_index = self.plus_tv_sm.currentIndex()
        if current_index.isValid():
            row = current_index.row()
            cell = self.plus_tv_im.item(row, 1)
            if cell:
                content = cell.text()
                self.plus_pte.setPlainText(content)
            else:
                self.plus_pte.setPlainText('')
        else:
            self.plus_pte.setPlainText('')

    def on_text_changed(self):
        text = self.plus_pte.toPlainText()
        current_index = self.plus_tv.currentIndex()
        if not current_index.isValid():
            return
        row = current_index.row()
        self.plus_tv_im.setItem(row, 1, QStandardItem(text))

    def on_delete(self):
        if self.plus_tv.selectionModel().hasSelection():
            selectedRow = self.plus_tv.selectionModel().selectedRows()[0]
            index = selectedRow.row()
            # logger.debug(f'{index=}')
            for i in range(index + 1, len(self.bubble_ellipses)):
                # logger.debug(f'{i=}')
                bubble_ellipse = self.bubble_ellipses[i]
                old_id = bubble_ellipse.id
                new_id = old_id - 1
                bubble_ellipse.change_id(new_id)
                # logger.debug(f'{old_id=}, {new_id=}')
                item = self.plus_tv_im.item(i, 0)  # ID column
                item.setText(str(new_id))
            # 删除场景中对应的气泡
            bubble_ellipse = self.bubble_ellipses.pop(index)
            self.cgs.removeItem(bubble_ellipse)
            self.plus_tv_im.removeRow(index)

            # 选择新行
            new_index = min(index, self.plus_tv_im.rowCount() - 1)
            if new_index >= 0:
                self.plus_tv.selectRow(new_index)

    def on_move(self, move_num):
        current_index = self.plus_tv.currentIndex()
        if not current_index.isValid():
            return

        row = current_index.row()
        if isinstance(move_num, str):
            if move_num == '1':
                # 移到开头
                target_row = 0
            else:  # move_num == '-1'
                # 移到结尾
                target_row = self.plus_tv_im.rowCount() - 1
            items = [self.plus_tv_im.takeItem(row, col) for col in range(self.plus_tv_im.columnCount())]
            self.plus_tv_im.removeRow(row)
            self.plus_tv_im.insertRow(target_row, items)

            # 移动气泡并更新ID
            bubble_to_move = self.bubble_ellipses.pop(row)
            self.bubble_ellipses.insert(target_row, bubble_to_move)
            # 更新所有行的ID和气泡ID
            for idx, bubble in enumerate(self.bubble_ellipses):
                bubble.change_id(idx + 1)
                item0 = QStandardItem(str(idx + 1))
                self.plus_tv_im.setItem(idx, 0, item0)
        else:
            target_row = row + move_num
            if 0 <= target_row < self.plus_tv_im.rowCount():
                items = [self.plus_tv_im.takeItem(row, col) for col in range(self.plus_tv_im.columnCount())]
                self.plus_tv_im.removeRow(row)
                self.plus_tv_im.insertRow(target_row, items)

                src_bubble_ellipse = self.bubble_ellipses[row]
                dst_bubble_ellipse = self.bubble_ellipses[target_row]

                item0 = QStandardItem(str(row + 1))
                target_item0 = QStandardItem(str(target_row + 1))
                self.plus_tv_im.setItem(row, 0, item0)
                self.plus_tv_im.setItem(target_row, 0, target_item0)

                src_bubble_ellipse.change_id(target_row + 1)
                dst_bubble_ellipse.change_id(row + 1)

                # 交换
                self.bubble_ellipses[row], self.bubble_ellipses[target_row] = self.bubble_ellipses[target_row], \
                    self.bubble_ellipses[row]

        # 更新选中
        self.plus_tv.selectRow(target_row)

    def on_create(self, x, y):
        new_id = len(self.bubble_ellipses) + 1
        new_bubble = DraggableGroup(x, y, new_id, 1, self)
        self.cgs.addItem(new_bubble)
        self.bubble_ellipses.append(new_bubble)

        row = self.plus_tv_im.rowCount()
        self.plus_tv_im.insertRow(row)

        item0 = QStandardItem(str(new_id))
        item1 = QStandardItem('')
        item2 = QStandardItem('G1框内')
        item3 = QStandardItem(f'{x},{y}')
        self.plus_tv_im.setItem(row, 0, item0)
        self.plus_tv_im.setItem(row, 1, item1)
        self.plus_tv_im.setItem(row, 2, item2)
        self.plus_tv_im.setItem(row, 3, item3)

        # 选择新行
        self.plus_tv.selectRow(self.plus_tv_im.rowCount() - 1)

    def on_group(self, group):
        current_index = self.plus_tv.currentIndex()
        if not current_index.isValid():
            return

        row = current_index.row()
        if group == 1:
            group_str = 'G1框内'
        elif group == 2:
            group_str = 'G2框外'
        else:
            group_str = ''
        item2 = QStandardItem(group_str)
        self.plus_tv_im.setItem(row, 2, item2)
        bubble_ellipse = self.bubble_ellipses[row]
        bubble_ellipse.change_group(group)

    def open_folder_by_path(self, folder_path):
        # 判断文件夹路径是否存在
        folder_path = Path(folder_path)
        if folder_path.exists():
            # 获取所有图片文件的路径
            img_list = get_valid_imgs(folder_path)
            if img_list and folder_path != self.img_folder:
                self.img_folder = folder_path
                self.auto_subdir = Auto / self.img_folder.name
                make_dir(self.auto_subdir)
                self.img_list = img_list
                self.all_masks = get_valid_imgs(self.img_folder, vmode='mask')
                self.filter_img_list = self.img_list
                self.img_ind = 0
                self.img_file = self.img_list[self.img_ind]
                # ================更新导航栏中的图片列表================
                self.cil.load_img_list()
                self.open_img_by_path(self.img_file)

    def open_folder_by_dialog(self):
        # 如果self.img_folder已经设置，使用其上一级目录作为起始目录，否则使用当前目录
        self.img_folder = Path(self.img_folder) if self.img_folder else None
        start_directory = self.img_folder.parent.as_posix() if self.img_folder else "."
        img_folder = QFileDialog.getExistingDirectory(self, self.tr('Open Folder'), start_directory)
        if img_folder:
            self.open_folder_by_path(img_folder)

    def nav_img(self, nav_step):
        cur_img_path = self.img_list[self.img_ind]
        # 检查当前图片路径是否在过滤后的图片列表中
        if cur_img_path not in self.filter_img_list:
            return
        cur_filter_ind = self.filter_img_list.index(cur_img_path)
        if nav_step == "first":
            new_filter_ind = 0
        elif nav_step == "last":
            new_filter_ind = len(self.filter_img_list) - 1
        else:
            new_filter_ind = cur_filter_ind + nav_step
        if 0 <= new_filter_ind < len(self.filter_img_list):
            new_img_file = self.filter_img_list[new_filter_ind]
            self.open_img_by_path(new_img_file)

    def filter_imgs(self, search_text: str):
        # 获取搜索框的三个条件：是否区分大小写、是否全词匹配、是否使用正则表达式
        case_sensitive = self.search_line.case_sensitive_button.isChecked()
        whole_word = self.search_line.whole_word_button.isChecked()
        use_regex = self.search_line.regex_button.isChecked()

        # 获取对应的正则表达式模式对象
        regex = get_search_regex(search_text, case_sensitive, whole_word, use_regex)
        if not regex:
            return

        # 根据正则表达式筛选图片列表
        self.filter_img_list = [img_file for img_file in self.img_list if regex.search(img_file.name)]

        # 更新缩略图列表：如果图片名匹配正则表达式，显示该项，否则隐藏
        for index in range(self.cil.count()):
            item = self.cil.item(index)
            item_text = item.text()
            item.setHidden(not bool(regex.search(item_text)))

    def update_zoom_label(self):
        self.le_scale_percent.setText(f'{self.cgv.zoom_level * 100:.2f}')

    def scale_by_percent(self):
        target_scale = float(self.le_scale_percent.text()) / 100
        current_scale = self.cgv.transform().m11()
        scale_factor = target_scale / current_scale
        self.cgv.scale(scale_factor, scale_factor)

    def undo(self):
        if not self.undo_stack:
            return

        self.undo_action.setEnabled(len(self.undo_stack) > 0)
        self.redo_action.setEnabled(True)

    def redo(self):
        if not self.redo_stack:
            return

        self.redo_action.setEnabled(len(self.redo_stack) > 0)
        self.undo_action.setEnabled(True)

    def get_table_data(self):
        data_list = []
        for row in range(self.plus_tv_im.rowCount()):
            row_data = []
            for col in range(self.plus_tv_im.columnCount()):
                item = self.plus_tv_im.item(row, col)
                if item is not None:
                    row_data.append(item.text())
                else:
                    row_data.append("")  # 如果某个单元格没有数据，则添加空字符串
            data_list.append(row_data)
        return data_list

    def save2lp(self):
        self.formatted_stem = get_formatted_stem(self.img_file)
        table_data = self.get_table_data()
        if print_type == 'pprint':
            pprint(table_data)
        else:
            table = PrettyTable()
            # 设置表格的标题
            table.field_names = [self.tr('ID'), self.tr('Content'), self.tr('Group'), self.tr('Coordinates'), ]
            # 添加数据到表格
            for t in range(len(table_data)):
                row_data = table_data[t]
                table.add_row(row_data)
            # 打印表格
            print(table)

        rlp_pic_bubbles = []
        for t in range(len(table_data)):
            row_data = table_data[t]
            id, content, group_str, coor_str = row_data
            if group_str == 'G1框内':
                group = 1
            elif group_str == 'G2框外':
                group = 2
            else:
                group = 1
            x_str, par, y_str = coor_str.partition(',')
            coor_x = int(x_str) / self.iw
            coor_y = int(y_str) / self.ih
            rlp_pic_bubble = {}
            rlp_pic_bubble['id'] = id
            rlp_pic_bubble['coor_x'] = coor_x
            rlp_pic_bubble['coor_y'] = coor_y
            rlp_pic_bubble['group'] = group
            rlp_pic_bubble['content'] = content
            rlp_pic_bubbles.append(rlp_pic_bubble)
        self.rlp_para_dic[self.formatted_stem] = rlp_pic_bubbles
        save_rlp(self.rlp_txt, self.rlp_para_dic, self.img_list)


class MistWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.type = self.__class__.__name__
        self.a0_para()
        self.a1_initialize()
        self.a2_status_bar()
        self.a3_docks()
        self.a4_actions()
        self.a5_menubar()
        self.a6_toolbar()
        self.a9_setting()

    def b1_window(self):
        return

    def a0_para(self):
        # ================初始化变量================
        self.screen_icon = qicon('ei.screen')
        self.setWindowIcon(self.screen_icon)
        self.cgs = CustGraphicsScene(self)  # 图形场景对象
        self.cgv = CustGraphicsView(self)  # 图形视图对象
        self.cgv.setScene(self.cgs)
        self.cgv.zoomChanged.connect(self.update_zoom_label)
        self.setWindowTitle(window_title_prefix)
        self.resize(window_w, window_h)

    def a1_initialize(self):
        # ================图片列表================
        self.img_folder = None
        self.auto_subdir = None
        self.img_list = []
        self.img_file = None
        self.image = QImage()
        self.pixmap_item = QGraphicsPixmapItem()
        self.img_ind = -1
        self.display_mode = 0
        self.order_mode = None
        # ================最近文件夹================
        self.recent_folders = []
        # ================设置================
        # 根据需要修改组织和应用名
        self.program_settings = QSettings("MOMO", "MomoTranslator")
        self.media_type_ind = self.program_settings.value('media_type_ind', 0)
        self.media_type = media_type_dict[self.media_type_ind]
        self.media_lang_ind = self.program_settings.value('media_lang_ind', 0)
        self.media_lang = media_lang_dict[self.media_lang_ind]

    def a2_status_bar(self):
        # ================状态栏================
        self.status_bar = QStatusBar()
        # 设置状态栏，类似布局设置
        self.setStatusBar(self.status_bar)

    def a3_docks(self):
        # ================缩略图列表================
        self.nav_tab = QTabWidget(self)
        self.cil = CustImageList(self)
        self.nav_tab.addTab(self.cil, self.tr('Thumbnails'))

        self.search_line = SearchLine(self)
        self.vb_search_nav = QVBoxLayout(self)
        self.vb_search_nav.addWidget(self.search_line)
        self.vb_search_nav.addWidget(self.nav_tab)
        self.pics_widget = QWidget()
        self.pics_widget.setLayout(self.vb_search_nav)

        # 用于显示图片列表
        self.pics_dock = QDockWidget(self.tr('Image List'), self)
        self.pics_dock.setObjectName("PicsDock")
        self.pics_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        self.pics_dock.setWidget(self.pics_widget)

        # ================创建设置工具================
        self.lb_setting = QLabel('翻译设置')
        self.lb_setting.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)

        self.hb_media_type = QHBoxLayout()
        self.lb_media_type = QLabel(self.tr('Layout'))
        self.lb_media_type.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.hb_media_type.addWidget(self.lb_media_type)

        # 创建选项组
        self.media_type_names = [
            self.tr('Comic'),
            self.tr('Manga'),
        ]
        self.bg_media_type = QButtonGroup()
        self.bg_media_type.buttonClicked.connect(self.update_media_type)
        for index, option in enumerate(self.media_type_names):
            self.rb_media_type = QRadioButton(option)
            self.hb_media_type.addWidget(self.rb_media_type)
            self.bg_media_type.addButton(self.rb_media_type, index)
            if index == self.media_type_ind:
                self.rb_media_type.setChecked(True)

        self.hb_media_type.addStretch(1)

        # 创建语言选择下拉框
        self.cb_media_lang = QComboBox()
        self.cb_media_lang.addItems([
            self.tr("English"),
            self.tr("Chinese Simplified"),
            self.tr("Chinese Traditional"),
            self.tr("Japanese"),
            self.tr("Korean"),
        ])
        self.cb_media_lang.setCurrentIndex(self.media_lang_ind)
        self.cb_media_lang.currentIndexChanged.connect(self.update_media_lang)

        # 在布局中添加下拉框
        self.hb_media_lang = QHBoxLayout()
        self.lb_media_lang = QLabel(self.tr('Language'))
        self.lb_media_lang.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.hb_media_lang.addWidget(self.lb_media_lang)
        self.hb_media_lang.addWidget(self.cb_media_lang)
        self.hb_media_lang.addStretch(1)

        self.vb_setting = QVBoxLayout()
        self.vb_setting.addWidget(self.lb_setting)
        self.vb_setting.addLayout(self.hb_media_type)
        self.vb_setting.addLayout(self.hb_media_lang)
        self.vb_setting.addStretch(1)
        self.setting_tool = QWidget()
        self.setting_tool.setLayout(self.vb_setting)

        # ================创建文本工具================
        self.lb_text = QLabel('为LabelPlus提供兼容')
        self.lb_text.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.vb_text = QVBoxLayout()
        self.vb_text.addWidget(self.lb_text)
        self.vb_text.addStretch(1)
        self.text_tool = QWidget()
        self.text_tool.setLayout(self.vb_text)

        # ================创建图层工具================
        self.lb_layer = QLabel('为PS提供兼容')
        self.lb_layer.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.vb_layer = QVBoxLayout()
        self.vb_layer.addWidget(self.lb_layer)
        self.vb_layer.addStretch(1)
        self.layer_tool = QWidget()
        self.layer_tool.setLayout(self.vb_layer)

        # ================创建步骤工具================
        self.hb_step = QHBoxLayout()
        self.pb_step = QPushButton(self.tr('Start'))
        self.pb_step.clicked.connect(self.istart_task)
        self.hb_step.addWidget(self.pb_step)

        # 创建选项组
        self.task_names = [
            self.tr('Analyze Frames'),
            self.tr('Analyze Bubbles'),
            self.tr('Order Bubbles'),
            self.tr('OCR'),
            self.tr('Translate'),
        ]
        self.bg_step = QButtonGroup()
        for index, option in enumerate(self.task_names):
            self.rb_task = QRadioButton(option)
            self.hb_step.addWidget(self.rb_task)
            self.bg_step.addButton(self.rb_task, index)
            # 默认选中第一个 QRadioButton
            if index == 0:
                self.rb_task.setChecked(True)

        # 创建进度条
        self.pb_task = QProgressBar()
        self.hb_step.addWidget(self.pb_task)
        self.step_tool = QWidget()
        self.step_tool.setLayout(self.hb_step)

        self.setting_dock = QDockWidget(self.tr('Setting'), self)
        self.setting_dock.setObjectName("SettingDock")
        self.setting_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        self.setting_dock.setWidget(self.setting_tool)
        self.setting_dock.setMinimumWidth(200)

        self.text_dock = QDockWidget(self.tr('Text'), self)
        self.text_dock.setObjectName("TextDock")
        self.text_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        self.text_dock.setWidget(self.text_tool)
        self.text_dock.setMinimumWidth(200)

        self.layer_dock = QDockWidget(self.tr('Layer'), self)
        self.layer_dock.setObjectName("LayerDock")
        self.layer_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        self.layer_dock.setWidget(self.layer_tool)
        self.layer_dock.setMinimumWidth(200)

        self.step_dock = QDockWidget(self.tr('Step'), self)
        self.step_dock.setObjectName("StepDock")
        self.step_dock.setAllowedAreas(Qt.DockWidgetArea.BottomDockWidgetArea | Qt.DockWidgetArea.TopDockWidgetArea)
        self.step_dock.setWidget(self.step_tool)
        self.step_dock.setMinimumWidth(200)

        self.addDockWidget(Qt.DockWidgetArea.LeftDockWidgetArea, self.pics_dock)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.setting_dock)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.text_dock)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.layer_dock)
        self.addDockWidget(Qt.DockWidgetArea.BottomDockWidgetArea, self.step_dock)

    def a4_actions(self):
        # 添加缩放百分比输入框
        self.le_scale_percent = QLineEdit(self)
        self.le_scale_percent.setFixedWidth(50)
        self.le_scale_percent.setValidator(QDoubleValidator(1, 1000, 2))

        self.open_folder_action = iact(self.tr('Open Folder'), 'ei.folder', QKeySequence.StandardKey.Open,
                                       trig=self.open_folder_by_dialog)
        self.save_img_action = iact(self.tr('Save Image'), 'ri.image-edit-fill', QKeySequence.StandardKey.Save,
                                    trig=self.save_img)
        self.zoom_in_action = iact(self.tr('Zoom In'), 'ei.zoom-in', QKeySequence.StandardKey.ZoomIn,
                                   trig=self.cgv.cust_zoom_in)
        self.zoom_out_action = iact(self.tr('Zoom Out'), 'ei.zoom-out', QKeySequence.StandardKey.ZoomOut,
                                    trig=self.cgv.cust_zoom_out)
        self.fit2screen_action = iact(self.tr('Fit to Screen'), 'mdi6.fit-to-screen-outline', "Alt+F",
                                      trig=lambda: self.cgv.fit2view("screen"))
        self.fit2width_action = iact(self.tr('Fit to Width'), 'ei.resize-horizontal', "Alt+W",
                                     trig=lambda: self.cgv.fit2view("width"))
        self.fit2height_action = iact(self.tr('Fit to Height'), 'ei.resize-vertical', "Alt+H",
                                      trig=lambda: self.cgv.fit2view("height"))
        self.reset_zoom_action = iact(self.tr('Reset Zoom'), 'mdi6.backup-restore', "Ctrl+0",
                                      trig=lambda: self.cgv.fit2view("original"))
        self.prev_img_action = iact(self.tr('Previous Image'), 'ei.arrow-left', "Ctrl+Left",
                                    trig=lambda: self.nav_img(-1))
        self.next_img_action = iact(self.tr('Next Image'), 'ei.arrow-right', "Ctrl+Right",
                                    trig=lambda: self.nav_img(1))
        self.first_img_action = iact(self.tr('First Image'), 'ei.step-backward', "Ctrl+Home",
                                     trig=lambda: self.nav_img("first"))
        self.last_img_action = iact(self.tr('Last Image'), 'ei.step-forward', "Ctrl+End",
                                    trig=lambda: self.nav_img("last"))
        self.about_action = iact(f"{self.tr('About')} {APP_NAME}", None,
                                 trig=self.show_about_dialog)
        self.about_qt_action = iact(f"{self.tr('About')} Qt", None,
                                    trig=QApplication.instance().aboutQt)
        self.help_document_action = iact(f"{APP_NAME} {self.tr('Help')}", None,
                                         trig=self.show_help_document)
        self.feedback_action = iact('Bug Report', None, trig=self.show_feedback_dialog)
        self.update_action = iact('Update Online', None, trig=self.check_for_updates)

        self.le_scale_percent.editingFinished.connect(self.scale_by_percent)
        self.update_zoom_label()

    def a5_menubar(self):
        # 文件菜单
        self.file_menu = self.menuBar().addMenu(self.tr('File'))
        self.file_menu.addAction(self.open_folder_action)
        self.file_menu.addAction(self.save_img_action)
        self.file_menu.addSeparator()
        # 添加最近打开文件夹菜单项
        self.recent_folders_menu = self.file_menu.addMenu(self.tr('Recent Folders'))
        self.update_recent_folders_menu()

        # 显示菜单
        self.view_menu = self.menuBar().addMenu(self.tr('View'))
        # 添加视图菜单选项
        self.view_menu.addAction(self.pics_dock.toggleViewAction())
        self.view_menu.addAction(self.setting_dock.toggleViewAction())
        self.view_menu.addAction(self.text_dock.toggleViewAction())
        self.view_menu.addAction(self.layer_dock.toggleViewAction())
        self.view_menu.addAction(self.step_dock.toggleViewAction())
        self.view_menu.addSeparator()

        # 添加缩放选项
        self.view_menu.addAction(self.zoom_in_action)
        self.view_menu.addAction(self.zoom_out_action)
        self.view_menu.addSeparator()
        self.view_menu.addAction(self.fit2screen_action)
        self.view_menu.addAction(self.fit2width_action)
        self.view_menu.addAction(self.fit2height_action)
        self.view_menu.addAction(self.reset_zoom_action)
        self.view_menu.addSeparator()

        # 添加显示选项
        self.display_modes = [(self.tr('Show Thumbnails'), 0),
                              (self.tr('Show Filenames'), 1),
                              (self.tr('Show Both'), 2)]
        self.display_mode_group = QActionGroup(self)
        for display_mode in self.display_modes:
            action = QAction(display_mode[0], self, checkable=True)
            action.triggered.connect(lambda _, dmode=display_mode[1]: self.cil.set_display_mode(dmode))
            self.view_menu.addAction(action)
            self.display_mode_group.addAction(action)

        # 默认选中 Show Both 选项
        self.display_mode_group.actions()[2].setChecked(True)

        # 编辑菜单
        self.edit_menu = self.menuBar().addMenu(self.tr('Edit'))

        # 导航菜单
        self.nav_menu = self.menuBar().addMenu(self.tr('Navigate'))
        self.nav_menu.addAction(self.prev_img_action)
        self.nav_menu.addAction(self.next_img_action)
        self.nav_menu.addAction(self.first_img_action)
        self.nav_menu.addAction(self.last_img_action)

        # 帮助菜单
        self.help_menu = self.menuBar().addMenu(self.tr('Help'))
        self.help_menu.addAction(self.about_action)
        self.help_menu.addAction(self.about_qt_action)
        self.help_menu.addAction(self.help_document_action)
        self.help_menu.addAction(self.feedback_action)
        self.help_menu.addAction(self.update_action)

    def a6_toolbar(self):

        # 添加顶部工具栏
        self.tool_bar = QToolBar(self.tr('Toolbar'), self)
        self.tool_bar.setObjectName("Toolbar")
        self.tool_bar.setIconSize(QSize(24, 24))
        self.tool_bar.setMovable(False)
        self.tool_bar.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextUnderIcon)
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea, self.tool_bar)
        self.tool_bar.addAction(self.open_folder_action)
        self.tool_bar.addAction(self.save_img_action)
        self.tool_bar.addSeparator()
        self.tool_bar.addAction(self.zoom_in_action)
        self.tool_bar.addAction(self.zoom_out_action)
        self.tool_bar.addAction(self.fit2screen_action)
        self.tool_bar.addAction(self.fit2width_action)
        self.tool_bar.addAction(self.fit2height_action)
        self.tool_bar.addAction(self.reset_zoom_action)
        self.tool_bar.addSeparator()
        self.tool_bar.addAction(self.first_img_action)
        self.tool_bar.addAction(self.prev_img_action)
        self.tool_bar.addAction(self.next_img_action)
        self.tool_bar.addAction(self.last_img_action)
        self.tool_bar.addWidget(self.le_scale_percent)
        self.tool_bar.addWidget(QLabel('%'))

    def a9_setting(self):
        # 读取上一次打开的文件夹、窗口位置和状态
        last_opened_folder = self.program_settings.value('last_opened_folder', '')
        geometry = self.program_settings.value('window_geometry')
        state = self.program_settings.value('window_state')
        # 如果上一次打开的文件夹存在，则打开它
        if last_opened_folder and exists(last_opened_folder) and isdir(last_opened_folder):
            self.open_folder_by_path(last_opened_folder)
        # 如果上一次有记录窗口位置，则恢复窗口位置
        if geometry is not None:
            self.restoreGeometry(geometry)
        # 如果上一次有记录窗口状态，则恢复窗口状态
        if state is not None:
            self.restoreState(state)

        # 将 QGraphicsView 设置为中心窗口部件
        self.setCentralWidget(self.cgv)

        # 显示窗口
        self.show()

    def open_img_by_path(self, img_file):
        img_file = Path(img_file)
        if img_file.exists() and img_file != self.cgs.img_file:
            self.img_file = img_file
            self.img_file_size = getsize(self.img_file)
            self.img_ind = self.img_list.index(self.img_file)
            # ================显示图片================
            self.img_raw = imdecode(fromfile(self.img_file, dtype=uint8), -1)
            self.ih, self.iw = self.img_raw.shape[0:2]
            self.cgs.load_qimg(self.img_raw, self.img_file)
            self.scale_by_percent()
            self.update_zoom_label()
            # ================将当前图片项设为选中状态================
            self.cil.blockSignals(True)  # 阻止信号
            self.cil.setCurrentRow(self.img_ind)
            self.cil.blockSignals(False)  # 恢复信号
            # ================更新状态栏信息================
            index_str = f'{self.img_ind + 1}/{len(self.img_list)}'
            meta_str = f'{self.tr("Width")}: {self.iw} {self.tr("Height")}: {self.ih} | {self.tr("File Size")}: {self.img_file_size} bytes'
            status_text = f'{index_str} | {self.tr("Filnename")}: {self.img_file.name} | {meta_str}'
            self.status_bar.showMessage(status_text)

    def open_folder_by_path(self, folder_path):
        # 打开上次关闭程序时用到的文件夹
        # 打开最近的文件夹
        # 打开文件夹

        # 判断文件夹路径是否存在
        folder_path = Path(folder_path)
        if folder_path.exists():
            # 获取所有图片文件的路径
            img_list = get_valid_imgs(folder_path)
            if img_list:
                self.img_folder = folder_path
                self.auto_subdir = Auto / self.img_folder.name
                self.img_list = img_list
                self.filter_img_list = self.img_list
                self.img_ind = 0
                self.img_file = self.img_list[self.img_ind]

                # ================更新导航栏中的图片列表================
                self.cil.load_img_list()
                self.open_img_by_path(self.img_file)

                self.setWindowTitle(f'{window_title_prefix} - {self.img_folder}')
                # 将该文件夹添加到最近使用文件夹列表
                self.add_recent_folder(self.img_folder)

    def open_folder_by_dialog(self):
        # 如果self.img_folder已经设置，使用其上一级目录作为起始目录，否则使用当前目录
        self.img_folder = Path(self.img_folder) if self.img_folder else None
        start_directory = self.img_folder.parent.as_posix() if self.img_folder else "."
        img_folder = QFileDialog.getExistingDirectory(self, self.tr('Open Folder'), start_directory)
        if img_folder:
            self.open_folder_by_path(img_folder)

    def update_recent_folders_menu(self):
        self.recent_folders_menu.clear()
        recent_folders = self.program_settings.value("recent_folders", [])
        for folder in recent_folders:
            action = QAction(str(folder), self)
            action.triggered.connect(lambda checked, p=folder: self.open_folder_by_path(p))
            self.recent_folders_menu.addAction(action)

    def add_recent_folder(self, folder_path):
        recent_folders = self.program_settings.value("recent_folders", [])
        recent_folders = reduce_list(recent_folders)
        if folder_path in recent_folders:
            recent_folders.remove(folder_path)
        recent_folders.insert(0, folder_path)
        # 保留最多10个最近文件夹
        recent_folders = recent_folders[:10]

        self.program_settings.setValue("recent_folders", recent_folders)
        self.update_recent_folders_menu()

    def nav_img(self, nav_step):
        cur_img_path = self.img_list[self.img_ind]
        # 检查当前图片路径是否在过滤后的图片列表中
        if cur_img_path not in self.filter_img_list:
            return
        cur_filter_ind = self.filter_img_list.index(cur_img_path)
        if nav_step == "first":
            new_filter_ind = 0
        elif nav_step == "last":
            new_filter_ind = len(self.filter_img_list) - 1
        else:
            new_filter_ind = cur_filter_ind + nav_step
        if 0 <= new_filter_ind < len(self.filter_img_list):
            new_img_file = self.filter_img_list[new_filter_ind]
            self.open_img_by_path(new_img_file)

    def filter_imgs(self, search_text: str):
        # 获取搜索框的三个条件：是否区分大小写、是否全词匹配、是否使用正则表达式
        case_sensitive = self.search_line.case_sensitive_button.isChecked()
        whole_word = self.search_line.whole_word_button.isChecked()
        use_regex = self.search_line.regex_button.isChecked()

        # 获取对应的正则表达式模式对象
        regex = get_search_regex(search_text, case_sensitive, whole_word, use_regex)
        if not regex:
            return

        # 根据正则表达式筛选图片列表
        self.filter_img_list = [img_file for img_file in self.img_list if regex.search(img_file.name)]

        # 更新缩略图列表：如果图片名匹配正则表达式，显示该项，否则隐藏
        for index in range(self.cil.count()):
            item = self.cil.item(index)
            item_text = item.text()
            item.setHidden(not bool(regex.search(item_text)))

    def save_img(self):
        file_name, _ = QFileDialog.getSaveFileName(self, self.tr('Save Image'), "", "Images (*.png *.xpm *.jpg)")
        if file_name:
            image = QImage(self.cgs.sceneRect().size().toSize(), QImage.Format.Format_ARGB32)
            image.fill(Qt.GlobalColor.transparent)
            painter = QPainter(image)
            self.cgs.render(painter)
            painter.end()
            image.save(file_name)

    def update_zoom_label(self):
        self.le_scale_percent.setText(f'{self.cgv.zoom_level * 100:.2f}')

    def scale_by_percent(self):
        target_scale = float(self.le_scale_percent.text()) / 100
        current_scale = self.cgv.transform().m11()
        scale_factor = target_scale / current_scale
        self.cgv.scale(scale_factor, scale_factor)

    def show_about_dialog(self):
        about_dialog = QDialog(self)
        about_dialog.setWindowTitle(f"{self.tr('About')} {APP_NAME}")

        app_name_label = QLabel(APP_NAME)
        app_name_label.setFont(QFont("Arial", 20, QFont.Bold))

        version_label = QLabel(f"版本: {APP_VERSION}")
        libraries_label = QLabel("使用的库：PyQt6")
        license_label = QLabel("许可证: MIT License")

        vb_dialog = QVBoxLayout()
        vb_dialog.addWidget(app_name_label)
        vb_dialog.addWidget(version_label)
        vb_dialog.addWidget(libraries_label)
        vb_dialog.addWidget(license_label)
        about_dialog.setLayout(vb_dialog)
        about_dialog.exec_()

    def show_help_document(self):
        # 在此处显示帮助文档
        pass

    def show_feedback_dialog(self):
        recipient_email = "annebing@qq.com"
        subject = "Bug Report - 漫画翻译工具"
        body = "请在此处详细描述您遇到的问题：\n\n\n\n软件版本：1.0.0"
        mailto_url = f"mailto:{recipient_email}?subject={subject}&body={body}"
        try:
            webbrowser.open(mailto_url)
        except webbrowser.Error as e:
            QMessageBox.warning(self, "错误", f"无法打开邮件客户端：{e}")

    def check_for_updates(self):
        # 在此处实现在线更新软件的功能
        pass

    @timer_decorator
    # @logger.catch
    def step0_analyze_frames(self):
        frame_data = iload_data(self.frame_yml)
        # ================分析画格================
        processed_imgs = 0
        for p, img_file in enumerate(self.img_list):
            img_file, frame_grid_strs = analyze1frame(img_file, frame_data, self.auto_subdir, media_type)
            frame_data[img_file.name] = frame_grid_strs

            processed_imgs += 1
            self.pb_task.setValue(int(processed_imgs / len(self.img_list) * 100))
            QApplication.processEvents()

            pprint(frame_grid_strs)
            if len(frame_grid_strs) >= 2:
                grid_masks, marked_frames = get_grid_masks(img_file, frame_grid_strs)
                marked_frames_jpg = self.auto_subdir / f'{img_file.stem}-画格.jpg'
                marked_frames = imdecode(fromfile(marked_frames_jpg, dtype=uint8), -1)
                self.cgs.load_qimg(marked_frames, marked_frames_jpg)
                QApplication.processEvents()

        self.pb_task.setValue(100)
        QApplication.processEvents()
        frame_data_sorted = {k: frame_data[k] for k in natsorted(frame_data)}
        with open(self.frame_yml, 'w') as yml_file:
            yaml.dump(frame_data_sorted, yml_file)

    @timer_decorator
    # @logger.catch
    def step1_analyze_bubbles(self):
        frame_data = iload_data(self.frame_yml)
        processed_imgs = 0
        all_masks_old = get_valid_imgs(self.img_folder, vmode='mask')

        for p, img_file in enumerate(self.img_list):
            logger.warning(f'{img_file=}')
            img_raw = imdecode(fromfile(img_file, dtype=uint8), -1)
            img_raw = toBGR(img_raw)
            ih, iw = img_raw.shape[0:2]
            # ================矩形画格信息================
            frame_grid_strs = frame_data.get(img_file.name, [f'0,0,{iw},{ih}~0,0,{iw},{ih}'])
            # ================模型检测文字，文字显示为白色================
            if use_torch and CTD_model is not None:
                CTD_mask = get_CTD_mask(img_file)
            else:
                CTD_mask = None
            # ================针对每一张图================
            for c in range(len(color_patterns)):
                # ================遍历每种气泡文字颜色组合================
                color_pattern = color_patterns[c]
                cp_mask_cnt_png = get_bubbles_by_cp(img_file, color_pattern, frame_grid_strs, CTD_mask, media_type,
                                                    self.auto_subdir)
                if cp_mask_cnt_png.exists():
                    transparent_img = imdecode(fromfile(cp_mask_cnt_png, dtype=uint8), -1)
                    self.cgs.load_qimg(transparent_img, cp_mask_cnt_png)
                QApplication.processEvents()

            processed_imgs += 1
            self.pb_task.setValue(int(processed_imgs / len(self.img_list) * 100))
            QApplication.processEvents()

        # ================搬运气泡蒙版================
        auto_all_masks = get_valid_imgs(self.auto_subdir, vmode='mask')
        # 如果步骤开始前没有气泡蒙版
        if not all_masks_old:
            for mask_src in auto_all_masks:
                mask_dst = self.img_folder / mask_src.name
                copy2(mask_src, mask_dst)

        # Update progress bar to 100%
        self.pb_task.setValue(100)
        QApplication.processEvents()

    @timer_decorator
    # @logger.catch
    def step2_order(self):
        frame_data = iload_data(self.frame_yml)
        order_data = iload_data(self.order_yml)
        all_masks = get_valid_imgs(self.img_folder, vmode='mask')
        processed_imgs = 0

        for p, img_file in enumerate(self.img_list):
            logger.warning(f'{img_file=}')
            img_raw = imdecode(fromfile(img_file, dtype=uint8), -1)
            ih, iw = img_raw.shape[0:2]

            # ================矩形画格信息================
            # 从 frame_data 中获取画格信息，默认为整个图像的矩形区域
            frame_grid_strs = frame_data.get(img_file.name, [f'0,0,{iw},{ih}~0,0,{iw},{ih}'])
            bubble_order_strs = order_data.get(img_file.name, [])
            grid_masks, marked_frames = get_grid_masks(img_file, frame_grid_strs)
            order_preview_jpg = self.auto_subdir / f'{img_file.stem}-气泡排序.jpg'

            # ================获取对应的文字图片================
            mask_pics = [x for x in all_masks if x.stem.startswith(img_file.stem)]
            if mask_pics:
                single_cnts = get_single_cnts(img_raw, mask_pics)
                logger.debug(f'{len(single_cnts)=}')
                single_cnts_grids_ordered = get_ordered_cnts(single_cnts, img_file, grid_masks, bubble_order_strs,
                                                             media_type)
                ordered_cnts = list(chain(*single_cnts_grids_ordered))
                order_preview = get_order_preview(marked_frames, single_cnts_grids_ordered)
                write_pic(order_preview_jpg, order_preview)
                self.cgs.load_qimg(order_preview, order_preview_jpg)

            processed_imgs += 1
            self.pb_task.setValue(int(processed_imgs / len(self.img_list) * 100))
            QApplication.processEvents()

        # Update progress bar to 100%
        self.pb_task.setValue(100)
        QApplication.processEvents()

    def step3_OCR(self):
        frame_data = iload_data(self.frame_yml)
        order_data = iload_data(self.order_yml)
        ocr_data = iload_data(self.ocr_yml)
        # ================分析画格================
        processed_imgs = 0
        # ================气泡蒙版================
        all_masks = get_valid_imgs(self.img_folder, vmode='mask')
        ocr_doc = Document()
        for i in range(len(self.img_list)):
            img_file = self.img_list[i]
            logger.warning(f'{img_file=}')
            pic_results = ocr1pic(img_file, frame_data, order_data, ocr_data, all_masks, media_type, media_lang, vert)
            ocr_doc, all_cropped_imgs = update_ocr_doc(ocr_doc, pic_results, self.img_folder, i, self.img_list,
                                                       media_type)

            self.pb_task.setValue(int(processed_imgs / len(self.img_list) * 100))
            QApplication.processEvents()
            processed_imgs += 1

        # ================保存为DOCX================
        write_docx(self.ocr_docx, ocr_doc)
        self.pb_task.setValue(100)

    def istart_task(self, checked):
        # 当使用 QPushButton 的 clicked 信号连接到槽时，它会传递一个布尔值表示按钮是否被选中
        # 检查哪个 QRadioButton 被选中并执行相应的任务
        self.task_name = self.bg_step.checkedButton().text()
        self.task_ind = self.bg_step.checkedId()
        logger.debug(f'[{self.task_ind}]{self.task_name} - {self.img_folder}')
        # 在这里执行您的任务逻辑
        if self.img_folder is not None:
            self.img_folder = Path(self.img_folder)
            if self.img_folder.exists():
                self.frame_yml = self.img_folder.parent / f'{self.img_folder.name}.yml'
                self.order_yml = self.img_folder.parent / f'{self.img_folder.name}-气泡排序.yml'
                self.ocr_yml = self.img_folder.parent / f'{self.img_folder.name}-文字识别.yml'
                self.ocr_docx = self.img_folder.parent / f'{self.img_folder.name}-1识别.docx'
                self.auto_subdir = Auto / self.img_folder.name
                make_dir(self.auto_subdir)

                if self.task_ind == 0:
                    self.step0_analyze_frames()
                elif self.task_ind == 1:
                    self.step1_analyze_bubbles()
                elif self.task_ind == 2:
                    self.step2_order()
                elif self.task_ind == 3:
                    self.step3_OCR()
                elif self.task_ind == 4:
                    step5_translate(self.img_folder)

                # 自动选中下一个步骤，除非当前步骤是最后一个步骤
                if self.task_ind < self.bg_step.id(self.bg_step.buttons()[-1]):
                    next_task_ind = self.task_ind + 1
                    next_task_button = self.bg_step.button(next_task_ind)
                    if next_task_button:
                        next_task_button.setChecked(True)
                        self.task_ind = next_task_ind

        QApplication.processEvents()

    def update_media_type(self, button):
        index = self.bg_media_type.id(button)
        self.media_type_ind = index
        self.media_type = media_type_dict[index]

    def update_media_lang(self, index):
        self.media_lang_ind = index
        self.media_lang = media_lang_dict[index]

    def closeEvent(self, event):
        # 如果有打开图片的文件夹，将其保存到程序设置中
        if self.img_folder:
            self.program_settings.setValue('last_opened_folder', self.img_folder)
        else:
            self.program_settings.setValue('last_opened_folder', '')
        # 保存窗口的几何和状态信息到程序设置中
        self.program_settings.setValue('window_geometry', self.saveGeometry())
        self.program_settings.setValue('window_state', self.saveState())
        self.program_settings.setValue('media_type_ind', self.media_type_ind)
        self.program_settings.setValue('media_lang_ind', self.media_lang_ind)
        event.accept()


@logger.catch
def order_qt(appgui):
    order_window = OrderWindow()
    sys.exit(appgui.exec())


@logger.catch
def lp_qt(appgui):
    lp_window = LabelPlusWindow()
    sys.exit(appgui.exec())


@logger.catch
def mist_qt(appgui):
    mist_window = MistWindow()
    sys.exit(appgui.exec())


@logger.catch
def get_line_segments(mask, target_color='ffffff'):
    """
    根据给定的掩码获取非目标颜色的线段。
    每个线段都表示为一个 (start, end) 元组。
    """
    # 根据目标颜色设置目标值
    if target_color == 'ffffff':
        target_value = 255
    elif target_color == '000000':
        target_value = 0
    else:
        raise ValueError("不支持的颜色。")

    segments = []
    start = None
    # 遍历掩码，找到非目标颜色的线段
    for i, value in enumerate(mask):
        if value != target_value:
            if start is None:
                start = i
        else:
            if start is not None:
                segments.append((start, i))
                start = None
    if start is not None:
        segments.append((start, len(mask)))
    return segments


@logger.catch
def get_comb_mask(lower_bound, upper_bound, slices):
    masks = [inRange(slice, lower_bound, upper_bound) for slice in slices]
    combined_mask = reduce(bitwise_and, masks)
    return combined_mask


# @logger.catch
def get_added_frames(frame_grid_strs, img_raw, color_name0):
    img_raw = toBGR(img_raw)
    directions = ['top', 'bottom', 'left', 'right']
    # 根据边框颜色获取绘制用颜色
    if color_name0 == 'white':
        pic_frame_color = 'ffffff'
        color = color_black
        lower_bound = lower_bound_white
        upper_bound = upper_bound_white
    else:
        pic_frame_color = '000000'
        color = color_white
        lower_bound = lower_bound_black
        upper_bound = upper_bound_black
    target_color = pic_frame_color
    logger.debug(f'{pic_frame_color=}')

    top_gap = gaps[0]
    bottom_gap = gaps[1]
    left_gap = gaps[2]
    right_gap = gaps[3]

    for frame_grid_str in frame_grid_strs:
        int_values = list(map(int, findall(r'\d+', frame_grid_str)))
        x, y, w, h, xx, yy, ww, hh = int_values

        if frame_level == 3:
            masks = {
                'top': get_comb_mask(lower_bound, upper_bound,
                                     [img_raw[yy + 2:yy + 3, xx:xx + ww],
                                      img_raw[yy + 3:yy + 4, xx:xx + ww],
                                      img_raw[yy + 4:yy + 5, xx:xx + ww],
                                      ]),
                'bottom': get_comb_mask(lower_bound, upper_bound,
                                        [img_raw[yy + hh - 3:yy + hh - 2, xx:xx + ww],
                                         img_raw[yy + hh - 2:yy + hh - 1, xx:xx + ww],
                                         img_raw[yy + hh - 1:yy + hh, xx:xx + ww],
                                         ]),
                'left': get_comb_mask(lower_bound, upper_bound,
                                      [img_raw[yy:yy + hh, xx + 1:xx + 2],
                                       img_raw[yy:yy + hh, xx + 2:xx + 3],
                                       img_raw[yy:yy + hh, xx + 3:xx + 4],
                                       ]),
                'right': get_comb_mask(lower_bound, upper_bound,
                                       [img_raw[yy:yy + hh, xx + ww - 3:xx + ww - 2],
                                        img_raw[yy:yy + hh, xx + ww - 2:xx + ww - 1],
                                        img_raw[yy:yy + hh, xx + ww - 1:xx + ww],
                                        ])
            }
        elif frame_level == 2:
            masks = {
                'top': get_comb_mask(lower_bound, upper_bound,
                                     [img_raw[yy + 2:yy + 3, xx:xx + ww],
                                      img_raw[yy + 3:yy + 4, xx:xx + ww],
                                      ]),
                'bottom': get_comb_mask(lower_bound, upper_bound,
                                        [img_raw[yy + hh - 2:yy + hh - 1, xx:xx + ww],
                                         img_raw[yy + hh - 1:yy + hh, xx:xx + ww],
                                         ]),
                'left': get_comb_mask(lower_bound, upper_bound,
                                      [img_raw[yy:yy + hh, xx + 1:xx + 2],
                                       img_raw[yy:yy + hh, xx + 2:xx + 3],
                                       ]),
                'right': get_comb_mask(lower_bound, upper_bound,
                                       [img_raw[yy:yy + hh, xx + ww - 2:xx + ww - 1],
                                        img_raw[yy:yy + hh, xx + ww - 1:xx + ww],
                                        ])
            }
        else:
            masks = {
                'top': get_comb_mask(lower_bound, upper_bound, [img_raw[yy:yy + 1, xx:xx + ww], ]),
                'bottom': get_comb_mask(lower_bound, upper_bound, [img_raw[yy + hh - 1:yy + hh, xx:xx + ww], ]),
                'left': get_comb_mask(lower_bound, upper_bound, [img_raw[yy:yy + hh, xx + 1:xx + 2], ]),
                'right': get_comb_mask(lower_bound, upper_bound, [img_raw[yy:yy + hh, xx + ww - 1:xx + ww], ])
            }

        masks_segments = {
            direction: get_line_segments(squeeze(mask), target_color=target_color) for direction, mask in masks.items()
        }

        for direction, segments in masks_segments.items():
            total_length = ww if direction in ['top', 'bottom'] else hh
            non_white_length = sum([(seg[1] - seg[0]) for seg in segments])
            percentage = non_white_length / total_length
            gap_value = gaps[directions.index(direction)]

            if percentage <= frame_thres:
                if gap_value == 0:
                    draw_points = {
                        'top': ((xx, yy), (xx + ww, yy)),
                        'bottom': ((xx, yy + hh), (xx + ww, yy + hh)),
                        'left': ((xx, yy), (xx, yy + hh)),
                        'right': ((xx + ww, yy), (xx + ww, yy + hh))
                    }
                    line(img_raw, draw_points[direction][0], draw_points[direction][1], color, 2)
                else:
                    for offset in range(-2 if direction in ['top', 'bottom'] else -1, gap_value + 1):
                        for start, end in segments:
                            draw_points = {
                                'top': ((xx + start, yy - offset), (xx + end, yy - offset)),
                                'bottom': ((xx + start, yy + hh + offset), (xx + end, yy + hh + offset)),
                                'left': ((xx - offset, yy + start), (xx - offset, yy + end)),
                                'right': ((xx + ww + offset, yy + start), (xx + ww + offset, yy + end))
                            }
                            if abs(start - end) >= min_outline_len:
                                line(img_raw, draw_points[direction][0], draw_points[direction][1], color, 1)

                    # 绘制外侧扩大的矩形，考虑gaps延长线段
                    draw_points = {
                        'top': ((xx - left_gap, yy - top_gap), (xx + ww + right_gap, yy - top_gap)),
                        'bottom': ((xx - left_gap, yy + hh + bottom_gap), (xx + ww + right_gap, yy + hh + bottom_gap)),
                        'left': ((xx - left_gap, yy - top_gap), (xx - left_gap, yy + hh + bottom_gap)),
                        'right': ((xx + ww + right_gap, yy - top_gap), (xx + ww + right_gap, yy + hh + bottom_gap))
                    }
                    line(img_raw, draw_points[direction][0], draw_points[direction][1], color, 2)
    return img_raw


# @logger.catch
def compute_frame_mask(img_raw, dominant_colors, tolerance):
    masks = []
    for i, dominant_color in enumerate(dominant_colors):
        dominant_color_int32 = array(dominant_color[:3]).astype(int32)
        lower_bound = maximum(0, dominant_color_int32 - tolerance)
        upper_bound = minimum(255, dominant_color_int32 + tolerance)
        mask = inRange(img_raw, lower_bound, upper_bound)
        masks.append(mask)
    # 将多个掩码相加
    combined_mask = np.sum(masks, axis=0)
    # 将 combined_mask 转换为 uint8 类型
    combined_mask_uint8 = combined_mask.astype(uint8)
    bubble_mask = combined_mask_uint8
    letter_mask = bitwise_not(bubble_mask)
    filter_cnts = get_raw_bubbles(bubble_mask, letter_mask, None, None, None)
    for f in range(len(filter_cnts)):
        bubble_cnt = filter_cnts[f]
        combined_mask_uint8 = drawContours(combined_mask_uint8, [bubble_cnt.contour], 0, 0, -1)
    return combined_mask_uint8


@logger.catch
def get_proj_img(img_array, direction, target_color="black"):
    """
    根据输入的二维图像数组和投影方向，创建投影图像并返回白色像素的投影。

    :param img_array: 一个表示二值化图像的二维数组，形状为 (height, width)。
    :param direction: 投影方向，可以是 "horizontal" 或 "vertical"。
    :param target_color: 目标颜色，可以是 "black" 或 "white"。
    :return: 投影图像，以及一个列表，列表的每个元素是每行（或每列）的白色像素数量。
    """
    # 如果是PIL图像，转换为NumPy数组
    if isinstance(img_array, Image.Image):
        img_array = array(img_array)
        # 如果图像有三个维度，并且颜色为三通道，则进行颜色空间的转换
        if img_array.ndim == 3 and img_array.shape[2] == 3:
            img_array = cvtColor(img_array, COLOR_RGB2BGR)

    if len(img_array.shape) == 3:
        # 转换为灰度
        img_array = cvtColor(img_array, COLOR_BGR2GRAY)

    height, width = img_array.shape[:2]
    is_target_black = target_color.lower() == "black"

    target_value = 0 if is_target_black else 255
    img_bg_color = 255 if is_target_black else 0
    img_fg_color = 0 if is_target_black else 255

    projection_data = []
    if direction.lower() == "horizontal":
        # 计算水平投影
        projection_data = [row.count(target_value) for row in img_array.tolist()]
        # 创建水平投影图像
        proj_pil = Image.new('L', (width, height), img_bg_color)
        for y, value in enumerate(projection_data):
            for x in range(value):
                proj_pil.putpixel((x, y), img_fg_color)
    else:  # if direction.lower() == "vertical"
        # 计算垂直投影
        projection_data = [col.count(target_value) for col in zip(*img_array.tolist())]
        # 创建垂直投影图像
        proj_pil = Image.new('L', (width, height), img_bg_color)
        for x, value in enumerate(projection_data):
            for y in range(height - value, height):
                proj_pil.putpixel((x, y), img_fg_color)
    return proj_pil, projection_data


# @logger.catch
def analyze1frame(img_file, frame_data, auto_subdir, media_type):
    with_frames_jpg = img_file.parent / f'{img_file.stem}-画框.jpg'
    pic_div_psd = img_folder / f'{img_file.stem}-分框.psd'
    added_frames_jpg = auto_subdir / f'{img_file.stem}-加框.jpg'
    logger.warning(f'{img_file=}')
    img_raw = imdecode(fromfile(img_file, dtype=uint8), -1)
    img_formal = img_raw.copy()
    if with_frames_jpg.exists():
        img_formal = imdecode(fromfile(with_frames_jpg, dtype=uint8), -1)
    img_formal = toBGR(img_formal)

    color_name0 = 'white'
    if img_file.name in frame_data:
        frame_grid_strs = frame_data[img_file.name]
    elif pic_div_psd.exists():
        frame_grid_strs = []
    else:
        if frame_color is None:
            # 获取RGBA格式的边框像素
            edge_pixels_rgba = get_edge_pxs(img_raw)
            # 只检查前两种主要颜色
            color_and_ratios, combined_color_and_ratios = find_dominant_colors(edge_pixels_rgba)
            dominant_color0, color_ratio0 = combined_color_and_ratios[0]
            color_name0 = get_color_name(dominant_color0)
            dominant_colors = [dominant_color0]
            imes = f"{img_file.stem}边框颜色中出现次数最多的颜色：{dominant_color0}, 颜色名称：{color_name0}, {color_ratio0=:.4f}"
            if check_2nd_color and len(combined_color_and_ratios) >= 2:
                # 第二种颜色较多且不为黑色
                dominant_color1, color_ratio1 = combined_color_and_ratios[1]
                color_name1 = get_color_name(dominant_color1)
                if color_name1 != 'black' and color_ratio1 >= 0.1:
                    dominant_colors = [dominant_color0, dominant_color1]
                    imes += f", 次多的颜色：{dominant_color1}, 颜色名称：{color_name1}, {color_ratio1=:.4f}"
        else:
            # ================从配置中读取颜色================
            dominant_color0 = color2rgb(f'#{frame_color}') + (255,)
            color_name0 = get_color_name(dominant_color0)
            dominant_colors = [dominant_color0]
            imes = f"{img_file.stem}边框颜色为：{dominant_color0}, 颜色名称：{color_name0}"
        logger.info(imes)

        # 计算主要颜色遮罩
        frame_mask = compute_frame_mask(img_formal, dominant_colors, tolerance=normal_tolerance)
        frame_mask_white = compute_frame_mask(img_formal, [rgba_white], tolerance=white_tolerance)
        frame_mask_black = compute_frame_mask(img_formal, [rgba_black], tolerance=black_tolerance)

        frame_mask_group = [frame_mask]
        if check_more_frame_color:
            if color_name0 != 'white':
                frame_mask_group.append(frame_mask_white)
            if color_name0 != 'black':
                frame_mask_group.append(frame_mask_black)

        # 创建水平投影图像
        hori_proj_bi, hori_proj_data = get_proj_img(frame_mask, 'horizontal', 'white')
        # 创建垂直投影图像
        ver_proj_bi, ver_proj_data = get_proj_img(frame_mask, 'vertical', 'white')

        if frame_color is not None:
            thres = 300
            # 计算一阶差分
            dif = diff(hori_proj_data)
            # 找到大于等于阈值的一阶差分的索引
            high_diff_indices = where(np.abs(dif) >= thres)[0]
            # 使用 argrelextrema 函数找到大于thres的一阶差分的所有局部最大值的索引
            high_diff_maxima = argrelextrema(dif[high_diff_indices], greater)[0]
            # 注意high_diff_maxima返回的是high_diff_indices的局部索引，我们需要把它转换成hori_proj_data的索引
            high_diff_maxima = high_diff_indices[high_diff_maxima]
            logger.info(f'{high_diff_indices=}')
            logger.info(f'{high_diff_maxima=}')

        if do_dev_pic:
            frame_mask_png = debug_dir / f'{media_type}_frame_mask.png'
            frame_mask_white_png = debug_dir / f'{media_type}_frame_mask_white.png'
            frame_mask_black_png = debug_dir / f'{media_type}_frame_mask_black.png'
            hori_proj_frame_mask_png = debug_dir / f'{media_type}_hori_proj_bi.png'
            ver_proj_frame_mask_png = debug_dir / f'{media_type}_ver_proj_bi.png'
            write_pic(frame_mask_png, frame_mask)
            write_pic(frame_mask_white_png, frame_mask_white)
            write_pic(frame_mask_black_png, frame_mask_black)
            write_pic(hori_proj_frame_mask_png, hori_proj_bi)
            write_pic(ver_proj_frame_mask_png, ver_proj_bi)

        grids = get_grids(frame_mask_group, media_type)
        frame_grid_strs = [rect.frame_grid_str for rect in grids]
        frame_data[img_file.name] = frame_grid_strs
    pprint(frame_grid_strs)
    # 对画格进行标记
    grid_masks, marked_frames = get_grid_masks(img_file, frame_grid_strs)

    if do_add_frame:
        added_frames = get_added_frames(frame_grid_strs, img_raw, color_name0)
        write_pic(added_frames_jpg, added_frames)
    return img_file, frame_grid_strs


# @timer_decorator
def step0_analyze_frames(img_folder, frame_yml, media_type, auto_subdir, img_inds):
    """
    分析画格，获取画格的位置和尺寸，将结果保存到YAML文件中。

    :param frame_yml: frame_yml文件路径
    :return: 包含已排序画格数据的字典。
    """
    # 这是一个用于分析图像画格并将结果保存到YAML文件中的Python函数。该函数以YAML文件路径作为输入，并在存在YAML文件时加载YAML文件中的数据。然后，它会分析给定图像列表中的每个图像，并将画格的位置和尺寸保存为字符串列表，并将其保存到YAML文件中。
    # 每个图像的分析包括找到画格边缘像素的主要颜色，根据主要颜色计算遮罩，并使用遮罩将画格分成子矩形。然后，函数将每个子矩形的位置和大小保存为字符串列表中的元素。
    # 如果某个图像中有两个或更多的子矩形，则函数可以使用get_marked_frames函数对其进行标记，但是此部分代码目前被注释掉了。
    # 最后，函数将img_data字典按图像文件名排序，并将排序后的字典保存到YAML文件中。
    img_list = get_valid_imgs(img_folder)
    # ================加载YAML文件================
    frame_data = iload_data(frame_yml)

    if img_inds:
        img_list_roi = [img_list[i] for i in img_inds]
        # 从img_data中移除这些图像的数据，以便后面重新生成
        for image in img_list_roi:
            frame_data.pop(image.name, None)
    # ================分析画格================
    if thread_method == 'queue':
        for p, img_file in enumerate(img_list):
            img_file, frame_grid_strs = analyze1frame(img_file, frame_data, auto_subdir, media_type)
            frame_data[img_file.name] = frame_grid_strs
    else:
        with ThreadPoolExecutor(os.cpu_count()) as executor:
            # 提交所有图片处理任务
            futures = [executor.submit(analyze1frame, img_file, frame_data, auto_subdir, media_type) for
                       img_file in img_list]
            for future in as_completed(futures):
                try:
                    img_file, frame_grid_strs = future.result()
                    frame_data[img_file.name] = frame_grid_strs
                except Exception as e:
                    printe(e)

    # 按图像文件名对 frame_data 字典进行排序
    frame_data_sorted = {k: frame_data[k] for k in natsorted(frame_data)}

    # 将 frame_data_sorted 字典保存到 yml 文件
    with open(frame_yml, 'w') as yml_file:
        yaml.dump(frame_data_sorted, yml_file)

    return frame_data_sorted


# @logger.catch
def water_seg(filled_contour, textblocks):
    """
    使用分水岭算法对给定的填充轮廓进行分割，以分离其中的多个对象。

    :param filled_contour: 已填充的轮廓。
    :param textblocks: 文本块。
    :return: 分割后的轮廓列表。
    """
    # 现在我们想要分离图像中的多个对象
    distance = ndimage.distance_transform_edt(filled_contour)
    # 创建一个与图像大小相同的标记数组
    markers_full = zeros_like(filled_contour, dtype=int32)

    # 在分水岭算法中，每个标记值代表一个不同的区域，算法会将每个像素点归属到与其距离最近的标记值所代表的区域。
    for g in range(len(textblocks)):
        textblock = textblocks[g]
        markers_full[textblock.br_n, textblock.br_m] = g + 1

    markers_full = ndimage.label(markers_full)[0]
    labels = watershed(-distance, markers_full, mask=filled_contour)

    # 创建一个空的彩色图像
    color_labels = zeros((labels.shape[0], labels.shape[1], 3), dtype=uint8)
    # 为每个标签分配一种颜色
    unique_labels = unique(labels)
    colors = (colormap_tab20(arange(len(unique_labels)))[:, :3] * 255).astype(uint8)
    for label, color in zip(unique_labels, colors):
        if label == 0:
            continue
        color_labels[labels == label] = color

    if do_dev_pic:
        color_labels_png = debug_dir / f'{media_type}_color_labels.png'
        write_pic(color_labels_png, color_labels)

    seg_cnts = []
    for u in range(len(unique_labels)):
        target_label = unique_labels[u]
        binary_img = zeros_like(labels, dtype=uint8)
        binary_img[labels == target_label] = 255
        # binary_img_path = debug_dir / f'binary_img{u}.png'
        # write_pic(binary_img_path, binary_img)

        # 查找轮廓
        sep_contours, hierarchy = findContours(binary_img, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
        for s in range(len(sep_contours)):
            contour = sep_contours[s]
            cnt = Contr(contour)
            # ================对轮廓数值进行初筛================
            condition_p1 = 300 <= cnt.area <= 1000000
            condition_p2 = 80 <= cnt.perimeter <= 50000
            condition_ps = [
                condition_p1,
                condition_p2,
            ]
            print(f'{condition_ps=}, {cnt.area=:.1f}, {cnt.perimeter=:.2f}')
            if all(condition_ps):
                seg_cnts.append(cnt)
    return seg_cnts


@timer_decorator
@logger.catch
def get_textwords(letter_cnts, letter_in_contour, kernal_word, cnt_ind):
    # ================单词================
    textwords = []
    letter_cnts4words = deepcopy(letter_cnts)

    if text_direction == 'Horizontal':
        # 从左到右，然后从上到下
        letter_cnts4words.sort(key=lambda x: (x.br_x, x.br_y))
    else:  # media_type in ['Manga']
        # 从上到下，然后从右到左
        letter_cnts4words.sort(key=lambda x: (x.br_y, -x.br_x))
    if use_dilate:
        # ================通过膨胀查找单词================
        word_in_contour = dilate(letter_in_contour, kernal_word, iterations=1)
        word_contours, word_hier = findContours(word_in_contour, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
        if do_dev_pic:
            if cnt_ind is not None:
                word_in_contour_png = debug_dir / f'letter_in_contour_{cnt_ind}_dilated.png'
            else:
                word_in_contour_png = debug_dir / f'letter_in_contour_dilated.png'
            write_pic(word_in_contour_png, word_in_contour)
        word_cnts = []
        for w in range(len(word_contours)):
            word_contour = word_contours[w]
            word_cnt = Contr(word_contour)
            if word_cnt.area >= 50 or len(word_contours) <= 3:
                word_cnts.append(word_cnt)
        for w in range(len(word_cnts)):
            word_cnt = word_cnts[w]
            letter_cnts_in_word = []
            # 枚举每个letter_cnt，检查是否在word_cnt内部
            i = 0
            while i < len(letter_cnts4words):
                letter_cnt = letter_cnts4words[i]
                if letter_cnt.is_inside(word_cnt):
                    letter_cnts_in_word.append(letter_cnt)
                    letter_cnts4words.pop(i)
                else:
                    i += 1
            textword = TextWord(letter_cnts_in_word, media_type)
            textword.add_letter_mask(letter_in_contour)
            textwords.append(textword)
    else:
        while letter_cnts4words:
            # 取最左边的单词，逐词往右查找
            letter_cnt = letter_cnts4words.pop(0)
            textword = TextWord(letter_cnt, media_type)
            while True:
                # 对原始的文本行进行深度复制，以便之后比较是否有变化
                old_textword = deepcopy(textword)
                for letter in letter_cnts4words[:]:
                    # 检查单字的质心和单词的中心点的高度差异是否在允许的范围内
                    if text_direction == 'Horizontal':
                        c_height_diff = abs(letter.cy - textword.br_center[1])
                    else:  # media_type in ['Manga']
                        c_height_diff = abs(letter.cx - textword.br_center[0])
                    # 如果高度差异在允许的范围内且单词和文本行的轮廓有交集
                    if c_height_diff <= max_hdiff and textword.ext_brp.intersects(letter.brp):
                        # 则认为单词属于这个文本行
                        textword.add_letter_cnt(letter)
                        letter_cnts4words.remove(letter)
                # 如果在这个循环中文本行没有变化（没有新的 letter 被加入），则退出循环
                if len(textword.letter_cnts) == len(old_textword.letter_cnts):
                    break
            textword.add_letter_mask(letter_in_contour)
            # logger.debug(f'{textword.letter_area=:.2f}')
            if textword.letter_area <= letter_area_max:
                textwords.append(textword)
        # ================字母i最上方的点可能被孤立================
        if media_type in ['Comic'] and check_dots:
            dots = [x for x in textwords if x.br_area <= 30 and len(x.letter_cnts) == 1]
            not_dots = [x for x in textwords if x not in dots]
            for d in range(len(dots)):
                dot = dots[d]
                # 筛选出在 dot 下方的 textword，并计算距离
                textwords_under = [textword for textword in not_dots if textword.br_y >= dot.br_y - 30]
                distance_tups = []
                for textword in textwords_under:
                    distance = dot.brp.distance(textword.brp)
                    # 如果textword在dot的上方，将距离乘以3
                    if textword.br_y < dot.br_y:
                        distance *= 3
                    distance_tups.append((distance, textword))
                # 按距离进行排序
                distance_tups.sort(key=lambda x: x[0])
                if distance_tups:
                    closest_dist, closest_textword = distance_tups[0]
                    index = not_dots.index(closest_textword)
                    not_dots[index].add_letter_cnt(dot.letter_cnts[0])
            textwords = not_dots
    textwords.sort(key=lambda x: (x.br_y, x.br_x))
    logger.debug(f'{len(textwords)=}')
    return textwords


@timer_decorator
@logger.catch
def get_textlines(textwords):
    # ================文本行================
    textlines = []
    textwords4lines = deepcopy(textwords)
    if text_direction == 'Horizontal':
        # 从左到右，然后从上到下
        textwords4lines.sort(key=lambda x: (x.br_x, x.br_y))
    else:  # media_type in ['Manga']
        # 从上到下，然后从右到左
        textwords4lines.sort(key=lambda x: (x.br_y, -x.br_x))
    while textwords4lines:
        # 取最左边的单词，逐词往右查找
        textword = textwords4lines[0]
        textwords4lines.remove(textword)
        textline = TextLine(textword, media_type)
        while True:
            # 对原始的文本行进行深度复制，以便之后比较是否有变化
            old_textline = deepcopy(textline)
            for textword in textwords4lines[:]:
                # 检查单词的质心和文本行的中心点的高度差异是否在允许的范围内
                if text_direction == 'Horizontal':
                    c_height_diff = abs(textword.br_center[1] - textword.br_center[1])
                else:  # media_type in ['Manga']
                    c_height_diff = abs(textword.br_center[0] - textword.br_center[0])
                # 如果高度差异在允许的范围内且单词和文本行的轮廓有交集
                if c_height_diff <= max_hdiff and textline.ext_brp.intersects(textword.brp):
                    # 则认为单词属于这个文本行
                    textline.add_textword(textword)
                    textwords4lines.remove(textword)
            # 如果在这个循环中文本行没有变化（没有新的 word 被加入），则退出循环
            if len(textline.textwords) == len(old_textline.textwords):
                break
        textlines.append(textline)
    textlines.sort(key=lambda x: x.br_y)
    logger.debug(f'{len(textlines)=}')
    return textlines


@timer_decorator
@logger.catch
def get_textblocks(textlines, cnt_ind, color_input, ih, iw):
    # ================根据文本行检测文本块================
    textblocks_raw = []
    textlines4blocks = deepcopy(textlines)

    if text_direction == 'Horizontal':
        # 从上到下，然后从左到右
        textlines4blocks.sort(key=lambda x: (x.br_y, x.br_x))
    else:  # media_type in ['Manga']
        # 从右到左，然后从上到下
        textlines4blocks.sort(key=lambda x: (-x.br_x, x.br_y))
    while textlines4blocks:
        if text_direction == 'Horizontal':
            # 取最下方的文本行，逐行往上查找
            textline = textlines4blocks[-1]
        else:  # media_type in ['Manga']
            # 取最右方的文本行，逐行往左查找
            textline = textlines4blocks[0]
        textlines4blocks.remove(textline)
        textblock = iTextBlock(textline, media_type)
        while True:
            old_textblock = deepcopy(textblock)
            for textline in textlines4blocks[:]:
                is_inside = textblock.core_brp.intersects(textline.brp) or textline.brp.contains(textblock.core_brp)
                if is_inside:
                    textblock.add_textline(textline)
                    textlines4blocks.remove(textline)
            if len(textblock.textlines) == len(old_textblock.textlines):
                break
        textblocks_raw.append(textblock)
    textblocks = deepcopy(textblocks_raw)
    textblocks.sort(key=lambda x: x.br_y + x.br_x)

    # ================检查底行================
    if text_direction == 'Horizontal' and check_bottom and len(textblocks) >= 2:
        textblocks.sort(key=lambda x: x.br_y)
        # ================底行文本块================
        bottom_textblock = textblocks[-1]
        bup_textblock = textblocks[-2]
        bottom_textlines = bottom_textblock.textlines
        textline = bottom_textlines[-1]
        br_m_diff = abs(bottom_textblock.br_m - bup_textblock.br_m)
        y_diff = bottom_textblock.br_y - bup_textblock.br_v
        logger.debug(f'{br_m_diff=}, {y_diff=}')
        if len(bottom_textlines) == 1 and br_m_diff <= 20 and y_diff <= 15:
            textblocks = textblocks[:-1]
            textblocks[-1].add_textline(textline)

    # ================去除相交且相交面积大于等于面积更小的文本块的面积的80%的小文本块================
    if exclude_smaller:
        i = 0
        while i < len(textblocks):
            tb1 = textblocks[i]
            removed = False
            for j, tb2 in enumerate(textblocks):
                if i != j and tb1.brp.intersects(tb2.brp):
                    intersection_area = tb1.brp.intersection(tb2.brp).area
                    min_area = min(tb1.br_area, tb2.br_area)
                    if intersection_area >= intersect_ratio * min_area:
                        if tb1.br_area < tb2.br_area:
                            textblocks.pop(i)
                            removed = True
                            break
            if not removed:
                i += 1
    # ================去除完全包含在另一个文本块中的小文本块================
    if exclude_inside:
        filter_textblocks = []
        for i, tb1 in enumerate(textblocks):
            is_inside_another = False
            for j, tb2 in enumerate(textblocks):
                if i != j and tb2.brp.contains(tb1.brp):
                    is_inside_another = True
                    break
            if not is_inside_another:
                filter_textblocks.append(tb1)
        textblocks = filter_textblocks
    else:
        filter_textblocks = deepcopy(textblocks)
    # ================其他筛选================
    textblocks_valid = [x for x in textblocks if textblock_letters_min <= x.letter_count <= textblock_letters_max]
    if textblocks_valid:
        textblocks = textblocks_valid

    textblocks_valid = [x for x in textblocks if
                        textblock_w_percent_min * iw <= x.br_x <= x.br_u <= textblock_w_percent_max * iw]
    if textblocks_valid:
        textblocks = textblocks_valid

    textblocks_valid = [x for x in textblocks if
                        textblock_h_percent_min * ih <= x.br_y <= x.br_v <= textblock_h_percent_max * ih]
    if textblocks_valid:
        textblocks = textblocks_valid
    logger.warning(f'{len(textblocks_raw)}=>{len(filter_textblocks)}=>{len(textblocks)}')

    # textblocks = textblocks_raw

    if cnt_ind is None and not color_input:
        # ================无框================
        textblocks = [x for x in textblocks if textblock_area_min <= x.block_cnt.area <= textblock_area_max]
        textblocks = [x for x in textblocks if textblock_wmin <= x.block_cnt.br_w <= textblock_wmax]
        textblocks = [x for x in textblocks if textblock_hmin <= x.block_cnt.br_h <= textblock_hmax]
        logger.debug(f'{len(textblocks)=}')
    # ================其他排序================
    if sort_by_y:
        if len(textblocks) in [3, 4, 5]:
            textblocks.sort(key=lambda x: x.br_y + lower_ratio * x.br_x)
            # textblocks.sort(key=lambda x: x.br_y)
    if sort_by_x:
        if len(textblocks) in [3, 4, 5]:
            textblocks.sort(key=lambda x: lower_ratio * x.br_y + x.br_x)
            # textblocks.sort(key=lambda x: x.br_x)
    logger.debug(f'{len(textblocks)=}')
    return textblocks


@timer_decorator
@logger.catch
def get_textblocks_full(letter_in_contour, media_type, cnt_ind=None):
    ih, iw = letter_in_contour.shape[:2]
    black_bg = zeros((ih, iw), dtype=uint8)

    color_input = False
    # 确保图像是灰度的
    if len(letter_in_contour.shape) == 3:
        # 如果图像是彩色的
        color_input = True
        # 转换为灰度
        gray_letter_in_contour = cvtColor(letter_in_contour, COLOR_BGR2GRAY)
        # 二值化图像
        ret, letter_in_contour = threshold(gray_letter_in_contour, bi_thres, 255, THRESH_BINARY)
        # 反色
        letter_in_contour = bitwise_not(letter_in_contour)

    if text_direction == 'Horizontal':
        kh = 1
        kw = kernel_depth
    else:
        kh = kernel_depth
        kw = 1
    kernal_word = kernel_hw(kh, kw)

    # letter_in_contour_inv = bitwise_not(letter_in_contour)
    if do_dev_pic:
        if cnt_ind is not None:
            letter_in_contour_name = f'letter_in_contour_{cnt_ind}.png'
        else:
            letter_in_contour_name = 'letter_in_contour.png'
        letter_in_contour_png = current_dir / letter_in_contour_name
        logger.debug(f'{letter_in_contour_name=}')
        write_pic(letter_in_contour_png, letter_in_contour)

    # ================单字轮廓================
    raw_letter_cnts = []
    # letter_contours, letter_hier = findContours(letter_in_contour, RETR_LIST, CHAIN_APPROX_SIMPLE)
    letter_contours, letter_hier = findContours(letter_in_contour, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
    logger.debug(f'{len(letter_contours)=}')
    if cnt_ind is None and not color_input:
        # ================无框================
        for l in range(len(letter_contours)):
            letter_contour = letter_contours[l]
            letter_cnt = Contr(letter_contour)
            # ================对数值进行初筛================
            condition_c1 = cnt_area_min <= letter_cnt.area <= cnt_area_max
            condition_c2 = brw_min <= letter_cnt.br_w <= brw_max
            condition_c3 = brh_min <= letter_cnt.br_h <= brh_max
            logger.debug(f"{letter_cnt.area=}, {letter_cnt.br_w=}, {letter_cnt.br_h=}")
            condition_cs = [
                condition_c1,
                condition_c2,
                condition_c3,
            ]
            if all(condition_cs):
                logger.debug(f"{letter_cnt=}")
                raw_letter_cnts.append(letter_cnt)
    else:
        # ================气泡内文本块================
        for l in range(len(letter_contours)):
            letter_contour = letter_contours[l]
            letter_cnt = Contr(letter_contour)
            raw_letter_cnts.append(letter_cnt)
    # ================通过图像再次筛选================
    if len(raw_letter_cnts) <= 100:
        letter_cnts = []
        for l in range(len(raw_letter_cnts)):
            letter_cnt = raw_letter_cnts[l]
            all_letters = drawContours(black_bg.copy(), letter_cnt.contour, -1, 255, FILLED)
            letter_in_word = bitwise_and(letter_in_contour, all_letters)
            px_pts = transpose(nonzero(letter_in_word))
            px_area = px_pts.shape[0]
            if letter_area_min <= px_area <= letter_area_max:
                letter_cnts.append(letter_cnt)
    else:
        letter_cnts = raw_letter_cnts
    logger.debug(f"{len(raw_letter_cnts)=}, {len(letter_contours)=}, {len(letter_cnts)=}")

    textwords = get_textwords(letter_cnts, letter_in_contour, kernal_word, cnt_ind)
    textlines = get_textlines(textwords)
    textblocks = get_textblocks(textlines, cnt_ind, color_input, ih, iw)
    return textblocks


# @logger.catch
def get_full_scnts(seg_cnts, filled_contour, filled_contour_split):
    ih, iw = filled_contour.shape[0:2]
    color_img = zeros((ih, iw, 3), dtype=uint8)
    for idx, cnt in enumerate(seg_cnts):
        color = colormap_tab20(idx % 20)[:3]
        color_rgb = tuple(int(c * 255) for c in color)
        color_bgr = color_rgb[::-1]
        drawContours(color_img, [cnt.contour], 0, color_bgr, -1)

    diff_mask = filled_contour - filled_contour_split
    for y, x in argwhere(diff_mask == 255):
        dists = []
        point = Point(x, y)
        for idx, contour in enumerate(seg_cnts):
            if contour.polygon is None:  # 安全检查
                print(f"索引为 {idx} 的轮廓没有polygon属性!")
                continue
            dist = contour.polygon.distance(point)
            dists.append((dist, idx))
        dists.sort()
        # 获取最小距离的索引
        min_dist_idx = dists[0][1]
        # 使用该索引获取颜色
        color = colormap_tab20(min_dist_idx % 20)[:3]
        color_rgb = tuple(int(c * 255) for c in color)
        nearest_color = color_rgb[::-1]
        color_img[y, x] = nearest_color

    new_scnts = []
    for idx in range(len(seg_cnts)):
        color = colormap_tab20(idx % 20)[:3]
        color_rgb = tuple(int(c * 255) for c in color)
        color_bgr = color_rgb[::-1]
        mask = (color_img == color_bgr).all(axis=-1).astype(uint8) * 255
        contours, _ = findContours(mask, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
        new_scnts.extend([Contr(contour) for contour in contours])
    seg_cnts = new_scnts
    return seg_cnts


def get_outline_pt(intersection, ref_pt1, ref_pt2, min_dist=1):
    """
    根据交点和参考点获取轮廓点。

    :param intersection: 交点，可以是点、线或多线
    :param ref_pt1: 参考点1
    :param ref_pt2: 参考点2
    :return: 轮廓点
    """
    # 将输入的参考点转换为Point对象
    ref_pt1 = Point(ref_pt1)
    ref_pt2 = Point(ref_pt2)

    potential_pts = []
    if isinstance(intersection, MultiLineString):
        for line in intersection.geoms:
            potential_pts.extend(line.coords)
    elif isinstance(intersection, LineString):
        potential_pts.extend(intersection.coords)
    elif isinstance(intersection, Point):
        potential_pts.append(intersection.coords[0])

    # 排除离ref_point1和ref_point2距离小于等于min_dist的点
    potential_pts = [pt for pt in potential_pts if
                     ref_pt1.distance(Point(pt)) > min_dist and ref_pt2.distance(Point(pt)) > min_dist]

    if not potential_pts:
        return None

    # 找到距离ref_point2最近的点
    closest_pt = min(potential_pts, key=lambda pt: ref_pt2.distance(Point(pt)))
    return Point(closest_pt)


@logger.catch
def pivot_proc(filter_cnt, filled_contour, letter_in_contour, textblocks, cnt_ind):
    filled_contour_split = filled_contour.copy()

    # 创建一个预览图像，用于显示分割效果
    preview_canvas = cvtColor(filled_contour.copy(), COLOR_GRAY2BGR)
    # 将轮廓中的字母部分设置为黑色，以便在示意图中清晰地看到分割线
    preview_canvas[letter_in_contour == 255] = 0
    for textblock in textblocks:
        # 在示意图上标记文本块的中心点，并绘制文本块的轮廓
        preview_canvas = circle(preview_canvas, pt2tup(textblock.block_poly.centroid), 5, color_green, -1)
        preview_canvas = drawContours(preview_canvas, [textblock.block_contour], 0, color_olive, 1)

    split_preview = preview_canvas.copy()
    for i in range(len(textblocks) - 1):
        textblock_comb = textblocks[i:i + 2]
        # 使用针转切割
        textblock1, textblock2 = textblock_comb[0:2]
        intersect_pts = nearest_points(textblock1.block_poly, textblock2.block_poly)
        intersect_points_geom = MultiPoint(intersect_pts)
        # 计算最近距离交点的中心（质心）
        dist_center = intersect_points_geom.centroid
        # 定义连线
        link_line = LineString([textblock1.br_center, textblock2.br_center])
        sect1 = textblock1.block_poly.convex_hull.intersection(link_line)
        sect2 = textblock2.block_poly.convex_hull.intersection(link_line)
        outline_pt1 = get_outline_pt(sect1, textblock1.br_center, textblock2.br_center)
        outline_pt2 = get_outline_pt(sect2, textblock2.br_center, textblock1.br_center)
        if outline_pt1 is None or outline_pt2 is None:
            img = filled_contour.copy()
            blended_img = get_textblock_bubbles(img, textblocks)
            if do_dev_pic:
                check_pts_png = debug_dir / f'{media_type}_check_pts.png'
                write_pic(check_pts_png, blended_img)
        if outline_pt1 is None and outline_pt2 is None:
            pt_center = None
        elif outline_pt1 is None:
            pt_center = outline_pt2
        elif outline_pt2 is None:
            pt_center = outline_pt1
        else:
            points = MultiPoint([outline_pt1, outline_pt2])
            mid_point = points.centroid
            pt_center = mid_point

        if pt_center is None:
            raise ValueError("pt_center是None，无法继续。")

        logger.info(f"{outline_pt1}, {outline_pt2}, {pt_center=}")
        # 在质心周围创建一个像素范围
        pixel_range = range(-px_step, px_step + 1)

        # ================计算合适的分割线长度================
        line_length = 500
        lengths = []
        for angle in range(0, 180, angle_step):
            rad_angle = radians(angle)
            # 计算线段的两个端点
            line_pt1 = (
                pt_center.x + line_length * cos(rad_angle),
                pt_center.y + line_length * sin(rad_angle)
            )
            line_pt2 = (
                pt_center.x - line_length * cos(rad_angle),
                pt_center.y - line_length * sin(rad_angle)
            )
            cline = LineString([line_pt1, line_pt2])
            intersection = cline.intersection(filter_cnt.polygon)
            if intersection.geom_type == "LineString":
                lengths.append(intersection.length)
        min_intersect_length = min(lengths)
        seg_line_len = min_intersect_length * 1.5
        seg_line_len = max(seg_line_min, min(seg_line_max, seg_line_len))
        if not filter_cnt.polygon.contains(pt_center):
            seg_line_len = (seg_line_min + seg_line_max) / 2

        # ================计算合适的分割================
        all_tups = []
        for p in pixel_range:
            # 创建一个新的质心点，偏移了i个像素
            cur_pt_center = Point(pt_center.x, pt_center.y + p)
            for angle in range(0, 180, angle_step):
                rad_angle = radians(angle)
                # 计算线段的两个端点
                sep_line_pt1 = (
                    cur_pt_center.x + seg_line_len * cos(rad_angle),
                    cur_pt_center.y + seg_line_len * sin(rad_angle)
                )
                sep_line_pt2 = (
                    cur_pt_center.x - seg_line_len * cos(rad_angle),
                    cur_pt_center.y - seg_line_len * sin(rad_angle)
                )
                sep_line_pts = [sep_line_pt1, sep_line_pt2]
                sep_line = LineString(sep_line_pts)

                # 计算分隔线与过滤后轮廓的交点
                cnt_line_sect = sep_line.intersection(filter_cnt.polygon)
                # 不会改变原始图形的形状，但会清除任何存在的几何不规则性，例如自交叉的线条或重叠的点。
                tb1_sect = sep_line.intersection(textblock1.block_poly.buffer(0))
                tb2_sect = sep_line.intersection(textblock2.block_poly.buffer(0))
                sum_sect = tb1_sect.length + tb2_sect.length
                comb_len = cnt_line_sect.length + 1000 * sum_sect
                tup = (cur_pt_center, angle, sep_line, cnt_line_sect, tb1_sect, tb2_sect, sum_sect, comb_len)
                all_tups.append(tup)

        all_tups.sort(key=lambda x: x[-1])
        best_tup = all_tups[0]
        cur_pt_center, angle, sep_line, cnt_line_sect, tb1_sect, tb2_sect, sum_sect, comb_len = best_tup
        if cnt_line_sect.geom_type == "LineString":
            coords_list = list(cnt_line_sect.coords)
        else:  # cnt_line_sect.geom_type == "MultiLineString":
            dists = []
            # 计算每个线段到pt_center的距离
            for linestring in cnt_line_sect.geoms:
                distance = linestring.distance(pt_center)
                dists.append(distance)
            # 寻找最小距离及其对应的线段
            min_dist = min(dists)
            # 获取最小距离的索引
            index_of_closest = dists.index(min_dist)
            closest_linestring = cnt_line_sect.geoms[index_of_closest]
            # 从最近的线段中获取坐标列表
            coords_list = list(closest_linestring.coords)

        if len(coords_list) >= 2 and sum_sect <= sum_sect_max:
            # 根据坐标列表获取起始点和终点
            p0, p1 = map(lambda coord: tuple(map(floor, coord)), coords_list[:2])
            logger.debug(f'{cur_pt_center=}, {p0=}, {p1=}, {outline_pt1=}, {outline_pt2=}')

            filled_contour_split = line(filled_contour_split, p0, p1, 0, 2)

            ct1 = pt2tup(textblock_comb[0].block_poly.centroid)
            ct2 = pt2tup(textblock_comb[1].block_poly.centroid)
            # 在示意图上绘制从两个文本块中心点连接的线条
            split_preview = line(split_preview, ct1, ct2, color_maroon, 2)
            # 在示意图上绘制实际的分割线
            split_preview = line(split_preview, p0, p1, color_red, 2)
            if outline_pt1 is not None and outline_pt2 is not None:
                split_preview = line(split_preview, pt2tup(outline_pt1), pt2tup(outline_pt2), color_slate_blue, 1)
            # 在示意图上标记分割线的中心点
            split_preview = circle(split_preview, pt2tup(pt_center), 3, color_yellow, -1)
            split_preview = circle(split_preview, pt2tup(cur_pt_center), 3, color_blue, -1)
        else:
            logger.error(f'{coords_list=}')
            # 生成示意图
            error_preview = preview_canvas.copy()
            # 在示意图上标记分割线的中心点
            error_preview = circle(error_preview, pt2tup(pt_center), 3, color_yellow, -1)
            error_preview = circle(error_preview, pt2tup(cur_pt_center), 3, color_blue, -1)
            # 保存示意图
            if do_dev_pic:
                error_preview_png = debug_dir / f'{media_type}_error_preview.png'
                write_pic(error_preview_png, error_preview)
            logger.debug(f'{cur_pt_center=}, {outline_pt1=}, {outline_pt2=}')

    contours, _ = findContours(filled_contour_split, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
    seg_cnts = [Contr(contour) for contour in contours if Contr(contour).area >= area_min]
    seg_cnts = get_full_scnts(seg_cnts, filled_contour, filled_contour_split)

    if do_dev_pic:
        split_preview_png = debug_dir / f'split_preview_{cnt_ind}.png'
        filled_contour_split_png = debug_dir / f'{media_type}_filled_contour_split.png'
        write_pic(split_preview_png, split_preview)
        write_pic(filled_contour_split_png, filled_contour_split)
    return seg_cnts, True


@logger.catch
def seg_bubbles(filter_cnts, bubble_mask, letter_mask, media_type):
    ih, iw = bubble_mask.shape[0:2]
    black_bg = zeros((ih, iw), dtype=uint8)
    single_cnts = []

    all_textblocks = []
    for cnt_ind in range(len(filter_cnts)):
        # ================每个初始对话气泡================
        filter_cnt = filter_cnts[cnt_ind]
        br_area_real = (filter_cnt.br_w - 2) * (filter_cnt.br_h - 2)
        fulfill_ratio = filter_cnt.area / br_area_real

        filled_contour = drawContours(black_bg.copy(), [filter_cnt.contour], 0, 255, -1)
        # bubble_in_contour = bitwise_and(bubble_mask, filled_contour)
        letter_in_contour = bitwise_and(letter_mask, filled_contour)

        raw_textblocks = get_textblocks_full(letter_in_contour, media_type, cnt_ind)
        # ================排除过小的文本块================
        big_textblocks = [x for x in raw_textblocks if textblock_area_min <= x.block_cnt.area <= textblock_area_max]
        small_textblocks = [x for x in raw_textblocks if x.block_cnt.area < textblock_area_min]
        note_textblocks = []
        if big_textblocks:
            max_block_cnt_area = max(x.block_cnt.area for x in big_textblocks)
            logger.warning(f'{max_block_cnt_area=}')
            # ================检查音符================
            dists_dic = {}
            if check_note and len(big_textblocks) >= 2 and max_block_cnt_area >= note_area_max:
                for idx, textblock in enumerate(big_textblocks[:]):
                    # if note_area_min <= textblock.block_cnt.area <= note_area_max:
                    dists = [filter_cnt.polygon.boundary.distance(Point(p)) for p in
                             textblock.block_poly.exterior.coords]
                    min_dist = min(dists)
                    dists_dic[idx] = dists
                    logger.warning(f'[{idx}]{textblock.block_cnt.area=}, {min_dist=}')
                    if min_dist <= max_note_dist:
                        note_textblocks.append(textblock)
                        big_textblocks.remove(textblock)
                # if dists_dic:
                #     print()
            textblocks = big_textblocks
        else:
            textblocks = raw_textblocks
        all_textblocks.extend(textblocks)

        # ================如果有音符================
        if note_textblocks:
            for textblock in note_textblocks:
                textblock_contour_mask = zeros(letter_mask.shape, dtype=uint8)
                drawContours(textblock_contour_mask, [textblock.block_contour], 0, 255, -1)
                letter_in_textblock = bitwise_and(letter_mask, textblock_contour_mask)
                contours, _ = findContours(letter_in_textblock, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)

                for cnt in contours:
                    # 找到从音符轮廓到边界的最短线段
                    dists = [(Point(p1[0]).distance(Point(p2[0])), (p1[0], p2[0])) for p1 in cnt for p2 in
                             filter_cnt.contour]
                    dists.sort(key=lambda x: x[0])
                    min_segment = dists[0][1]
                    # 音符轮廓及最短线段
                    mask = zeros(letter_mask.shape, dtype=uint8)
                    drawContours(mask, [cnt], 0, 255, -1)
                    line(mask, tuple(min_segment[0]), tuple(min_segment[1]), 255, 2)
                    dilated_mask = dilate(mask, kernel3, iterations=1)
                    filled_contour[dilated_mask == 255] = 0

            if do_dev_pic:
                filled_contour_preview_png = debug_dir / f'{media_type}_filled_contour_preview.png'
                write_pic(filled_contour_preview_png, filled_contour)

            contours_in_filled_contour, _ = findContours(filled_contour, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
            largest_contour = max(contours_in_filled_contour, key=contourArea)
            filter_cnt = Contr(largest_contour)

        # ================不是矩形且有两个以上文本框================
        if len(textblocks) >= 2 and fulfill_ratio <= 0.95:
            # 有2个以上的文字块轮廓则进行切割
            # 先用分水岭切割
            for t in range(len(textblocks)):
                textblock = textblocks[t]
                logger.debug(f'[{t + 1}]{textblock=}')
            seg_cnts = water_seg(filled_contour, textblocks)

            # 先检查数量一致性
            seg_success = (len(seg_cnts) == len(textblocks))
            logger.debug(f'{seg_success=}')
            # 再检查分割是否不破坏文本
            if seg_success and check_intact:
                for cnt in seg_cnts:
                    for textblock in textblocks:
                        if not cnt.polygon.is_valid:
                            cnt.polygon = cnt.polygon.buffer(0)
                        if not textblock.block_poly.is_valid:
                            textblock.block_poly = textblock.block_poly.buffer(0)
                        intersection = cnt.polygon.intersection(textblock.block_poly)
                        if intersection.area not in [0, textblock.block_poly.area]:
                            seg_success = False
                            break
                    if not seg_success:
                        break

            if not seg_success and use_pivot_split:
                if do_dev_pic:
                    textblock_bubbles = get_textblock_bubbles(bubble_mask, textblocks)
                    cp_textblock_jpg = current_dir / f'TextBlock-{cnt_ind}.jpg'
                    logger.warning(f'{cp_textblock_jpg=}')
                    write_pic(cp_textblock_jpg, textblock_bubbles)
                logger.warning(f'{len(textblocks)=}, 分水岭切割失败')
                # 再用轴点针转切割
                seg_cnts, seg_success = pivot_proc(filter_cnt, filled_contour, letter_in_contour, textblocks, cnt_ind)
            if not seg_success:
                logger.error(f'{len(textblocks)=}, 轴点针转切割失败')
                seg_cnts = [filter_cnt]
            single_cnts.extend(seg_cnts)
        else:
            single_cnts.append(filter_cnt)
    return single_cnts, all_textblocks


@logger.catch
def get_bubbles_by_cp(img_file, color_pattern, frame_grid_strs, CTD_mask, media_type, auto_subdir, bboxes=None):
    """
    分析气泡并根据指定的颜色模式对其进行处理。

    :param img_file: 输入的漫画或其他带文字的图像。
    :param color_pattern: 用于分析和处理气泡的颜色模式。
    :param frame_grid_strs: 框架网格字符串列表，用于确定气泡在图像中的位置和顺序。
    :param CTD_mask: 漫画文本检测器生成的气泡掩码。
    """

    added_frames_jpg = img_file.parent / f'{img_file.stem}-加框.jpg'

    bubble_mask_png = debug_dir / f'{media_type}_bubble_mask.png'
    letter_mask_png = debug_dir / f'{media_type}_letter_mask.png'

    img_raw = imdecode(fromfile(img_file, dtype=uint8), -1)
    if added_frames_jpg.exists():
        img_raw = imdecode(fromfile(added_frames_jpg, dtype=uint8), -1)
    img_raw = toBGR(img_raw)
    ih, iw = img_raw.shape[0:2]
    black_bg = zeros((ih, iw), dtype=uint8)

    cp_bubble, cp_letter = color_pattern
    logger.debug(f'{color_pattern=}, {cp_bubble=}, {cp_letter=}')
    if cp_bubble == '':
        color_bubble = None
    elif isinstance(cp_bubble, list):
        # 气泡为渐变色
        color_bubble = ColorGradient(cp_bubble)
    else:
        # 气泡为单色
        color_bubble = Color(cp_bubble)

    if isinstance(cp_letter, list):
        # 文字为双色
        color_letter = ColorDouble(cp_letter)
    else:
        # 文字为单色
        color_letter = Color(cp_letter)
    letter_mask = color_letter.get_range_img(img_raw)
    # ================无框================
    if color_bubble is None:
        # ================排除已经识别的气泡================
        mask_pics = [x for x in all_masks if x.stem.startswith(img_file.stem)]
        if mask_pics:
            single_cnts = get_single_cnts(img_raw, mask_pics)
            bubble_contours = [x.contour for x in single_cnts]
            ready_bubble_mask = drawContours(black_bg.copy(), bubble_contours, -1, 255, FILLED)
            dilated_mask = dilate(ready_bubble_mask, kernel5, iterations=1)
            dilated_mask_inv = bitwise_not(dilated_mask)
            letter_mask = bitwise_and(letter_mask, dilated_mask_inv)
        if CTD_mask is not None:
            ret, CTD_mask = threshold(CTD_mask, bi_thres, 255, THRESH_BINARY)
            letter_mask = bitwise_and(letter_mask, CTD_mask)

        if do_dev_pic:
            write_pic(letter_mask_png, letter_mask)
        all_textblocks = get_textblocks_full(letter_mask, media_type)
        ordered_cnts = [x.expanded_cnt for x in all_textblocks]

        cp_raw_jpg = auto_subdir / f'{img_file.stem}-Raw-{color_letter.rgb_str}-{color_letter.color_name}.jpg'
        cp_textblock_jpg = auto_subdir / f'{img_file.stem}-TextBlock-{color_letter.rgb_str}-{color_letter.color_name}.jpg'
        cp_preview_jpg = auto_subdir / f'{img_file.stem}-{color_letter.rgb_str}-{color_letter.color_name}.jpg'
        cp_mask_cnt_png = auto_subdir / f'{img_file.stem}-Mask-{color_letter.rgb_str}-{color_letter.color_name}.png'
    else:
        if color_bubble.type == 'Color':
            bubble_mask = color_bubble.get_range_img(img_raw)
            left_sample, right_sample = None, None
        else:  # ColorGradient
            img_tuple = color_bubble.get_range_img(img_raw)
            bubble_mask, left_sample, right_sample = img_tuple
        if do_dev_pic:
            write_pic(bubble_mask_png, bubble_mask)
            write_pic(letter_mask_png, letter_mask)
        cp_raw_jpg = auto_subdir / f'{img_file.stem}-Raw-{color_bubble.rgb_str}-{color_bubble.color_name}~{color_letter.rgb_str}-{color_letter.color_name}.jpg'
        cp_textblock_jpg = auto_subdir / f'{img_file.stem}-TextBlock-{color_bubble.rgb_str}-{color_bubble.color_name}~{color_letter.rgb_str}-{color_letter.color_name}.jpg'
        cp_preview_jpg = auto_subdir / f'{img_file.stem}-{color_bubble.rgb_str}-{color_bubble.color_name}~{color_letter.rgb_str}-{color_letter.color_name}.jpg'
        cp_mask_cnt_png = auto_subdir / f'{img_file.stem}-Mask-{color_bubble.rgb_str}-{color_bubble.color_name}~{color_letter.rgb_str}-{color_letter.color_name}.png'

        filter_cnts = get_raw_bubbles(bubble_mask, letter_mask, left_sample, right_sample, CTD_mask, bboxes)
        colorful_raw_bubbles = get_colorful_bubbles(img_raw, filter_cnts)
        # ================切割相连气泡================
        single_cnts, all_textblocks = seg_bubbles(filter_cnts, bubble_mask, letter_mask, media_type)
        # ================通过画格重新排序气泡框架================
        single_cnts_grids_ordered = []
        for f in range(len(frame_grid_strs)):
            frame_grid_str = frame_grid_strs[f]
            # 使用正则表达式提取字符串中的所有整数
            int_values = list(map(int, findall(r'\d+', frame_grid_str)))
            # 按顺序分配值
            x, y, w, h, xx, yy, ww, hh = int_values
            single_cnts_grid = []
            for s in range(len(single_cnts)):
                single_cnt = single_cnts[s]
                if x <= single_cnt.cx <= x + w and y <= single_cnt.cy <= y + h:
                    single_cnts_grid.append(single_cnt)
                else:
                    # 计算该single_cnt到所有frame_grid的距离
                    dists = []
                    for frame_grid_str_opt in frame_grid_strs:
                        int_values_opt = list(map(int, findall(r'\d+', frame_grid_str_opt)))
                        x_opt, y_opt, w_opt, h_opt, _, _, _, _ = int_values_opt
                        rect_opt = (x_opt, y_opt, w_opt, h_opt)
                        if single_cnt.polygon:
                            dist2rect = get_dist2rect(single_cnt.polygon.centroid, rect_opt)
                            dists.append(dist2rect)
                    if dists:
                        # 寻找最小距离及其对应的frame_grid的索引
                        closest_dist = min(dists)
                        closest_frame_idx = dists.index(closest_dist)
                        if closest_frame_idx == f:
                            single_cnts_grid.append(single_cnt)

            if media_type == 'Manga':
                ax = -1
            else:
                ax = 1
            single_cnts_grid.sort(key=lambda x: ax * x.cx + x.cy)
            single_cnts_grids_ordered.append(single_cnts_grid)
        ordered_cnts = list(chain(*single_cnts_grids_ordered))

    if color_bubble is None and bboxes:
        filtered_cnts = []
        for o in range(len(ordered_cnts)):
            cnt = ordered_cnts[o]
            if cnt.polygon:
                # 确保轮廓足以构成多边形
                for bbox in bboxes:
                    bbox_polygon = box(*bbox)
                    # 计算两个多边形的交集
                    intersection = cnt.polygon.intersection(bbox_polygon)
                    # 返回重叠区域的面积
                    overlap_area = intersection.area
                    min_area = min(100, int(0.5 * bbox_polygon.area))
                    if overlap_area >= min_area:
                        filtered_cnts.append(cnt)
                        # 如果找到至少一个满足条件的bbox，就不再继续检查这个cnt
                        break
    else:
        filtered_cnts = ordered_cnts

    logger.debug(f'{len(bboxes)=}')
    logger.debug(f'{len(ordered_cnts)=}')
    logger.debug(f'{len(filtered_cnts)=}')
    if len(filtered_cnts) >= 1 or (color_bubble is None and do_dev_pic):
        colorful_single_bubbles = get_colorful_bubbles(img_raw, filtered_cnts)
        textblock_bubbles = get_textblock_bubbles(img_raw, all_textblocks)

        # 创建一个带有bg_alpha透明度原图背景的图像
        transparent_img = zeros((ih, iw, 4), dtype=uint8)
        transparent_img[..., :3] = img_raw
        transparent_img[..., 3] = int(255 * bg_alpha)
        # 在透明图像上绘制contours，每个contour使用不同颜色

        for s in range(len(filtered_cnts)):
            bubble_cnt = filtered_cnts[s]
            # 从 tab20 颜色映射中选择一个颜色
            color = colormap_tab20(s % 20)[:3]
            color_rgb = tuple(int(c * 255) for c in color)
            color_bgra = color_rgb[::-1] + (255,)
            drawContours(transparent_img, [bubble_cnt.contour], -1, color_bgra, -1)

        # write_pic(cp_raw_jpg, colorful_raw_bubbles)
        write_pic(cp_textblock_jpg, textblock_bubbles)
        write_pic(cp_preview_jpg, colorful_single_bubbles)
        write_pic(cp_mask_cnt_png, transparent_img)
    return cp_mask_cnt_png


@logger.catch
def analyze1pic(img_file, frame_data, color_patterns, media_type, auto_subdir):
    logger.warning(f'{img_file=}')
    img_raw = imdecode(fromfile(img_file, dtype=uint8), -1)
    img_raw = toBGR(img_raw)
    ih, iw = img_raw.shape[0:2]
    # ================矩形画格信息================
    frame_grid_strs = frame_data.get(img_file.name, [f'0,0,{iw},{ih}~0,0,{iw},{ih}'])
    # ================模型检测文字，文字显示为白色================
    blk_bboxes = []
    line_bboxes = []
    if use_torch and CTD_model is not None:
        CTD_mask = get_CTD_mask(img_file)
        if better_data:  # and do_dev_pic
            mask, blks, lines_map, resize_ratio = get_CTD_data(img_raw)
            bccs, lines = CTD2data(blks, lines_map, resize_ratio)
            line_img = img_raw.copy()
            for bbox, cls, conf in zip(*bccs):
                x_min, y_min, x_max, y_max = bbox
                blk_bboxes.append(bbox)
                # 在标注图像上绘制边界框，使用红色线条
                rectangle(line_img, (x_min, y_min), (x_max, y_max), (255, 0, 0), 1)
            for ii, line in enumerate(lines):
                # 计算线条的边界
                bx1, bx2 = line[:, 0].min(), line[:, 0].max()
                by1, by2 = line[:, 1].min(), line[:, 1].max()
                bbox = (bx1, by1, bx2, by2)
                line_bboxes.append(bbox)
                rectangle(line_img, (bx1, by1), (bx2, by2), (0, 0, 255), 1)
            if do_dev_pic:
                line_path = debug_dir / f'{media_type}_line_anno.png'
                write_pic(line_path, line_img)
    else:
        CTD_mask = None
        mask, bccs, lines = None, None, None

    if bbox_type == 'block':
        bboxes = blk_bboxes
    elif bbox_type == 'line':
        bboxes = line_bboxes
    else:
        bboxes = blk_bboxes + line_bboxes
    # ================针对每一张图================
    for c in range(len(color_patterns)):
        # ================遍历每种气泡文字颜色组合================
        color_pattern = color_patterns[c]
        cp_mask_cnt_png = get_bubbles_by_cp(img_file, color_pattern, frame_grid_strs, CTD_mask, media_type, auto_subdir,
                                            bboxes)


@timer_decorator
# @logger.catch
def step1_analyze_bubbles(img_folder, media_type, auto_subdir):
    """
    分析气泡，获取气泡位置和尺寸，将结果可视化并保存到文件中。
    """
    # 这是一个用于分析漫画气泡并将结果保存到文件中的Python函数。该函数首先从YAML文件中加载画格信息，然后对给定的漫画图像列表进行处理。
    # 函数的主要工作是遍历每种气泡文字颜色组合，并分析每张图像中的气泡和文字。对于每种颜色组合，函数使用该颜色来创建气泡蒙版，并使用白色文字蒙版来检测文本。函数使用这些蒙版来过滤出符合条件的气泡，并对相邻的气泡进行切割和排序，以获得正确的气泡框架。
    # 接下来，函数为每个气泡框架绘制一个不同的颜色，并将结果保存为JPEG文件和PNG文件。这些文件用于可视化结果和后续的处理。
    # 最后，函数将处理后的气泡蒙版复制到原始图像所在的文件夹中。
    img_list = get_valid_imgs(img_folder)
    frame_yml = img_folder.parent / f'{img_folder.name}.yml'
    frame_data = iload_data(frame_yml)
    all_masks_old = get_valid_imgs(img_folder, vmode='mask')

    if pic_thread_method == 'queue':
        for p, img_file in enumerate(img_list):
            analyze1pic(img_file, frame_data, color_patterns, media_type, auto_subdir)
    else:
        with ThreadPoolExecutor(os.cpu_count()) as executor:
            futures = [
                executor.submit(analyze1pic, img_file, frame_data, color_patterns, media_type, auto_subdir) for
                img_file in img_list]

    # ================搬运气泡蒙版================
    auto_all_masks = get_valid_imgs(auto_subdir, vmode='mask')
    # 如果步骤开始前没有气泡蒙版
    if not all_masks_old:
        for mask_src in auto_all_masks:
            mask_dst = img_folder / mask_src.name
            copy2(mask_src, mask_dst)
    return auto_all_masks


# ================获取气泡内的文字================
def get_right_mask(letter_mask, bit_white_bubble, color_bubble, single_cnt):
    ih, iw = letter_mask.shape[0:2]
    black_bg = zeros((ih, iw), dtype=uint8)
    mask = bit_white_bubble
    letter_in_bubble_raw = bitwise_and(letter_mask, letter_mask, mask=bit_white_bubble)
    if has_decoration:
        # ================有装饰线================
        if color_bubble is None or single_cnt.area <= 900:
            # 无框或气泡较小
            pass
        else:
            if single_cnt.area <= 30000:
                kernel = kernel30
            elif single_cnt.area <= 40000:
                kernel = kernel35
            else:
                kernel = kernel40
            bit_white_bubble_ero = erode(bit_white_bubble, kernel, iterations=1)
            bit_white_bubble_ero_inv = bitwise_not(bit_white_bubble_ero)
            outline_mask = bitwise_and(bit_white_bubble, bit_white_bubble, mask=bit_white_bubble_ero_inv)
            letter_outline = bitwise_and(letter_mask, letter_mask, mask=outline_mask)
            # 获取所有非零像素的坐标
            px_pts = transpose(nonzero(letter_outline))
            # 计算非零像素的数量
            px_area = px_pts.shape[0]
            if px_area <= 10:
                # 可能文字靠近边缘
                pass
            else:
                all_x = px_pts[:, 1]
                all_y = px_pts[:, 0]
                raw_min_x, raw_max_x = np.min(all_x), np.max(all_x)
                raw_min_y, raw_max_y = np.min(all_y), np.max(all_y)
                raw_w = raw_max_x - raw_min_x
                raw_h = raw_max_y - raw_min_y
                if raw_w <= 0.6 * single_cnt.br_w or raw_h <= 0.6 * single_cnt.br_h:
                    # 可能文字靠近边缘
                    pass
                else:
                    mask = bit_white_bubble_ero.copy()
                    letter_contours_raw, _ = findContours(letter_in_bubble_raw, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
                    for contour in letter_contours_raw:
                        cnt = Contr(contour)
                        if cnt.area >= 2:
                            contour_mask = drawContours(black_bg.copy(), [cnt.contour], -1, 255, -1)
                            # 计算当前轮廓与outline_mask的交集
                            intersection_mask = bitwise_and(contour_mask, outline_mask)
                            has_white = np.any(intersection_mask == 255)
                            if has_white:
                                # 有交集，将当前轮廓涂黑到mask中
                                drawContours(mask, [cnt.contour], -1, 0, FILLED)
                                # pass
    return mask


# @logger.catch
# @timer_decorator
def get_single_cnts(img_raw, mask_pics):
    """
    从原始图像和对应的掩码图像中提取单个气泡的轮廓及其裁剪后的图像。

    :param img_raw: 原始图像，通常为漫画或其他带文字的图像。
    :param mask_pics: 包含掩码图像的列表，这些掩码用于在原始图像中找到气泡。
    :return 单个气泡轮廓及其裁剪后的图像的列表
    """
    img_raw = toBGR(img_raw)
    ih, iw = img_raw.shape[0:2]
    black_bg = zeros((ih, iw), dtype=uint8)
    single_cnts = []

    # 一次性读取所有的透明图像
    transparent_imgs = [imdecode(fromfile(mask_pic, dtype=uint8), -1) for mask_pic in mask_pics]

    for m, transparent_img in enumerate(transparent_imgs):
        # ================获取轮廓================
        # 将半透明像素变为透明
        alpha_channel = transparent_img[..., 3]
        transparent_mask = alpha_channel < 255
        opaque_mask = alpha_channel == 255
        transparent_img[transparent_mask] = 0
        # 获取所有完全不透明像素的颜色
        non_transparent_pxs = transparent_img[opaque_mask, :3]
        unique_colors = np.unique(non_transparent_pxs, axis=0)

        # 对于每种颜色，提取对应的contour
        contour_list = []
        for color in unique_colors:
            mask = np.all(transparent_img[..., :3] == color, axis=-1)
            mask = (mask & opaque_mask).astype(uint8) * 255
            contours, _ = findContours(mask, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
            contour_list.extend(contours)

        # ================遍历每种气泡文字颜色组合================
        color_pattern = get_color_pattern(mask_pics[m])
        cp_bubble, cp_letter = color_pattern

        if cp_bubble == '':
            # 无框
            color_bubble = None
        elif isinstance(cp_bubble, list):
            # 气泡为渐变色
            color_bubble = ColorGradient(cp_bubble)
            img_tuple = color_bubble.get_range_img(img_raw)
            bubble_mask, left_sample, right_sample = img_tuple
        else:
            # 气泡为单色
            color_bubble = Color(cp_bubble)
            bubble_mask = color_bubble.get_range_img(img_raw)

        if isinstance(cp_letter, list):
            # 文字为双色
            color_letter = ColorDouble(cp_letter)
        else:
            # 文字为单色
            color_letter = Color(cp_letter)
        letter_mask = color_letter.get_range_img(img_raw)

        if do_dev_pic:
            letter_mask_png = debug_dir / f'{media_type}_letter_mask.png'
            write_pic(letter_mask_png, letter_mask)

        # ================获取对应的气泡文字蒙版================
        for c, contour in enumerate(contour_list):
            single_cnt = Contr(contour)
            if single_cnt.area >= 30:
                # 在黑色背景上绘制单个轮廓的白色填充图像
                bit_white_bubble = drawContours(black_bg.copy(), [single_cnt.contour], 0, 255, -1)
                bit_white_bubble_inv = bitwise_not(bit_white_bubble)

                # 获取原始图像在bit_white_bubble范围内的图像，其他部分为气泡颜色
                img_bubble_only = bitwise_and(img_raw, img_raw, mask=bit_white_bubble)
                if color_bubble is None:
                    # 无框，掩模中黑色部分用白色
                    img_bubble_only[bit_white_bubble == 0] = (255, 255, 255)
                else:
                    # 掩模中黑色部分用气泡原本的颜色
                    img_bubble_only[bit_white_bubble == 0] = color_bubble.rgb
                img_bubble_only_inv = bitwise_not(img_bubble_only)

                mask = get_right_mask(letter_mask, bit_white_bubble, color_bubble, single_cnt)
                letter_in_bubble = bitwise_and(letter_mask, letter_mask, mask=mask)
                letter_in_bubble_inv = bitwise_not(letter_in_bubble)

                # 获取所有非零像素的坐标
                px_pts = transpose(nonzero(letter_in_bubble))
                # 计算非零像素的数量
                px_area = px_pts.shape[0]
                # 获取所有非零像素的x和y坐标
                all_x = px_pts[:, 1]
                all_y = px_pts[:, 0]
                # logger.debug(f'{px_area=}')

                if px_area >= 1:
                    # 计算不带padding的最小和最大x、y坐标
                    raw_min_x, raw_max_x = np.min(all_x), np.max(all_x)
                    raw_min_y, raw_max_y = np.min(all_y), np.max(all_y)

                    # 计算中心点坐标
                    center_x = (raw_min_x + raw_max_x) / 2
                    center_y = (raw_min_y + raw_max_y) / 2
                    center_pt = (int(center_x), int(center_y))

                    # 使用clip方法简化坐标计算
                    min_x = clip(raw_min_x - padding, 0, iw)
                    max_x = clip(raw_max_x + padding, 0, iw)
                    min_y = clip(raw_min_y - padding, 0, ih)
                    max_y = clip(raw_max_y + padding, 0, ih)

                    # 创建包含所有坐标的元组
                    letter_coors = (raw_min_x, raw_min_y, raw_max_x, raw_max_y, min_x, min_y, max_x, max_y, center_pt)

                    if isinstance(cp_bubble, list):
                        # 渐变框
                        mod_img = letter_in_bubble_inv.copy()
                    elif color_bubble is None:
                        # 无框
                        if color_letter.rgb_str == 'ffffff':
                            # 文字颜色为白色
                            mod_img = letter_in_bubble_inv.copy()
                        elif color_letter.rgb_str == '000000':
                            # 文字颜色为黑色
                            mod_img = img_bubble_only.copy()
                        else:
                            mod_img = letter_in_bubble_inv.copy()
                    else:
                        # 从带有白色背景的图像中裁剪出带有padding的矩形区域
                        if color_bubble.rgb_str == 'ffffff':
                            # 气泡颜色为白色则使用原图
                            mod_img = img_bubble_only.copy()
                        elif color_bubble.rgb_str == '000000':
                            # 气泡颜色为黑色则使用原图反色
                            mod_img = img_bubble_only_inv.copy()
                        elif color_letter.rgb_str == 'ffffff':
                            # 文字颜色为白色则使用原图反色把气泡颜色替换为白色后的转换成的灰度图
                            mod_img = img_bubble_only_inv.copy()
                            mod_img[bubble_mask == 255] = color_white
                            mod_img[bit_white_bubble_inv == 255] = color_white
                            if do_dev_pic:
                                mod_img_png = debug_dir / f'{media_type}_mod_img.png'
                                write_pic(mod_img_png, mod_img)
                        elif color_letter.rgb_str == '000000':
                            # 文字颜色为黑色则使用原图把气泡颜色替换为白色后的转换成的灰度图
                            mod_img = img_bubble_only.copy()
                            mod_img[bubble_mask == 255] = color_white
                            mod_img[bit_white_bubble_inv == 255] = color_white
                        else:
                            mod_img = letter_in_bubble_inv.copy()

                    if len(mod_img.shape) == 3:
                        mod_img = cvtColor(mod_img, COLOR_BGR2GRAY)
                    if len(mod_img.shape) == 2:
                        mod_img = cvtColor(mod_img, COLOR_GRAY2BGR)
                    cropped_img = mod_img[min_y:max_y, min_x:max_x]

                    single_cnt.add_cropped_img(cropped_img, letter_coors, color_pattern)
                    single_cnts.append(single_cnt)
    return single_cnts


def parse_color_parts(parts):
    """
    解析颜色部分，生成颜色信息。

    :param parts: 分割后的颜色字符串列表。
    :return: 颜色信息字符串或列表。
    """
    if len(parts) == 4:
        # 如果有四个部分，说明有两种颜色
        return [f'{parts[0]}-{parts[2]}', f'{parts[1]}-{parts[3]}']
    elif len(parts) == 2:
        # 如果有两个部分，说明只有一种颜色
        return f'{parts[0]}-{parts[1]}'
    return ''  # 如果没有颜色信息，返回空字符串


@logger.catch
def get_color_pattern(mask_pic):
    mask_prefix, par, cp_str = mask_pic.stem.partition('-Mask-')
    # 先切分气泡颜色和文字颜色
    # 白底黑字
    # ffffff-15-white~000000-60-black
    # 渐变底黑字
    # 75a1ba-c3d7d8-cadetblue-lightgray~000000-black
    # 无框双色字
    # 595ca9-a8dcea-slateblue-lightblue
    left_str, partition, right_str = cp_str.partition('~')
    left_parts = left_str.split('-')
    if right_str == '':
        # 无框
        right_parts = deepcopy(left_parts)
        left_parts = []
    else:
        right_parts = right_str.split('-')
    c_bubble = parse_color_parts(left_parts)
    c_letter = parse_color_parts(right_parts)
    color_pattern = [c_bubble, c_letter]
    return color_pattern


def get_start_panels(frame_grid_strs):
    panels = []
    for f in range(len(frame_grid_strs)):
        frame_grid_str = frame_grid_strs[f]
        # 使用正则表达式提取字符串中的所有整数
        int_values = list(map(int, findall(r'\d+', frame_grid_str)))
        if len(int_values) == 8:
            # 按顺序分配值
            x, y, w, h, xx, yy, ww, hh = int_values
            outer_br = (x, y, w, h)
            inner_br = (xx, yy, ww, hh)
            br_tup = (outer_br, inner_br)
            panels.append(br_tup)
    return panels


# @timer_decorator
@logger.catch
def get_panel_layers(pic_div_psd):
    panel_layers = []
    if pic_div_psd.exists():
        psd = PSDImage.open(pic_div_psd)
        layers = [layer for layer in psd]
        # ================读取 PSD 文件中所有可见的非背景图层================
        visible_layers = [layer for layer in layers if layer.visible]
        visible_layers = [layer for layer in visible_layers if layer.name != '背景']
        for n in range(len(visible_layers)):
            # ================每个图层================
            layer = visible_layers[n]
            layer_name = layer.name
            if layer_name.startswith('图层'):
                layer_name = layer_name.removeprefix('图层').strip()
            # 当前图层名称为数字时，计入画格图层
            if layer_name.isdigit():
                panel_layers.append(layer)
    return panel_layers


# @logger.catch
@timer_decorator
def get_grid_masks(img_file, frame_grid_strs):
    """
    根据给定的矩形框架字符串列表和 PSD 文件，生成每个框架的二值掩膜。

    :param img_file: 输入的漫画或其他带文字的图像。
    :param frame_grid_strs: 字符串列表，包含每个矩形框架的位置、宽度和高度信息。

    :return: 掩膜图像列表，包含每个框架的二值掩膜。
    """

    img_folder = img_file.parent
    auto_subdir = Auto / img_folder.name
    pic_div_psd = img_folder / f'{img_file.stem}-分框.psd'
    marked_frames_jpg = auto_subdir / f'{img_file.stem}-画格.jpg'

    img_raw = imdecode(fromfile(img_file, dtype=uint8), -1)
    panel_pil = fromarray(cvtColor(img_raw, COLOR_BGR2RGB))
    panel_draw = ImageDraw.Draw(panel_pil, 'RGBA')

    ih, iw = img_raw.shape[0:2]
    black_bg = zeros((ih, iw), uint8)
    logger.warning(f'{frame_grid_strs=}')

    # ================获取画格内外方框================
    panels = get_start_panels(frame_grid_strs)
    # ================读取psd画格================
    panel_layers = get_panel_layers(pic_div_psd)
    # ================图层按数字序数排序================
    panel_layers.sort(key=lambda x: int(x.name.removeprefix('图层').strip()))
    # ================以正确顺序插入到panels================
    for p in range(len(panel_layers)):
        # ================每个画格图层================
        layer = panel_layers[p]
        layer_index = int(layer.name.removeprefix('图层').strip())
        panels.insert(layer_index - 1, layer)

    # ================生成grid_masks================
    grid_masks = []
    for p in range(len(panels)):
        # ================每个画格图层================
        panel = panels[p]
        # 从 tab20 颜色映射中选择一个颜色
        color = colormap_tab20(p % 20)[:3]
        color_rgb = tuple(int(c * 255) for c in color)
        color_bgra = color_rgb[::-1] + (255,)
        color_bgr127 = color_rgb[::-1] + (127,)

        if isinstance(panel, tuple):
            # ================内外方框================
            outer_br, inner_br = panel
            # logger.warning(f'{inner_br=}')
            x, y, w, h = outer_br
            xx, yy, ww, hh = inner_br
            # ================绘制外圈矩形================
            pt1 = (x, y)
            pt2 = (x + w, y + h)
            # ================内圈矩形================
            pt3 = (xx, yy)
            pt4 = (xx + ww, yy + hh)
            panel_draw.rectangle([pt1, pt2], outline=color_bgra, width=3)
            panel_draw.rectangle([pt3, pt4], outline=color_bgr127, width=3)
            # 以外圈矩形为基点计算序数坐标
            text_pos = (x + 5, y - 5)
            # 生成矩形画格信息对应的二值图像，画格内部为 255（白色），其他部分为 0（黑色）
            filled_frame_outer = rectangle(black_bg.copy(), pt1, pt2, 255, -1)
            filled_frame_inner = rectangle(black_bg.copy(), pt3, pt4, 255, -1)
            grid_masks.append(filled_frame_outer)
        else:
            # ================psd图层================
            layer = panel
            layer_index = int(layer.name.removeprefix('图层').strip())

            # 获取图层的边界框坐标
            lef, topp, righ, bott = layer.bbox
            # ================numpy================
            layer_img_np = layer.numpy()
            # 将图层数据转换为 0-255 范围的 uint8 类型
            layer_img_np = (layer_img_np * 255).astype(uint8)

            # 对图层的 alpha 通道应用阈值处理，生成二值化掩膜
            ret, layer_mask_raw = threshold(layer_img_np[:, :, 3], 0, 255, THRESH_BINARY)

            layer_mask = fromarray(black_bg.copy())
            layer_mask_raw = fromarray(layer_mask_raw)
            # 将二值化掩膜粘贴到黑色背景上，以生成完整的掩膜图像
            layer_mask.paste(layer_mask_raw, (lef, topp))
            # 将图像对象转换为 numpy 数组
            layer_mask = asarray(layer_mask)

            px_pts = transpose(nonzero(layer_mask))
            px_area = px_pts.shape[0]
            all_x = px_pts[:, 1]
            all_y = px_pts[:, 0]
            all_x.sort()
            all_y.sort()

            # 计算图层边界框的位置、宽度和高度
            layer_x = int(min(all_x))
            layer_y = int(min(all_y))
            layer_w = int(max(all_x) - min(all_x))
            layer_h = int(max(all_y) - min(all_y))
            layer_br = (layer_x, layer_y, layer_w, layer_h)
            logger.warning(f'{layer_br=}')
            layer_area = layer_w * layer_h
            # 像素面积与图层边界框面积之比
            px_br_ratio = px_area / layer_area
            if px_br_ratio >= 1.2:
                # 使用黑色背景创建填充的矩形框
                pt1 = (layer_x, layer_y)
                pt2 = (layer_x + layer_w, layer_y + layer_h)
                filled_frame = rectangle(black_bg.copy(), pt1, pt2, 255, -1)
            else:
                # 否则，对掩膜进行腐蚀和膨胀操作，去除小点，以生成填充的框架
                filled_frame = erode(layer_mask, kernel5)
                filled_frame = dilate(filled_frame, kernel5)
            grid_masks.append(filled_frame)

            px_pts = transpose(nonzero(filled_frame))
            px_area = px_pts.shape[0]
            all_x = px_pts[:, 1]
            all_y = px_pts[:, 0]
            all_x.sort()
            all_y.sort()

            grid_x = min(all_x)
            grid_y = min(all_y)
            grid_w = max(all_x) - min(all_x)
            grid_h = max(all_y) - min(all_y)
            grid_br = (grid_x, grid_y, grid_w, grid_h)
            lef, topp = grid_x, grid_y
            # 偏左了，向右移
            lef += 10
            text_pos = (lef, topp)

            filled_frame_erode10 = erode(filled_frame, kernel10)
            filled_frame_erode10_inv = bitwise_not(filled_frame_erode10)
            outline_mask = bitwise_and(filled_frame, filled_frame, mask=filled_frame_erode10_inv)
            filled_outline = cvtColor(outline_mask, COLOR_GRAY2BGRA)
            # 黑色转为透明
            filled_outline[np.all(filled_outline == rgba_black, axis=2)] = rgba_zero
            # 白色转为颜色
            filled_outline[np.all(filled_outline == rgba_white, axis=2)] = color_bgra
            filled_outline = fromarray(cvtColor(filled_outline, COLOR_BGRA2RGBA))
            panel_pil.paste(filled_outline, pos_0, filled_outline)
        panel_draw.text(text_pos, f'{p + 1}', index_color, font=msyh_font60)
    marked_frames = cvtColor(array(panel_pil), COLOR_RGB2BGR)
    write_pic(marked_frames_jpg, marked_frames)
    return grid_masks, marked_frames


def ocr_by_tesseract_simple(image, media_lang, vert=False):
    # 获取所选语言的配置
    language_config = tesseract_language_options[media_lang]
    config = f'-c preserve_interword_spaces=1 --psm 6 -l {language_config}'
    # 如果识别竖排文本
    if vert:
        config += " -c textord_old_baselines=0"
    text = image_to_string(image, config=config)
    return text.strip()


# 'level': 层次结构中的级别。1 表示页面级别，2 表示区域级别，3 表示段落级别，4 表示文本行级别，5 表示单词级别。
# 'page_num': 页码。
# 'block_num': 区域编号。
# 'par_num': 段落编号。
# 'line_num': 行号。
# 'word_num': 单词编号。
# 'left': 边界框左上角的 x 坐标。
# 'top': 边界框左上角的 y 坐标。
# 'width': 边界框的宽度。
# 'height': 边界框的高度。
# 'conf': 置信度，表示 Tesseract 识别结果的可信度，范围为 0 到 100。
# 'text': 识别的文本。


def ocr_by_tesseract(image, media_lang, vert=False):
    # 获取所选语言的配置
    language_config = tesseract_language_options[media_lang]
    config = f'-c preserve_interword_spaces=1 --psm 6 -l {language_config}'
    # 如果识别竖排文本
    if vert:
        config += " -c textord_old_baselines=0"
    data = image_to_data(image, config=config, output_type='dict')
    # 将各个键的值组合成元组
    tess_zdata = zip(
        data['level'],
        data['page_num'],
        data['block_num'],
        data['par_num'],
        data['line_num'],
        data['word_num'],
        data['left'],
        data['top'],
        data['width'],
        data['height'],
        data['conf'],
        data['text']
    )
    return list(tess_zdata)


# @logger.catch
def ocr_by_vision(image, media_lang):
    languages = [vision_language_options[media_lang]]  # 设置需要识别的语言

    image = conv_img(image, target_format='CV')
    height, width = image.shape[:2]
    image_data = image.tobytes()
    provider = CGDataProviderCreateWithData(None, image_data, len(image_data), None)
    cg_image = CGImageCreate(width, height, 8, 8 * 3, width * 3, CGColorSpaceCreateDeviceRGB(), 0, provider, None,
                             False, kCGRenderingIntentDefault)

    handler = VNImageRequestHandler.alloc().initWithCGImage_options_(cg_image, None)

    request = VNRecognizeTextRequest.new()
    # 使用快速识别（0-快速，1-精确）
    request.setRecognitionLevel_(0)
    # request.setRecognitionLevel_(1)
    request.setUsesLanguageCorrection_(True)  # 使用语言纠正
    request.setRecognitionLanguages_(languages)  # 设置识别的语言

    error = None
    success = handler.performRequests_error_([request], error)
    if success:
        results = request.results()
        text_results = []
        for result in results:
            text = result.text()
            confidence = result.confidence()
            bounding_box = result.boundingBox()
            x_raw = bounding_box.origin.x
            y_raw = bounding_box.origin.y
            w_raw = bounding_box.size.width
            h_raw = bounding_box.size.height
            x = round(x_raw * width)
            y = round(y_raw * height)
            w = round(w_raw * width)
            h = round(h_raw * height)
            # line_br = (x, y, w, h)
            tup = (x, y, w, h, confidence, text)
            text_results.append(tup)
        return text_results
    else:
        print("Error: ", error)
        return []


async def recognize_text(img, lang):
    image_pil = fromarray(cvtColor(img, COLOR_BGR2RGB))
    result = await recognize_pil(image_pil, lang)
    return result


def ocr_by_winocr(img, media_lang):
    lang = winocr_language_options[media_lang]
    image_pil = fromarray(cvtColor(img, COLOR_BGR2RGB))
    result = asyncio.run(recognize_text(img, lang))
    text_results = []
    for l in range(len(result.lines)):
        line = result.lines[l]
        for w in range(len(line.words)):
            word = line.words[w]
            x = word.bounding_rect.x
            y = word.bounding_rect.y
            width = word.bounding_rect.width
            height = word.bounding_rect.height
            # print(f'[{x}, {y}, {width}, {height}]{word.text}')
            tup = (x, y, width, height, l, w, word.text)
            text_results.append(tup)
    return text_results


# @logger.catch
def ocr_by_paddle(image, language):
    lang_code = paddleocr_language_map.get(language, language)
    if lang_code == 'en' and init_ocr:
        ocr = en_ocr
    else:
        ocr = PaddleOCR(use_gpu=False, lang=lang_code)
    result = ocr.ocr(image)
    result0 = result[0]
    text_results = []
    if result0:
        for res in result0:
            para_pts = res[0]
            text, confidence = res[1]
            pts = list(chain(*para_pts))
            pts = [int(x) for x in pts]
            tup = tuple(pts) + (confidence, text)
            text_results.append(tup)
    return text_results


# @logger.catch
def ocr_by_easy(image, language):
    lang_code = easyocr_language_map.get(language, language)
    if lang_code == 'English' and init_ocr:
        reader = en_reader
    else:
        reader = Reader([lang_code])
    result = reader.readtext(image)
    text_results = []
    for res in result:
        # logger.debug(f'{res=}')
        para_pts = res[0]
        text = res[1]
        confidence = res[2]
        pts = list(chain(*para_pts))
        pts = [int(x) for x in pts]
        tup = tuple(pts) + (confidence, text)
        # logger.debug(f'{tup=}')
        text_results.append(tup)
    return text_results


# language_type
# 识别语言类型，默认为CHN_ENG
# 可选值包括：
# - CHN_ENG：中英文混合
# - ENG：英文
# - JAP：日语
# - KOR：韩语
# - FRE：法语
# - SPA：西班牙语
# - POR：葡萄牙语
# - GER：德语
# - ITA：意大利语
# - RUS：俄语

# detect_direction
# 是否检测图像朝向，默认不检测，即：false。朝向是指输入图像是正常方向、逆时针旋转90/180/270度。可选值包括:
# - true：检测朝向；
# - false：不检测朝向。

# detect_language
# 是否检测语言，默认不检测。当前支持（中文、英语、日语、韩语）

# paragraph
# 是否输出段落信息

# probability
# 是否返回识别结果中每一行的置信度
def ocr_by_baidu_base(image, language, req_type):
    _, buffer = imencode('.jpg', image)
    image_as_bytes = buffer.tobytes()

    obd_lang = 'ENG'
    options = {
        # 'detect_direction': 'true',
        'language_type': obd_lang,
        # 'detect_language': 'true',
        'paragraph': 'true',
        'probability': 'true',
    }

    # 调用通用文字识别接口
    if req_type == 'basic':
        result = aip_client.basicGeneral(image_as_bytes, options)
    else:
        result = aip_client.basicAccurate(image_as_bytes, options)

    text_results = []
    paras_result = result['paragraphs_result']
    words_result = result['words_result']
    for p in range(len(paras_result)):
        para = paras_result[p]
        words_result_idx = para['words_result_idx']
        for ind in words_result_idx:
            words_text = words_result[ind]['words']
            probability = words_result[ind]['probability']
            prob_average = probability['average']
            prob_min = probability['min']
            prob_variance = probability['variance']
            tup = (p, ind, prob_average, prob_min, prob_variance, words_text)
            text_results.append(tup)
    return text_results


def ocr_by_baidu(image, language):
    return ocr_by_baidu_base(image, language, req_type='basic')


def ocr_by_baidu_accu(image, language):
    return ocr_by_baidu_base(image, language, req_type='accurate')


def find_intervals(projection, threshold):
    intervals = []  # 用于存储符合条件的区间
    start = None  # 区间的起始点
    length = 0  # 区间的长度

    # 遍历投影数据
    for i, value in enumerate(projection):
        # 如果当前值大于等于阈值
        if value >= threshold:
            # 如果区间尚未开始，设置起始点为当前位置
            if start is None:
                start = i
            # 区间长度加1
            length += 1
        else:
            # 如果区间已经开始，将区间添加到结果列表中
            if start is not None:
                intervals.append((start, length))
                start = None
                length = 0

    # 如果循环结束时区间仍然有效，将区间添加到结果列表中
    if start is not None:
        intervals.append((start, length))

    return intervals


# @logger.catch
def sort_bubble(cnt, ax, origin_x, origin_y):
    val = ax * (cnt.cx - origin_x) + (cnt.cy - origin_y)
    if y_1st:
        if cnt.core_br_y - origin_y <= 60:
            val = 0.1 * ax * (cnt.cx - origin_x) + (cnt.cy - origin_y)
        elif cnt.core_br_y - origin_y <= 300:
            val = 0.3 * ax * (cnt.cx - origin_x) + (cnt.cy - origin_y)
        elif cnt.core_br_y - origin_y <= 600:
            val = 0.5 * ax * (cnt.cx - origin_x) + (cnt.cy - origin_y)
    return val


# @logger.catch
@timer_decorator
def get_ordered_cnts(single_cnts, img_file, grid_masks, bubble_order_strs, media_type):
    img_raw = imdecode(fromfile(img_file, dtype=uint8), -1)
    ih, iw = img_raw.shape[0:2]
    black_bg = zeros((ih, iw), uint8)

    if media_type == 'Manga':
        ax = -1
    else:
        ax = 1
    all_contours = [single_cnt.contour for single_cnt in single_cnts]
    white_bubbles = drawContours(black_bg.copy(), all_contours, -1, 255, FILLED)
    # ================通过画格重新排序气泡框架================
    single_cnts_grids_ordered = []
    single_cnts_grids = []
    for g in range(len(grid_masks)):
        grid_mask = grid_masks[g]
        grid_mask_px_pts = transpose(nonzero(grid_mask))
        grid_mask_px_area = grid_mask_px_pts.shape[0]
        # ================获取外接矩形坐标作为参考================
        all_x = grid_mask_px_pts[:, 1]
        all_y = grid_mask_px_pts[:, 0]
        all_x.sort()
        all_y.sort()
        grid_x = min(all_x)
        grid_y = min(all_y)
        grid_w = max(all_x) - min(all_x)
        grid_h = max(all_y) - min(all_y)
        logger.debug(f'[画格{g + 1}]外接矩形({grid_x}, {grid_y}, {grid_w}, {grid_h})')
        grid_white_bubbles = bitwise_and(white_bubbles, grid_mask)
        grid_white_bubbles_px_pts = transpose(nonzero(grid_white_bubbles))
        grid_white_bubbles_px_area = grid_white_bubbles_px_pts.shape[0]
        single_cnts_grid = []
        # ================如果画格和气泡有重叠区域================
        if grid_white_bubbles_px_area >= 100:
            # ================获取当前画格内所有气泡轮廓================
            for s in range(len(single_cnts)):
                single_cnt = single_cnts[s]
                cxy_color = grid_mask[single_cnt.cy, single_cnt.cx]
                if cxy_color == 255:
                    # ================如果轮廓质心在此画格内部================
                    single_cnts_grid.append(single_cnt)
        single_cnts_grids.append(single_cnts_grid)

    # ================边界外的轮廓================
    catered_cnts = list(chain(*single_cnts_grids))
    uncatered_cnts = [x for x in single_cnts if x not in catered_cnts]
    for u in range(len(uncatered_cnts)):
        uncatered_cnt = uncatered_cnts[u]
        logger.info(f'未分类轮廓{uncatered_cnt.cxy=}')
        for g in range(len(grid_masks)):
            grid_mask = grid_masks[g]
            if grid_mask[uncatered_cnt.br_v, uncatered_cnt.cx] == 255:
                logger.info(f'{g=}')
                # ================如果轮廓质心在此画格内部================
                single_cnts_grids[g].append(uncatered_cnt)
                break

    for g in range(len(grid_masks)):
        grid_mask = grid_masks[g]
        grid_mask_px_pts = transpose(nonzero(grid_mask))
        grid_mask_px_area = grid_mask_px_pts.shape[0]
        # ================获取外接矩形坐标作为参考================
        all_x = grid_mask_px_pts[:, 1]
        all_y = grid_mask_px_pts[:, 0]
        all_x.sort()
        all_y.sort()
        grid_x = min(all_x)
        grid_y = min(all_y)
        grid_w = max(all_x) - min(all_x)
        grid_h = max(all_y) - min(all_y)
        logger.info(f'[{g + 1}]({grid_x}, {grid_y}, {grid_w}, {grid_h})')
        grid_white_bubbles = bitwise_and(white_bubbles, grid_mask)
        grid_white_bubbles_px_pts = transpose(nonzero(grid_white_bubbles))
        grid_white_bubbles_px_area = grid_white_bubbles_px_pts.shape[0]

        single_cnts_grid = single_cnts_grids[g]
        logger.debug(f'[画格{g + 1}]{len(single_cnts_grid)=}')
        # ================根据高度块重排序================
        bulk_cnts_grid = []
        if single_cnts_grid:
            if len(single_cnts_grid) == 1:
                # ================只有一个轮廓================
                single_cnts_grid_ordered = single_cnts_grid
                single_cnt = single_cnts_grid_ordered[0]
                single_cnt.get_core_br(black_bg)
            else:
                # ================多个轮廓================
                single_cnts_grid = sorted(single_cnts_grid, key=lambda x: x.br_y)
                # ================获取高度块================
                # 将在Y轴方向上有重叠的部分划分到一起
                grouped_bulks = []
                current_group = [single_cnts_grid[0]]
                for s in range(1, len(single_cnts_grid)):
                    curr_cnt = single_cnts_grid[s]
                    group_bottom = max([cnt.br_v for cnt in current_group])
                    group_top = min([cnt.br_y for cnt in current_group])
                    # 如果当前cnt的顶部在组内任何一个cnt的底部之上（考虑阈值）并且
                    # 当前cnt的底部在组内任何一个cnt的顶部之下（考虑阈值）
                    if curr_cnt.br_y <= group_bottom + bulk_thres and curr_cnt.br_v >= group_top - bulk_thres:
                        current_group.append(curr_cnt)
                    else:
                        grouped_bulks.append(current_group)
                        current_group = [curr_cnt]
                if current_group:
                    grouped_bulks.append(current_group)
                # ================如果grid_mask又宽又矮================
                if (grid_mask_px_area <= 0.25 * ih * iw and grid_w >= 0.5 * iw) or fully_framed:
                    # ================所有高度块强制设置为同一组================
                    grouped_bulks = [list(chain(*grouped_bulks))]
                # ================对已按高度分组的轮廓继续进行分析================
                bulk_cnts_group = []
                for b in range(len(grouped_bulks)):
                    grouped_bulk = grouped_bulks[b]
                    logger.debug(f'[高度块{b + 1}]{len(grouped_bulk)=}')
                    y_min = min([cnt.br_y for cnt in grouped_bulk])
                    y_max = max([cnt.br_v for cnt in grouped_bulk])
                    # ================获取团块================
                    mask_y = black_bg.copy()
                    mask_y[y_min:y_max, :] = 255
                    masked_grid_white_part = bitwise_and(grid_white_bubbles, mask_y)
                    masked_grid_mask = bitwise_and(grid_mask, mask_y)
                    dilated_masked = dilate(masked_grid_white_part.copy(), kernel60, iterations=1)
                    dilated_masked = bitwise_and(dilated_masked, masked_grid_mask)
                    bulk_contours, _ = findContours(dilated_masked, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
                    bulk_cnts = []
                    for b in range(len(bulk_contours)):
                        bulk_contour = bulk_contours[b]
                        bulk_cnt = Contr(bulk_contour)
                        if bulk_cnt.area >= 30:
                            bulk_cnt.get_core_br(black_bg)
                            bulk_cnts.append(bulk_cnt)
                    all_br_x = [cnt.core_br_x for cnt in bulk_cnts]
                    all_br_y = [cnt.core_br_y for cnt in bulk_cnts]
                    origin_x, origin_y = min(all_br_x), min(all_br_y)
                    bulk_cnts.sort(key=lambda x: ax * (x.cx - origin_x) + (x.cy - origin_y))
                    bulk_cnts_group.append(bulk_cnts)
                bulk_cnts_grid = list(chain(*bulk_cnts_group))
                if bulk_cnts_grid:
                    logger.debug(f'{len(bulk_cnts_grid)=}')
                    cnts_in_bulk_grid = []
                    for b in range(len(bulk_cnts_grid)):
                        bulk_cnt = bulk_cnts_grid[b]
                        # ================质心在团块轮廓中的气泡轮廓================
                        cnts_in_bulk = [x for x in single_cnts if
                                        pointPolygonTest(bulk_cnt.contour, x.order_pt, False) > 0]
                        if cnts_in_bulk:
                            # ================获取整体的外接矩形左上角坐标================
                            for c in range(len(cnts_in_bulk)):
                                cnt = cnts_in_bulk[c]
                                cnt.get_core_br(black_bg)
                            all_br_x = [cnt.core_br_x for cnt in cnts_in_bulk]
                            all_br_y = [cnt.core_br_y for cnt in cnts_in_bulk]
                            origin_x, origin_y = min(all_br_x), min(all_br_y)
                            # ================画格内按坐标排序================
                            cnts_in_bulk.sort(key=lambda cnt: sort_bubble(cnt, ax, origin_x, origin_y))
                            cnts_in_bulk_grid.append(cnts_in_bulk)
                    single_cnts_grid_ordered = list(chain(*cnts_in_bulk_grid))
                # ================使用手动排序得到的记录排序================
                if bubble_order_strs:
                    try:
                        # 根据bubble_order_strs中的cxy_str顺序对single_cnts_grid_ordered进行排序
                        single_cnts_grid_ordered.sort(key=lambda cnt: bubble_order_strs.index(cnt.cxy_str))
                    except ValueError as e:
                        logger.error(f'{img_file=}')
            single_cnts_grids_ordered.append(single_cnts_grid_ordered)
    # ================对排序后的轮廓去重================
    cxys = set()
    new_grids = []
    for o in range(len(single_cnts_grids_ordered)):
        single_cnts_grid = single_cnts_grids_ordered[o]
        logger.debug(f'[画格{o + 1}]{len(single_cnts_grid)=}')
        new_grid = [single_cnt for single_cnt in single_cnts_grid if single_cnt.cxy not in cxys]
        cxys.update(single_cnt.cxy for single_cnt in new_grid)
        new_grids.append(new_grid)
    return new_grids


@lru_cache
def get_sd_img_folder(img_folder):
    # ================获取漫画名================
    series4file = img_folder.name
    m_issue_w_dot = p_issue_w_dot.match(folder_name)
    issue = None
    if m_issue_w_dot:
        series4file = m_issue_w_dot.group(1)
        issue = m_issue_w_dot.group(2)

    sd_folder_name = f'{series4file}-SD'
    if issue:
        sd_folder_name += f' {issue}'
    sd_img_folder = img_folder.parent / sd_folder_name
    return sd_img_folder


# @logger.catch
def figure1pic(img_folder, i):
    auto_subdir = Auto / img_folder.name
    img_list = get_valid_imgs(img_folder)

    # ================气泡蒙版================
    all_masks = get_valid_imgs(img_folder, vmode='mask')
    img_file = img_list[i]
    logger.warning(f'{img_file=}')
    added_frames_jpg = auto_subdir / f'{img_file.stem}-加框.jpg'
    added_frames_jpg_dst = img_folder / f'{img_file.stem}-加框.jpg'
    img_raw = imdecode(fromfile(img_file, dtype=uint8), -1)
    ih, iw = img_raw.shape[0:2]

    # ================获取对应的文字图片================
    mask_pics = [x for x in all_masks if x.stem.startswith(img_file.stem)]
    if mask_pics:
        single_cnts = get_single_cnts(img_raw, mask_pics)
        logger.debug(f'{len(single_cnts)=}')

        blk_bboxes = []
        line_bboxes = []
        mask, blks, lines_map, resize_ratio = get_CTD_data(img_raw)
        bccs, lines = CTD2data(blks, lines_map, resize_ratio)
        line_img = img_raw.copy()
        for bbox, cls, conf in zip(*bccs):
            x_min, y_min, x_max, y_max = bbox
            blk_bboxes.append(bbox)
            # 在标注图像上绘制边界框，使用红色线条
            rectangle(line_img, (x_min, y_min), (x_max, y_max), (255, 0, 0), 1)
        for ii, line in enumerate(lines):
            # 计算线条的边界
            bx1, bx2 = line[:, 0].min(), line[:, 0].max()
            by1, by2 = line[:, 1].min(), line[:, 1].max()
            bbox = (bx1, by1, bx2, by2)
            line_bboxes.append(bbox)
            rectangle(line_img, (bx1, by1), (bx2, by2), (0, 0, 255), 1)
        bboxes = blk_bboxes

        filtered_bboxes = []
        for b in range(len(bboxes)):
            bbox = bboxes[b]
            bbox_polygon = box(*bbox)
            for s in range(len(single_cnts)):
                cnt = single_cnts[s]
                # 计算两个多边形的交集
                intersection = cnt.polygon.intersection(bbox_polygon)
                # 返回重叠区域的面积
                overlap_area = intersection.area
                min_area = min(100, int(0.5 * bbox_polygon.area))
                if overlap_area >= min_area:
                    filtered_bboxes.append(bbox)
                    # 如果找到至少一个满足条件的cnt，就不再继续检查这个bbox
                    break

        if len(filtered_bboxes) < len(bboxes):
            logger.error(f'{img_file.name=}')
            copy2(added_frames_jpg, added_frames_jpg_dst)


@timer_decorator
@logger.catch
def order1pic(img_folder, i, media_type):
    auto_subdir = Auto / img_folder.name
    img_list = get_valid_imgs(img_folder)

    frame_yml = img_folder.parent / f'{img_folder.name}.yml'
    order_yml = img_folder.parent / f'{img_folder.name}-气泡排序.yml'
    frame_data = iload_data(frame_yml)
    order_data = iload_data(order_yml)

    lp_txt = img_folder.parent / f'{img_folder.name}-8萌翻.txt'
    rlp_txt = img_folder.parent / f'{img_folder.name}翻译_0.txt'
    rlp_para_dic = None
    if rlp_txt.exists():
        rlp_text = read_txt(rlp_txt)
        rlp_lines = rlp_text.splitlines()
        rlp_index_dict, rlp_inds = create_index_dict(img_list, rlp_lines, 'txt')
        rlp_para_dic = create_para_dic(img_list, rlp_index_dict, rlp_inds, rlp_lines, 'txt')

    # ================气泡蒙版================
    all_masks = get_valid_imgs(img_folder, vmode='mask')

    img_file = img_list[i]
    logger.warning(f'{img_file=}')
    order_preview_jpg = auto_subdir / f'{img_file.stem}-气泡排序.jpg'
    lp_preview_jpg = auto_subdir / f'{img_file.stem}-气泡萌翻.jpg'
    img_raw = imdecode(fromfile(img_file, dtype=uint8), -1)
    ih, iw = img_raw.shape[0:2]

    # ================矩形画格信息================
    # 从 frame_data 中获取画格信息，默认为整个图像的矩形区域
    frame_grid_strs = frame_data.get(img_file.name, [f'0,0,{iw},{ih}~0,0,{iw},{ih}'])
    bubble_order_strs = order_data.get(img_file.name, [])
    grid_masks, marked_frames = get_grid_masks(img_file, frame_grid_strs)
    # ================获取对应的文字图片================
    mask_pics = [x for x in all_masks if x.stem.startswith(img_file.stem)]

    sd_img_folder = get_sd_img_folder(img_folder)
    sd_bubble_order_strs = []
    if sd_img_folder.exists():
        sd_img_list = get_valid_imgs(sd_img_folder)
        sd_order_yml = img_folder.parent / f'{sd_img_folder.name}-气泡排序.yml'
        sd_order_data = iload_data(sd_order_yml)
        sd_img_raw = imdecode(fromfile(img_file, dtype=uint8), -1)
        if i <= len(sd_img_list):
            sd_img_file = sd_img_list[i]
            sd_bubble_order_strs = sd_order_data.get(sd_img_file.name, [])
            sd_img_raw = imdecode(fromfile(sd_img_file, dtype=uint8), -1)
        sd_ih, sd_iw = sd_img_raw.shape[0:2]

    ordered_cnts = []
    if mask_pics:
        single_cnts = get_single_cnts(img_raw, mask_pics)
        single_cnts_grids_ordered = get_ordered_cnts(single_cnts, img_file, grid_masks, bubble_order_strs, media_type)
        ordered_cnts = list(chain(*single_cnts_grids_ordered))
        logger.warning(f'{len(single_cnts)=}')
        logger.warning(f'{len(single_cnts_grids_ordered)=}')
        logger.warning(f'{len(ordered_cnts)=}')
        if not rlp_para_dic:
            order_preview = get_order_preview(marked_frames, single_cnts_grids_ordered)
            write_pic(order_preview_jpg, order_preview)

        if rlp_para_dic:
            formatted_stem = get_formatted_stem(img_file)
            rlp_pic_bubbles = rlp_para_dic.get(formatted_stem, [])
            if len(rlp_pic_bubbles) == len(single_cnts) == 1:
                new_cnts = single_cnts
            else:
                new_cnts = []
                roi_cnts = deepcopy(single_cnts)
                for b in range(len(rlp_pic_bubbles)):
                    rlp_pic_bubble = rlp_pic_bubbles[b]
                    coor_x = rlp_pic_bubble['coor_x']
                    coor_y = rlp_pic_bubble['coor_y']
                    x = int(iw * coor_x)
                    y = int(ih * coor_y)
                    pt_xy = (x, y)
                    # 收集所有轮廓到点的距离
                    distances = []
                    for c, roi_cnt in enumerate(roi_cnts):
                        distance = pointPolygonTest(roi_cnt.contour, pt_xy, True)
                        distances.append((distance, roi_cnt))
                    # 根据距离排序，找到最近的轮廓
                    distances.sort(key=lambda x: -x[0])
                    if distances:
                        distances0 = distances[0]
                        distance, roi_cnt = distances0
                        if distance > 0 or abs(distance) <= dist_thres:
                            new_cnts.append(roi_cnt)
                            roi_cnts = [x for x in roi_cnts if x != roi_cnt]
                new_cnts.extend(roi_cnts)
            single_cnts_grids_ordered = [new_cnts]
            order_preview = get_order_preview(marked_frames, single_cnts_grids_ordered)
            write_pic(order_preview_jpg, order_preview)

            img_pil = fromarray(cvtColor(marked_frames, COLOR_BGR2RGB))
            num_overlay = Image.new('RGBA', img_pil.size, rgba_zero)
            num_draw = ImageDraw.Draw(num_overlay)
            # ================绘制示意图================
            for b in range(len(rlp_pic_bubbles)):
                rlp_pic_bubble = rlp_pic_bubbles[b]
                group = rlp_pic_bubble['group']
                coor_x = rlp_pic_bubble['coor_x']
                coor_y = rlp_pic_bubble['coor_y']
                x = int(iw * coor_x)
                y = int(ih * coor_y)
                pt_xy = (x, y)
                text = str(b + 1)  # 序数
                # 在pt_xy位置画一个紫色的点
                radius = 5  # 点的半径
                anchor_pt = (x - radius, y - radius, x + radius, y + radius)
                text_bbox = num_draw.textbbox(pt_xy, text, font=msyh_font100, anchor="mm")
                text_width = text_bbox[2] - text_bbox[0]
                text_height = text_bbox[3] - text_bbox[1]
                # xy_pos = (x - text_width // 2, y - text_height // 2)
                # 在点的右下角
                xy_pos = (x + radius, y + radius)
                if group == 1:
                    num_draw.ellipse(anchor_pt, fill=color_purple)
                    num_draw.text(xy_pos, text, font=msyh_font100, fill=trans_red)
                else:
                    num_draw.ellipse(anchor_pt, fill=color_olive)
                    num_draw.text(xy_pos, text, font=msyh_font100, fill=trans_blue)
            img_pil.paste(num_overlay, mask=num_overlay)
            lp_preview = cvtColor(array(img_pil), COLOR_RGB2BGR)
            write_pic(lp_preview_jpg, lp_preview)

        elif sd_img_folder.exists() and not bubble_order_strs:
            if len(sd_bubble_order_strs) == len(ordered_cnts):
                # 把低清版图源的排序转换并录入到高清版图源
                custom_cnts = []
                for s in range(len(sd_bubble_order_strs)):
                    sd_bubble_order_str = sd_bubble_order_strs[s]
                    x, par, y = sd_bubble_order_str.partition(',')
                    x = int(x)
                    y = int(y)
                    tar_x = x * iw / sd_iw
                    tar_y = y * ih / sd_ih
                    tar_x = int(tar_x)
                    tar_y = int(tar_y)
                    tar_xy = (tar_x, tar_y)
                    # 找到ordered_cnts中cxy最靠近tar_xy的cnt
                    tar_pt = Point(tar_xy)
                    nearest_cnts = deepcopy(ordered_cnts)
                    nearest_cnts.sort(key=lambda x: x.cxy_pt.distance(tar_pt))
                    nearest_cnt = nearest_cnts[0]
                    dist = nearest_cnt.cxy_pt.distance(tar_pt)
                    logger.warning(f'[{dist}]{tar_pt=}, {nearest_cnt.cxy_pt=}')
                    custom_cnts.append(nearest_cnt)

                order_data = iload_data(order_yml)
                content0 = [cnt.cxy_str for cnt in ordered_cnts]
                content = [cnt.cxy_str for cnt in custom_cnts]
                highlight_diffs(content0, content)
                if content != content0:
                    order_data[img_file.name] = content
                    bubble_data_sorted = {k: order_data[k] for k in natsorted(order_data)}
                    write_yml(order_yml, bubble_data_sorted)
                    logger.debug(f"已保存到{order_yml}")
            elif sd_bubble_order_strs:
                logger.error(
                    f'[{img_file.stem}][{len(sd_bubble_order_strs)}/{len(ordered_cnts)}]{sd_bubble_order_strs=}')
    return img_file, ordered_cnts


@timer_decorator
def get_order_preview(marked_frames, single_cnts_grids_ordered):
    ordered_cnts = list(chain(*single_cnts_grids_ordered))
    img_pil = fromarray(cvtColor(marked_frames, COLOR_BGR2RGB))
    # 创建与原图大小相同的透明图像
    bubble_overlay = Image.new('RGBA', img_pil.size, rgba_zero)
    core_br_overlay = Image.new('RGBA', img_pil.size, rgba_zero)
    line_overlay = Image.new('RGBA', img_pil.size, rgba_zero)
    bubble_draw = ImageDraw.Draw(bubble_overlay)
    core_br_draw = ImageDraw.Draw(core_br_overlay)
    line_draw = ImageDraw.Draw(line_overlay)
    line_alpha = 192
    # ================绘制示意图================
    for f in range(len(ordered_cnts)):
        bubble_cnt = ordered_cnts[f]
        contour_points = [tuple(point[0]) for point in bubble_cnt.contour]
        x, y, w, h = bubble_cnt.core_br
        # 从 tab20 颜色映射中选择一个颜色
        color = colormap_tab20(f % 20)[:3]
        bubble_color_rgba = tuple(int(c * 255) for c in color) + (int(255 * bubble_alpha),)
        bubble_draw.polygon(contour_points, fill=bubble_color_rgba)
        color_rgba = tuple(int(c * 255) for c in color) + (255,)
        # 绘制矩形
        core_br_draw.rectangle([(x, y), (x + w, y + h)], outline=color_rgba, width=1)

    for f in range(len(ordered_cnts)):
        bubble_cnt = ordered_cnts[f]
        # 在轮廓中心位置添加序数
        text = str(f + 1)
        text_bbox = bubble_draw.textbbox(bubble_cnt.cxy, text, font=msyh_font100, anchor="mm")
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]
        xy_pos = (bubble_cnt.cx - text_width // 2, bubble_cnt.cy - text_height // 2)
        bubble_draw.text(xy_pos, text, font=msyh_font100, fill=trans_purple)

    # ================按照画格内对话框顺序添加红色连线================
    for g in range(len(single_cnts_grids_ordered)):
        single_cnts_grid = single_cnts_grids_ordered[g]
        # 确定该画格内需要绘制的红线数量
        total_lines = len(single_cnts_grid) - 1
        if len(single_cnts_grid) >= 2:
            for b in range(len(single_cnts_grid)):
                bubble_cnt = single_cnts_grid[b]
                # 绿色圆圈表示质心
                radius = 10  # 锚点半径
                anchor_pt = (
                    bubble_cnt.cx - radius, bubble_cnt.cy - radius, bubble_cnt.cx + radius, bubble_cnt.cy + radius)
                trans_green = (0, 255, 0, line_alpha)
                line_draw.ellipse(anchor_pt, fill=trans_green)
                # 如果不是该画格内的第一个气泡，与上一个气泡的质心画连线
                if b > 0:
                    prev_cnt = single_cnts_grid[b - 1]
                    line_color = (255, 0, 0, line_alpha)
                    if total_lines >= 2:
                        # 根据当前红线的序号逐渐增加
                        blue_value = int(255 * (b - 1) / total_lines)
                        # 使用计算出的蓝色值来调整红线的颜色
                        line_color = (255, 0, blue_value, line_alpha)
                    # 在新的透明图层上用调整后的颜色绘制连线
                    line_draw.line([prev_cnt.cxy, bubble_cnt.cxy], fill=line_color, width=5)

    # 将带有半透明颜色轮廓的透明图像与原图混合
    img_pil.paste(bubble_overlay, mask=bubble_overlay)
    img_pil.paste(core_br_overlay, mask=core_br_overlay)
    img_pil.paste(line_overlay, mask=line_overlay)

    # 将 PIL 图像转换回 OpenCV 图像
    blended_img = cvtColor(array(img_pil), COLOR_RGB2BGR)
    return blended_img


@timer_decorator
# @logger.catch
def step2_order(img_folder, media_type):
    img_list = get_valid_imgs(img_folder)
    cnts_dic_pkl = img_folder.parent / f'{img_folder.name}.pkl'
    cnts_dic = iload_data(cnts_dic_pkl)
    lp_txt = img_folder.parent / f'{img_folder.name}-8萌翻.txt'
    rlp_txt = img_folder.parent / f'{img_folder.name}翻译_0.txt'

    if thread_method == 'queue':
        for i in range(len(img_list)):
            img_file, ordered_cnts = order1pic(img_folder, i, media_type)
            if ordered_cnts:
                cnts_dic[img_file] = ordered_cnts
    else:
        with ThreadPoolExecutor(os.cpu_count()) as executor:
            futures = [executor.submit(order1pic, img_folder, i, media_type)
                       for i in range(len(img_list))]
            for future in as_completed(futures):
                try:
                    img_file, ordered_cnts = future.result()
                    if ordered_cnts:
                        cnts_dic[img_file] = ordered_cnts
                except Exception as e:
                    printe(e)

    with open(cnts_dic_pkl, 'wb') as f:
        pickle.dump(cnts_dic, f)


def tesseract2text(tess_zdata):
    tess_zdata5 = [x for x in tess_zdata if x[0] == 5]
    par_nums = [x[3] for x in tess_zdata5]
    par_nums = reduce_list(par_nums)
    par_nums.sort()
    line_nums = [x[4] for x in tess_zdata5]
    line_nums = reduce_list(line_nums)
    line_nums.sort()
    lines_tess = []
    for par_num in par_nums:
        for line_num in line_nums:
            line_data = [x for x in tess_zdata5 if x[3] == par_num and x[4] == line_num]
            line_data.sort(key=lambda x: x[5])
            for l in range(len(line_data)):
                word_data = line_data[l]
                level, page_num, block_num, par_num, line_num, word_num, left, top, width, height, conf, text = word_data
            line_words = [x[-1] for x in line_data]
            line_text = ' '.join(line_words)
            if line_text != '':
                lines_tess.append(line_text)
    return lines_tess


def get_bubble_meta(bubble_meta_str):
    """
    从元数据字符串中提取气泡的各项属性。

    :param bubble_meta_str (str): 包含气泡元数据的字符串。

    :return: dict: 一个包含气泡属性的字典。
    """
    # 分割元数据字符串
    bubble_meta_list = bubble_meta_str.split('~')

    # 解析气泡的各项属性
    coors_str = bubble_meta_list[0]
    src_font_size = int(bubble_meta_list[1].removeprefix('S'))
    dst_font_size = int(bubble_meta_list[2].removeprefix('D'))
    bubble_shape = bubble_meta_list[3]
    text_direction = bubble_meta_list[4]
    text_alignment = bubble_meta_list[5]

    bubble_color_str = ''
    if len(bubble_meta_list) == 9:
        bubble_color_str = bubble_meta_list[6]

    letter_color_str = bubble_meta_list[-2]
    letter_font_name = bubble_meta_list[-1]

    # 解析坐标
    coors = [int(x) for x in coors_str.split(',')]
    min_x, min_y, br_w, br_h = coors

    # 计算中心点坐标
    center_x = min_x + br_w / 2
    center_y = min_y + br_h / 2
    center_pt = (int(center_x), int(center_y))

    # 将解析的属性数据放入一个字典中
    bubble_metadata = {
        'coordinates': {
            'min_x': min_x,
            'min_y': min_y,
            'br_w': br_w,
            'br_h': br_h,
            'center_x': center_x,
            'center_y': center_y,
            'center_pt': center_pt,
        },
        'src_font_size': src_font_size,
        'dst_font_size': dst_font_size,
        'bubble_shape': bubble_shape,
        'text_direction': text_direction,
        'text_alignment': text_alignment,
        'bubble_color_str': bubble_color_str,
        'letter_color_str': letter_color_str,
        'letter_font_name': letter_font_name
    }
    return bubble_metadata


def get_ocr_data(ocr_engine, pic_ocr_data, img_np, media_lang, elements_num):
    # 检查指定的OCR引擎是否已经用于当前图像
    if ocr_engine in pic_ocr_data:
        # 如果已经识别过，则直接从已有数据中获取OCR结果
        ocr_results = form2data(pic_ocr_data[ocr_engine])
    else:
        # 如果尚未识别，则调用相应的OCR引擎进行识别
        true_engine = ocr_engine.lower()
        if true_engine == 'tes':
            true_engine = 'tesseract'
        elif true_engine == 'vis':
            true_engine = 'vision'
        elif true_engine == 'win':
            true_engine = 'winocr'
        # ocr_by_{true_engine} 是一个函数，用于执行OCR操作
        ocr_results = globals()[f'ocr_by_{true_engine}'](img_np, media_lang)

        # 用于存储格式化后的OCR结果
        ocr_results_form = []

        # 遍历识别结果
        for row in ocr_results:
            # 分离出数字和文本部分
            row_nums = row[:-elements_num]  # 数字部分
            row_str = ','.join(map(str, row_nums))  # 数字转为字符串，并用逗号连接
            text = row[-1]  # 文本部分
            confs = row[-elements_num:-1]  # 置信度部分

            # 如果有置信度数据，则添加到字符串中
            if confs:
                row_str += ','
                row_str += ','.join([f'{x:.2f}' for x in confs])

            # 将文本部分添加到字符串末尾
            row_str += f'|{text}'

            # 将格式化的字符串添加到结果列表中
            ocr_results_form.append(row_str)

        # 将格式化后的结果保存到pic_ocr_data字典中，以便以后可以直接使用
        pic_ocr_data[ocr_engine] = ocr_results_form

    # 返回OCR结果和更新后的pic_ocr_data
    return ocr_results, pic_ocr_data


def rec2text(rec_results):
    lines_vision = []
    current_line = []
    last_y = None
    last_h = None
    for r in range(len(rec_results)):
        rec_result = rec_results[r]
        x, y, w, h, conf, text_seg = rec_result
        good_text = "".join([similar_chars_map.get(c, c) for c in text_seg])
        # 检查 y 值
        if last_y is not None and abs(y - last_y) <= y_thres and 0.45 * last_h <= h <= 1.6 * last_h:
            # 如果 y 值相近，添加到当前行
            current_line.append(good_text)
        else:
            # 否则，开始新的一行
            if current_line:
                lines_vision.append(' '.join(current_line))
            current_line = [good_text]
            last_y = y
            last_h = h
    if current_line:
        # 添加最后一行
        lines_vision.append(' '.join(current_line))
    return lines_vision


def get_upper_ratio(input_text):
    # 计算文本中大写和小写字母的数量
    letters = [char for char in input_text if char.isalpha()]
    letter_count = Counter([char.isupper() for char in letters])
    # 计算大写字母的百分比
    up_ratio = 0
    if letters:
        up_ratio = letter_count[True] / len(letters)
    return up_ratio


def correct_word(word_text):
    for old, new in corrections:
        if old in word_text:
            rep_word = word_text.replace(old, new)
            if rep_word.lower() in good_words:
                return rep_word

    # if word_text.endswith('IS'):
    #     new_word_text = word_text.removesuffix('IS')
    #     if new_word_text.lower() in good_words or new_word_text.capitalize() in good_words:
    #         return f"{new_word_text}'S"

    if word_text.endswith(('Z', '2')):
        new_word_text = word_text.removesuffix('Z').removesuffix('2')
        if new_word_text.lower() in good_words or new_word_text.capitalize() in good_words:
            return f"{new_word_text}?"

    return word_text


# @logger.catch
def better_text(input_text, ocr_type):
    fmt_text = input_text.strip()
    # 根据替换规则表进行替换
    for old, new in replace_rules.items():
        fmt_text = fmt_text.replace(old, new)
    # ================处理省略号================
    fmt_text = sub(r'(\. ?){2,}', '…', fmt_text)
    fmt_text = sub(r'-{3,}', '--', fmt_text)
    # ================处理引号================
    for old, new in replacements.items():
        fmt_text = fmt_text.replace(old, new)
    # ================多个0改成多个O================
    fmt_text = sub(r'0{3,}', lambda x: 'O' * len(x.group()), fmt_text)
    # ================处理有缩写符号的词================
    fmt_text = sub(r'\b15\b(?!\s*percent\b)', "IS", fmt_text, flags=IGNORECASE)
    for old, new in better_abbrs.items():
        fmt_text = sub(r'\b' + old + r'\b', new, fmt_text)

    letters = [char for char in input_text if char.isalpha()]
    up_ratio = get_upper_ratio(fmt_text)
    check_upper = False
    if up_ratio > 0.5 and len(letters) >= 6:
        check_upper = True

    if ocr_type in ['tesseract', 'vision', 'baidu', 'baidu_accu']:
        lines = fmt_text.splitlines()
        processed_lines = []
        for l in range(len(lines)):
            line = lines[l]
            words = list(finditer(r'\b\w+\b', line))
            new_words = []
            for w in range(len(words)):
                word = words[w]
                start, end = word.span()
                word_text = line[start:end]
                # if not word_text.lower() in good_words:
                #     word_text = correct_word(word_text)
                # 处理大小写
                if word_text.isupper() or word_text.lower().startswith(interjections):
                    new_word = word_text
                elif word_text[0].isdigit():
                    new_word = word_text
                elif check_upper:
                    new_word = word_text.upper()
                else:
                    new_word = word_text
                new_words.append(new_word)
            if new_words:
                # 转义单个大括号
                line = line.replace('{', '{{').replace('}', '}}')
                line = sub(r'\b\w+\b', '{}', line).format(*new_words)
            else:
                line = sub(r'\b\w+\b', '', line)
            processed_lines.append(line)
        fmt_text = lf.join(processed_lines)

    # ================处理标点格式================
    # 确保"…"前后没有空格或特殊字符
    fmt_text = sub(r' +… +| +…|… +', '…', fmt_text)
    fmt_text = sub(r'-+…|…-+|\*+…|…\*+', '…', fmt_text)
    # 确保"-"前后没有空格
    fmt_text = sub(r'- +| +-', '-', fmt_text)
    # 其他替换
    fmt_text = sub(r' *• *', ' ', fmt_text)
    fmt_text = sub(r' +([?!.,])', r'\1', fmt_text)
    # 在标点后的单词前添加空格
    fmt_text = sub(r'([?!.,])(?![0-9])(\w)', r'\1 \2', fmt_text)
    fmt_text = fmt_text.replace('#$0%', '#$@%')
    fmt_text = fmt_text.replace('D. C.', 'D.C.')
    return fmt_text


def find_punct(text):
    """找到文本的段首和段尾标点"""
    start_punct = match(r'^\W+', text)
    end_punct = search(r'\W+$', text)
    start_p = start_punct.group() if start_punct else ''
    end_p = end_punct.group() if end_punct else ''
    if start_p == '. ':
        start_p = '…'
    return start_p, end_p


def is_valid_tense(word):
    doc = nlp(word)
    for token in doc:
        if token.tag_ in ['VBD', 'VBG']:  # VBD: past tense, VBG: gerund/present participle
            return True
    return False


@logger.catch
def get_final_token(token1, token2, up_ratio, is_sent_start, i1, i2, text1, line1):
    # 找到这个字母所在的完整单词
    word_start, word_end = i1, i2
    while word_start > 0 and line1[word_start - 1].isalpha():
        word_start -= 1
    while word_end < len(line1) and line1[word_end].isalpha():
        word_end += 1
    ori_word = line1[word_start:word_end]
    rep_word = ori_word[:i1 - word_start] + token2 + ori_word[i2 - word_start:]
    ori_word_lower = ori_word.lower()
    rep_word_lower = rep_word.lower()
    final_token = token1
    # 检查是否是单个字母替换成单个字母的情况
    if len(token1) == 1 and len(token2) == 1 and token1.lower() == token2.lower():
        if up_ratio >= 0.9 or is_sent_start or ori_word in proper_nouns:
            # 全大写或句首
            final_token = token1.upper()
        else:
            final_token = token1.lower()
        if final_token != token1:
            print(f'{token1}->{final_token}')
    elif token1 in punct_rep_rules and token2 in punct_rep_rules[token1]:
        final_token = token2
    elif token1.isdigit() and len(token1) == 1 and all(ch in valid_chars for ch in token2):
        final_token = token2
    elif token1.isdigit() and len(token1) == 1 and token2.isdigit() and len(token2) == 1:
        final_token = token2
    elif token1 == 'I.' and token2 == 'T…':
        final_token = 'I…'
    elif len(token1) == 1 and len(token2) == 1 and token1.isalpha():
        # 尝试获取单数形式
        ori_word_singular = lemmatizer.lemmatize(ori_word_lower, pos='n')
        rep_word_singular = lemmatizer.lemmatize(rep_word_lower, pos='n')
        # 检查这个单词是否是一个“合理的”单词
        if ori_word_lower in interjections:
            # 如果是语气词，保留原词
            final_token = token1
        elif ori_word_lower in good_words or ori_word_singular in good_words or ori_word_lower in fine_words:
            final_token = token1
        else:
            if token2.isalpha():
                if ori_word_lower.endswith(('ing', 'ed')):
                    if use_nlp and is_valid_tense(ori_word):
                        # 如果是过去时或进行时的合理词，保留原词
                        final_token = token1
                    elif rep_word_lower in good_words or rep_word_lower in fine_words:
                        print(f'{ori_word}->{rep_word}')
                        # 否则，如果替换后的词是合理的，进行替换
                        final_token = token2
                    else:
                        # 否则，保留原词
                        final_token = token1
                elif rep_word_lower in good_words or rep_word_singular in good_words or rep_word_lower in fine_words:
                    final_token = token2
                    print(f'{ori_word}->{rep_word}')
                else:
                    final_token = token1
            else:
                final_token = token2
                print(f'{ori_word}->{rep_word}')
    elif token1 == ',' and token2 == '.':
        # 获取原始文本和添加标点后的文本中的句子
        adj_text1 = text1[:i1] + token2 + text1[i2:]
        final_token = adjust_sent(token1, token2, text1, adj_text1)
    else:
        final_token = token1
    return final_token


@logger.catch
def get_sents(text):
    """获取文本中的所有句子的起始和结束位置"""
    # 考虑句号、问号、感叹号、省略号等作为句子的结束
    sentence_boundaries = finditer(r'(?<=[.!?…])\s+', text)
    start = 0
    sentences = []
    for match in sentence_boundaries:
        end = match.start()
        sentences.append((start, end))
        start = match.end()
    sentences.append((start, len(text)))
    return sentences


@logger.catch
def adjust_sent(token1, token2, text1, adj_text1):
    final_token = token1
    ori_sents = get_sents(text1)
    mod_sents = get_sents(adj_text1)

    # 比较两个版本的句子结构，找到句子二和句子三的位置
    split_index = None
    for idx, (orig_sent, mod_sent) in enumerate(zip(ori_sents, mod_sents)):
        if orig_sent != mod_sent:
            split_index = idx
            break

    # 如果找到了被分割的句子
    if split_index is not None and split_index + 1 < len(mod_sents):
        # 获取句子三的起始和结束位置
        next_sent_i, next_sent_j = mod_sents[split_index + 1]
        next_sent = adj_text1[next_sent_i:next_sent_j]
        print(f'{next_sent=}')
        # 检查句子是否首字母大写且不是整句话全部大写
        if next_sent and next_sent[0].isupper() and not next_sent.isupper():
            final_token = token2
    return final_token


@logger.catch
def fix_w_tess(text1, text2):
    # 将输入文本按行分割
    lines1 = text1.strip().splitlines()
    lines2 = text2.strip().splitlines()
    up_ratio = get_upper_ratio(text1)
    text1_merged = ' '.join(lines1)
    if use_nlp:
        doc = nlp(text1_merged)
        sent_starts = [sent.start_char for sent in doc.sents]
    else:
        tokens = sent_tokenize(text1_merged)
        sent_starts = [text1_merged.find(sent) for sent in tokens]
    new_lines1 = []

    # 如果 lines1 比 lines2 少一行
    if len(lines1) + 1 == len(lines2):
        max_similarity = 0
        max_sim_index = -1
        # 枚举每次去掉 lines2 的一行
        for i in range(len(lines2)):
            temp_lines2 = lines2[:i] + lines2[i + 1:]
            total_similarity = 0
            for line1, line2 in zip(lines1, temp_lines2):
                s = SequenceMatcher(None, line1, line2)
                total_similarity += s.ratio()
            sim = total_similarity / len(lines1) if lines1 else 0
            if sim > max_similarity:
                max_similarity = sim
                max_sim_index = i
        # 根据相似度最高的情况调整 lines1
        lines1 = lines1[:max_sim_index] + [None] + lines1[max_sim_index:]
        # print()

    # 当前处理到的字符位置
    pin = 0
    for line1, line2 in zip_longest(lines1, lines2, fillvalue=None):
        # 遍历两个文本的每一行
        # 如果 line2 为 None，保持 line1 不变
        if line2 is None:
            new_lines1.append(line1)
            print(f'{line1=}')
        elif line1 is None:
            new_lines1.append(line2)
            print(f'{line2=}')
        else:
            s = SequenceMatcher(None, line1, line2)
            # logger.info(f"{line1},{line2}")
            new_line1 = ''
            # 获取操作码并遍历，以找出两行文本之间的差异
            for tag, i1, i2, j1, j2 in s.get_opcodes():
                token1 = line1[i1:i2]
                token2 = line2[j1:j2]
                last_sent_i = max([x for x in sent_starts if x <= pin], default=0)
                is_sent_start = all(not c.isalpha() for c in text1_merged[last_sent_i:pin])
                # 默认情况下，final_token 是 token1
                final_token = token1

                if tag == 'equal':
                    final_token = token1
                elif tag == 'replace':
                    final_token = get_final_token(token1, token2, up_ratio, is_sent_start, i1, i2, text1, line1)
                elif tag == 'insert':
                    # 使用正则表达式提取行中的所有单词
                    words = findall(r"[\w']+", line1)
                    # 使用finditer()获取每个单词的位置
                    word_positions = [(m.start(), m.end()) for m in finditer(r"[\w']+", line1)]
                    # 根据i1的位置找到正确的单词及其位置
                    ori_word = None
                    word_start = None
                    for word, (start, end) in zip(words, word_positions):
                        if start <= i1 < end:
                            ori_word = word
                            word_start = start
                            break
                    if ori_word is not None and word_start is not None:
                        rep_word = ori_word[:i1 - word_start] + token2 + ori_word[i1 - word_start:]
                    else:
                        word_start = i1
                        word_end = i1
                        while word_start > 0 and line1[word_start - 1].isalpha():
                            word_start -= 1
                        while word_end < len(line1) and line1[word_end].isalpha():
                            word_end += 1
                        ori_word = line1[word_start:word_end]
                        rep_word = ori_word[:i1 - word_start] + token2 + ori_word[i1 - word_start:]
                    if ori_word == "C'ON" and token2 == 'M':
                        # 检查是否是C'ON变成C'MON的情况
                        final_token = token2
                        logger.debug(f'{final_token=}')
                    elif token2 == '-' and i1 > 0 and line1[i1 - 1] == '-':
                        # 如果在-之后插入一个-，接受这个插入
                        final_token = token2
                        logger.debug(f'{final_token=}')
                    elif len(token2) == 1 and token2.isalpha() and len(ori_word) >= 1:
                        # 检查这个单词是否是一个“合理的”单词
                        if ori_word.lower() not in good_words and rep_word.lower() in good_words or rep_word.lower() in fine_words:
                            final_token = token2
                            logger.debug(f'{final_token=}')
                    elif match(r'[a-zA-Z]', token2) or match(r'\s*\d+$', token2) or match(r'\s+', token2):
                        # 忽略纯字母、纯数字和纯空格的插入
                        pass
                    elif token2 in (':', '¥', '?', '! ', '[', ']', '(', ')', '-', '"', 'O'):
                        pass
                    elif token2 in ['.', '. ']:
                        # 获取原始文本和添加标点后的文本中的句子
                        adj_text1 = text1[:i1] + token2 + text1[i1:]
                        final_token = adjust_sent(token1, token2, text1, adj_text1)
                    else:
                        final_token = token2
                        logger.debug(f'{final_token=}')
                elif tag == 'delete':
                    final_token = token1

                pin += len(token1)
                if final_token is None:
                    logger.error(f'{final_token=}')
                    final_token = ''
                new_line1 += f'{final_token}'
            # 将处理后的新行添加到结果列表中
            new_lines1.append(new_line1)
            # 计算空格
            pin += 1
    # 将所有新行连接成一个字符串并返回
    new_text1 = lf.join(new_lines1)
    if new_text1 != text1:
        # print(f'{text1=}')
        # print(f'{new_text1=}')
        pass
    return new_text1


@logger.catch
def better_punct(vision_fmt_text, refer_ocrs):
    """优化 vision_fmt_text 文本"""

    # 计算文本中的单词数量
    word_count = len(list(finditer(r'\b\w+\b', vision_fmt_text)))

    start_puncts = []
    end_puncts = []

    # 收集所有参考文本的段首和段尾标点
    for text in refer_ocrs:
        start, end = find_punct(text)
        if start:
            start_puncts.append(start)
        if end:
            end_puncts.append(end)

    # 计算出现次数
    start_counter = Counter(start_puncts)
    end_counter = Counter(end_puncts)

    # 找到出现至少两次的段首和段尾标点
    common_start = [k for k, v in start_counter.items() if v >= 2]
    common_end = [k for k, v in end_counter.items() if v >= 2]

    # 获取 vision_fmt_text 的段首和段尾标点
    vision_start, vision_end = find_punct(vision_fmt_text)

    # 如果 vision 的段首标点不在共同段首标点中，则替换
    if vision_start not in common_start and common_start:
        vision_fmt_text = common_start[0] + vision_fmt_text[len(vision_start):]

    certains = ['to be continued']
    # 如果 vision 的段尾标点不在共同段尾标点中，则替换
    if vision_end:
        if vision_end not in common_end and common_end:
            vision_fmt_text = vision_fmt_text[:-len(vision_end)] + common_end[0]
    else:
        if common_end:
            vision_fmt_text += common_end[0]
        else:
            # 如果没有共同的段尾标点，添加最长的段尾标点
            longest_end_punct = max(end_puncts, key=len, default='')
            if longest_end_punct:
                vision_fmt_text += longest_end_punct
            elif vision_fmt_text.lower().endswith(interjections):
                # 语气词
                pass
            elif vision_fmt_text.lower() in certains:
                # 未完待续等特殊段落
                pass
            elif word_count >= 8:
                # 如果单词数量大于或等于8，添加句号
                vision_fmt_text += '.'

    if end_counter['-'] >= 1 and end_counter['--'] >= 1:
        vision_fmt_text = vision_fmt_text.rstrip('-') + '--'

    # 如果文本以逗号结尾，替换为句号
    vision_fmt_text = sub(r',$', '.', vision_fmt_text)
    vision_fmt_text = vision_fmt_text.replace('…,', '…')
    tess_fmt_text = refer_ocrs[-1]
    if '¿' in vision_fmt_text:
        vision_fmt_text = tess_fmt_text
    vision_fmt_text = sub(r'\'OR', 'OR', vision_fmt_text)
    return vision_fmt_text


@logger.catch
def run2ansi(run):
    text = run.text
    if text.strip() == '':
        return text
    elif run.bold and run.italic:
        return f"\033[1;3;35m{text}\033[0m"  # 粗斜体紫色
    elif run.bold:
        return f"\033[1;34m{text}\033[0m"  # 粗体蓝色
    elif run.italic:
        return f"\033[3;32m{text}\033[0m"  # 斜体绿色
    else:
        return text


@logger.catch
def get_dst_font_size(src_font_size, bubble_color_str, letter_color_str):
    if font_size_ratio_min <= src_font_size / global_font_size <= font_size_ratio_max:
        dst_font_size = global_font_size
    else:
        dst_font_size = round(src_font_size * 1.2 / base_num) * base_num
    color_locate = f"{bubble_color_str}-{letter_color_str}"
    # logger.debug(f'{color_locate=}')
    if color_locate in font_size_range_dic:
        # 按照气泡和文字颜色
        cfont_size_min, cfont_size_max = parse_range(font_size_range_dic[color_locate])
        # 小的拉到最低值，大的不变
        if dst_font_size <= cfont_size_max:
            dst_font_size = clip(dst_font_size, cfont_size_min, cfont_size_max)
    else:
        # 默认
        if src_font_size * 1.1 <= font_size_max:
            dst_font_size = clip(dst_font_size, font_size_min, font_size_max)
    return int(dst_font_size)


@logger.catch
def draw_para(img_tups, ocr_tups, page_ind, bubble_ind):
    img_np, gray_limg = img_tups
    tess_zdata, tes_zdata, rec_results_vision, rec_results_vis, rec_results_winocr, rec_results_win = ocr_tups

    nh, nw = img_np.shape[:2]  # 获取原始图像的尺寸
    white_bg = ones((nh, nw, 3), dtype=uint8) * 255
    black_bg = zeros((nh, nw), dtype=uint8)

    img_pil = fromarray(cvtColor(img_np, COLOR_BGR2RGB))
    img_gray = cvtColor(img_np, COLOR_BGR2GRAY)
    ret, img_bi = threshold(img_gray, bi_thres, 255, THRESH_BINARY)
    # 白色二值化文字
    img_bi_inv = bitwise_not(img_bi)
    img_bi_pil = fromarray(cvtColor(img_bi, COLOR_BGR2RGB))

    gh, gw = gray_limg.shape[:2]
    limg_pil = fromarray(cvtColor(gray_limg, COLOR_BGR2RGB))

    # 创建一个新的大图像来容纳六个子图像
    combined_img = Image.new('RGBA', (nw * 2, nh * 3 + 2 * gh), rgba_white)
    proj_img = Image.new('RGBA', (nw * 2, nh * 2), rgba_white)

    combined_img.paste(img_pil, (0, 0))
    proj_img.paste(img_pil, (0, nh))
    proj_img.paste(img_bi_pil, (nw, 0))

    tess_zdata4 = [x for x in tess_zdata if x[0] == 4]
    tess_zdata5 = [x for x in tess_zdata if x[0] == 5]
    tes_zdata4 = [x for x in tes_zdata if x[0] == 4]
    tes_zdata5 = [x for x in tes_zdata if x[0] == 5]

    cp_bubble_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}.jpg'
    cp_br_tess_word_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_tess_word.jpg'
    cp_br_tes_word_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_tes_word.jpg'
    cp_br_tess_line_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_tess_line.jpg'
    cp_br_tes_line_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_tes_line.jpg'
    cp_br_vision_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_vision.jpg'
    cp_br_vis_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_vis.jpg'
    cp_br_winocr_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_winocr.jpg'
    cp_br_win_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_win.jpg'
    cp_br_textword_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_textword.jpg'
    cp_br_textline_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_textline.jpg'
    combined_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_combined.jpg'
    write_pic(cp_bubble_jpg, img_np)

    # ================vision坐标绘制================
    if rec_results_vision:
        br_vision_img_pil = img_pil.convert("RGBA")
        br_vision_overlay = Image.new('RGBA', img_pil.size, rgba_zero)
        br_vision_draw = ImageDraw.Draw(br_vision_overlay)
        for r in range(len(rec_results_vision)):
            rec_result = rec_results_vision[r]
            vx, vy, vw, vh, vconf, text_seg = rec_result
            br_vision_draw.rectangle([(vx, nh - vy - vh), (vx + vw, nh - vy)], outline=trans_olive, width=1)
        br_vision_img_pil.paste(br_vision_overlay, mask=br_vision_overlay)
        # write_pic(cp_br_vision_jpg, br_vision_img_pil)
        combined_img.paste(br_vision_img_pil, (nw, 0))

    # ================winocr坐标绘制================
    if rec_results_winocr:
        br_winocr_img_pil = img_pil.convert("RGBA")
        br_winocr_overlay = Image.new('RGBA', img_pil.size, rgba_zero)
        br_winocr_draw = ImageDraw.Draw(br_winocr_overlay)
        for r in range(len(rec_results_winocr)):
            rec_result = rec_results_winocr[r]
            wx, wy, ww, wh, l, w, word_text = rec_result
            logger.debug(f'winocr_br: {wx}, {wy}, {ww}, {wh}')
            br_winocr_draw.rectangle([(wx, wy), (wx + ww, wy + wh)], outline=trans_olive, width=1)
        br_winocr_img_pil.paste(br_winocr_overlay, mask=br_winocr_overlay)
        # write_pic(cp_br_winocr_jpg, br_winocr_img_pil)
        combined_img.paste(br_winocr_img_pil, (nw, 0))

    # ================tesseract坐标绘制================
    br_tess_word_img_pil = img_pil.convert("RGBA")
    br_tess_word_overlay = Image.new('RGBA', img_pil.size, rgba_zero)
    br_tess_word_draw = ImageDraw.Draw(br_tess_word_overlay)
    for l in range(len(tess_zdata5)):
        word_data = tess_zdata5[l]
        level, page_num, block_num, par_num, line_num, word_num, tleft, ttop, twidth, theight, tconf, text = word_data
        br_tess_word_draw.rectangle([(tleft, ttop), (tleft + twidth, ttop + theight)], outline=trans_red, width=1)
    br_tess_word_img_pil.paste(br_tess_word_overlay, mask=br_tess_word_overlay)
    # write_pic(cp_br_tess_word_jpg, br_tess_word_img_pil)

    br_tess_line_img_pil = img_pil.convert("RGBA")
    br_tess_line_overlay = Image.new('RGBA', img_pil.size, rgba_zero)
    br_tess_line_draw = ImageDraw.Draw(br_tess_line_overlay)
    for l in range(len(tess_zdata4)):
        word_data = tess_zdata4[l]
        level, page_num, block_num, par_num, line_num, word_num, tleft, ttop, twidth, theight, tconf, text = word_data
        br_tess_line_draw.rectangle([(tleft, ttop), (tleft + twidth, ttop + theight)], outline=trans_red, width=1)
    br_tess_line_img_pil.paste(br_tess_line_overlay, mask=br_tess_line_overlay)
    # write_pic(cp_br_tess_line_jpg, br_tess_line_img_pil)

    combined_img.paste(br_tess_word_img_pil, (0, nh))
    combined_img.paste(br_tess_line_img_pil, (nw, nh))

    # ================cv2坐标绘制================
    rec_textblocks = get_textblocks_full(img_np, media_type)
    if rec_textblocks:
        rec_textblock = rec_textblocks[0]
        textlines = rec_textblock.textlines
        textwords = list(chain(*[x.textwords for x in textlines]))

        br_textword_img_pil = img_pil.convert("RGBA")
        br_textword_overlay = Image.new('RGBA', img_pil.size, rgba_zero)
        br_textword_draw = ImageDraw.Draw(br_textword_overlay)
        for t in range(len(textwords)):
            textword = textwords[t]
            br_textword_draw.rectangle([textword.left_top, textword.right_bottom], outline=trans_purple, width=1)
        br_textword_img_pil.paste(br_textword_overlay, mask=br_textword_overlay)
        # write_pic(cp_br_textword_jpg, br_textword_img_pil)

        br_textline_img_pil = img_pil.convert("RGBA")
        br_textline_overlay = Image.new('RGBA', img_pil.size, rgba_zero)
        br_textline_draw = ImageDraw.Draw(br_textline_overlay)
        for t in range(len(textlines)):
            textline = textlines[t]
            br_textline_draw.rectangle([textline.left_top, textline.right_bottom], outline=trans_purple, width=1)
        br_textline_img_pil.paste(br_textline_overlay, mask=br_textline_overlay)
        # write_pic(cp_br_textline_jpg, br_textline_img_pil)

        textblock_bubbles = get_textblock_bubbles(img_np, rec_textblocks)
        cp_textblock_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_TextBlock[{len(textlines)}].jpg'
        write_pic(cp_textblock_jpg, textblock_bubbles)

        combined_img.paste(br_textword_img_pil, (0, nh * 2))
        combined_img.paste(br_textline_img_pil, (nw, nh * 2))

    # ================长图tesseract坐标绘制================
    br_tess_word_limg_pil = limg_pil.convert("RGBA")
    br_tess_word_overlay = Image.new('RGBA', limg_pil.size, rgba_zero)
    br_tess_word_draw = ImageDraw.Draw(br_tess_word_overlay)
    for l in range(len(tes_zdata5)):
        word_data = tes_zdata5[l]
        level, page_num, block_num, par_num, line_num, word_num, tleft, ttop, twidth, theight, tconf, text = word_data
        br_tess_word_draw.rectangle([(tleft, ttop), (tleft + twidth, ttop + theight)], outline=trans_red, width=1)
    br_tess_word_limg_pil.paste(br_tess_word_overlay, mask=br_tess_word_overlay)
    # write_pic(cp_br_tess_word_jpg, br_tess_word_limg_pil)

    br_tess_line_limg_pil = limg_pil.convert("RGBA")
    br_tess_line_overlay = Image.new('RGBA', limg_pil.size, rgba_zero)
    br_tess_line_draw = ImageDraw.Draw(br_tess_line_overlay)
    for l in range(len(tes_zdata4)):
        word_data = tes_zdata4[l]
        level, page_num, block_num, par_num, line_num, word_num, tleft, ttop, twidth, theight, tconf, text = word_data
        br_tess_line_draw.rectangle([(tleft, ttop), (tleft + twidth, ttop + theight)], outline=trans_red, width=1)
    br_tess_line_limg_pil.paste(br_tess_line_overlay, mask=br_tess_line_overlay)
    # write_pic(cp_br_tess_line_jpg, br_tess_line_limg_pil)

    combined_img.paste(br_tess_word_limg_pil, (0, 3 * nh))
    combined_img.paste(br_tess_line_limg_pil, (nw, 3 * nh))

    # ================长图vision坐标绘制================
    if rec_results_vis:
        br_vis_limg_pil = limg_pil.convert("RGBA")
        br_vis_overlay = Image.new('RGBA', limg_pil.size, rgba_zero)
        br_vis_draw = ImageDraw.Draw(br_vis_overlay)
        for r in range(len(rec_results_vis)):
            rec_result = rec_results_vis[r]
            vx, vy, vw, vh, vconf, text_seg = rec_result
            br_vis_draw.rectangle([(vx, gh - vy - vh), (vx + vw, gh - vy)], outline=trans_olive, width=1)
        br_vis_limg_pil.paste(br_vis_overlay, mask=br_vis_overlay)
        # write_pic(cp_br_vis_jpg, br_vis_img_pil)
        combined_img.paste(br_vis_limg_pil, (0, 3 * nh + gh))

    # ================长图winocr坐标绘制================
    if rec_results_win:
        br_win_limg_pil = limg_pil.convert("RGBA")
        br_win_overlay = Image.new('RGBA', limg_pil.size, rgba_zero)
        br_win_draw = ImageDraw.Draw(br_win_overlay)
        for r in range(len(rec_results_win)):
            rec_result = rec_results_win[r]
            wx, wy, ww, wh, l, w, word_text = rec_result
            logger.debug(f'win_br: {wx}, {wy}, {ww}, {wh}')
            br_win_draw.rectangle([(wx, wy), (wx + ww, wy + wh)], outline=trans_olive, width=1)
        br_win_limg_pil.paste(br_win_overlay, mask=br_win_overlay)
        # write_pic(cp_br_win_jpg, br_win_img_pil)
        combined_img.paste(br_win_limg_pil, (0, 3 * nh + gh))

    write_pic(combined_jpg, combined_img)
    return proj_img


# ================图形化参数================
@logger.catch
def cal_hori_paras(hori_proj_bi_inv, upper, lower):
    row_segment = hori_proj_bi_inv[int(upper):int(lower), :]
    # 每行宽度
    widths = np.sum(row_segment == 255, axis=1)
    # 计算方差
    vari = np.var(widths)
    area = sum(widths)
    avg_width = np.mean(widths)
    logger.debug(f'{area=}, {len(widths)=}, {avg_width=:.2f}, {vari=:.2f}')
    return area, widths, avg_width, vari


@logger.catch
def better_line_proj_cnts(line_proj_cnts, hori_proj_bi_inv):
    # ================处理只有点的文本行================
    if len(line_proj_cnts) >= 2:
        # ================点在最后一行则不处理================
        br_hs = [x.br_h for x in line_proj_cnts[:-1]]
        min_br_h = min(br_hs)
        if min_br_h <= dot_h:
            # ================找出所有点型文本行================
            dot_inds = []
            ind = 0
            while ind < len(line_proj_cnts) - 1:
                line_proj_cnt = line_proj_cnts[ind]
                if line_proj_cnt.br_h <= dot_h:
                    dot_inds.append(ind)
                ind += 1

            # ================找出所有点型文本行================
            new_line_proj_cnts = []
            pin = 0
            while pin < len(line_proj_cnts):
                line_proj_cnt = line_proj_cnts[pin]
                if pin in dot_inds:
                    line_proj_cnt = line_proj_cnts[pin]
                    line_proj_cnt_next = line_proj_cnts[pin + 1]
                    # 涂白最右边1像素宽
                    hori_proj_dot_roi = hori_proj_bi_inv.copy()
                    hori_proj_dot_roi[:line_proj_cnt.br_y - 1, :] = 0
                    hori_proj_dot_roi[line_proj_cnt.br_v:line_proj_cnt_next.br_y, -1:] = 255
                    hori_proj_dot_roi[line_proj_cnt_next.br_v + 1:, :] = 0
                    proj_dot_contours, proj_dot_hier = findContours(hori_proj_dot_roi, RETR_EXTERNAL,
                                                                    CHAIN_APPROX_SIMPLE)
                    proj_dot_cnts = [Contr(x) for x in proj_dot_contours]
                    proj_dot_cnts.sort(key=lambda x: x.br_y)
                    proj_dot_cnts_ok = [x for x in proj_dot_cnts if x.area >= min_cnt_area]
                    if proj_dot_cnts_ok:
                        alt_cnt = proj_dot_cnts_ok[0]
                    else:  # if proj_dot_cnts
                        alt_cnt = proj_dot_cnts[0]
                    new_line_proj_cnts.append(alt_cnt)
                elif pin - 1 in dot_inds:
                    pass
                else:
                    new_line_proj_cnts.append(line_proj_cnt)
                pin += 1
            line_proj_cnts = new_line_proj_cnts
    return line_proj_cnts


@logger.catch
def get_sub_line_proj_cnts(proj_cnt_raw, hori_proj_bi_inv, refer_h, proj_mask):
    sub_line_proj_cnts = []
    logger.debug(f'{proj_cnt_raw=}, {proj_cnt_raw.avg_w=:.2f}')
    upper = proj_cnt_raw.br_y
    lower = proj_cnt_raw.br_v
    upper_raw = deepcopy(upper)
    lower_raw = deepcopy(lower)
    area, widths, avg_width, vari = cal_hori_paras(hori_proj_bi_inv, upper, lower)
    area0, widths0, avg_width0, vari0 = cal_hori_paras(hori_proj_bi_inv, upper, upper + refer_h)
    area1, widths1, avg_width1, vari1 = cal_hori_paras(hori_proj_bi_inv, lower - refer_h, lower)
    # ================估测所需截断宽度================
    if proj_cnt_raw.br_w <= 60:
        ratio = 0.12
    elif proj_cnt_raw.br_w <= 120:
        ratio = 0.1
    elif proj_cnt_raw.br_w <= 180:
        ratio = 0.08
    elif proj_cnt_raw.br_w <= 240:
        ratio = 0.06
    else:
        ratio = 0.04
    line_edge = proj_cnt_raw.br_w * ratio
    line_edge = clamp(line_edge, line_edge_min, line_edge_max)
    line_edge = int(line_edge)
    # ================仅对对话框优化================
    if refer_h <= 20:
        if avg_width0 <= line_edge + 2:
            logger.error(f'{avg_width0=:.2f}, {line_edge=}, 首行过短')
            upper += refer_h
        if avg_width1 <= line_edge + 2:
            logger.error(f'{avg_width1=:.2f}, {line_edge=}, 尾行过短')
            lower -= refer_h
    # ================涂黑白色轮廓中最右边line_edge像素宽================
    hori_proj_roi = hori_proj_bi_inv.copy()
    hori_proj_roi[:proj_cnt_raw.br_y - 1, :] = 0
    hori_proj_roi[int(upper):int(lower), -line_edge:] = 0
    hori_proj_roi[proj_cnt_raw.br_v + 1:, :] = 0

    proj_contours_fit, fit_hier = findContours(hori_proj_roi, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
    proj_cnts_fit = [Contr(x) for x in proj_contours_fit]
    proj_cnts_fit.sort(key=lambda x: x.br_y)
    proj_cnts_fit_ok = [x for x in proj_cnts_fit if x.area >= min_cnt_area]
    logger.warning(f'[{line_edge}]{len(proj_cnts_fit)=}, {len(proj_cnts_fit_ok)=}')

    if len(proj_cnts_fit) <= 2 and len(proj_cnts_fit_ok) < 2:
        # 切分不成功，可能不是粘连
        sub_line_proj_cnts.append(proj_cnt_raw)
    else:
        # 多个文本行粘连切分成功
        split_lines = []
        for f in range(len(proj_cnts_fit_ok) - 1):
            cnt_fit = proj_cnts_fit_ok[f]
            cnt_fit_next = proj_cnts_fit_ok[f + 1]
            # 切割线应该在白色区域的最小值处
            start_y = cnt_fit.br_v
            end_y = cnt_fit_next.br_y
            row_segment = hori_proj_bi_inv[start_y:end_y, :]
            white_px_count = np.sum(row_segment == 255, axis=1)
            if len(white_px_count) > 0:
                split_line = argmin(white_px_count)
            else:
                split_line = 0
            split_line += start_y
            split_lines.append(split_line)
        all_lines = [proj_cnts_fit_ok[0].br_y - 1] + split_lines + [proj_cnts_fit_ok[-1].br_v + 1]
        # ================建立基准行================
        for a in range(len(all_lines) - 1):
            start_line = all_lines[a]
            end_line = all_lines[a + 1]
            proj_roi_line = hori_proj_bi_inv.copy()
            if a == 0:
                # 有粘连的行的第一个
                proj_roi_line[:start_line - 1, :] = 0
            else:
                proj_roi_line[:start_line, :] = 0
            if a == len(all_lines) - 1:
                # 有粘连的行的最后一个
                proj_roi_line[end_line + 1:, :] = 0
            else:
                proj_roi_line[end_line:, :] = 0

            contours_line, line_hier = findContours(proj_roi_line, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
            cnts_line = [Contr(x) for x in contours_line]
            cnts_line.sort(key=lambda x: x.br_y)
            cnts_line_ok = [x for x in cnts_line if x.area >= min_cnt_area]
            if cnts_line_ok:
                line_proj_cnt = cnts_line_ok[0]
            else:  # if cnts_line:
                line_proj_cnt = cnts_line[0]
            sub_line_proj_cnts.append(line_proj_cnt)
    # 示意图
    proj_mask[int(upper):int(lower), -line_edge:] = 127
    return sub_line_proj_cnts, proj_mask


@logger.catch
def get_line_proj_cnts(img_np, hori_proj_bi, ver_proj_bi, direction):
    # ================图像参数================
    nh, nw = img_np.shape[:2]  # 获取原始图像的尺寸
    white_bg = ones((nh, nw, 3), dtype=uint8) * 255
    black_bg = zeros((nh, nw), dtype=uint8)

    img_pil = fromarray(cvtColor(img_np, COLOR_BGR2RGB))
    img_gray = cvtColor(img_np, COLOR_BGR2GRAY)
    ret, img_bi = threshold(img_gray, bi_thres, 255, THRESH_BINARY)
    # 白色二值化文字
    img_bi_inv = bitwise_not(img_bi)
    img_bi_pil = fromarray(cvtColor(img_bi, COLOR_BGR2RGB))
    # 将img_np中特别浅的部分变为白色
    img_opt = img_gray.copy()
    img_opt[img_opt >= white_thres] = 255

    hori_proj_bi_np = array(hori_proj_bi)
    hori_proj_bi_inv = bitwise_not(hori_proj_bi_np)

    ver_proj_bi_np = array(ver_proj_bi)
    ver_proj_bi_inv = bitwise_not(ver_proj_bi_np)

    # ================示意图================
    if direction == 'Horizontal':
        proj_mask = hori_proj_bi_np.copy()
    else:
        proj_mask = ver_proj_bi_np.copy()

    # ================笔画和原始投影轮廓================
    stroke_contours, stroke_hier = findContours(img_bi_inv, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)

    stroke_cnts = [Contr(x) for x in stroke_contours]
    stroke_cnts = [x for x in stroke_cnts if x.area >= 1]
    stroke_cnts.sort(key=lambda x: (x.br_x, x.br_y))

    refer_w = get_refer_dimension(stroke_cnts, 'width')
    refer_h = get_refer_dimension(stroke_cnts, 'height')
    # ================多行文本的最低高度================
    line_multi_h_min = int(max(line_multi_h, 2 * refer_h - 1))
    line_multi_w_min = int(max(line_multi_w, 2 * refer_w - 1))

    if direction == 'Horizontal':
        proj_contours_raw, raw_hier = findContours(hori_proj_bi_inv, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
        proj_cnts_raw = [Contr(x) for x in proj_contours_raw]
        proj_cnts_raw.sort(key=lambda x: x.br_y)
        line_multi_roi_min = line_multi_h_min
    else:
        proj_contours_raw, raw_hier = findContours(ver_proj_bi_inv, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
        proj_cnts_raw = [Contr(x) for x in proj_contours_raw]
        proj_cnts_raw.sort(key=lambda x: -x.br_x)
        line_multi_roi_min = line_multi_w_min

    proj_cnts_raw_ok = [x for x in proj_cnts_raw if x.area >= min_cnt_area]
    br_hs = [x.br_h for x in proj_cnts_raw_ok]
    br_ws = [x.br_w for x in proj_cnts_raw_ok]

    if direction == 'Horizontal':
        br_rois = br_hs
    else:
        br_rois = br_ws

    line_proj_cnts = []
    if media_type in ['Manga']:
        line_proj_cnts = deepcopy(proj_cnts_raw)
    elif len(proj_cnts_raw) == 1 and not proj_cnts_raw_ok:
        # ================短小的一行，只有一个符号（省略号之类）================
        line_proj_cnts = deepcopy(proj_cnts_raw)
    elif br_rois and max(br_rois) <= line_multi_roi_min:
        # ================行高均符合规则，不存在多行粘连================
        line_proj_cnts = deepcopy(proj_cnts_raw)
    else:
        # ================可能存在多行粘连================
        for p in range(len(proj_cnts_raw)):
            proj_cnt_raw = proj_cnts_raw[p]
            if direction == 'Horizontal':
                br_roi = proj_cnt_raw.br_h
            else:
                br_roi = proj_cnt_raw.br_w
            if br_roi <= line_multi_roi_min:
                # ================单行================
                line_proj_cnts.append(proj_cnt_raw)
            elif direction == 'Horizontal':
                # ================多行================
                sub_line_proj_cnts, proj_mask = get_sub_line_proj_cnts(proj_cnt_raw, hori_proj_bi_inv, refer_h,
                                                                       proj_mask)
                line_proj_cnts.extend(sub_line_proj_cnts)

    if direction == 'Horizontal':
        line_proj_cnts = better_line_proj_cnts(line_proj_cnts, hori_proj_bi_inv)
    # ================添加文本行矩形================
    for line_proj_cnt in line_proj_cnts:
        line_proj_cnt.add_line_polygon(nw, nh)

    # ================进行高度筛选================
    if len(line_proj_cnts) >= 2:
        if direction == 'Horizontal':
            line_proj_cnts = [x for x in line_proj_cnts if x.br_h >= textline_h_min]
        # else:
        #     line_proj_cnts = [x for x in line_proj_cnts if x.br_w >= textline_w_min]

    if direction == 'Horizontal':
        for l in range(len(stroke_cnts)):
            # ================每个笔画轮廓================
            stroke_cnt = stroke_cnts[l]
            # 绘制stroke_cnt的遮罩
            stroke_mask = drawContours(black_bg.copy(), [stroke_cnt.contour], -1, (255), FILLED)
            # write_pic(stroke_jpg, stroke_mask)
            overlap_ratios = []
            for t in range(len(line_proj_cnts)):
                # ================每个基准行================
                line_proj_cnt = line_proj_cnts[t]
                if stroke_cnt.polygon.intersects(line_proj_cnt.line_polygon):
                    # 获取交集区域
                    overlap = stroke_cnt.polygon.intersection(line_proj_cnt.line_polygon)
                    line_mask = rectangle(black_bg.copy(), line_proj_cnt.line_br_xy, line_proj_cnt.line_br_uv, (255),
                                          FILLED)
                    # 创建交集遮罩
                    overlap_mask = bitwise_and(stroke_mask, line_mask)
                    # write_pic(overlap_jpg, overlap_mask)
                    # 计算两个区域的非白色像素数量
                    overlap_non_white_area = np.sum((img_opt[overlap_mask == 255] != 255).astype(int))
                    stroke_non_white_area = np.sum((img_opt[stroke_mask == 255] != 255).astype(int))
                    if stroke_non_white_area == 0:
                        overlap_ratio = 0
                    else:
                        overlap_ratio = overlap_non_white_area / stroke_non_white_area
                else:
                    overlap_ratio = 0
                overlap_ratios.append(overlap_ratio)

            ind = overlap_ratios.index(max(overlap_ratios))
            # logger.debug(f'[{ind}]{max(overlap_ratios)=}')
            value_2nd, index_2nd = find_nth_largest(overlap_ratios, 2)
            line_proj_cnt_1st = line_proj_cnts[ind]
            # 笔画相对于行高的比例
            h_ratio = stroke_cnt.br_h / line_proj_cnt_1st.br_h

            if max(overlap_ratios) >= 0.75 and h_ratio <= 1.25:
                belong = 'A'
                # 绝大部分都在一个文本行内
                line_proj_cnt_1st.add_stroke(stroke_cnt)
            elif len(overlap_ratios) >= 2 and value_2nd >= 0.1:
                logger.debug(f'{len(line_proj_cnts)=}, {len(overlap_ratios)=}')
                belong = 'B'
                # 找到与笔画重叠的文本行
                overlapping_lines = [
                    line_proj_cnts[ind],
                    line_proj_cnts[index_2nd],
                ]
                for line_proj_cnt in overlapping_lines:
                    # 创建新的交集遮罩
                    line_mask = rectangle(black_bg.copy(), line_proj_cnt.line_br_xy, line_proj_cnt.line_br_uv, (255),
                                          FILLED)
                    intersected_mask = bitwise_and(stroke_mask, line_mask)
                    # 创建新的笔画轮廓
                    sub_stroke_contours, sub_stroke_hier = findContours(intersected_mask, RETR_EXTERNAL,
                                                                        CHAIN_APPROX_SIMPLE)
                    sub_stroke_cnts = [Contr(x) for x in sub_stroke_contours]
                    sub_stroke_cnts = [x for x in sub_stroke_cnts if x.area >= 1]
                    for stroke_cnt in sub_stroke_cnts:
                        line_proj_cnt.add_stroke(stroke_cnt)
            elif max(overlap_ratios) >= 0.01:
                belong = 'C'
                line_proj_cnt_1st.add_stroke(stroke_cnt)
            else:
                belong = 'D'
                # 无归属且无重叠面积则归于最近的line_proj_cnt
                nearest_line_proj_cnt = min(line_proj_cnts, key=lambda cnt: abs(cnt.br_y - stroke_cnt.br_y))
                nearest_line_proj_cnt.add_stroke(stroke_cnt)

            imes = f'[{l}]{belong}[{len(line_proj_cnts)}行之{ind + 1}]{max(overlap_ratios)=:.2f}, {h_ratio=:.2f}'
            if len(overlap_ratios) >= 2:
                imes += f', {value_2nd=}, {index_2nd=}'
    if not line_proj_cnts:
        if do_dev_pic:
            hori_proj_bi_png = debug_dir / f'{media_type}_hori_proj_bi.png'
            img_np_png = debug_dir / f'{media_type}_img_np.png'
            write_pic(hori_proj_bi_png, hori_proj_bi)
            write_pic(img_np_png, img_np)
    if len(line_proj_cnts) == 0:
        print()
    return line_proj_cnts, proj_mask


@logger.catch
def draw_line_proj_cnts(img_np, line_proj_cnts):
    ih, iw = img_np.shape[0:2]
    black_bg = zeros((ih, iw), dtype=uint8)

    img_pil = fromarray(cvtColor(img_np, COLOR_BGR2RGB))
    img_gray = cvtColor(img_np, COLOR_BGR2GRAY)
    ret, img_bi = threshold(img_gray, bi_thres, 255, THRESH_BINARY)
    # 白色二值化文字
    img_bi_inv = bitwise_not(img_bi)
    img_bi_pil = fromarray(cvtColor(img_bi, COLOR_BGR2RGB))
    # 将img_np中特别浅的部分变为白色
    img_opt = img_gray.copy()
    img_opt[img_opt >= white_thres] = 255

    # ================绘制示意图的彩色文本行边框================
    img_w_lines = deepcopy(img_pil)
    lines_draw = ImageDraw.Draw(img_w_lines, 'RGBA')
    for p in range(len(line_proj_cnts)):
        cnt = line_proj_cnts[p]
        color = colormap_tab20(p % 20)[:3]
        color_rgb = tuple(int(c * 255) for c in color)
        color_rgba = color_rgb + (200,)
        if text_direction == 'Horizontal':
            start_y = cnt.br_y
            end_y = cnt.br_v
            row_segment = img_bi[start_y:end_y + 1, :]
            if np.any(row_segment == 0):
                left = np.min(where(row_segment == 0)[1])
                right = np.max(where(row_segment == 0)[1])
                lines_draw.rectangle([(left, start_y), (right, end_y)], outline=color_rgba)
        else:
            start_x = cnt.br_x
            end_x = cnt.br_u
            row_segment = img_bi[:, start_x:end_x + 1]
            if np.any(row_segment == 0):
                up = np.min(where(row_segment == 0)[0])
                down = np.max(where(row_segment == 0)[0])
                lines_draw.rectangle([(start_x, up), (end_x, down)], outline=color_rgba)

    # ================获取每行文本对应的笔画轮廓================
    line_binarys = []
    for t in range(len(line_proj_cnts)):
        line_proj_cnt = line_proj_cnts[t]
        contours = [x.contour for x in line_proj_cnt.strokes]
        # line_binary = drawContours(black_bg.copy(), contours, -1, 255, FILLED)
        # 初始化一个黑色背景
        combined_mask = black_bg.copy()
        # 对每个轮廓单独绘制
        for contour in contours:
            mask = black_bg.copy()  # 每个轮廓都用一个新的遮罩
            drawContours(mask, [contour], -1, 255, FILLED)  # 绘制轮廓
            combined_mask = bitwise_or(combined_mask, mask)  # 将当前轮廓的遮罩合并到总遮罩上
        line_binary = combined_mask
        # 最终的遮罩将所有轮廓合并，即使有重叠，重叠部分也是白色
        line_binarys.append(line_binary)

    # ================获取未归属区域================
    mask_belong = black_bg.copy()
    # 叠加所有文本行的掩模
    for line_binary in line_binarys:
        mask_belong = maximum(mask_belong, line_binary)
    mask_unbelong = bitwise_not(mask_belong)

    # ================计算周边区域归属并更新未归属区域================
    line_binary_borders = []
    for l in range(len(line_binarys)):
        line_binary = line_binarys[l]
        line_binary_d3 = dilate(line_binary, kernel3, iterations=1)
        line_binary_border = bitwise_and(line_binary_d3, mask_unbelong)
        line_binary_borders.append(line_binary_border)
        # 更新 mask_unbelong 以去掉已归属的边界区域
        mask_unbelong = bitwise_and(mask_unbelong, bitwise_not(line_binary_border))

    # ================将未归属区域按基准矩形分配到每个文本行================
    real_masks = []
    textline_imgs = []
    for t in range(len(line_proj_cnts)):
        line_proj_cnt = line_proj_cnts[t]
        line_binary = line_binarys[t]

        line_binary_border = line_binary_borders[t]
        # 创建每一行文本的基准矩形掩模
        line_belong_mask = zeros_like(img_bi, dtype=uint8)
        x, y, w, h = line_proj_cnt.line_br
        line_belong_mask[y:y + h, x:x + w] = 255
        # 将基准矩形掩模与mask_unbelong重叠的部分视为每一行的真正区域
        line_belong_mask = minimum(line_belong_mask, mask_unbelong)
        # 每一行的真正掩模为 line_binary 加上 line_binary_border 加上 line_belong_mask
        real_mask = maximum(maximum(line_binary, line_binary_border), line_belong_mask)
        real_masks.append(real_mask)
        # ================获取每一行真正掩模对应的非白色像素区域================
        # 创建一个白色背景图像
        textline_img = ones_like(img_opt) * 255
        # 从原始灰度图像提取非白色像素，并放置在白色背景图像上
        textline_img[real_mask == 255] = img_opt[real_mask == 255]
        textline_imgs.append(textline_img)
        if do_dev_pic:
            line_binary_png = debug_dir / f'{media_type}_line_binary.png'
            textline_img_png = debug_dir / f'{media_type}_textline_img_{t + 1}.png'
            write_pic(line_binary_png, line_binary)
            write_pic(textline_img_png, textline_img)
    return img_w_lines, line_binarys, textline_imgs


@timer_decorator
@logger.catch
def ocr1pic(img_file, frame_data, order_data, ocr_data, all_masks, media_type, media_lang, vert):
    logger.warning(f'{img_file=}')
    pic_results = []
    stem_tup = ('stem', img_file.stem)
    pic_results.append(stem_tup)
    img_ind = img_list.index(img_file)
    page_ind = img_ind

    img_raw = imdecode(fromfile(img_file, dtype=uint8), -1)
    ih, iw = img_raw.shape[0:2]
    # ================矩形画格信息================
    # 从 frame_data 中获取画格信息，默认为整个图像的矩形区域
    frame_grid_strs = frame_data.get(img_file.name, [f'0,0,{iw},{ih}~0,0,{iw},{ih}'])
    bubble_order_strs = order_data.get(img_file.name, [])
    grid_masks, marked_frames = get_grid_masks(img_file, frame_grid_strs)

    # ================获取对应的文字图片================
    mask_pics = [x for x in all_masks if x.stem.startswith(img_file.stem)]
    if mask_pics:
        single_cnts = get_single_cnts(img_raw, mask_pics)
        logger.debug(f'{len(single_cnts)=}')

        single_cnts_grids_ordered = get_ordered_cnts(single_cnts, img_file, grid_masks, bubble_order_strs, media_type)
        ordered_cnts = list(chain(*single_cnts_grids_ordered))

        for c in range(len(ordered_cnts)):
            single_cnt = ordered_cnts[c]
            bubble_ind = c + 1
            color_pattern = single_cnt.color_pattern
            cp_bubble, cp_letter = color_pattern
            if isinstance(cp_letter, list):
                # 文字为双色
                color_letter = ColorDouble(cp_letter)
            else:
                # 文字为单色
                color_letter = Color(cp_letter)
            if isinstance(cp_bubble, list):
                # 气泡为渐变色
                color_bubble = ColorGradient(cp_bubble)
            elif cp_bubble == '':
                # 无框字
                color_bubble = None
            else:
                # 气泡为单色
                color_bubble = Color(cp_bubble)
            bubble_color_str = cp_bubble
            if '-' in cp_bubble:
                bubble_color_str = cp_bubble.split('-')[0]
            letter_color_str = cp_letter
            if '-' in cp_letter:
                letter_color_str = cp_letter.split('-')[0]

            img_np = single_cnt.cropped_img
            img_md5 = generate_md5(img_np)
            # ================图像参数================
            nh, nw = img_np.shape[:2]  # 获取原始图像的尺寸
            white_bg = ones((nh, nw, 3), dtype=uint8) * 255
            black_bg = zeros((nh, nw), dtype=uint8)

            img_pil = fromarray(cvtColor(img_np, COLOR_BGR2RGB))
            img_gray = cvtColor(img_np, COLOR_BGR2GRAY)
            ret, img_bi = threshold(img_gray, bi_thres, 255, THRESH_BINARY)
            # 白色二值化文字
            img_bi_inv = bitwise_not(img_bi)
            img_bi_pil = fromarray(cvtColor(img_bi, COLOR_BGR2RGB))
            # 将img_np中特别浅的部分变为白色
            img_opt = img_gray.copy()
            img_opt[img_opt >= white_thres] = 255

            # ================投影================
            hori_proj_bi, hori_proj_data = get_proj_img(img_bi_pil, 'horizontal', 'white')
            ver_proj_bi, ver_proj_data = get_proj_img(img_bi_pil, 'vertical', 'white')

            # ================绘制分行示意图================
            direction = text_direction
            if nw >= 1.4 * nh:
                direction = 'Horizontal'
            line_proj_cnts, proj_mask = get_line_proj_cnts(img_np, hori_proj_bi, ver_proj_bi, direction)
            proj_png = auto_subdir / f'P{page_ind}_B{bubble_ind}_proj.png'
            write_pic(proj_png, proj_mask)

            img_w_lines, line_binarys, textline_imgs = draw_line_proj_cnts(img_np, line_proj_cnts)

            # ================每行的标记图================
            textlines_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_TextLines[{len(line_proj_cnts)}].jpg'
            write_pic(textlines_jpg, img_w_lines)

            if direction == 'Horizontal':
                # ================制作最终效果图================
                stitched_bi_img = stitch_imgs(line_binarys, 'B')
                gray_limg = stitch_imgs(textline_imgs, 'W')
                # gray_limg = stitch_imgs_normal(textline_imgs)
                # 转为黑白
                gray_limg = cvtColor(gray_limg, COLOR_BGR2GRAY)
                gray_limg = cvtColor(gray_limg, COLOR_GRAY2BGR)
                stitched_bi_png = auto_subdir / f'P{page_ind}_B{bubble_ind}_bi.png'
                stitched_gray_png = auto_subdir / f'P{page_ind}_B{bubble_ind}_gray.png'
                write_pic(stitched_bi_png, stitched_bi_img)
                write_pic(stitched_gray_png, gray_limg)
            else:
                gray_limg = None

            # ================笔画轮廓================
            stroke_contours, stroke_hier = findContours(img_bi_inv, RETR_EXTERNAL, CHAIN_APPROX_SIMPLE)
            stroke_cnts = [Contr(x) for x in stroke_contours]
            stroke_cnts = [x for x in stroke_cnts if x.area >= 1]
            stroke_cnts.sort(key=lambda x: (x.br_x, x.br_y))
            refer_w = get_refer_dimension(stroke_cnts, 'width')
            refer_h = get_refer_dimension(stroke_cnts, 'height')

            # ================气泡颜色参数================
            cp_bubble, cp_letter = color_pattern
            pic_locate = f'{img_file.stem}, {c}, {img_md5}'
            pic_ocr_data = ocr_data.get(pic_locate, {})
            rec_textblocks = get_textblocks_full(img_np, media_type)

            # ================对裁剪后的图像进行文字识别================
            if text_direction == 'Horizontal':
                tess_zdata, pic_ocr_data = get_ocr_data('tesseract', pic_ocr_data, img_np, media_lang, 1)
                tess_zdata4 = [x for x in tess_zdata if x[0] == 4]
                tess_zdata5 = [x for x in tess_zdata if x[0] == 5]

                tes_zdata, pic_ocr_data = get_ocr_data('tes', pic_ocr_data, gray_limg, media_lang, 1)
                tes_zdata4 = [x for x in tes_zdata if x[0] == 4]
                tes_zdata5 = [x for x in tes_zdata if x[0] == 5]

                if custom_ocr_engines and 'Paddle' in custom_ocr_engines:
                    rec_results_paddle, pic_ocr_data = get_ocr_data('paddle', pic_ocr_data, img_np, media_lang, 2)
                if custom_ocr_engines and 'Easy' in custom_ocr_engines:
                    rec_results_easy, pic_ocr_data = get_ocr_data('easy', pic_ocr_data, img_np, media_lang, 2)
                if custom_ocr_engines and 'Baidu' in custom_ocr_engines:
                    rec_results_baidu, pic_ocr_data = get_ocr_data('baidu', pic_ocr_data, img_np, media_lang, 4)
                if custom_ocr_engines and 'BaiduAccu' in custom_ocr_engines:
                    rec_results_baidu_accu, pic_ocr_data = get_ocr_data('baidu_accu', pic_ocr_data, img_np, media_lang,
                                                                        4)
                ocr_data[pic_locate] = pic_ocr_data

                lines_tess = tesseract2text(tess_zdata)
                lines_tes = tesseract2text(tes_zdata)
                tess_text = lf.join(lines_tess)
                tes_text = lf.join(lines_tes)

                if SYSTEM in ['MAC', 'M1']:
                    rec_results_vision, pic_ocr_data = get_ocr_data('vision', pic_ocr_data, img_np, media_lang, 2)
                    lines_vision = rec2text(rec_results_vision)
                    vision_text = lf.join(lines_vision)

                    rec_results_vis, pic_ocr_data = get_ocr_data('vis', pic_ocr_data, gray_limg, media_lang, 2)
                    lines_vis = rec2text(rec_results_vis)
                    vis_text = lf.join(lines_vis)

                    rec_results_winocr = None
                    rec_results_win = None
                else:
                    rec_results_winocr, pic_ocr_data = get_ocr_data('winocr', pic_ocr_data, img_np, media_lang, 1)
                    vision_text = deepcopy(tess_text)

                    rec_results_win, pic_ocr_data = get_ocr_data('win', pic_ocr_data, gray_limg, media_lang, 1)

                    rec_results_vision = None
                    rec_results_vis = None

                # ================其他OCR================
                refer_ocrs = []
                if 'Baidu' in custom_ocr_engines:
                    lines_baidu = [x[-1] for x in rec_results_baidu]
                    baidu_text = lf.join(lines_baidu)
                    baidu_fmt_text = better_text(baidu_text, 'baidu')
                    refer_ocrs.append(baidu_fmt_text)
                if 'BaiduAccu' in custom_ocr_engines:
                    lines_baidu_accu = [x[-1] for x in rec_results_baidu_accu]
                    baidu_accu_text = lf.join(lines_baidu_accu)
                    baidu_accu_fmt_text = better_text(baidu_accu_text, 'baidu_accu')
                    refer_ocrs.append(baidu_accu_fmt_text)
                if 'Paddle' in custom_ocr_engines:
                    lines_paddle = [x[-1] for x in rec_results_paddle]
                    paddle_text = lf.join(lines_paddle)
                    paddle_fmt_text = better_text(paddle_text, 'paddle')
                    refer_ocrs.append(paddle_fmt_text)
                if 'Easy' in custom_ocr_engines:
                    lines_easy = [x[-1] for x in rec_results_easy]
                    easy_text = lf.join(lines_easy)
                    easy_fmt_text = better_text(easy_text, 'easy')
                    refer_ocrs.append(easy_fmt_text)

                tess_fmt_text = better_text(tess_text, 'tesseract')
                tes_fmt_text = better_text(tes_text, 'tesseract')
                refer_ocrs.append(tess_fmt_text)
                if SYSTEM in ['MAC', 'M1']:
                    vision_fmt_text = better_text(vision_text, 'vision')
                    vision_fmt_text = fix_w_tess(vision_fmt_text, tess_fmt_text)
                    vision_fmt_text = better_punct(vision_fmt_text, refer_ocrs)

                    vis_fmt_text = better_text(vis_text, 'vis')
                    vis_fmt_text = fix_w_tess(vis_fmt_text, tess_fmt_text)
                    vis_fmt_text = better_punct(vis_fmt_text, refer_ocrs)
                else:
                    vision_fmt_text = deepcopy(tess_fmt_text)
                    vis_fmt_text = deepcopy(tes_fmt_text)
            else:
                tess_zdata = None
                tes_zdata = None
                rec_results_vision = None
                rec_results_vis = None
                rec_results_winocr = None
                rec_results_win = None
                vis_fmt_text = None
                if 'tess' in pic_ocr_data:
                    line_text = pic_ocr_data['tess']
                else:
                    textlines = []
                    for t in range(len(textline_imgs)):
                        textline_img = textline_imgs[t]
                        tess_zdata = ocr_by_tesseract(textline_img, media_lang)  # , vert=True
                        tess_zdata4 = [x for x in tess_zdata if x[0] == 4]
                        tess_zdata5 = [x for x in tess_zdata if x[0] == 5]
                        words = [x[-1] for x in tess_zdata5]
                        textline = ''.join(words)
                        textlines.append(textline)
                    line_text = lf.join(textlines)
                    pic_ocr_data['tess'] = line_text
                if python_ver in ['3.11.8']:
                    if 'mangaocr' in pic_ocr_data:
                        line_text = pic_ocr_data['mangaocr']
                    else:
                        textlines = []
                        for t in range(len(textline_imgs)):
                            textline_img = textline_imgs[t]
                            textline_pil = fromarray(textline_img)
                            res = mocr(textline_pil)
                            textline = res
                            textlines.append(textline)
                        line_text = lf.join(textlines)
                        pic_ocr_data['mangaocr'] = line_text
                vision_fmt_text = line_text
                ocr_data[pic_locate] = pic_ocr_data

            # ================绘制投影示意图================
            proj_jpg = auto_subdir / f'P{page_ind}_B{bubble_ind}_proj.jpg'
            img_tups = (img_np, gray_limg)
            ocr_tups = (tess_zdata, tes_zdata, rec_results_vision, rec_results_vis, rec_results_winocr, rec_results_win)
            if text_direction == 'Horizontal':
                proj_img = draw_para(img_tups, ocr_tups, page_ind, bubble_ind)
                proj_img.paste(hori_proj_bi, (0, 0))
                proj_img.paste(ver_proj_bi, (nw, nh))
                write_pic(proj_jpg, proj_img)

            # ================获取气泡内文字的基本信息================
            raw_min_x, raw_min_y, raw_max_x, raw_max_y, min_x, min_y, max_x, max_y, center_pt = single_cnt.letter_coors
            br_w = raw_max_x - raw_min_x
            br_h = raw_max_y - raw_min_y
            src_font_size = int(0.8 * global_font_size)
            dst_font_size = global_font_size

            if text_direction == 'Horizontal' and len(stroke_cnts) >= 10 and refer_h >= dot_line_h:
                src_font_size = int(refer_h)
                dst_font_size = get_dst_font_size(src_font_size, bubble_color_str, letter_color_str)

            br_area_real = (single_cnt.br_w - 2) * (single_cnt.br_h - 2)
            fulfill_ratio = single_cnt.area / br_area_real
            bubble_shape = '未知'
            if fulfill_ratio >= 0.95:
                bubble_shape = '矩形'
            font_name = global_font_name
            if bubble_shape == '矩形' and global_rec_font_name:
                font_name = global_rec_font_name

            # 如果cp_bubble是列表，表示气泡是渐变色的
            if isinstance(cp_bubble, list):
                b0 = cp_bubble[0].split('-')[0]
                b1 = cp_bubble[1].split('-')[0]
                bubble_color_str = f'{b0}-{b1}'
            elif cp_bubble == '':
                bubble_color_str = ''
            else:
                bubble_color_str = cp_bubble.split('-')[0]
            # 如果cp_letter是列表，表示文字是双色的
            if isinstance(cp_letter, list):
                b0 = cp_letter[0].split('-')[0]
                b1 = cp_letter[1].split('-')[0]
                letter_color_str = f'{b0}-{b1}'
            else:
                letter_color_str = cp_letter.split('-')[0]
            bubble_meta_list = [
                f'{raw_min_x},{raw_min_y},{br_w},{br_h}',
                f'S{src_font_size}',
                f'D{dst_font_size}',
                bubble_shape,
                text_direction,
                text_alignment,
                bubble_color_str,  # 白色气泡
                letter_color_str,  # 黑色文字
                font_name,
            ]
            bubble_meta_list = [x for x in bubble_meta_list if x != '']
            bubble_meta_str = '~'.join(bubble_meta_list)

            # ================将图片添加到docx文件中================
            pic_tup = ('picture', img_np, pic_locate, pic_ocr_data)
            pic_results.append(pic_tup)
            # ================将识别出的文字添加到docx文件中================
            para_tup = (
                'paragraph', bubble_meta_str,
                tess_zdata, tes_zdata,
                rec_results_vision, rec_results_vis,
                vision_fmt_text, vis_fmt_text,
            )
            pic_results.append(para_tup)
    return pic_results


def most_common_color(image):
    """返回图像中出现最多的颜色"""
    pixels = list(image.getdata())
    most_common_pixel = Counter(pixels).most_common(1)[0][0]
    return most_common_pixel


def add_pad2img(image, padding, bg_color=None):
    """给图像加上指定的padding，并使用指定的背景色或默认背景色"""
    if bg_color is None:
        bg_color = most_common_color(image)

    new_width = image.width + 2 * padding
    new_height = image.height + 2 * padding
    padded_img = Image.new('RGB', (new_width, new_height), bg_color)
    padded_img.paste(image, (padding, padding))
    return padded_img


@logger.catch
def get_line_infos(tess_zdata, img_np):
    line_infos = []
    tess_zdata5 = [x for x in tess_zdata if x[0] == 5]
    par_nums = [x[3] for x in tess_zdata5]
    par_nums = reduce_list(par_nums)
    par_nums.sort()
    line_nums = [x[4] for x in tess_zdata5]
    line_nums = reduce_list(line_nums)
    line_nums.sort()
    for par_num in par_nums:
        # 每个段落
        for line_num in line_nums:
            # 每一行
            line_data = [x for x in tess_zdata5 if x[3] == par_num and x[4] == line_num]
            line_data.sort(key=lambda x: x[5])
            word_imgs = []
            for l in range(len(line_data)):
                # 每个单词
                word_data = line_data[l]
                level, page_num, block_num, par_num, line_num, word_num, left, top, width, height, conf, text = word_data
                word_br = (left, top, width, height)
                br_area = width * height
                # 裁剪出单词
                crop_br = (top, top + height, left, left + width)
                word_img = img_np[top - 1:top + height + 1, left - 1:left + width + 1]
                # 检查图像是否为空或维度是否为0
                if word_img.size == 0 or word_img.shape[0] == 0 or word_img.shape[1] == 0:
                    logger.error(f'{crop_br=}')
                elif text != '':
                    pos_meta = (text, par_num, line_num, word_num, left, top, width, height, conf, word_img)
                    word_imgs.append(pos_meta)
            if word_imgs:
                line_infos.append(word_imgs)
    return line_infos


@logger.catch
def stitch_imgs_normal(input_imgs):
    # 转换所有输入图像为 PIL 图像
    pil_imgs = []
    for img in input_imgs:
        if isinstance(img, ndarray):
            # 转换 NumPy 数组为 PIL 图像
            if len(img.shape) == 2:  # 灰度或黑白图像
                pil_img = Image.fromarray(img, 'L')
            else:  # if len(img.shape) == 3:  # 彩色图像
                pil_img = Image.fromarray(img, 'RGB')
        else:  # isinstance(img, Image.Image)
            pil_img = img
        pil_imgs.append(pil_img)

    # 垂直拼接
    total_height = sum(img.height for img in pil_imgs)
    max_width = max(img.width for img in pil_imgs)
    stitched_img = Image.new('RGB' if pil_imgs[0].mode == 'RGB' else 'L', (max_width, total_height), color=255)
    y_offset = 0
    for img in pil_imgs:
        stitched_img.paste(img, (0, y_offset))
        y_offset += img.height
    return stitched_img


@logger.catch
def stitch_imgs(input_imgs, bg_color='B'):
    tolerance = 30
    pil_imgs = []
    for img in input_imgs:
        if isinstance(img, ndarray):
            # 转换 NumPy 数组为 PIL 图像
            if len(img.shape) == 2:  # 灰度或黑白图像
                pil_img = Image.fromarray(img, 'L')
            else:  # if len(img.shape) == 3:  # 彩色图像
                pil_img = Image.fromarray(img, 'RGB')
        else:  # isinstance(img, Image.Image)
            pil_img = img
        # 转换背景为透明
        if pil_img.mode != 'RGBA':
            pil_img = pil_img.convert('RGBA')
        data = array(pil_img)
        if bg_color == 'B':
            # 纯黑色变为透明
            BC = [0, 0, 0]
        else:  # if bg_color == 'W'
            # 纯白色变为透明
            BC = [255, 255, 255]
        mask = np.all(data[:, :, :3] == BC, axis=-1)
        data[mask, 3] = 0
        pil_img = Image.fromarray(data)
        pil_imgs.append(pil_img)

    # 计算拼接后的图像总高度
    total_height = 0
    if pil_imgs:
        total_height = pil_imgs[0].height
    total_height += line_spacing * (len(pil_imgs) - 1)
    max_width = max(img.width for img in pil_imgs)

    # 根据bg_color设置背景颜色
    bg_fill_color = rgba_black if bg_color == 'B' else rgba_white

    # 创建新图像并指定背景颜色
    stitched_img = Image.new('RGBA', (max_width, total_height), bg_fill_color)
    y_offset = 0
    for img in pil_imgs:
        stitched_img.paste(img, (0, y_offset), img)
        y_offset += line_spacing
    stitched_img = conv_img(stitched_img, 'CV')
    return stitched_img


def get_refer_dimension(stroke_cnts, dimension='width'):
    threshold = dot_line_w if dimension == 'width' else dot_line_h
    attr = 'br_w' if dimension == 'width' else 'br_h'

    # 过滤出符合条件的轮廓
    stroke_cnts_fit = [x for x in stroke_cnts if getattr(x, attr) >= threshold]
    if not stroke_cnts_fit:
        stroke_cnts_fit = stroke_cnts

    # 计算基准尺寸
    dimensions = [getattr(x, attr) for x in stroke_cnts_fit]
    mean_dim = mean(dimensions)
    if len(stroke_cnts_fit) >= 10:
        # 去掉10%的最大值
        cutoff = percentile(dimensions, 90)
        filtered_dims = [d for d in dimensions if d <= cutoff]
        avg_dim = mean(filtered_dims)
        # 找出大于平均值的尺寸
        greater_than_avg = [d for d in filtered_dims if d > avg_dim]
        if greater_than_avg:
            # 尝试找出众数，如果没有众数，则取中值
            modes = multimode(greater_than_avg)
            refer_dim = modes[0] if modes else median(greater_than_avg)
        else:
            # 如果没有大于平均值的尺寸，使用过滤后的平均值
            refer_dim = avg_dim
    else:
        # 如果数量小于10，直接取平均值
        refer_dim = mean_dim

    logger.warning(f'{len(stroke_cnts)=}, {mean_dim=:.2f}, {refer_dim=}')
    return refer_dim


@logger.catch
def process_para(ocr_doc, img_folder, pic_result, img_np, page_ind, bubble_ind, media_type):
    logger.warning(f'[{page_ind=}]{bubble_ind=}')
    auto_subdir = Auto / img_folder.name

    # ================图像参数================
    nh, nw = img_np.shape[:2]  # 获取原始图像的尺寸
    white_bg = ones((nh, nw, 3), dtype=uint8) * 255
    black_bg = zeros((nh, nw), dtype=uint8)

    img_pil = fromarray(cvtColor(img_np, COLOR_BGR2RGB))
    img_gray = cvtColor(img_np, COLOR_BGR2GRAY)
    ret, img_bi = threshold(img_gray, bi_thres, 255, THRESH_BINARY)
    # 白色二值化文字
    img_bi_inv = bitwise_not(img_bi)
    img_bi_pil = fromarray(cvtColor(img_bi, COLOR_BGR2RGB))

    # ================文本================
    bubble_meta_str = pic_result[1]
    tess_zdata, tes_zdata = pic_result[2:4]
    rec_results_vision, rec_results_vis = pic_result[4:6]
    vision_fmt_text, vis_fmt_text = pic_result[6:8]
    vision_lines = vision_fmt_text.splitlines()
    if len(vision_lines) >= 2 and vision_lines[-1] == '':
        vision_lines = vision_lines[:-1]
    # vision_lines = vis_fmt_text.splitlines()
    # ================获取气泡元数据================
    bubble_metadata = get_bubble_meta(bubble_meta_str)
    bubble_shape = bubble_metadata['bubble_shape']
    bubble_color_str = bubble_metadata['bubble_color_str']
    letter_color_str = bubble_metadata['letter_color_str']
    color_locate = f"{bubble_color_str}-{letter_color_str}"
    bold_thres = bold_thres_dic['default']
    if color_locate in bold_thres_dic:
        bold_thres = bold_thres_dic[color_locate]
    logger.warning(f'{color_locate=}, {bold_thres=}')
    char_area_dict, word_actual_areas_dict = {}, {}
    if color_locate in area_dic:
        char_area_dict, word_actual_areas_dict = area_dic[color_locate]
    elif letter_color_str in ['000000', 'ffffff'] and 'ffffff-000000' in area_dic:
        # ================黑字或白字强制使用白底黑字================
        char_area_dict, word_actual_areas_dict = area_dic['ffffff-000000']

    # ================添加段落================
    new_para = ocr_doc.add_paragraph()
    # ================气泡数据================
    new_run = new_para.add_run(bubble_meta_str)
    new_run = new_para.add_run(lf)  # 软换行

    if text_direction == 'Horizontal':
        # ================每一个对话框================
        line_infos = get_line_infos(tess_zdata, img_np)

        run_infos = []
        if len(vision_lines) == len(line_infos):
            # 行数相等时
            for l in range(len(line_infos)):
                # 每一行
                word_imgs = line_infos[l]
                ocr_line = vision_lines[l]
                tess_words = [x[0] for x in word_imgs]
                line_words = ocr_line.split(' ')
                # ================如果词数有差异则进行修正================
                if len(tess_words) != len(line_words):
                    new_line_words = []
                    for l in range(len(line_words)):
                        line_word = line_words[l]
                        if '…' in line_word:
                            a, par, b = line_word.partition('…')
                            new_line_words.append(f'{a}{par}')
                            new_line_words.append(b)
                        else:
                            new_line_words.append(line_word)
                    line_words = new_line_words

                # ================词数相等时================
                if len(word_imgs) == len(line_words):
                    for w in range(len(word_imgs)):
                        # 每一词
                        pos_meta = word_imgs[w]
                        line_word = line_words[w]
                        # text, par_num, line_num, word_num, left, top, width, height, conf, word_img = pos_meta
                        word_img = pos_meta[-1]

                        # 检查图像的维度
                        if word_img.shape[0] == 0 or word_img.shape[1] == 0:
                            logger.error(f'图片错误: {line_word}')
                            logger.error(f'{pos_meta[:-1]=}')
                            continue

                        # 计算裁剪出的单词图片的黑色像素面积
                        # 检查是否为 NumPy 数组，如果是，将其转换为 PIL 图像
                        if isinstance(word_img, ndarray):
                            word_img = fromarray(word_img)
                        # 转换为灰度图像
                        gray_img = word_img.convert('L')
                        # 转换为二值图像
                        binary_img = gray_img.point(lambda x: 0 if x <= bit_thres else 255, '1')
                        # 计算黑色像素的数量（假设黑色像素的值为0）
                        black_px_area = binary_img.histogram()[0]
                        pchars = [char for char in line_word if char in char_area_dict]
                        npchars = [char for char in line_word if char not in char_area_dict]
                        run_info = (line_word, '')
                        if italic_color_locates and color_locate in italic_color_locates:
                            run_info = (line_word, 'i')
                        elif bubble_shape == '矩形' and rec2italic:
                            run_info = (line_word, 'i')
                        # 使用正则表达式对line_word进行处理，匹配前面有字母或者后面有字母的"I"
                        line_word = sub(r'(?<=[a-zA-Z])I|I(?=[a-zA-Z])', '|', line_word)
                        if not npchars or len(npchars) / len(line_word) <= 0.2:
                            if line_word in word_actual_areas_dict:
                                word_actual_areas = word_actual_areas_dict[line_word]
                                expect_area = sum(word_actual_areas) / len(word_actual_areas)
                            else:
                                exp_areas = []
                                for char in line_word:
                                    # 移除重音符号
                                    normalized_char = normalize('NFD', char).encode('ascii', 'ignore').decode('utf-8')
                                    if char in char_area_dict:
                                        # 如果字符在字典中，直接使用其面积
                                        exp_area = char_area_dict[char]
                                    else:
                                        # 如果标准化后的字符在字典中，使用其面积
                                        if normalized_char in char_area_dict:
                                            exp_area = char_area_dict[normalized_char]
                                        elif char.lower() in char_area_dict:
                                            exp_area = char_area_dict[char.lower()]
                                        else:
                                            logger.warning(f'{char=}')
                                            exp_area = 5
                                    exp_areas.append(exp_area)
                                expect_area = sum(exp_areas)
                            if expect_area == 0:
                                logger.warning(f'{expect_area=}')
                            else:
                                refer_ratio = black_px_area / expect_area
                                if refer_ratio >= bold_thres:
                                    logger.warning(f'[{refer_ratio:.2f}]{line_word}')
                                    run_info = (line_word, 'bi')
                        else:
                            if media_type in ['Comic']:
                                imes = f'[{line_word}]'
                                if pchars and npchars:
                                    imes += f'{pchars=}, {npchars=}'
                                elif pchars and not npchars:
                                    imes += f'{pchars=}'
                                elif not pchars:
                                    imes += f'{npchars=}'
                                logger.error(imes)
                        run_infos.append(run_info)
                        if w != len(word_imgs) - 1:
                            # 如果不是最后一词就添加空格
                            run_info = (' ', '')
                            run_infos.append(run_info)
                else:
                    run_info = (ocr_line, '')
                    run_infos.append(run_info)
                    logger.error(f'词数不相等,{tess_words=},{line_words=}')
                if l != len(line_infos) - 1:
                    # 如果不是最后一行就添加换行
                    run_info = (lf, '')
                    run_infos.append(run_info)
        else:
            run_info = (vision_fmt_text, '')
            run_infos.append(run_info)
            logger.error(f'行数不相等, [{len(vision_lines)}-{len(line_infos)}]')
            logger.error(f'{vision_fmt_text}')
            for l in range(len(line_infos)):
                word_imgs = line_infos[l]
                words = [x[0] for x in word_imgs]
                logger.debug(f'{words=}')

        for r in range(len(run_infos)):
            run_info = run_infos[r]
            line_word, form = run_info
            if line_word not in [' ', lf]:
                logger.debug(f'[{line_word}]')
            line_word = line_word.replace('|', 'I')
            line_word = line_word.replace('—', '-')
            new_run = new_para.add_run(line_word)
            if 'b' in form:
                new_run.bold = True
            if 'i' in form:
                new_run.italic = True
    else:
        new_run = new_para.add_run(vision_fmt_text)  # 软换行
    ocr_doc.add_paragraph('')
    return ocr_doc


@logger.catch
def update_ocr_doc(ocr_doc, pic_results, img_folder, page_ind, img_list, media_type):
    ocr_yml = img_folder.parent / f'{img_folder.name}-文字识别.yml'
    ocr_data = iload_data(ocr_yml)
    # 示例单词，长度为34
    font_color = "black"

    img_stems = [x.stem for x in img_list]
    cpre = common_prefix(img_stems)
    csuf = common_suffix(img_stems)
    img_file = img_list[page_ind]
    simple_stem = img_file.stem.removeprefix(cpre).removesuffix(csuf)
    pic_text_png = auto_subdir / f'{img_folder.name}-1拼接-{simple_stem}.png'
    # 图与图之间的空行（像素）
    all_cropped_imgs = []
    cur_img_np = None
    if pic_results:
        bubble_ind = 1
        for p, pic_result in enumerate(pic_results):
            result_type = pic_result[0]
            if result_type == 'stem':
                # ================页码================
                pic_stem = pic_result[1]
                if 1 <= stem_level <= 9:
                    ocr_doc.add_heading(pic_stem, level=stem_level)
                else:
                    ocr_doc.add_paragraph(pic_stem)
                ocr_doc.add_paragraph('')
            elif result_type == 'picture':
                # ================图片================
                img_np, pic_locate, pic_ocr_data = pic_result[1:4]
                cur_img_np = img_np
                ocr_data[pic_locate] = pic_ocr_data
                img_pil = fromarray(img_np)
                all_cropped_imgs.append(img_pil)
                # 获取图片的dpi，如果没有dpi信息，则使用默认值
                img_dpi = img_pil.info.get('dpi', (custom_dpi, custom_dpi))[0]
                pic_width_inches = img_pil.width / img_dpi
                with BytesIO() as temp_buffer:
                    img_pil.save(temp_buffer, format=docx_img_format.upper())
                    temp_buffer.seek(0)
                    ocr_doc.add_picture(temp_buffer, width=Inches(pic_width_inches))
            elif result_type == 'paragraph':
                ocr_doc = process_para(ocr_doc, img_folder, pic_result, cur_img_np, page_ind, bubble_ind, media_type)
                bubble_ind += 1
        write_yml(ocr_yml, ocr_data)
        # ================生成并保存长图================
        if all_cropped_imgs:
            # 使用Pillow绘制单词并获取其尺寸
            text_img = Image.new("RGBA", (1000, 1000), rgba_zero)
            text_draw = ImageDraw.Draw(text_img)
            text_draw.text((0, 0), sep_word, font=msyh_font20, fill=font_color)

            # 裁剪文本图像
            text_bbox = text_img.getbbox()
            cropped_text_img = text_img.crop(text_bbox)
            word_width, word_height = cropped_text_img.size

            # 考虑最低宽度
            max_bubble_width = max(img.width for img in all_cropped_imgs)
            max_width = max(max_bubble_width, word_width)

            # 考虑高度
            total_img_h = sum(img.height for img in all_cropped_imgs)
            total_spacing = bubble_spacing * 2 * (len(all_cropped_imgs) - 1)
            total_word_h = word_height * (len(all_cropped_imgs) - 1)
            total_height = total_img_h + total_spacing + total_word_h

            long_img = Image.new('RGB', (max_width, total_height), color_white)
            y_offset = 0
            for img in all_cropped_imgs[:-1]:  # 除了最后一个图像
                long_img.paste(img, ((max_width - img.width) // 2, y_offset))
                y_offset += img.height + bubble_spacing
                # 粘贴分隔词图像
                word_x = (max_width - word_width) // 2
                long_img.paste(cropped_text_img, (word_x, y_offset), cropped_text_img)
                y_offset += word_height + bubble_spacing
            # 添加最后一个图像
            long_img.paste(all_cropped_imgs[-1], ((max_width - all_cropped_imgs[-1].width) // 2, y_offset))
            long_img = add_pad2img(long_img, 20, color_white)
            if not pic_text_png.exists() or renew_stitch:
                write_pic(pic_text_png, long_img)
    else:
        logger.error(f'{page_ind=}')
    return ocr_doc, all_cropped_imgs


@timer_decorator
@logger.catch
def merge_update_doc(src_docx, sd_src_docx, img_stems, sd_img_stems):
    ocr_doc = Document(src_docx)
    sd_ocr_doc = Document(sd_src_docx)
    ocr_full_paragraphs = [para.text for para in ocr_doc.paragraphs]
    ocr_docx_index_dict, ocr_docx_inds = create_index_dict(img_stems, ocr_full_paragraphs)
    sd_ocr_full_paragraphs = [para.text for para in sd_ocr_doc.paragraphs]
    sd_ocr_docx_index_dict, sd_ocr_docx_inds = create_index_dict(sd_img_stems, sd_ocr_full_paragraphs)

    for i in range(len(ocr_full_paragraphs)):
        hd_para = ocr_doc.paragraphs[i]
        sd_para = sd_ocr_doc.paragraphs[i]
        if 'Horizontal' in hd_para.text or 'Vertical' in hd_para.text:
            print(hd_para.text)
            new_run = hd_para.add_run(lf)
            for sd_run in sd_para.runs:
                new_run = hd_para.add_run(sd_run.text)
                new_run.font.bold = sd_run.font.bold
                new_run.font.italic = sd_run.font.italic
                new_run.font.underline = sd_run.font.underline
                new_run.font.strike = sd_run.font.strike
                new_run.font.name = sd_run.font.name
                new_run.font.size = sd_run.font.size
                new_run.font.color.rgb = sd_run.font.color.rgb

        if 'Horizontal' in hd_para.text or 'Vertical' in hd_para.text:
            print(hd_para.text)
            bubble_meta_str = hd_para.text.splitlines()[0]
            hd_para.clear()
            # ================气泡数据================
            new_run = hd_para.add_run(bubble_meta_str)
            new_run = hd_para.add_run(lf)  # 软换行

            for r in range(len(sd_para.runs)):
                sd_run = sd_para.runs[r]
                sd_text = sd_run.text
                sd_textlines = sd_text.splitlines()
                if r == 0:
                    if sd_textlines:
                        first_line = sd_textlines[0]
                    else:
                        first_line = ''
                    text2add = sd_text.removeprefix(f'{first_line}\n')
                else:
                    text2add = sd_text

                if text2add:
                    new_run = hd_para.add_run(text2add)
                    new_run.font.bold = sd_run.font.bold
                    new_run.font.italic = sd_run.font.italic
                    new_run.font.underline = sd_run.font.underline
                    new_run.font.strike = sd_run.font.strike
                    new_run.font.name = sd_run.font.name
                    new_run.font.size = sd_run.font.size
                    new_run.font.color.rgb = sd_run.font.color.rgb

    ocr_doc.save('updated_document.docx')
    return


@timer_decorator
@logger.catch
def step3_OCR(img_folder, media_type, media_lang, vert):
    """
    对给定文件夹中的图像进行OCR识别，并将结果保存到一个docx文件中。

    :param img_folder: 包含要识别的图像的文件夹。
    :param frame_yml: 包含框架配置的YAML文件路径。
    :param media_type: 媒体类型，如'Manga'。
    :param media_lang: 用于OCR的媒体语言。
    :param vertical: 是否进行垂直文本识别。
    :param ocr_doc: 生成的DOCX文档
    """
    # step3_OCR 函数接收一个包含要识别的图像的文件夹、一个包含框架配置的YAML文件路径、媒体类型、OCR识别使用的媒体语言和一个指示是否进行垂直文本识别的布尔值。
    # 然后，对文件夹中的每个图像进行OCR识别，并将识别结果保存到一个docx文件中。
    img_list = get_valid_imgs(img_folder)
    frame_yml = img_folder.parent / f'{img_folder.name}.yml'
    order_yml = img_folder.parent / f'{img_folder.name}-气泡排序.yml'
    ocr_yml = img_folder.parent / f'{img_folder.name}-文字识别.yml'
    frame_data = iload_data(frame_yml)
    order_data = iload_data(order_yml)
    ocr_data = iload_data(ocr_yml)
    # ================气泡蒙版================
    all_masks = get_valid_imgs(img_folder, vmode='mask')

    all_pic_results = []
    if thread_method == 'queue':
        for i in range(len(img_list)):
            img_file = img_list[i]
            pic_results = ocr1pic(img_file, frame_data, order_data, ocr_data, all_masks, media_type, media_lang, vert)
            all_pic_results.append(pic_results)
    else:
        with ThreadPoolExecutor(os.cpu_count()) as executor:
            futures = [
                executor.submit(ocr1pic, img_file, frame_data, order_data, ocr_data, all_masks, media_type,
                                media_lang, vert)
                for img_file in img_list]
            for future in as_completed(futures):
                pic_results = future.result()
                all_pic_results.append(pic_results)
        for i, item in enumerate(all_pic_results):
            if item is None or not isinstance(item, list) or not item:
                logger.error(f"错误的元素在索引 {i}: {item}")
            elif len(item[0]) < 2:
                logger.error(f"元素结构不正确在索引 {i}: {item}")
        all_pic_results = [x for x in all_pic_results if x is not None]
        all_pic_results.sort(key=lambda x: x[0][1])

    ocr_doc = Document()
    for page_ind, pic_results in enumerate(all_pic_results):
        ocr_doc, all_cropped_imgs = update_ocr_doc(ocr_doc, pic_results, img_folder, page_ind, img_list, media_type)
    return ocr_doc


def get_formatted_stem(file_stem, format='txt'):
    if format == 'doc':
        formatted_stem = file_stem
    elif format == 'html':
        formatted_stem = f'<p>{file_stem}</p>'
    else:  # format == 'txt':
        formatted_stem = f'>>>>>>>>[{file_stem.name}]<<<<<<<<'
    formatted_stem = normalize('NFC', formatted_stem)
    # logger.debug(f'{formatted_stem=}')
    return formatted_stem


@logger.catch
def get_format_code(run):
    format_code = 0
    if run.italic:
        format_code += 1
    if run.bold:
        format_code += 2
    if run.underline:
        format_code += 4
    return format_code


@logger.catch
def get_para_text_line(valid_para_lines):
    if media_lang in ['English']:
        # 将每句话首字母大写
        single_line_text = ' '.join(valid_para_lines)
        # 将文本分割成句子并保留句末标点符号
        sentence_parts = split(r'( *[.?!…]+[\'")\]，。？！]* *)', single_line_text)
        if do_cap:
            # 对每个句子进行首字母大写处理
            format_parts = [x.capitalize() for x in sentence_parts]
        else:
            format_parts = sentence_parts
        # 将处理后的句子连接成一个字符串
        para_text_line = ''.join(format_parts).strip()
    else:
        para_text_line = ''.join(valid_para_lines)
    return para_text_line


@logger.catch
def copy_run(run, new_run, format=True):
    new_run.font.size = run.font.size
    new_run.font.name = run.font.name
    if format:
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.strike = run.font.strike
        new_run.font.highlight_color = run.font.highlight_color
        new_run.font.color.rgb = run.font.color.rgb


@timer_decorator
def step4_readable(img_folder):
    """
    该函数用于处理OCR识别后的docx文档，使其更易阅读。它将多行文本的段落转换为单行文本，
    并保留原始格式和样式。

    :param src_docx: OCR识别后的docx文件路径。
    :param img_list: 与docx文件关联的图像列表。
    :param read_docx: 可读化docx文件路径。
    :param read_html: 可读化html文件路径。
    :return: new_doc: 处理后的易阅读的docx文档对象。
    """
    stitch_docx = img_folder.parent / f'{img_folder.name}-1拼接.docx'
    ocr_docx = img_folder.parent / f'{img_folder.name}-1识别.docx'
    src_docx = img_folder.parent / f'{img_folder.name}-2校对.docx'
    read_docx = img_folder.parent / f'{img_folder.name}-3段落.docx'
    read_html = img_folder.parent / f'{img_folder.name}-3段落.html'
    read_txt = img_folder.parent / f'{img_folder.name}-3段落.txt'
    read_stitch_txt = img_folder.parent / f'{img_folder.name}-3拼接段落.txt'
    img_list = get_valid_imgs(img_folder)
    img_stems = [x.stem for x in img_list]
    cpre = common_prefix(img_stems)
    csuf = common_suffix(img_stems)
    simple_stems = [x.removeprefix(cpre).removesuffix(csuf) for x in img_stems]

    if not src_docx.exists():
        copy2(ocr_docx, src_docx)

    if not src_docx.exists():
        logger.error(f'{src_docx}不存在')
        return

    # 打开docx文件
    ocr_doc = Document(src_docx)
    # 读取并连接文档中的所有段落文本
    ocr_full_paragraphs = [para.text for para in ocr_doc.paragraphs]
    ocr_docx_index_dict, ocr_docx_inds = create_index_dict(img_stems, ocr_full_paragraphs)

    readable_doc = Document()
    # 把每个段落的多行文本改成单行文本，同时保留原始格式和样式
    for p in range(len(ocr_doc.paragraphs)):
        # ================每个段落================
        paragraph = ocr_doc.paragraphs[p]
        runs_text = []
        runs_format = []
        format_codes = []
        if p in ocr_docx_inds:
            # 图片名称所在的段落不处理
            pass
        elif paragraph.text.strip() == '':
            # 空段落及图片段落不处理
            pass
        else:
            # OCR识别段落
            para_text = paragraph.text
            para_lines = para_text.splitlines()
            if len(para_lines) >= 2:
                meta_line = para_lines[0]
                valid_para_lines = para_lines[1:]
                para_text_line = get_para_text_line(valid_para_lines)
                if para_text_line != '':
                    start_pos = 0
                    runs = paragraph.runs
                    for r in range(len(runs)):
                        run = runs[r]
                        run_text_raw = run.text
                        # ================换行改掉================
                        run_text = run_text_raw.replace('\n', new_break)
                        if r == 0:
                            # ================移除Meta数据行================
                            new_run_text_raw = run_text.removeprefix(meta_line)
                            new_run_text_raw = new_run_text_raw.removeprefix(' ')
                            end_pos = start_pos + len(new_run_text_raw)
                        else:
                            end_pos = start_pos + len(run_text)
                        # ================截取文本并更新起始位置================
                        new_run_text = para_text_line[start_pos:end_pos]
                        start_pos = end_pos

                        if new_run_text != '':
                            format_code = get_format_code(run)
                            runs_text.append(new_run_text)
                            runs_format.append(run)
                            format_codes.append(format_code)

        if runs_text:
            if len(runs_text) >= 3 and ' ' in runs_text[1:-1] and max(format_codes) > 0:
                # 找到所有单个空格字符串的位置
                space_indices = [index for index, item in enumerate(runs_text) if item == " "]
                space_indices = [x for x in space_indices if 1 <= x <= len(runs_text) - 1]
                merged_runs_text = []
                merged_runs_format = []
                merged_format_codes = []
                roi_indices = []
                for s in range(len(space_indices)):
                    space_indice = space_indices[s]
                    format_before = format_codes[space_indice - 1]
                    format_after = format_codes[space_indice + 1]
                    if format_before == format_after > 0:
                        roi_indices.append(space_indice)

                for r in range(len(runs_text)):
                    new_run_text = runs_text[r]
                    run = runs_format[r]
                    format_code = format_codes[r]
                    if r + 1 in roi_indices:
                        new_run_text += runs_text[r + 1]
                    elif r in roi_indices:
                        new_run_text = ''
                    if new_run_text != '':
                        merged_runs_text.append(new_run_text)
                        merged_runs_format.append(run)
                        merged_format_codes.append(format_code)

                if roi_indices:
                    logger.debug(f'{roi_indices=}')
            else:
                merged_runs_text = runs_text
                merged_runs_format = runs_format
                merged_format_codes = format_codes

            new_para = readable_doc.add_paragraph()
            for new_run_text, run, format_code in zip(merged_runs_text, merged_runs_format, merged_format_codes):
                new_run = new_para.add_run(new_run_text)
                copy_run(run, new_run)
    write_docx(read_docx, readable_doc)
    logger.warning(f'{read_docx=}')

    # ================转为英文段落================
    with open(read_docx, 'rb') as docx_file:
        result = convert_to_html(docx_file)
        result_html = result.value
    result_html = result_html.replace(r'</p><p>', '</p>\n<p>')
    write_txt(read_html, result_html)

    # ================提取纯文本================
    soup = BeautifulSoup(result_html, 'html.parser')
    plain_text = soup.get_text()
    write_txt(read_txt, plain_text)
    # ================检查是否有俄罗斯字母================
    lines = plain_text.splitlines()
    for line_number, line in enumerate(lines, start=1):
        russian_chars = findall(r'[а-яА-ЯёЁ]', line)
        if search(r'[а-яА-ЯёЁ]', line):
            logger.warning(f"第{line_number}行: {line.strip()}")
            # 查找并显示这一行中的俄罗斯字母
            print(f"{''.join(russian_chars)}")

    # ================生成拼接段落================
    if stitch_docx.exists():
        stitch_doc = Document(stitch_docx)
        # 读取并连接文档中的所有段落文本
        stitch_full_paragraphs = [para.text for para in stitch_doc.paragraphs]

        index_dict = {}
        last_ind = 0
        indexes = []
        for i, simple_stem in enumerate(simple_stems):
            formatted_stem = f'{img_folder.name}-1拼接-{simple_stem}'
            if formatted_stem in stitch_full_paragraphs[last_ind:]:
                ind = stitch_full_paragraphs[last_ind:].index(formatted_stem) + last_ind
                index_dict[formatted_stem] = ind
                indexes.append(ind)
                last_ind = ind
        indexes.append(len(stitch_full_paragraphs))

        bubbles = []
        bubble = []
        for p in range(len(stitch_doc.paragraphs)):
            paragraph = stitch_doc.paragraphs[p]
            para_text = paragraph.text
            if p in indexes:
                # 图片名所在的行
                if bubble:
                    bubbles.append(bubble)
                    bubble = []
                logger.warning(f'{para_text=}')
            elif para_text != '':
                if para_text == sep_word:
                    if not bubble:
                        bubble = ['']
                    bubbles.append(bubble)
                    bubble = []
                else:
                    bubble.append(para_text)
                logger.debug(f'{para_text=}')
        if bubble:
            bubbles.append(bubble)
        stitch_lines = []
        for b in range(len(bubbles)):
            bubble = bubbles[b]
            para_text_line = get_para_text_line(bubble)
            stitch_lines.append(para_text_line)
        stitch_full_text = lf.join(stitch_lines)
        stitch_full_text = sub(r'\.{2,}', '…', stitch_full_text)
        write_txt(read_stitch_txt, stitch_full_text)


# @logger.catch
@timer_decorator
def google_translate(simple_lines, target_lang, strip_empty=True):
    """
    将 .docx 文档翻译成指定的目标语言，并将翻译后的文本保存到一个 .txt 文件中。

    :param simple_lines: 要翻译的文本列表
    :param target_lang: 目标语言的代码，例如 'zh-CN' 或 'en'
    """
    chunks = []
    current_chunk = ""

    # 将文本分成多个块，以便在翻译时遵守最大字符数限制
    for line in simple_lines:
        # 检查将当前行添加到当前块后的长度是否超过最大字符数
        if len(current_chunk) + len(line) + 1 > google_max_chars:  # 加1是为了考虑换行符
            chunks.append(current_chunk.strip())
            current_chunk = ""
        current_chunk += line + "\n"

    # 添加最后一个块（如果有内容的话）
    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    # ================分段翻译================
    translated_chunks = []
    # 对每个块使用谷歌翻译 API 进行翻译
    for chunk in chunks:
        translated_chunk = GoogleTranslator(source='auto', target=target_lang).translate(chunk)
        translated_chunks.append(translated_chunk)
    # 将翻译后的块连接成一个字符串
    translated_text = lf.join(translated_chunks)
    if strip_empty:
        translated_lines = translated_text.splitlines()
        translated_lines = [x for x in translated_lines if x.strip() != '']
        translated_text = lf.join(translated_lines)
    return translated_text


@timer_decorator
def fill_textarea(browser, input_text, activate_browser=True, hit_enter=True):
    """
    使用 AppleScript 在指定浏览器中查找并填充页面上的第一个 textarea 元素，然后模拟回车键的按下。

    :param browser (str): 浏览器名称，支持 'Safari' 和 'Google Chrome'。
    :param input_text (str): 需要填入 textarea 的文本。
    :param activate_browser (bool): 是否激活（前台显示）浏览器。

    :return: AppleScript 执行结果，如果有错误返回 None。
    """
    input_lines = input_text.splitlines()
    # 将输入文本复制到剪贴板
    pyperclip.copy(input_text)

    # 如果需要激活浏览器，则设置 activate_command 为 "activate"，否则为空字符串。
    activate_command = "activate" if activate_browser else ""

    js_command = press_enter_js
    as_command = as_paste_n_enter
    if not hit_enter:
        js_command = ''
        as_command = as_paste

    if len(input_lines) == 1:
        # 单行提问
        if browser == 'Safari':
            apple_script = f'''
            tell application "{browser}"
                {activate_command}
                do JavaScript "Array.from(document.querySelectorAll('textarea'))[0].value = '{input_text}'; {js_command}" in front document
            end tell
            '''
        elif browser.startswith('Google Chrome'):
            apple_script = f'''
            tell application "{browser}"
                {activate_command}
                set js_code to "Array.from(document.querySelectorAll('textarea'))[0].value = '{input_text}'; {js_command}"
                execute active tab of front window javascript js_code
            end tell
            '''
        else:
            print(f"Error: Unsupported browser {browser}.")
            return None
    else:
        # 多行提问
        if browser == 'Safari':
            apple_script = f'''
            tell application "{browser}"
                {activate_command}
                do JavaScript "Array.from(document.querySelectorAll('textarea'))[0].focus();" in front document
                {as_command}
            end tell
            '''
        elif browser.startswith('Google Chrome'):
            apple_script = f'''
            tell application "{browser}"
                {activate_command}
                set js_code to "Array.from(document.querySelectorAll('textarea'))[0].focus();"
                execute active tab of front window javascript js_code
                {as_command}
            end tell
            '''
        else:
            print(f"Error: Unsupported browser {browser}.")
            return None
    return run_apple_script(apple_script)


@timer_decorator
@logger.catch
def get_QA(browser, local_name=None, div_type='code'):
    """
    从当前浏览器中提取问答对，并进行格式化和清理。

    :param browser: 当前使用的浏览器名称
    :return: 包含已清理和格式化问答对的列表
    """
    # ================保存当前浏览器内容================
    if local_name:
        chatgpt_html = ChatGPT / f'{local_name}.html'
    else:
        chatgpt_html = save_from_browser(browser)
    chatgpt_text = read_txt(chatgpt_html)
    # ================解析当前浏览器内容================
    soup = BeautifulSoup(chatgpt_text, 'html.parser')
    # 删除所有不需要的标签，如 meta、script、link 和 style
    extra_tags = [
        'meta',
        'script',
        'link',
        'style',
    ]
    for extra_tag in extra_tags:
        for meta in soup.find_all(extra_tag):
            meta.decompose()

    # 查找并删除具有特定 class 的 <div> 标签
    extra_classes = [
        'overflow-x-hidden',  # 历史聊天记录
        'bottom-0',  # 侧边对话
    ]
    for extra_class in extra_classes:
        for div in soup.find_all('div', class_=extra_class):
            div.decompose()

    # 格式化HTML并去除多余的空行
    pretty_html = soup.prettify()
    simple_chatgpt_text = '\n'.join([line for line in pretty_html.splitlines() if line.strip()])
    simple_chatgpt_html = chatgpt_html.parent / f'{chatgpt_html.stem}_simple.html'
    write_txt(simple_chatgpt_html, simple_chatgpt_text)

    # ================进行问答解析================
    simple_soup = BeautifulSoup(simple_chatgpt_text, 'html.parser')
    target_tups = []
    # class_ = "min-h-[20px]"
    class_ = '[.text-message+&]:mt-5'
    message_divs = simple_soup.find_all('div', class_=class_)

    # 模型名称
    class_ = "line-clamp-1 text-sm"
    style = "opacity: 0; padding-left: 0px; width: 0px;"
    model_spans = simple_soup.find_all('span', class_=class_, style=style)
    # 打印找到的所有符合条件的<span>标签
    model_names = []
    for m in range(len(model_spans)):
        model_span = model_spans[m]
        model_name = model_span.get_text().strip()
        model_names.append(model_name)

    handler_normal = HTML2Text()
    handler_no_link = HTML2Text()
    # 如果不需要处理Markdown中的链接可以设置为True
    handler_no_link.ignore_links = True

    model_i = 0
    for m in range(len(message_divs)):
        message_div = message_divs[m]
        raw_div_str = str(message_div)
        # 根据文本特征判断发送者身份
        if 'max-w-[70%]' in raw_div_str:
            text_role = '用户'
            model_name = ''
        else:
            text_role = 'chatGPT'
            if model_i < len(model_names):
                model_name = model_names[model_i]
                model_i += 1
            else:
                model_name = ''

        # 查找 code 标签
        code_tag = message_div.find('code')
        if code_tag and div_type == 'code':
            # 如果找到 code 标签，提取其内容
            # 用于文稿翻译
            target_div = str(code_tag).strip()
        elif text_role == '用户':
            # 用户提问
            # 查找所有的<code>标签
            code_tags = message_div.find_all('code')
            # 替换<code>标签内容，用反引号包围其文本
            for code in code_tags:
                code.string = f"`{code.get_text(strip=True)}`"
            # 提取整个文档的文本，此时<code>中的文本已被修改
            target_div = message_div.get_text(strip=True)

            # 提问转为Markdown
            target_md = handler_normal.handle(raw_div_str).strip()
            if gpt4o_spec_str in target_md:
                pic_part = target_md.split(gpt4o_spec_str)[0].strip()
                target_div = f'{pic_part}{lf}{target_div}'
        else:
            # 回答转为Markdown
            target_div = handler_no_link.handle(raw_div_str).strip()
        target_tup = (text_role, model_name, target_div)
        target_tups.append(target_tup)
    return target_tups


@logger.catch
def get_target_tups(browser, div_type='code'):
    all_target_tups = []
    if read_local:
        target_tups = get_QA(browser, local_name, div_type=div_type)
    else:
        target_tups = get_QA(browser, div_type=div_type)
    all_target_tups.extend(target_tups)
    if read_history:
        for h in range(len(history_names)):
            history_name = history_names[h]
            target_tups = get_QA(browser, history_name, div_type=div_type)
            all_target_tups.extend(target_tups)
    return all_target_tups


def get_split_lines(html_text):
    # 将输入的HTML文本按行分割存入html_lines列表
    html_lines = html_text.splitlines()

    # ================对网页内容进行分段================
    # 初始化split_lines列表，用于存储最终的分段结果，每个元素是一个包含多行文本的列表
    split_lines = []
    # 初始化current_lines列表，用于暂存当前处理的段落
    input_lines = []
    # 初始化当前段落的行数计数器
    current_line_count = 0
    # 初始化当前段落的字符数计数器
    current_char_count = 0

    # 遍历每一行HTML文本
    for i in range(len(html_lines)):
        html_line = html_lines[i]
        # 获取当前行的长度
        line_len = len(html_line)

        # 判断是否可以将当前行添加到current_lines中
        if current_line_count + 1 <= gpt_line_max and current_char_count + line_len <= gpt_char_max:
            # 添加当前行到段落中
            input_lines.append(html_line)
            # 行数计数器加一
            current_line_count += 1
            # 字符数计数器增加当前行的字符数
            current_char_count += line_len
        else:
            # 如果当前行不能添加到current_lines中，则将current_lines作为一个完成的段落添加到split_lines中
            split_lines.append(input_lines)
            # 重置current_lines，开始新段落
            input_lines = [html_line]
            # 重置行数计数器
            current_line_count = 1
            # 重置字符数计数器
            current_char_count = line_len

    # 循环结束后，如果current_lines中有数据，也添加到split_lines中
    if input_lines:
        split_lines.append(input_lines)

    # 返回分段后的所有段落列表
    return split_lines


def html2docx(html_text, new_para):
    # 生成一个唯一的临时文件名
    temp_filename = f"{uuid4()}.docx"

    # 将HTML转换为docx并将其保存到磁盘上
    convert_text(html_text, 'docx', format='html', outputfile=temp_filename, extra_args=['--wrap=none'])

    # 从磁盘读取生成的docx文件
    with open(temp_filename, 'rb') as f:
        temp_doc = Document(f)

    # 将转换后的docx内容添加到新段落
    for para in temp_doc.paragraphs:
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            copy_run(run, new_run)

    # 删除临时文件
    os.remove(temp_filename)


def get_innermost_tag(tag):
    """递归地获取最内层的标签"""
    if tag.contents and isinstance(tag.contents[-1], Tag):
        return get_innermost_tag(tag.contents[-1])
    return tag


# 将所有英文标点替换为对应的中文标点
punct_map = {
    ',': '，',
    '.': '。',
    '?': '？',
    '!': '！',
    # ':': '：',
    # ';': '；',
}


# @logger.catch
@timer_decorator
def get_dst_doc(src_docx, img_list, raw_html, dst_html):
    # 打开docx文件
    ocr_doc = Document(src_docx)
    new_doc = Document()
    dst_html_text = read_txt(raw_html)
    # ================省略号处理================
    dst_html_text = dst_html_text.replace('…', '...')
    dst_html_text = sub(r'\.{2,} ?', '…', dst_html_text)
    # ================左右括号处理================
    dst_html_text = dst_html_text.replace('‹', '<').replace('›', '>')
    # ================其他符号处理================
    dst_html_text = dst_html_text.replace('——', '--')

    soup = BeautifulSoup(dst_html_text, 'html.parser')

    if target_lang == 'zh-CN':
        # 去除中文句子中的空格
        dst_html_text = sub(r'(?<=[\u4e00-\u9fa5])\s+(?=[\u4e00-\u9fa5])', '', dst_html_text)
        for eng_punc, chi_punc in punct_map.items():
            dst_html_text = dst_html_text.replace(eng_punc, chi_punc)
        # 去除前后都不是英文单词的空格
        dst_html_text = sub(r'(?<![a-zA-Z])\s|\s(?![a-zA-Z])', '', dst_html_text)

        cn_puncts = "，。？！：；（）【】｛｝《》…“‘”’"
        for p in soup.find_all('p'):
            for idx, child in enumerate(p.contents):
                if (isinstance(child, NavigableString) and child[0] in cn_puncts and
                        child.parent.name == 'p'):  # 确保该子节点是<p>标签的直接文本内容
                    # 如果当前子节点是字符串并且以中文标点开头
                    if idx > 0:  # 确保它不是第一个子节点
                        prev_sibling = p.contents[idx - 1]
                        if isinstance(prev_sibling, Tag) and prev_sibling.name in ['strong', 'em']:
                            innermost_tag = get_innermost_tag(prev_sibling)
                            if innermost_tag.string and innermost_tag.string[-1] != child[0]:
                                innermost_tag.string += child[0]
                                child.replace_with(child[1:])

    if dst_html.exists():
        dst_html_text = read_txt(dst_html)
    else:
        # 更新dst_html_text为soup的HTML内容
        dst_html_text = str(soup)
        dst_html_text = dst_html_text.replace('</p><p>', '</p>\n<p>')
        write_txt(dst_html, dst_html_text)
    dst_html_lines = dst_html_text.splitlines()

    img_stems = [x.stem for x in img_list]
    cpre = common_prefix(img_stems)
    csuf = common_suffix(img_stems)

    # 读取并连接文档中的所有段落文本
    full_text = []
    for para in ocr_doc.paragraphs:
        full_text.append(para.text)

    # 创建一个字典，将图像文件名与其在文档中的位置关联起来
    index_dict = {}
    last_ind = 0
    indexes = []
    for i in range(len(img_list)):
        img_file = img_list[i]
        if img_file.stem in full_text[last_ind:]:
            ind = full_text[last_ind:].index(img_file.stem) + last_ind
            index_dict[img_file.stem] = ind
            indexes.append(ind)
            last_ind = ind
    indexes.append(len(full_text))

    dst_pin = 0
    for p in range(len(ocr_doc.paragraphs)):
        paragraph = ocr_doc.paragraphs[p]
        if p in indexes:
            # 图片名称所在的行
            pic_name = paragraph.text
            simple_stem = pic_name.removeprefix(cpre).removesuffix(csuf)
            if p != 0:
                new_para = new_doc.add_paragraph('')
            new_para = new_doc.add_paragraph(simple_stem)
        else:
            para_text = paragraph.text
            para_lines = para_text.splitlines()
            if len(para_lines) >= 2:
                meta_line = para_lines[0]
                valid_para_lines = para_lines[1:]
                para_text_line = lf.join(valid_para_lines)
                logger.info(f'[{len(para_lines)}]{para_text_line=}')
                if para_text_line != '':
                    # 放上对应的翻译
                    trans_html_line = dst_html_lines[dst_pin]
                    # 将HTML文本转换为docx格式
                    new_para = new_doc.add_paragraph()
                    html2docx(trans_html_line, new_para)
                    dst_pin += 1
    return new_doc


@logger.catch
def str2dic(dic_str):
    spec_dic = {}
    manual_lines = dic_str.strip().splitlines()
    for m in range(len(manual_lines) - 1):
        manual_line = manual_lines[m]
        next_manual_line = manual_lines[m + 1]
        if m % 2 == 0:
            spec_dic[manual_line] = next_manual_line
    return spec_dic


@logger.catch
def get_code_text(user_html):
    user_soup = BeautifulSoup(user_html, 'html.parser')
    # 查找 code 标签
    code_tag = user_soup.find('code')
    if code_tag is None:
        code_text = user_html
    else:
        code_text = code_tag.get_text().strip()
    # 对文本内容进行反向转义
    unescaped_text = unescape(code_text)
    return code_text


# @timer_decorator
@logger.catch
def get_gpt_dic(target_tups):
    gpt_dic = {}
    gpt_dic_user = str2dic(gpt_user_str)
    gpt_dic.update(gpt_dic_user)
    for t in range(len(target_tups) - 1):
        target_tup = target_tups[t]
        next_target_tup = target_tups[t + 1]
        text_role, model_name, target_div = target_tup
        next_text_role, next_model_name, next_target_div = next_target_tup
        if text_role == '用户' and next_text_role == 'chatGPT':
            user_html = target_div
            gpt_html = next_target_div
            if '<code' in user_html:
                user_code_text = get_code_text(user_html)
                gpt_code_text = get_code_text(gpt_html)
                if user_code_text.startswith('html'):
                    user_code_text = user_code_text.removeprefix('html').strip()
                input_lines = user_code_text.strip().splitlines()
                output_lines = gpt_code_text.strip().splitlines()
                if len(input_lines) == len(output_lines) >= 1:
                    ilines = input_lines
                    olines = output_lines
                    for c in range(len(ilines)):
                        input_line = ilines[c]
                        output_line = olines[c]
                        gpt_dic[input_line] = output_line
    return gpt_dic


@logger.catch
def step5_translate(img_folder):
    src_docx = img_folder.parent / f'{img_folder.name}-2校对.docx'
    read_html = img_folder.parent / f'{img_folder.name}-3段落.html'
    dst_docx = img_folder.parent / f'{img_folder.name}-4翻译.docx'
    raw_html = img_folder.parent / f'{img_folder.name}-4原翻.html'
    dst_html = img_folder.parent / f'{img_folder.name}-4翻译.html'
    googletrans_txt = img_folder.parent / f'{img_folder.name}-9-api-谷歌翻译.txt'
    img_list = get_valid_imgs(img_folder)
    activate_browser = True
    if not read_html.exists():
        step4_readable(img_folder)
    if read_html.exists():
        read_html_text = read_txt(read_html)
        soup = BeautifulSoup(read_html_text, 'html.parser')
        text = soup.get_text()
        simple_lines = text.splitlines()
        read_html_mtime = getmtime(read_html)
        if not googletrans_txt.exists():
            # ================谷歌翻译================
            logger.debug('not googletrans_txt.exists()')
            translated_text = google_translate(simple_lines, target_lang)
            write_txt(googletrans_txt, translated_text)
        else:
            # 获取文件的修改时间
            googletrans_txt_mtime = getmtime(googletrans_txt)
            # 比较修改时间
            if read_html_mtime > googletrans_txt_mtime:
                # ================如有更新则更新谷歌翻译================
                logger.debug('read_html_mtime > googletrans_txt_mtime')
                translated_text = google_translate(simple_lines, target_lang)
                write_txt(googletrans_txt, translated_text)
            else:
                # ================否则进行GPT4翻译================
                logger.debug('进行GPT4翻译')
                if not raw_html.exists():
                    read_html_text = read_txt(read_html)
                    roi_htmls = read_html_text.splitlines()
                    if SYSTEM in ['MAC', 'M1']:
                        target_tups = get_target_tups(browser)
                        gpt_dic = get_gpt_dic(target_tups)
                        # ================排除已经翻译的部分================
                        need2trans_lines = [x for x in roi_htmls if x not in gpt_dic]
                        need2trans_lines = reduce_list(need2trans_lines)
                        html_text = lf.join(need2trans_lines)
                        split_lines = get_split_lines(html_text)
                        split_lines = [x for x in split_lines if x != []]
                        # ================添加提示词================
                        for s in range(len(split_lines)):
                            input_lines = split_lines[s]
                            input_text = lf.join(input_lines)
                            full_prompt = f'{prompt_prefix}{lf}```html{lf}{input_text}{lf}```'
                            possible_divs = [x for x in target_tups if x[0] == '用户' and full_prompt in x[-1]]
                            # ================尚未提问的段落================
                            if not possible_divs:  # or True
                                logger.warning(
                                    f'[{s + 1}/{len(split_lines)}], {len(input_lines)=}, {len(input_text)=}')
                                logger.info(full_prompt)
                                if do_automate:
                                    fill_textarea(browser, full_prompt, activate_browser, hit_enter)
                                    if s != len(split_lines):
                                        stime = sleep_time
                                    else:
                                        # 最后一次等待时间
                                        stime = int(0.6 * sleep_time)
                                    # 等待回答完成
                                    sleep(stime)
                        logger.warning(f'chatGPT翻译完成, {len(split_lines)=}')

                        target_tups = get_target_tups(browser)
                        gpt_dic = get_gpt_dic(target_tups)
                        dst_lines = []
                        trans_len = 0
                        for r in range(len(roi_htmls)):
                            src_line = roi_htmls[r]
                            if src_line in gpt_dic:
                                dst_line = gpt_dic[src_line]
                                trans_len += 1
                            else:
                                dst_line = src_line
                            dst_lines.append(dst_line)
                        if trans_len >= 0.996 * len(roi_htmls):
                            dst_html_text = lf.join(dst_lines)
                            print(dst_html_text)
                            write_txt(raw_html, dst_html_text)
                    else:
                        translated_text = read_txt(googletrans_txt)
                        unescaped_texts = translated_text.splitlines()
                        unescaped_texts = [f'<p>{x}</p>' for x in unescaped_texts]
                        dst_html_text = lf.join(unescaped_texts)
                        write_txt(raw_html, dst_html_text)
                dst_doc = get_dst_doc(src_docx, img_list, raw_html, dst_html)
                write_docx(dst_docx, dst_doc)


# @logger.catch
def colored_text(text, color):
    colors = {
        'blue': '\033[94m',
        'green': '\033[92m',
        'purple': '\033[95m',
        'highlight': '\033[91m',
        'end': '\033[0m'
    }
    return colors[color] + text + colors['end']


def filter_w_bounds(data, key=lambda x: x):
    """
    过滤掉列表中指定比例的最高和最低值。

    :param data: 列表数据。
    :param key: 用于提取比较值的函数。
    :param discard_ratio: 要丢弃的数据比例。
    :return: 过滤后的列表。
    """
    if not data:
        return []

    # 根据提供的键函数排序
    sorted_data = sorted(data, key=key)
    discard_count = int(discard_ratio * len(sorted_data))

    # 计算上下界
    lower_bound = key(sorted_data[discard_count])
    upper_bound = key(sorted_data[-discard_count - 1])

    # 过滤数据
    return [x for x in data if lower_bound - 1 <= key(x) <= upper_bound + 1]


@timer_decorator
@logger.catch
def get_area_dic(area_folder_names):
    # 设置最小有效数据长度
    min_data_len = 28
    area_data = defaultdict(list)
    area_dic = {}
    all_data = []
    if area_folder_names is None:
        area_folder_names = []

    # ================加载文件夹数据================
    # 遍历每个文件夹名称
    for area_folder_name in area_folder_names:
        # 构造文件夹路径
        area_img_folder = ComicProcess / area_folder_name
        sub_area_yml = area_img_folder.parent / f'{area_img_folder.name}-文字面积.yml'
        sub_area_data = iload_data(sub_area_yml)
        for key, val in sub_area_data.items():
            area_data[key].extend(val)

    # ================加载气泡文字颜色对和对应单词的数据================
    for color_locate in area_data:
        words_w_format_str = area_data[color_locate]
        words_w_format_str.sort()
        raw_words_data = []
        for w in range(len(words_w_format_str)):
            word_w_format_str = words_w_format_str[w]
            parts = word_w_format_str.rsplit('|', 2)
            if len(parts) == 2:
                word = parts[0]
                word_format = ''
                nums_str = parts[1]
            elif len(parts) == 3:
                word = parts[0]
                word_format = parts[1]
                nums_str = parts[2]
            else:
                # 如果parts的长度大于3，那么word中可能包含'|'字符
                word = '|'.join(parts[:-2])
                word_format = parts[-2]
                nums_str = parts[-1]
            # 解析高度和黑色像素面积
            height, black_px_area = map(int, nums_str.split(','))
            raw_word_data = (word, word_format, height, black_px_area)
            raw_words_data.append(raw_word_data)
        data_tup = (color_locate, raw_words_data)
        all_data.append(data_tup)

    for data_tup in all_data:
        color_locate, raw_words_data = data_tup
        print(color_locate)
        # 如果数据长度满足最小要求
        if len(raw_words_data) >= min_data_len:
            # 对数据进行过滤和排序
            words_data = filter_w_bounds(raw_words_data, key=lambda x: x[2])
            logger.debug(f'{len(raw_words_data)=}, {len(words_data)=}')
            # 对words_data进行处理，把表示我的I替换成|
            for a, (word, word_format, height, area) in enumerate(words_data):
                # 使用正则表达式对line_word进行处理，匹配前面有字母或者后面有字母的"I"
                word = sub(r'(?<=[a-zA-Z])I|I(?=[a-zA-Z])', '|', word)
                words_data[a] = (word, word_format, height, area)

            # 获取所有独特的字符
            appeared_chars = set()
            unique_chars = set(''.join([word for word, word_format, height, black_px_area in words_data]))
            formatted_chars = set()
            # 为每个字符添加格式变体
            for char in unique_chars:
                formatted_chars.add(char)  # 无格式字符
                formatted_chars.add(f"{char}_b")  # 粗体字符
                formatted_chars.add(f"{char}_i")  # 斜体字符
                formatted_chars.add(f"{char}_bi")  # 粗斜体字符
            # 创建字符到索引的映射
            char2index = {char: index for index, char in enumerate(formatted_chars)}

            # 初始化矩阵和向量
            A = zeros((len(words_data), len(formatted_chars)))
            b = zeros(len(words_data))
            weights = zeros(len(words_data))
            # 填充矩阵和向量
            for n, (word, word_format, height, area) in enumerate(words_data):
                for char in word:
                    char_key = char
                    if word_format != '':
                        char_key += f'_{word_format}'
                        # logger.debug(f'{char_key=}')
                    A[n, char2index[char_key]] += 1
                    appeared_chars.add(char_key)
                b[n] = area
                weights[n] = len(word)

            # 构建加权矩阵
            W = np.diag(np.sqrt(weights))
            WA = W @ A
            Wb = W @ b
            logger.warning(f'{cal_method=}')
            if cal_method == '加权最小二乘法':
                # 使用加权最小二乘法解方程组
                x, residuals, rank, s = np.linalg.lstsq(WA, Wb, rcond=None)
                # x 是一个向量，其中每个元素是一个字符的期望黑色像素面积
            elif cal_method == '非负最小二乘法':
                # 使用非负最小二乘法解方程组
                x, residuals = nnls(WA, Wb)
                # x 是一个向量，其中每个元素是一个字符的期望黑色像素面积
            else:
                # 加权非负最小二乘法
                # 设置非负约束
                bounds = (min_char_area, np.inf)
                result = lsq_linear(WA, Wb, bounds=bounds)
                x = result.x
            # 构建字符到面积的映射
            char_to_area = {char: max(min_char_area, x[char2index[char]]) for char in formatted_chars}

            # 输出结果
            char_table = PrettyTable()
            char_table.field_names = ["Character", "Expected Area", "Bold", "Italic", "Bold Italic"]
            for char in sorted(unique_chars):
                row = [
                    char,
                    round(char_to_area[char], 2) if char in appeared_chars else '',
                    round(char_to_area[f"{char}_b"], 2) if f"{char}_b" in appeared_chars else '',
                    round(char_to_area[f"{char}_i"], 2) if f"{char}_i" in appeared_chars else '',
                    round(char_to_area[f"{char}_bi"], 2) if f"{char}_bi" in appeared_chars else ''
                ]
                char_table.add_row(row)
            if print_tables:
                print(char_table)

            word_table = PrettyTable()
            word_table.field_names = ['Word', 'Actual Area', 'Expected Area', 'Difference']
            for word, word_format, height, area in sorted(words_data, key=lambda x: x[0]):
                color_map = {
                    '': word,
                    'b': colored_text(word, 'blue'),
                    'i': colored_text(word, 'green'),
                    'bi': colored_text(word, 'purple')
                }
                colored_word = f"{color_map[word_format]}[{height}]"
                expect_area = sum(
                    [char_to_area[char + ('' if word_format == '' else f'_{word_format}')] for char in word])
                difference = round(expect_area - area, 2)
                if abs(difference) >= 0.2 * area:
                    difference = colored_text(str(difference), 'highlight')
                word_table.add_row([colored_word, area, round(expect_area, 2), difference])
            if print_tables:
                print(word_table)
            # 创建字典来存储字符和其期望的面积
            char_area_dict = {char: round(char_to_area[char], 2) for char in appeared_chars}

            # 创建字典来存储单词（包括格式后缀）和其期望的面积，以及另一个字典来存储单词的实际面积列表
            word_area_dict = {}
            word_actual_areas_dict = {}

            for word, word_format, height, actual_area in words_data:
                word_key = word + ('' if word_format == '' else f'_{word_format}')
                expect_area = sum(
                    [char_to_area[char + ('' if word_format == '' else f'_{word_format}')] for char in word])
                if word_key not in word_area_dict:
                    word_area_dict[word_key] = [expect_area]
                    word_actual_areas_dict[word_key] = [actual_area]
                else:
                    word_area_dict[word_key].append(expect_area)
                    word_actual_areas_dict[word_key].append(actual_area)

            # 计算字符出现次数
            char_count = {}
            for word, word_format, height, area in words_data:
                for char in word:
                    char_key = char + ('' if word_format == '' else f'_{word_format}')
                    char_count[char_key] = char_count.get(char_key, 0) + 1

            # 输出结果
            seg_char_table = PrettyTable()
            seg_char_table.field_names = ["#", "Character", "Expected Area"]
            for idx, (char, area) in enumerate(sorted(char_area_dict.items()), 1):
                count_suffix = f"[{char_count[char]}]" if char in char_count and char_count[char] > 1 else ""
                seg_char_table.add_row([idx, char, f"{area}{count_suffix}"])
            if print_tables:
                print(seg_char_table)
            seg_word_table = PrettyTable()
            seg_word_table.field_names = ['#', 'Word', 'Expected Area', 'Actual Areas']
            for idx, (word_key, areas) in enumerate(sorted(word_area_dict.items()), 1):
                avg_expect_area = round(sum(areas) / len(areas), 2)
                # ================简化Actual Areas的表示================
                areas = word_actual_areas_dict[word_key]
                areas.sort()
                simplified = []
                start = areas[0]
                end = start
                for i in range(1, len(areas)):
                    if areas[i] == end + 1:
                        end = areas[i]
                    else:
                        if start == end:
                            simplified.append(str(start))
                        else:
                            simplified.append(f"{start}~{end}")
                        start = areas[i]
                        end = start
                if start == end:
                    simplified.append(str(start))
                else:
                    simplified.append(f"{start}~{end}")
                simp = ', '.join(simplified)
                seg_word_row = [idx, word_key, avg_expect_area, simp]
                seg_word_table.add_row(seg_word_row)
            if print_tables:
                print(seg_word_table)
            area_dic[color_locate] = (char_area_dict, word_actual_areas_dict)
    return area_dic


@timer_decorator
# @logger.catch
def stepA_figure(img_folder):
    img_list = get_valid_imgs(img_folder)
    if thread_method == 'queue':
        for i in range(len(img_list)):
            figure1pic(img_folder, i)
    else:
        with ThreadPoolExecutor(os.cpu_count()) as executor:
            futures = [executor.submit(figure1pic, img_folder, i)
                       for i in range(len(img_list))]


def as_proc(as_code):
    try:
        result = applescript.run(as_code)
        print(f"Script output: {result.out}")
    except Exception as e:
        print(f"Script error: {e}")


def mouse_click(x, y):
    mouse_event = CGEventCreateMouseEvent(None, kCGEventLeftMouseDown, CGPoint(x, y), kCGMouseButtonLeft)
    CGEventPost(kCGHIDEventTap, mouse_event)
    CGEventSetType(mouse_event, kCGEventLeftMouseUp)
    CGEventPost(kCGHIDEventTap, mouse_event)


@logger.catch
def find_n_click(roi_logo, click_it=False):
    roi_location = None
    try:
        roi_location = locateOnScreen(roi_logo.as_posix(), confidence=0.98)
        if roi_location:
            logger.info(f'{roi_location=}')
            # 获取图片中心点坐标
            ct = center(roi_location)
            logger.info(f'{ct=}')
            pos_x = int(ct.x / 2)
            pos_y = int(ct.y / 2)
            pos = (pos_x, pos_y)
            logger.info(f'{pos=}')
            if click_it:
                if click_type == 'pyautogui':
                    # click(ct)
                    click(pos)
                else:
                    mouse_click(pos_x, pos_y)
    except Exception as e:
        pass
    return roi_location


@timer_decorator
@logger.catch
def choose1pic(down_num, ask_mode):
    # ================点击附加文件按钮================
    if ask_mode == 'web':
        # 退出输入框
        press('esc')
        sleep(0.2)
        press('esc')
        sleep(0.2)
        press('esc')
        sleep(0.3)

        # 使用Vimium显示链接
        press('f')
        sleep(0.5)

        # 点击附加文件按钮
        press('s')
        sleep(0.1)
        press('s')
    else:
        # 找到屏幕上的图片位置
        upload_location = find_n_click(upload_logo, click_it=True)
        sleep(0.5)

        # 模拟下箭头键
        press('down')
        sleep(0.2)

        # 模拟回车键
        press('enter')
        sleep(0.2)

    # ================选取图片================
    sleep(2)
    for p in range(down_num):
        # 模拟下箭头键
        press('down')
        sleep(0.2)

    # 模拟回车键
    press('enter')
    sleep(0.2)


@logger.catch
def warn_user(warn_str):
    logger.error(warn_str)
    as_proc(as_funk)
    voice_str = warn_str.replace('重', '虫')
    as_Tingting = f"""
    say "{voice_str}" speaking rate 180
    """
    as_proc(as_Tingting)


@timer_decorator
@logger.catch
def stepB_ask(img_folder):
    series4file = folder_name
    m_issue_w_dot = p_issue_w_dot.match(folder_name)
    issue = ''
    logger.debug(f'{series4file=}')
    if m_issue_w_dot:
        series4file = m_issue_w_dot.group(1)
        issue = m_issue_w_dot.group(2)
        logger.debug(f'{issue=}')

    img_list = get_valid_imgs(img_folder)
    all_psds = get_files(img_folder, 'psd', True)
    all_dirs = get_dirs(img_folder)

    img_stems = [x.stem for x in img_list]
    cpre = common_prefix(img_stems)
    csuf = common_suffix(img_stems)

    target_tups = get_target_tups(browser, div_type='pic')
    logger.debug(f'{len(target_tups)=}')
    question_tups = [x for x in target_tups if x[0] == '用户']

    if ask_mode == 'web':
        app_name = browser
    else:
        app_name = 'ChatGPT'
    as_activate_app = f'tell application "{app_name}" to activate'

    activate_browser = False

    # 起始页
    offset_i = 0
    start_i = round(len(target_tups) / 2)

    down_num = 0
    down_num += len(all_psds)
    down_num += len(all_dirs)

    has_error = False
    unfinished_img_files = []
    for i in range(len(img_list)):
        img_file = img_list[i]
        simple_stem = img_file.stem.removeprefix(cpre).removesuffix(csuf)
        locate_prompt = f'这张图片是来自漫画`{img_folder.name}`的`{img_file.name}`。'
        roi_question_tups = [x for x in question_tups if locate_prompt in x[-1]]
        if not roi_question_tups:
            # 尚未进行对此图片的提问
            logger.info(f'{img_file.name}')
            unfinished_img_files.append(img_file)
            if i == 0:
                # 需要选取正确的文件夹
                warn_user('封面需要手动放入')
                has_error = True
                break
    logger.warning(f'{len(unfinished_img_files)=}')

    if do_automate and not has_error:
        for i in range(len(img_list)):
            img_file = img_list[i]
            simple_stem = img_file.stem.removeprefix(cpre).removesuffix(csuf)
            folder_added_frames_jpg = img_folder / f'{img_file.stem}-加框.jpg'
            locate_prompt = f'这张图片是来自漫画`{img_folder.name}`的`{img_file.name}`。'
            full_prompt = f'{locate_prompt}{lf}{chatGPT4o_prompt}'
            roi_question_tups = [x for x in question_tups if locate_prompt in x[-1]]

            down_num += 1
            if folder_added_frames_jpg.exists():
                down_num += 1
                # logger.warning(f'{folder_added_frames_jpg=}')

            if not roi_question_tups:
                # 尚未进行对此图片的提问
                logger.warning(f'[{down_num}]{img_file.name=}')
                as_notification = f'display notification "{img_file.name}" with title "需上传的图片" subtitle "{down_num}"'
                sleep(2)
                # 调取浏览器或应用到前台
                as_proc(as_activate_app)
                # 显示通知
                # as_proc(as_notification)
                sleep(2)
                # ================开始每张图的流程================
                if ask_mode == 'web':
                    current_url = get_browser_current_tab_url(browser)
                    title = get_browser_current_tab_title(browser)
                    logger.debug(f'{current_url=}')
                    logger.warning(f'{title=}')
                    sleep(1)

                    choose1pic(down_num, ask_mode)

                    # 等待图片上传
                    sleep(20)
                    # 粘贴prompt
                    fill_textarea(browser, full_prompt, activate_browser, hit_enter)
                    # 等待回答
                    sleep(60)
                    # break
                else:
                    retry_location = find_n_click(retry_logo)
                    if retry_location:
                        logger.error(f'{retry_location=}')
                        warn_user('需重试')
                        break
                    reconnect_location = find_n_click(reconnect_logo)
                    if reconnect_location:
                        logger.error(f'{reconnect_location=}')
                        warn_user('需重连')
                        break

                    choose1pic(down_num, ask_mode)
                    # 等待图片上传
                    sleep(14)
                    for a in range(2):
                        # 找到屏幕上的图片位置来确定图片上传完成
                        gray_up_arrow_location = find_n_click(gray_up_arrow_logo)
                        if gray_up_arrow_location:
                            sleep(1)
                        else:
                            sleep(0.2)
                    for b in range(25):
                        # 找到屏幕上的图片位置来确定图片上传完成
                        up_arrow_location = find_n_click(up_arrow_logo)
                        if up_arrow_location:
                            logger.info(f'{up_arrow_location=}')
                            break
                        else:
                            sleep(1)
                    sleep(1)

                    # 粘贴prompt
                    pyperclip.copy(full_prompt)
                    sleep(0.05)
                    as_proc(as_paste)
                    sleep(1)

                    try:
                        poses = []
                        for pos in locateAllOnScreen(white_x_logo.as_posix()):
                            poses.append(pos)
                        if len(poses) == 1:
                            if hit_enter:
                                reconnect_location = find_n_click(reconnect_logo)
                                if reconnect_location:
                                    logger.error(f'{reconnect_location=}')
                                    warn_user('需重连')
                                    break

                                sleep(2)

                                # 按下回车键
                                # as_proc(as_enter)
                                # press('return')

                                keyDown('enter')
                                sleep(0.1)
                                keyUp('enter')

                                sleep(2)

                                up_arrow_location = find_n_click(up_arrow_logo)
                                if up_arrow_location:
                                    logger.error(f'{up_arrow_location=}')
                                    warn_user('没按出回车')
                                    break

                                sleep(3)
                                retry_location = find_n_click(retry_logo)
                                if retry_location:
                                    logger.error(f'{retry_location=}')
                                    warn_user('需重试')
                                    break

                                if i == len(img_list) - 1:
                                    as_proc(as_submarine)
                                    as_proc(as_Tingting_uploaded)
                                    break

                                # 等待回答
                                sleep(60)

                                # 到达最下
                                down_arrow_location = find_n_click(down_arrow_logo, True)
                                sleep(2)
                                if down_arrow_location:
                                    # 点击输入框
                                    send_message_ocation = find_n_click(send_message_logo, True)
                                    sleep(2)

                                # ================确保答案已经生成完毕================
                                headphone_location = None
                                for a in range(30):
                                    # 找到屏幕上的图片位置
                                    headphone_location = find_n_click(headphone_logo)
                                    if headphone_location:
                                        # 如果找到就不再等待
                                        logger.info(f'{headphone_location=}')
                                        break
                                    else:
                                        sleep(1)
                                if not headphone_location:
                                    warn_user('答案未生成完毕')
                                    break
                            else:
                                warn_user('用户设定不再继续')
                                break
                        else:
                            logger.error(f'{len(poses)=}')
                            warn_user('图片不是一张')
                            break
                    except Exception as e:
                        printe(e)
                        warn_user('找图出错')
                        break


@timer_decorator
@logger.catch
def stepC_conclude(img_folder):
    img_list = get_valid_imgs(img_folder)

    desc_md = img_folder.parent / f'{img_folder.name}-C描述.md'
    img_stems = [x.stem for x in img_list]
    cpre = common_prefix(img_stems)
    csuf = common_suffix(img_stems)

    target_tups = get_target_tups(browser, div_type='pic')
    logger.debug(f'{len(target_tups)=}')
    question_tups = [x for x in target_tups if x[0] == '用户']

    desc_text = ''
    for i in range(len(img_list)):
        img_file = img_list[i]
        simple_stem = img_file.stem.removeprefix(cpre).removesuffix(csuf)
        locate_prompt = f'这张图片是来自漫画`{img_folder.name}`的`{img_file.name}`。'
        roi_question_tups = [x for x in question_tups if locate_prompt in x[-1]]
        if roi_question_tups:
            roi_question_tup = roi_question_tups[-1]
            roi_ind = target_tups.index(roi_question_tup)
            logger.debug(f'{roi_ind=}')
            if roi_ind < len(target_tups) - 1:
                roi_answer_tup = target_tups[roi_ind + 1]
                # 图片名作为大标题
                desc_text += f'# {img_file.stem}{lf}{lf}'
                # 用户提问
                question_div = roi_question_tup[-1]
                # desc_text += f'{question_div}{lf}{lf}'
                # GPT回答
                answer_div = roi_answer_tup[-1]
                desc_text += f'{answer_div}{lf}{lf}'

    desc_text = desc_text.strip()
    write_txt(desc_md, desc_text)


@timer_decorator
@logger.catch
def stepD_fill(img_folder):
    img_list = get_valid_imgs(img_folder)
    roi_folder = img_folder / '成品'
    for i in range(len(img_list)):
        img_file = img_list[i]
        pic_path1 = roi_folder / img_file.name
        pic_path2 = roi_folder / f'{img_file.stem} 拷贝.jpg'
        if pic_path1.exists() or pic_path2.exists():
            pass
        else:
            copy2(img_file, pic_path1)
            logger.warning(f'{img_file}->{pic_path1}')


@timer_decorator
@logger.catch
def stepE_group(img_folder):
    series4file = folder_name
    m_issue_w_dot = p_issue_w_dot.match(folder_name)
    if m_issue_w_dot:
        series4file = m_issue_w_dot.group(1)
        issue = m_issue_w_dot.group(2)

    logger.debug(f'{series4file=}')
    auto_group_dir = Auto / series4file
    make_dir(auto_group_dir)

    hanhua_folders = get_dirs(Hanhua)
    roi_folders = [x for x in hanhua_folders if x.name.startswith(series4file)]
    for r in range(len(roi_folders)):
        hanhua_folder = roi_folders[r]
        final_folder = hanhua_folder / '成品'
        all_pics = get_files(final_folder, 'pic', True)
        final_jpgs = [x for x in all_pics if x.suffix.lower() in ('.jpg', '.jpeg')]
        if final_jpgs:
            cover_jpg = final_jpgs[0]
            dst_jpg = auto_group_dir / f'{hanhua_folder.name}-封面.jpg'
            copy2(cover_jpg, dst_jpg)
            logger.warning(f'[{r + 1}/{len(roi_folders)}]{cover_jpg}->{dst_jpg}')


@logger.catch
def folder_proc(img_folder, step_str, img_inds):
    frame_yml = img_folder.parent / f'{img_folder.name}.yml'
    ocr_docx = img_folder.parent / f'{img_folder.name}-1识别.docx'
    src_docx = img_folder.parent / f'{img_folder.name}-2校对.docx'
    read_docx = img_folder.parent / f'{img_folder.name}-3段落.docx'
    read_html = img_folder.parent / f'{img_folder.name}-3段落.html'
    dst_docx = img_folder.parent / f'{img_folder.name}-4翻译.docx'
    raw_html = img_folder.parent / f'{img_folder.name}-4原翻.html'
    dst_html = img_folder.parent / f'{img_folder.name}-4翻译.html'
    cut_txt = img_folder.parent / f'{img_folder.name}-5断句.txt'
    lettering_txt = img_folder.parent / f'{img_folder.name}-6填字.txt'
    mark_txt = img_folder.parent / f'{img_folder.name}-7标记.txt'
    lp_txt = img_folder.parent / f'{img_folder.name}-8萌翻.txt'
    rlp_txt = img_folder.parent / f'{img_folder.name}翻译_0.txt'
    googletrans_txt = img_folder.parent / f'{img_folder.name}-9-api-谷歌翻译.txt'
    chara_docx = img_folder.parent / f'{img_folder.name}-人物.docx'
    chara_M_yml = img_folder.parent / f'{img_folder.name}-男.yml'
    chara_F_yml = img_folder.parent / f'{img_folder.name}-女.yml'

    sd_img_folder = get_sd_img_folder(img_folder)
    sd_src_docx = img_folder.parent / f'{sd_img_folder.name}-2校对.docx'

    auto_subdir = Auto / img_folder.name
    make_dir(auto_subdir)

    img_list = get_valid_imgs(img_folder)
    img_stems = [x.stem for x in img_list]
    cpre = common_prefix(img_stems)
    csuf = common_suffix(img_stems)
    logger.debug(f'{cpre=}, {csuf=}')

    sd_img_list = get_valid_imgs(sd_img_folder)
    sd_img_stems = [x.stem for x in sd_img_list]
    sd_cpre = common_prefix(sd_img_stems)
    sd_csuf = common_suffix(sd_img_stems)

    if '0' in step_str:
        frame_data_sorted = step0_analyze_frames(img_folder, frame_yml, media_type, auto_subdir, img_inds)
    if '1' in step_str:
        auto_all_masks = step1_analyze_bubbles(img_folder, media_type, auto_subdir)
    if '2' in step_str:
        step2_order(img_folder, media_type)
    if '3' in step_str:
        if src_docx.exists() and sd_src_docx.exists() and merge_update:
            logger.warning(f'{sd_src_docx=}')
            merge_update_doc(src_docx, sd_src_docx, img_stems, sd_img_stems)
        else:
            ocr_doc = step3_OCR(img_folder, media_type, media_lang, vert)
            write_docx(ocr_docx, ocr_doc)
            logger.warning(f'{ocr_docx=}')
    if '4' in step_str:
        step4_readable(img_folder)
    if '5' in step_str:
        step5_translate(img_folder)
    if 'A' in step_str:
        stepA_figure(img_folder)
    if 'B' in step_str:
        stepB_ask(img_folder)
    if 'C' in step_str:
        stepC_conclude(img_folder)
    if 'D' in step_str:
        stepD_fill(img_folder)
    if 'E' in step_str:
        stepE_group(img_folder)


def z():
    pass


if __name__ == "__main__":
    Comic = homedir / 'Comic'
    Hanhua = Comic / '汉化'

    MomoHanhua = DOCUMENTS / '默墨汉化'
    Auto = MomoHanhua / 'Auto'
    Log = MomoHanhua / 'Log'
    DataOutput = MomoHanhua / 'DataOutput'
    ComicProcess = MomoHanhua / 'ComicProcess'  # 美漫
    MangaProcess = MomoHanhua / 'MangaProcess'  # 日漫
    ManhuaProcess = MomoHanhua / 'ManhuaProcess'  # 国漫
    ManhwaProcess = MomoHanhua / 'ManhwaProcess'  # 韩漫
    Storage = MomoHanhua / 'Storage'

    MomoYolo = DOCUMENTS / '默墨智能'
    ChatGPT = MomoYolo / 'ChatGPT'

    AutomateUserDataFolder = ProgramFolder / 'MomoAutomateUserData'
    ChatGPTApp = AutomateUserDataFolder / 'ChatGPTApp'
    upload_logo = ChatGPTApp / '回形针上传文件.png'
    gray_up_arrow_logo = ChatGPTApp / '灰色上箭头.png'
    up_arrow_logo = ChatGPTApp / '上箭头.png'
    down_arrow_logo = ChatGPTApp / '下箭头.png'
    stop_logo = ChatGPTApp / '停止.png'
    send_message_logo = ChatGPTApp / '发送消息.png'
    white_x_logo = ChatGPTApp / '白叉.png'
    microphone_logo = ChatGPTApp / '话筒.png'
    headphone_logo = ChatGPTApp / '耳机.png'
    retry_logo = ChatGPTApp / '重试.png'
    reconnect_logo = ChatGPTApp / '重连.png'
    continue_logo = ChatGPTApp / '继续生成.png'

    make_dir(Comic)
    make_dir(Hanhua)

    make_dir(MomoHanhua)
    make_dir(Auto)
    make_dir(Log)
    make_dir(DataOutput)
    make_dir(ComicProcess)
    make_dir(MangaProcess)
    make_dir(ManhuaProcess)
    make_dir(ManhwaProcess)
    make_dir(Storage)

    make_dir(MomoYolo)
    make_dir(ChatGPT)

    date_str = strftime('%Y_%m_%d')
    log_path = Log / f'日志-{date_str}.log'
    logger.add(
        log_path.as_posix(),
        rotation='500MB',
        encoding='utf-8',
        enqueue=True,
        compression='zip',
        retention='10 days',
        # backtrace=True,
        # diagnose=True,
        # colorize=True,
        # format="<green>{time}</green> <level>{message}</level>",
    )

    CTD_onnx = Storage / 'comictextdetector.pt.onnx'
    CTD_model = None
    uoln = None

    # ================选择语言================
    # 不支持使用软件期间重选语言
    # 因为界面代码是手写的，没有 retranslateUi
    # lang_code = 'en_US'
    lang_code = 'zh_CN'
    # lang_code = 'zh_TW'
    # lang_code = 'ja_JP'
    qm_path = UserDataFolder / f'{APP_NAME}_{lang_code}.qm'

    do_qt_translate = False
    do_qt = False
    do_dev_folder = False
    do_dev_pic = False
    do_requirements = False
    do_structure = False
    do_roi = False

    mode_list = [
        'do_qt_translate',
        'do_qt',
        'do_dev_folder',
        'do_dev_pic',
        'do_requirements',
        'do_structure',
        'do_roi',
    ]


    def steps():
        pass


    read_local = True
    read_local = False

    read_history = True
    read_history = False

    history_names = []

    local_name = ''

    debug_dir = current_dir

    system_fonts = font_manager.findSystemFonts()
    system_font_dic = {}
    system_fonts = sorted(system_fonts)
    for i in range(len(system_fonts)):
        font = system_fonts[i]
        font_path = Path(font)
        if font_path.exists():
            system_font_dic[font_path.name] = font_path.as_posix()

    app_config = AppConfig.load()

    do_mode = app_config.config_data['do_mode']
    img_ind = app_config.config_data['img_ind']
    step_str = app_config.config_data['step_str']
    media_type = app_config.config_data['media_type']
    text_direction = app_config.config_data['text_direction']
    if text_direction is None:
        if media_type in ['Comic', 'Manhua', 'Manhwa']:
            text_direction = 'Horizontal'
        else:
            text_direction = 'Vertical'
    if text_direction == 'Horizontal':
        text_alignment = 'Center'
    else:  # text_direction == 'Vertical'
        text_alignment = 'Left'
    hit_enter = app_config.config_data['hit_enter']
    use_pivot_split = app_config.config_data['use_pivot_split']
    bbox_type = app_config.config_data['bbox_type']
    better_data = app_config.config_data['better_data']
    force_trad = app_config.config_data['force_trad']
    punc_type = app_config.config_data['punc_type']
    enable_lettering = app_config.config_data['enable_lettering']
    lettering_type = app_config.config_data['lettering_type']
    default_chara = app_config.config_data['default_chara']
    default_dpi = app_config.config_data['default_dpi']
    custom_dpi = app_config.config_data['custom_dpi']
    hide_extra = app_config.config_data['hide_extra']
    print_type = app_config.config_data['print_type']
    dist_thres = app_config.config_data['dist_thres']
    docx_img_format = app_config.config_data['docx_img_format']
    folder_name = app_config.config_data['folder_name']
    area_folder_names = app_config.config_data['area_folder_names']
    WLB_ext_px_dic = app_config.config_data['WLB_ext_px_dic']
    WLB_ext_px = WLB_ext_px_dic['Comic']
    if ',' in str(WLB_ext_px):
        ext_pxs = WLB_ext_px.split(',')
        ext_pxs = [int(x) for x in ext_pxs]
        word_ext_px, line_ext_px, block_ext_px = ext_pxs
    else:
        word_ext_px = line_ext_px = block_ext_px = int(WLB_ext_px)

    lettering = app_config.config_data['lettering']
    global_font_size = lettering['global_font_size']
    global_font_name = lettering['global_font_name']
    global_rec_font_name = lettering['global_rec_font_name']
    global_spacing = lettering['global_spacing']
    check_len = lettering['check_len']
    rec2italic = lettering['rec2italic']

    browser = app_config.config_data['browser']
    thumb_size = app_config.config_data['thumb_size']
    window_size = app_config.config_data['window_size']
    if ',' in str(window_size):
        window_w, par, window_h = window_size.partition(',')
        window_w = int(window_w)
        window_h = int(window_h)
    else:
        window_w = window_h = window_size
    WBLDN_padding = app_config.config_data['WBLDN_padding']
    if ',' in str(WBLDN_padding):
        paddings = WBLDN_padding.split(',')
        paddings = [int(x) for x in paddings]
        white_padding, black_padding, light_padding, dark_padding, normal_padding = paddings
    else:
        white_padding = black_padding = light_padding = dark_padding = normal_padding = WBLDN_padding
    font_size_range_dic = app_config.config_data['font_size_range_dic']
    bold_thres_dic = app_config.config_data['bold_thres_dic']
    stem_level = app_config.config_data['stem_level']
    thread_method = app_config.config_data['thread_method']
    pic_thread_method = app_config.config_data['pic_thread_method']
    if use_torch:
        pic_thread_method = 'queue'

    font_dic = app_config.config_data['font_dic']

    chatgpt = app_config.config_data['chatgpt']
    do_automate = chatgpt['do_automate']
    ask_mode = chatgpt['ask_mode']
    click_type = chatgpt['click_type']
    sleep_minute = chatgpt['sleep_minute']
    sleep_time = sleep_minute * 60
    gpt_line_max = chatgpt['gpt_line_max']
    gpt_char_max = chatgpt['gpt_char_max']

    long_pic = app_config.config_data['long_pic']
    single_width_min = long_pic['single_width_min']
    group_len = long_pic['group_len']

    frame_settings = app_config.config_data['frame_settings']
    frame_color = frame_settings['frame_color']
    min_frame_thickness = frame_settings['min_frame_thickness']
    WBN_tolerance = frame_settings['WBN_tolerance']
    if ',' in str(WBN_tolerance):
        tolerances = WBN_tolerance.split(',')
        tolerances = [int(x) for x in tolerances]
        white_tolerance, black_tolerance, normal_tolerance = tolerances
    else:
        white_tolerance = black_tolerance = normal_tolerance = WBN_tolerance
    edge_size = frame_settings['edge_size']
    edge_ratio = frame_settings['edge_ratio']
    grid_width_min = frame_settings['grid_width_min']
    grid_height_min = frame_settings['grid_height_min']
    max_white_ratio = frame_settings['max_white_ratio']
    do_add_frame = frame_settings['do_add_frame']
    frame_level = frame_settings['frame_level']
    frame_thres = frame_settings['frame_thres']
    min_outline_len = frame_settings['min_outline_len']
    gaps = frame_settings['gaps']
    if isinstance(gaps, int):
        gaps = [gaps] * 4
    check_2nd_color = frame_settings['check_2nd_color']
    check_more_frame_color = frame_settings['check_more_frame_color']
    y_1st = frame_settings['y_1st']
    fully_framed = frame_settings['fully_framed']
    optimize_inner = frame_settings['optimize_inner']
    black_pxs_thres = frame_settings['black_pxs_thres']
    merge_thres = frame_settings['merge_thres']
    merge_nearby = frame_settings['merge_nearby']

    bubble_condition = app_config.config_data['bubble_condition']
    area_range = bubble_condition['area_range']
    perimeter_range = bubble_condition['perimeter_range']
    thickness_range = bubble_condition['thickness_range']
    br_w_range = bubble_condition['br_w_range']
    br_h_range = bubble_condition['br_h_range']
    br_wh_range = bubble_condition['br_wh_range']
    br_wnh_range = bubble_condition['br_wnh_range']
    br_w_rarange = bubble_condition['br_w_rarange']
    br_h_rarange = bubble_condition['br_h_rarange']
    br_rarange = bubble_condition['br_rarange']
    br_w_prange = bubble_condition['br_w_prange']
    br_h_prange = bubble_condition['br_h_prange']
    portion_rarange = bubble_condition['portion_rarange']
    area_perimeter_rarange = bubble_condition['area_perimeter_rarange']
    bubble_px_range = bubble_condition['bubble_px_range']
    letter_px_range = bubble_condition['letter_px_range']
    bubble_px_rarange = bubble_condition['bubble_px_rarange']
    letter_px_rarange = bubble_condition['letter_px_rarange']
    BnL_px_rarange = bubble_condition['BnL_px_rarange']
    CTD_mask_px_rarange = bubble_condition['CTD_mask_px_rarange']
    CTD_thres = bubble_condition['CTD_thres']
    seg_thresh = bubble_condition['seg_thresh']
    box_thresh = bubble_condition['box_thresh']
    max_candidates = bubble_condition['max_candidates']
    unclip_ratio = bubble_condition['unclip_ratio']
    nms_thresh = bubble_condition['nms_thresh']
    conf_thresh = bubble_condition['conf_thresh']
    mask_thresh = bubble_condition['mask_thresh']
    edge_px_count_range = bubble_condition['edge_px_count_range']
    border_thickness_range = bubble_condition['border_thickness_range']
    letter_area_range = bubble_condition['letter_area_range']
    textblock_area_range = bubble_condition['textblock_area_range']
    textblock_wrange = bubble_condition['textblock_wrange']
    textblock_hrange = bubble_condition['textblock_hrange']
    textblock_w_prange = bubble_condition['textblock_w_prange']
    textblock_h_prange = bubble_condition['textblock_h_prange']
    letter_cnts_range = bubble_condition['letter_cnts_range']
    textblock_letters_range = bubble_condition['textblock_letters_range']
    note_area_range = bubble_condition['note_area_range']
    max_note_dist = bubble_condition['max_note_dist']
    intersect_ratio = bubble_condition['intersect_ratio']
    seg_line_range = bubble_condition['seg_line_range']
    angle_px_step = bubble_condition['angle_px_step']
    if ',' in str(angle_px_step):
        steps = angle_px_step.split(',')
        steps = [int(x) for x in steps]
        angle_step, px_step = steps
    else:
        angle_step = px_step = angle_px_step
    textline_hrange = bubble_condition['textline_hrange']
    textline_wrange = bubble_condition['textline_wrange']
    line_multi_h = bubble_condition['line_multi_h']
    line_multi_w = bubble_condition['line_multi_w']
    line_edge_range = bubble_condition['line_edge_range']
    dot_line_h = bubble_condition['dot_line_h']
    dot_line_w = bubble_condition['dot_line_w']
    dot_h = bubble_condition['dot_h']
    bulk_thres = bubble_condition['bulk_thres']
    overlap_ratio = bubble_condition['overlap_ratio']

    area_min, area_max = parse_range(area_range)
    perimeter_min, perimeter_max = parse_range(perimeter_range)
    thickness_min, thickness_max = parse_range(thickness_range)
    br_w_min, br_w_max = parse_range(br_w_range)
    br_h_min, br_h_max = parse_range(br_h_range)
    br_wh_min, br_wh_max = parse_range(br_wh_range)
    br_wnh_min, br_wnh_max = parse_range(br_wnh_range)

    br_w_ratio_min, br_w_ratio_max = parse_range(br_w_rarange)
    br_h_ratio_min, br_h_ratio_max = parse_range(br_h_rarange)
    br_ratio_min, br_ratio_max = parse_range(br_rarange)
    br_w_percent_min, br_w_percent_max = parse_range(br_w_prange)
    br_h_percent_min, br_h_percent_max = parse_range(br_h_prange)
    portion_ratio_min, portion_ratio_max = parse_range(portion_rarange)
    area_perimeter_ratio_min, area_perimeter_ratio_max = parse_range(area_perimeter_rarange)
    bubble_px_min, bubble_px_max = parse_range(bubble_px_range)
    letter_px_min, letter_px_max = parse_range(letter_px_range)
    bubble_px_ratio_min, bubble_px_ratio_max = parse_range(bubble_px_rarange)
    letter_px_ratio_min, letter_px_ratio_max = parse_range(letter_px_rarange)
    BnL_px_ratio_min, BnL_px_ratio_max = parse_range(BnL_px_rarange)
    CTD_mask_px_ratio_min, CTD_mask_px_ratio_max = parse_range(CTD_mask_px_rarange)
    edge_px_count_min, edge_px_count_max = parse_range(edge_px_count_range)
    border_thickness_min, border_thickness_max = parse_range(border_thickness_range)
    letter_area_min, letter_area_max = parse_range(letter_area_range)
    textblock_area_min, textblock_area_max = parse_range(textblock_area_range)
    textblock_wmin, textblock_wmax = parse_range(textblock_wrange)
    textblock_hmin, textblock_hmax = parse_range(textblock_hrange)
    textblock_w_percent_min, textblock_w_percent_max = parse_range(textblock_w_prange)
    textblock_h_percent_min, textblock_h_percent_max = parse_range(textblock_h_prange)
    letter_cnts_min, letter_cnts_max = parse_range(letter_cnts_range)
    textblock_letters_min, textblock_letters_max = parse_range(textblock_letters_range)
    note_area_min, note_area_max = parse_range(note_area_range)
    seg_line_min, seg_line_max = parse_range(seg_line_range)
    textline_h_min, textline_h_max = parse_range(textline_hrange)
    textline_w_min, textline_w_max = parse_range(textline_wrange)
    line_edge_min, line_edge_max = parse_range(line_edge_range)

    # 根据 DPI 调整所有相关的尺寸参数
    if custom_dpi != default_dpi:
        br_w_min = int(br_w_min * custom_dpi / default_dpi)
        br_w_max = int(br_w_max * custom_dpi / default_dpi)
        br_h_min = int(br_h_min * custom_dpi / default_dpi)
        br_h_max = int(br_h_max * custom_dpi / default_dpi)
        br_wh_min = int(br_wh_min * custom_dpi / default_dpi)
        br_wh_max = int(br_wh_max * custom_dpi / default_dpi)
        br_wnh_min = int(br_wnh_min * custom_dpi / default_dpi)
        br_wnh_max = int(br_wnh_max * custom_dpi / default_dpi)

        textblock_wmin = int(textblock_wmin * custom_dpi / default_dpi)
        textblock_wmax = int(textblock_wmax * custom_dpi / default_dpi)
        textblock_hmin = int(textblock_hmin * custom_dpi / default_dpi)
        textblock_hmax = int(textblock_hmax * custom_dpi / default_dpi)

        bubble_px_min = int(bubble_px_min * custom_dpi ** 2 / default_dpi ** 2)
        bubble_px_max = int(bubble_px_max * custom_dpi ** 2 / default_dpi ** 2)
        letter_px_min = int(letter_px_min * custom_dpi ** 2 / default_dpi ** 2)
        letter_px_max = int(letter_px_max * custom_dpi ** 2 / default_dpi ** 2)

    char_condition = app_config.config_data['char_condition']
    px_area_range = char_condition['px_area_range']
    cnt_area_range = char_condition['cnt_area_range']
    thick_range = char_condition['thick_range']
    brw_range = char_condition['brw_range']
    brh_range = char_condition['brh_range']

    px_area_min, px_area_max = parse_range(px_area_range)
    cnt_area_min, cnt_area_max = parse_range(cnt_area_range)
    thick_min, thick_max = parse_range(thick_range)
    brw_min, brw_max = parse_range(brw_range)
    brh_min, brh_max = parse_range(brh_range)

    bubble_recognize = app_config.config_data['bubble_recognize']
    max_hdiff = bubble_recognize['max_hdiff']
    max_font_size = bubble_recognize['max_font_size']
    kernel_depth = bubble_recognize['kernel_depth']
    block_ratio = bubble_recognize['block_ratio']
    bubble_alpha = bubble_recognize['bubble_alpha']
    bg_alpha = bubble_recognize['bg_alpha']
    WLB_alpha = bubble_recognize['WLB_alpha']
    if ',' in str(WLB_alpha):
        alphas = WLB_alpha.split(',')
        alphas = [int(x) / 100 for x in alphas]
        textword_alpha, textline_alpha, textblock_alpha = alphas
    else:
        textword_alpha = textline_alpha = textblock_alpha = int(WLB_alpha) / 100
    padding = bubble_recognize['padding']
    r_dot = bubble_recognize['r_dot']
    rec_pad_h = bubble_recognize['rec_pad_h']
    rec_pad_w = bubble_recognize['rec_pad_w']
    sum_sect_max = bubble_recognize['sum_sect_max']
    lower_ratio = bubble_recognize['lower_ratio']

    bubble_seg = app_config.config_data['bubble_seg']
    use_rec = bubble_seg['use_rec']
    use_dilate = bubble_seg['use_dilate']
    adapt_big_letter = bubble_seg['adapt_big_letter']
    check_note = bubble_seg['check_note']
    check_dots = bubble_seg['check_dots']
    check_intact = bubble_seg['check_intact']
    sort_by_y = bubble_seg['sort_by_y']
    sort_by_x = bubble_seg['sort_by_x']

    ocr_settings = app_config.config_data['ocr_settings']
    discard_ratio = ocr_settings['discard_ratio']
    y_thres = ocr_settings['y_thres']
    min_char_area = ocr_settings['min_char_area']
    min_cnt_area = ocr_settings['min_cnt_area']
    bit_thres = ocr_settings['bit_thres']
    white_thres = ocr_settings['white_thres']
    bi_thres = ocr_settings['bi_thres']
    ext_padding = ocr_settings['ext_padding']
    base_num = ocr_settings['base_num']
    cal_method = ocr_settings['cal_method']
    print_tables = ocr_settings['print_tables']
    init_ocr = ocr_settings['init_ocr']
    use_textwords = ocr_settings['use_textwords']
    line_spacing = ocr_settings['line_spacing']
    bubble_spacing = ocr_settings['bubble_spacing']
    merge_update = ocr_settings['merge_update']
    all_caps = ocr_settings['all_caps']
    do_cap = ocr_settings['do_cap']
    has_decoration = ocr_settings['has_decoration']
    renew_stitch = ocr_settings['renew_stitch']
    renew_CTD = ocr_settings['renew_CTD']
    CTD_cut = ocr_settings['CTD_cut']
    better_info = ocr_settings['better_info']
    check_bottom = ocr_settings['check_bottom']
    exclude_smaller = ocr_settings['exclude_smaller']
    exclude_inside = ocr_settings['exclude_inside']
    ellipsis_type = ocr_settings['ellipsis_type']
    use_sd = ocr_settings['use_sd']

    baidu_ocr = app_config.config_data['baidu_ocr']
    obd_app_id = baidu_ocr['APP_ID']
    obd_app_key = baidu_ocr['API_KEY']
    obd_secret_key = baidu_ocr['SECRET_KEY']
    aip_client = None
    if obd_app_id and obd_app_key and obd_secret_key:
        aip_client = AipOcr(obd_app_id, obd_app_key, obd_secret_key)

    custom_ocr_engines = app_config.config_data['custom_ocr_engines']
    prompt_prefix_dic = app_config.config_data['prompt_prefix_dic']
    grid_ratio_dic = app_config.config_data['grid_ratio_dic']
    italic_color_locates = app_config.config_data['italic_color_locates']
    color_pattern_dic = app_config.config_data['color_pattern_dic']

    if lettering_type == '332':
        global_font_size = 30
        global_font_name = '方正兰亭准黑'
        ellipsis_type = '...'
        punc_type = 'simple'
        font_size_range_dic = {
            'default': '30~30',
            'ffffff-000000': '30~30',
            '000000-6ac7cf': '35~35',
        }
    elif lettering_type == '不朽':
        global_font_size = 28
        global_font_name = '思源黑体'
        ellipsis_type = '...'
        font_size_range_dic = {
            'default': '28~28',
            'ffffff-000000': '28~28',
        }
    elif lettering_type == '永恒':
        global_font_name = '方正兰亭中黑'
        font_size_range_dic = {
            'default': '30~30',
            'ffffff-000000': '30~30',
            '9fce5a-000000': '35~35',
        }
    elif lettering_type == 'SD':
        global_font_size = 20
        font_size_range_dic = {
            'default': '20~20',
            'ffffff-000000': '20~20',
        }
        single_width_min = 1900
        custom_dpi = 96
    elif lettering_type == 'HD':
        global_font_size = 40
        font_size_range_dic = {
            'default': '40~40',
        }
        single_width_min = 3200
        custom_dpi = 200
    elif media_type in ['Manga']:
        punc_type = 'none'
        global_font_name = '華康儷中黑'
        global_font_size = 40
        font_size_range_dic = {
            'default': '40~60',
        }

        global_font_size = 80
        font_size_range_dic = {
            'default': '80~120',
        }

    font_size_min, font_size_max = parse_range(font_size_range_dic['default'])
    font_size_rarange = app_config.config_data['font_size_rarange']
    font_size_ratio_min, font_size_ratio_max = parse_range(font_size_rarange)

    font_meta_csv = UserDataFolder / f'字体信息_{processor_name}_{ram}GB.csv'
    font_head = ['字体文件名', '字体名', '字体PostScript名']
    font_meta_list = update_font_metadata(font_meta_csv, font_head)

    font_filename = font_dic.get(global_font_name)
    msyh_font_ttc = system_font_dic.get(font_filename)
    logger.debug(f'{msyh_font_ttc=}')
    msyh_font20 = ImageFont.truetype(msyh_font_ttc, 20)
    msyh_font30 = ImageFont.truetype(msyh_font_ttc, 30)
    msyh_font60 = ImageFont.truetype(msyh_font_ttc, 60)
    msyh_font100 = ImageFont.truetype(msyh_font_ttc, 100)

    for dmode in mode_list:
        globals()[dmode] = (do_mode == dmode)

    ProcessDir = MomoHanhua / f'{media_type}Process'
    img_folder = ProcessDir / folder_name

    auto_subdir = Auto / img_folder.name
    make_dir(auto_subdir)

    if media_type in WLB_ext_px_dic:
        WLB_ext_px = WLB_ext_px_dic[media_type]
        if ',' in str(WLB_ext_px):
            ext_pxs = WLB_ext_px.split(',')
            ext_pxs = [int(x) for x in ext_pxs]
            word_ext_px, line_ext_px, block_ext_px = ext_pxs
        else:
            word_ext_px = line_ext_px = block_ext_px = int(WLB_ext_px)

    media_lang = 'English'
    target_lang = 'zh-CN'

    if media_type == 'Comic':
        media_lang = 'English'
    elif media_type == 'Manga':
        media_lang = 'Japanese'
    elif media_type == 'Manhua':
        media_lang = 'Chinese Simplified'
        target_lang = 'en'
    elif media_type == 'Manhwa':
        media_lang = 'Korean'

    vert = False  # 设置为True以识别竖排文本
    if media_lang == 'Japanese':
        vert = True

    new_break = ''
    if media_lang in ['English']:
        new_break = ' '

    prompt_prefix = prompt_prefix_dic['Comic']
    if media_type in prompt_prefix_dic:
        prompt_prefix = prompt_prefix_dic[media_type]

    logger.warning(f'{media_type=}, {vert=}, {new_break=}')

    img_list = get_valid_imgs(img_folder)
    img_stems = [x.stem for x in img_list]
    cpre = common_prefix(img_stems)
    csuf = common_suffix(img_stems)
    all_masks = get_valid_imgs(img_folder, vmode='mask')

    # ================获取漫画名================
    series4file = folder_name
    m_issue_w_dot = p_issue_w_dot.match(folder_name)
    if m_issue_w_dot:
        series4file = m_issue_w_dot.group(1)
        issue = m_issue_w_dot.group(2)

    # ================气泡颜色模式================
    logger.debug(f'{series4file=}')
    color_patterns_raw = app_config.config_data.get('color_patterns', [])
    media_type_color_patterns = color_pattern_dic.get(media_type, [])
    series4file_color_patterns = color_pattern_dic.get(series4file, [])
    if color_patterns_raw is None:
        color_patterns_raw = []
    if media_type_color_patterns:
        color_patterns_raw.extend(media_type_color_patterns)
    if series4file_color_patterns:
        color_patterns_raw.extend(series4file_color_patterns)
    color_patterns_raw = reduce_list(color_patterns_raw)

    color_patterns = []
    for c in range(len(color_patterns_raw)):
        cp_str = color_patterns_raw[c]
        cp_bubble, par, cp_letter = cp_str.partition('~')
        if '=' in cp_bubble:
            cp1, par, cp2 = cp_bubble.partition('=')
            cp_bubble = [cp1, cp2]
        if '=' in cp_letter:
            cp3, par, cp4 = cp_letter.partition('=')
            cp_letter = [cp3, cp4]
        if cp_letter == '':
            cp_bubble, cp_letter = cp_letter, cp_bubble
        color_pattern = [cp_bubble, cp_letter]

        color_patterns.append(color_pattern)
    logger.debug(f'{color_patterns=}')

    if use_torch and CTD_onnx.exists():
        CTD_model = dnn.readNetFromONNX(CTD_onnx.as_posix())
        uoln = CTD_model.getUnconnectedOutLayersNames()

    if '3' in step_str or '9' in step_str:
        en_reader = None
        en_ocr = None
        if init_ocr:
            en_reader = Reader(['en'])
            en_ocr = PaddleOCR(use_gpu=False, lang='en')

    logger.warning(f'{do_mode=}, {thread_method=}, {pic_thread_method=}, {folder_name=}')
    if '3' in step_str:
        area_dic = get_area_dic(area_folder_names)
    else:
        area_dic = {}

    if python_ver in ['3.11.8']:
        mocr = MangaOcr()

    if do_qt:
        appgui = QApplication(sys.argv)
        translator = QTranslator()
        translator.load(str(qm_path))
        QApplication.instance().installTranslator(translator)
        appgui.installTranslator(translator)
        screen = QApplication.primaryScreen()
        if sys.platform == 'darwin':
            # 如果是 MacOS 系统
            scaling_factor = screen.devicePixelRatio()
        else:
            scaling_factor = 1
        scaling_factor_reci = 1 / scaling_factor
        mist_qt(appgui)
    elif do_dev_folder:
        logger.debug(f'{folder_name=}')
        logger.info(f'{step_str=}')
        img_inds = [
            # 12,
            # 14,
            # 16,
        ]
        folder_proc(img_folder, step_str, img_inds)
    elif do_dev_pic:
        img_file = img_list[img_ind]
        img_raw = imdecode(fromfile(img_file, dtype=uint8), -1)
        ih, iw = img_raw.shape[0:2]

        frame_yml = img_folder.parent / f'{img_folder.name}.yml'
        order_yml = img_folder.parent / f'{img_folder.name}-气泡排序.yml'
        ocr_yml = img_folder.parent / f'{img_folder.name}-文字识别.yml'
        ocr_docx = img_folder.parent / f'{img_folder.name}-1识别.docx'
        src_docx = img_folder.parent / f'{img_folder.name}-2校对.docx'
        read_docx = img_folder.parent / f'{img_folder.name}-3段落.docx'
        read_html = img_folder.parent / f'{img_folder.name}-3段落.html'
        dst_docx = img_folder.parent / f'{img_folder.name}-4翻译.docx'
        raw_html = img_folder.parent / f'{img_folder.name}-4原翻.html'
        dst_html = img_folder.parent / f'{img_folder.name}-4翻译.html'
        cut_txt = img_folder.parent / f'{img_folder.name}-5断句.txt'
        lettering_txt = img_folder.parent / f'{img_folder.name}-6填字.txt'
        mark_txt = img_folder.parent / f'{img_folder.name}-7标记.txt'
        lp_txt = img_folder.parent / f'{img_folder.name}-8萌翻.txt'
        rlp_txt = img_folder.parent / f'{img_folder.name}翻译_0.txt'
        googletrans_txt = img_folder.parent / f'{img_folder.name}-9-api-谷歌翻译.txt'

        img_stems = [x.stem for x in img_list]
        cpre = common_prefix(img_stems)
        csuf = common_suffix(img_stems)
        simple_stems = [x.removeprefix(cpre).removesuffix(csuf) for x in img_stems]

        logger.info(f'{step_str=}')

        simple_stem = img_file.stem.removeprefix(cpre).removesuffix(csuf)
        pic_ocr_docx = img_folder.parent / f'{img_folder.name}-1识别-{simple_stem}.docx'

        frame_data = iload_data(frame_yml)
        if '0' in step_str:
            img_file, frame_grid_strs = analyze1frame(img_file, frame_data, auto_subdir, media_type)
        frame_data = iload_data(frame_yml)
        order_data = iload_data(order_yml)
        ocr_data = iload_data(ocr_yml)
        if '1' in step_str:
            analyze1pic(img_file, frame_data, color_patterns, media_type, auto_subdir)
        if '2' in step_str or 'L' in step_str:
            appgui = QApplication(sys.argv)
            translator = QTranslator()
            translator.load(str(qm_path))
            QApplication.instance().installTranslator(translator)
            appgui.installTranslator(translator)
            screen = QApplication.primaryScreen()
            if sys.platform == 'darwin':
                # 如果是 MacOS 系统
                scaling_factor = screen.devicePixelRatio()
            else:
                scaling_factor = 1
            scaling_factor_reci = 1 / scaling_factor
            if '2' in step_str:
                order_qt(appgui)
            else:
                lp_qt(appgui)
        if '3' in step_str:
            pic_results = ocr1pic(img_file, frame_data, order_data, ocr_data, all_masks, media_type, media_lang, vert)
            ocr_doc = Document()
            ocr_doc, all_cropped_imgs = update_ocr_doc(ocr_doc, pic_results, img_folder, img_ind, img_list, media_type)
            write_docx(pic_ocr_docx, ocr_doc)
            logger.warning(f'{pic_ocr_docx=}')
